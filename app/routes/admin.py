from io import BytesIO
from types import SimpleNamespace
from flask import Blueprint, render_template, redirect, url_for, flash, request, jsonify, send_file, Response
from flask_login import login_required, current_user
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from openpyxl.utils import get_column_letter
from app.extensions import db
from app.models.user import User, Role, ROLE_DISPLAY
from app.models.unit import DonVi, LoaiDonVi
from app.models.personnel import QuanNhan, CapBac, HocHam, HocVi, DoiTuong
from app.models.certificate import ChungChi, LoaiChungChi
from app.models.nomination import DeXuat, DeXuatChiTiet, TrangThaiDeXuat, LoaiDanhHieu, DanhHieu, TieuChi, TrangThaiChiTiet
from app.models.evaluation import NhomTieuChi, DanhGiaHangNam, DiemQuyDinhDanhHieu
from app.models.approval import PheDuyet, PhongDuyet, KetQuaDuyet, KetQuaDuyetChiTiet
from app.models.reward import KhenThuong
from app.models.hoi_dong import HoiDongBieuQuyet, HOI_DONG_VAI_TRO, HOI_DONG_VAI_TRO_DISPLAY
from app.models.catalog import ChucVuOption, CapBacOption, DoiTuongOption
from app.models.notification import ThongBao
from app.utils.decorators import admin_required, admin_or_reward_viewer_required
from app.utils.file_upload import save_upload, delete_upload
from app.utils.activity_logger import log_action
from datetime import datetime
from html import escape
from sqlalchemy import case
from sqlalchemy.exc import ProgrammingError, OperationalError

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
from flask import send_file
from datetime import date

import hashlib, binascii, os
admin_bp = Blueprint('admin', __name__)

# The six reviewing departments (excluding admin)
DEPT_NAMES = [
    'Phòng Khoa học', 'Phòng Đào tạo',
    'Thủ trưởng Phòng TM-HC',
    'Ban Cán bộ', 'Ban Tổ chức', 'Ban Tuyên huấn', 'Ban Công tác quần chúng',
    'Ban Công nghệ thông tin', 'Ban Tác huấn', 'Ban Khảo thí', 'Ban Bảo vệ an ninh',
    'Ủy ban Kiểm tra', 'Ban Quân lực', 'Phòng Hậu cần - Kỹ thuật', 'Ban Sau đại học',
]

# Display order for approval columns in tracking screens
TRACKING_DEPT_COLUMNS = [
    {'key': 'Ban Cán bộ', 'label': 'Ban Cán bộ'},
    {'key': 'Ban Tổ chức', 'label': 'Ban Tổ chức'},
    {'key': 'Ban Tuyên huấn', 'label': 'Ban Tuyên huấn'},
    {'key': 'Ban Công tác quần chúng', 'label': 'Ban Công tác quần chúng'},
    {'key': 'Ban Bảo vệ an ninh', 'label': 'Ban Bảo vệ an ninh'},
    {'key': 'Ban Công nghệ thông tin', 'label': 'Ban Công nghệ thông tin'},
    {'key': 'Ban Tác huấn', 'label': 'Ban Tác huấn'},
    {'key': 'Ban Quân lực', 'label': 'Ban Quân lực'},
    {'key': 'Thủ trưởng Phòng TM-HC', 'label': 'TT phòng TM-HC'},
    {'key': 'Phòng Đào tạo', 'label': 'Phòng Đào tạo'},
    {'key': 'Phòng Khoa học', 'label': 'Phòng Khoa học quân sự'},
    {'key': 'Ban Khảo thí', 'label': 'Ban Khảo thí'},
    {'key': 'Ủy ban Kiểm tra', 'label': 'Ủy ban Kiểm tra'},
    {'key': 'Phòng Hậu cần - Kỹ thuật', 'label': 'Phòng Hậu cần - Kỹ thuật'},
    {'key': 'Ban Sau đại học', 'label': 'Ban Sau đại học'},
]

# Display order for approval columns in tracking screens
TRACKING_DEPT_COLUMNS = [
    {'key': 'Ban Cán bộ', 'label': 'Ban Cán bộ'},
    {'key': 'Ban Tổ chức', 'label': 'Ban Tổ chức'},
    {'key': 'Ban Tuyên huấn', 'label': 'Ban Tuyên huấn'},
    {'key': 'Ban Công tác quần chúng', 'label': 'Ban Công tác quần chúng'},
    {'key': 'Ban Bảo vệ an ninh', 'label': 'Ban Bảo vệ an ninh'},
    {'key': 'Ban Công nghệ thông tin', 'label': 'Ban Công nghệ thông tin'},
    {'key': 'Ban Tác huấn', 'label': 'Ban Tác huấn'},
    {'key': 'Ban Quân lực', 'label': 'Ban Quân lực'},
    {'key': 'Thủ trưởng Phòng TM-HC', 'label': 'TT phòng TM-HC'},
    {'key': 'Phòng Đào tạo', 'label': 'Phòng Đào tạo'},
    {'key': 'Phòng Khoa học', 'label': 'Phòng Khoa học quân sự'},
    {'key': 'Ban Khảo thí', 'label': 'Ban Khảo thí'},
    {'key': 'Ủy ban Kiểm tra', 'label': 'Ủy ban Kiểm tra'},
    {'key': 'Phòng Hậu cần - Kỹ thuật', 'label': 'Phòng Hậu cần - Kỹ thuật'},
    {'key': 'Ban Sau đại học', 'label': 'Ban Sau đại học'},
]


def _is_auto_scope_approved(dept_name, doi_tuong):
    if dept_name == 'Ban Quân lực':
        return doi_tuong not in BAN_QUANLUC_DOI_TUONG
    if dept_name == 'Ban Cán bộ':
        return doi_tuong in BAN_QUANLUC_DOI_TUONG
    return False


# Gate columns each Thủ trưởng role waits for
_TTR_GATE_COLUMNS = {
    'Thủ trưởng Phòng TM-HC': [
        'Ban Công nghệ thông tin', 'Ban Tác huấn', 'Ban Quân lực',
    ],
}


def _is_thu_truong_gate_all_approved(de_xuat_id, ct, dept_name):
    """For Thủ trưởng roles: auto-pass if all gate sub-departments approved this individual."""
    gate_cols = _TTR_GATE_COLUMNS.get(dept_name)
    if not gate_cols:
        return False
    for gate_dept in gate_cols:
        if not _is_individual_dept_approved(de_xuat_id, ct, gate_dept):
            return False
    return True


def _is_individual_dept_approved(de_xuat_id, ct, dept_name):
    """Check one individual's approval for one department, honoring auto-scope rules."""
    is_auto = _is_auto_scope_approved(dept_name, ct.doi_tuong)
    if is_auto:
        return True

    # Thủ trưởng roles: approved if all their gate sub-depts approved
    if dept_name in _TTR_GATE_COLUMNS:
        pd = PheDuyet.query.filter_by(de_xuat_id=de_xuat_id, phong_duyet=dept_name).first()
        if pd and pd.ket_qua == KetQuaDuyet.DONG_Y.value:
            return True
        # Check gate columns
        return _is_thu_truong_gate_all_approved(de_xuat_id, ct, dept_name)

    pd = PheDuyet.query.filter_by(de_xuat_id=de_xuat_id, phong_duyet=dept_name).first()

    if not pd:
        return False

    if pd.ket_qua == KetQuaDuyet.TU_CHOI.value:
        return False

    # Nếu phòng đã đồng ý ở cấp đề xuất (pd-level) → tính là đồng ý cho mọi cá nhân
    if pd.ket_qua == KetQuaDuyet.DONG_Y.value:
        return True

    kq = KetQuaDuyetChiTiet.query.filter_by(phe_duyet_id=pd.id, chi_tiet_id=ct.id).first()
    if not kq:
        return False

    return kq.ket_qua == KetQuaDuyet.DONG_Y.value

# Đối tượng thuộc diện Ban Quân lực quản lý
BAN_QUANLUC_DOI_TUONG = ['Công nhân viên', 'Quân nhân chuyên nghiệp','Hạ sĩ quan chiến sĩ']

# All criteria field labels (for detailed view)
ALL_FIELD_LABELS = {
    'muc_do_hoan_thanh': 'Hoàn thành NV',
    'phieu_tin_nhiem': 'Phiếu tín nhiệm',
    'kiem_tra_dieu_lenh': 'KT Điều lệnh',
    'ban_sung': 'Bắn súng',
    'the_luc': 'Thể lực',
    'kiem_tra_chinh_tri': 'KT Chính trị',
    'kiem_tra_tin_hoc': 'Kỹ năng số',
    'dia_ly_quan_su': 'ĐHQS',
    'xep_loai_dang_vien': 'XL đảng viên',
    'ket_qua_doan_the': 'KQ đoàn thể',
    'xep_loai_doan_vien': 'XL đoàn viên',
    'hinh_thuc_khen_thuong_qc': 'KT quần chúng',
    'ket_qua_phu_nu': 'XL phụ nữ',
    'hinh_thuc_khen_thuong_pn': 'KT phụ nữ',
    'chu_tri_don_vi_danh_hieu': 'Chủ trì ĐV',
    'danh_hieu_gv_gioi': 'DH GV giỏi',
    'dinh_muc_giang_day': 'Định mức GD',
    'ket_qua_kiem_tra_giang': 'KT giảng',
    'thoi_gian_lao_dong_kh': 'LĐ khoa học',
    'tien_do_pgs': 'Tiến độ PGS',
    'danh_hieu_hv_gioi': 'DH HV giỏi',
    'diem_tong_ket': 'Điểm TK',
    'ket_qua_thuc_hanh': 'Thực hành',
    'ket_qua_ren_luyen': 'KQ rèn luyện',
    'hinh_thuc_tot_nghiep': 'HT thi TN',
    'diem_tn_ctd': 'Điểm CTĐ (TN)',
    'diem_tn_ct': 'Điểm CT (TN)',
    'diem_tn_ta': 'Điểm TA (TN)',
    'diem_tn_mon4': 'Điểm môn 4 (TN)',
    'diem_tn_chuyennganh': 'Điểm CN (TN)',
    'diem_tn_baove': 'Điểm BV KL (TN)',
    'ket_qua_doan_the': 'KQ Đoàn thể',
    'xep_loai_doan_vien': 'Xếp loại đoàn viên',
    'hinh_thuc_khen_thuong_qc': 'KT phong trào QC',
    'ket_qua_phu_nu': 'Xếp loại phụ nữ',
    'hinh_thuc_khen_thuong_pn': 'KT hội phụ nữ',
    'chu_tri_don_vi_danh_hieu': 'Chủ trì ĐV',
    'diem_nckh': 'Điểm NCKH',
    'nckh_noi_dung': 'ND NCKH',
    'nckh_minh_chung': 'MC NCKH',
    'mo_ta_khoa_hoc': 'Mô tả thành tích KH',
    'diem_tot_nghiep': 'Điểm tốt nghiệp',
    'minh_chung_thanh_tich_khac': 'MC thành tích khác',
    'thanh_tich_ca_nhan_khac': 'Thành tích khác',
}

# Criteria fields in display order
ALL_FIELDS = [
    'muc_do_hoan_thanh', 'phieu_tin_nhiem',
    'kiem_tra_chinh_tri', 'kiem_tra_dieu_lenh', 'kiem_tra_tin_hoc',
    'dia_ly_quan_su', 'ban_sung', 'the_luc',
    'xep_loai_dang_vien',
    'ket_qua_doan_the', 'xep_loai_doan_vien',
    'hinh_thuc_khen_thuong_qc', 'ket_qua_phu_nu', 'hinh_thuc_khen_thuong_pn',
    'chu_tri_don_vi_danh_hieu',
    'danh_hieu_gv_gioi', 'dinh_muc_giang_day', 'ket_qua_kiem_tra_giang',
    'tien_do_pgs', 'thoi_gian_lao_dong_kh',
    'danh_hieu_hv_gioi', 'diem_tong_ket', 'ket_qua_thuc_hanh', 'ket_qua_ren_luyen',
    'hinh_thuc_tot_nghiep',
    'diem_tn_ctd', 'diem_tn_ct', 'diem_tn_ta', 'diem_tn_mon4',
    'diem_tn_chuyennganh', 'diem_tn_baove',
    'diem_nckh', 'nckh_noi_dung', 'nckh_minh_chung', 'mo_ta_khoa_hoc',
    'diem_tot_nghiep', 'minh_chung_thanh_tich_khac',
    'thanh_tich_ca_nhan_khac',
]

# Score fields that can have minimum score requirements
DIEM_FIELDS = [
    'diem_kiem_tra_tin_hoc',
    'diem_kiem_tra_dieu_lenh',
    'diem_dia_ly_quan_su',
    'diem_ban_sung',
    'diem_the_luc',
    'diem_kiem_tra_chinh_tri',
    'diem_tong_ket',
    'diem_nckh',
    'diem_tot_nghiep',
    'diem_tn_ctd',
    'diem_tn_ct',
    'diem_tn_ta',
    'diem_tn_mon4',
    'diem_tn_chuyennganh',
    'diem_tn_baove',
]

DIEM_FIELD_LABELS = {
    'diem_kiem_tra_tin_hoc': 'Điểm kỹ năng số',
    'diem_kiem_tra_dieu_lenh': 'Điểm điều lệnh',
    'diem_dia_ly_quan_su': 'Điểm địa hình quân sự',
    'diem_ban_sung': 'Điểm bắn súng',
    'diem_the_luc': 'Điểm thể lực',
    'diem_kiem_tra_chinh_tri': 'Điểm chính trị',
    'diem_tong_ket': 'Điểm tổng kết',
    'diem_nckh': 'Điểm NCKH',
    'diem_tot_nghiep': 'Điểm tốt nghiệp',
    'diem_tn_ctd': 'Điểm CTĐ',
    'diem_tn_ct': 'Điểm CT',
    'diem_tn_ta': 'Điểm Tiếng Anh',
    'diem_tn_mon4': 'Điểm môn 4',
    'diem_tn_chuyennganh': 'Điểm chuyên ngành',
    'diem_tn_baove': 'Điểm bảo vệ',
}


def _get_doi_tuong_option_list():
    rows = DoiTuongOption.query.filter_by(is_active=True).order_by(DoiTuongOption.thu_tu, DoiTuongOption.ten).all()
    return [x.ten for x in rows] or [e.value for e in DoiTuong]


from collections import defaultdict
from sqlalchemy.orm import joinedload, subqueryload, contains_eager
from sqlalchemy import case, func

@admin_bp.route('/tracking')
@login_required
@admin_required
def approval_tracking():
    status_filter    = request.args.get('status', '')
    unit_filter      = request.args.get('unit', '', type=str)
    danh_hieu_filter = request.args.get('danh_hieu', '')
    search_query     = request.args.get('q', '').strip()
    scope_filter     = request.args.get('scope', '')
    view_mode        = request.args.get('view', 'compact')
    nam_hoc_filter   = request.args.get('nam_hoc', '')

    # ════════════════════════════════════════════════════════
    # 1. QUERY ĐỀ XUẤT — eager load
    # ════════════════════════════════════════════════════════
    query = DeXuat.query.filter(
        DeXuat.trang_thai != TrangThaiDeXuat.NHAP.value,
        DeXuat.trang_thai != TrangThaiDeXuat.PHE_DUYET_CUOI.value,
    ).options(
        joinedload(DeXuat.don_vi),
        subqueryload(DeXuat.chi_tiets).joinedload(DeXuatChiTiet.quan_nhan),
        subqueryload(DeXuat.chi_tiets).subqueryload(DeXuatChiTiet.minh_chungs),
        subqueryload(DeXuat.phe_duyets).subqueryload(PheDuyet.chi_tiet_duyet),
    )

    if nam_hoc_filter:
        query = query.filter(DeXuat.nam_hoc == nam_hoc_filter)
    if status_filter:
        query = query.filter(DeXuat.trang_thai == status_filter)

    query = query.join(DonVi, DeXuat.don_vi_id == DonVi.id)
    if unit_filter:
        query = query.filter(DonVi.ten_don_vi == unit_filter)

    query = query.order_by(
        DonVi.thu_tu.asc(),
        DeXuat.nam_hoc.desc()
       
    )
    nominations = query.all()

    # ════════════════════════════════════════════════════════
    # 2. STATS — 1 query GROUP BY
    # ════════════════════════════════════════════════════════
    stats_rows = db.session.query(
        DeXuat.trang_thai,
        func.count(DeXuat.id).label('cnt')
    ).filter(
        DeXuat.trang_thai != TrangThaiDeXuat.NHAP.value
    ).group_by(DeXuat.trang_thai).all()

    stats_map = {row.trang_thai: row.cnt for row in stats_rows}
    stats = {
        'total':          sum(stats_map.values()),
        'pending':        stats_map.get(TrangThaiDeXuat.CHO_DUYET.value, 0),
        'reviewing':      stats_map.get(TrangThaiDeXuat.DANG_DUYET.value, 0),
        'dept_approved':  stats_map.get(TrangThaiDeXuat.HOI_DONG.value, 0),
        'final_approved': stats_map.get(TrangThaiDeXuat.PHE_DUYET_CUOI.value, 0),
        'rejected':       stats_map.get(TrangThaiDeXuat.TU_CHOI.value, 0),
    }

    # ════════════════════════════════════════════════════════
    # 3. APPROVED CT IDS
    # ════════════════════════════════════════════════════════
    approved_ct_ids = set(
        row[0] for row in db.session.query(KhenThuong.chi_tiet_id).all()
    )

    # ════════════════════════════════════════════════════════
    # 4. NAM HOC LIST
    # ════════════════════════════════════════════════════════
    nam_hoc_list = [
        n[0] for n in db.session.query(DeXuat.nam_hoc)
        .filter(DeXuat.trang_thai != TrangThaiDeXuat.NHAP.value)
        .distinct().order_by(DeXuat.nam_hoc.desc()).all()
    ]

    # ════════════════════════════════════════════════════════
    # 5. UNIT NAMES
    # ════════════════════════════════════════════════════════
    unit_names = [
        u[0] for u in db.session.query(DonVi.ten_don_vi, DonVi.thu_tu)
        .join(DeXuat, DeXuat.don_vi_id == DonVi.id)
        .filter(DeXuat.trang_thai != TrangThaiDeXuat.NHAP.value)
        .distinct()
        .order_by(DonVi.thu_tu.asc(), DonVi.ten_don_vi.asc()).all()
    ]
    
    # ════════════════════════════════════════════════════════
    # 6. PRE-BUILD TTR GATE MAP — bulk load 1 lần, tra cứu O(1)
    #    Tránh _is_thu_truong_gate_all_approved() query trong loop
    # ════════════════════════════════════════════════════════
    ttr_gate_map = {}

    if _TTR_GATE_COLUMNS:
        all_ct_ids_for_gate = [
            ct.id
            for dx in nominations
            for ct in dx.chi_tiets
            if not ct.bi_loai and ct.id not in approved_ct_ids
        ]

        if all_ct_ids_for_gate:
            from app.models import KetQuaDuyetChiTiet as _KQ

            # ★ Join PheDuyet để lấy phong_duyet
            gate_rows = (
                db.session.query(
                    _KQ.chi_tiet_id,
                    PheDuyet.phong_duyet,
                    _KQ.ket_qua,
                )
                .join(PheDuyet, _KQ.phe_duyet_id == PheDuyet.id)
                .filter(
                    _KQ.chi_tiet_id.in_(all_ct_ids_for_gate),
                    PheDuyet.phong_duyet.in_(list(_TTR_GATE_COLUMNS)),
                )
                .all()
            )

            from collections import defaultdict
            gate_raw = defaultdict(list)
            for row in gate_rows:
                gate_raw[(row.chi_tiet_id, row.phong_duyet)].append(row.ket_qua)

            for (ct_id, dept_name), results in gate_raw.items():
                ttr_gate_map[(ct_id, dept_name)] = (
                    len(results) > 0 and
                    all(r == KetQuaDuyet.DONG_Y.value for r in results)
                )

    # ════════════════════════════════════════════════════════
    # 7. PROCESS NOMINATIONS
    # ════════════════════════════════════════════════════════
    tracking_dept_names = [c['key'] for c in TRACKING_DEPT_COLUMNS]
    unit_groups_dict    = {}
    total_individuals   = 0

    for dx in nominations:
        unit_name  = dx.don_vi.ten_don_vi
        dept_lookup = {}
        for pd in dx.phe_duyets:
            dept_lookup[pd.phong_duyet] = {
                'phe_duyet': pd,
                'items': {kq.chi_tiet_id: kq for kq in pd.chi_tiet_duyet},
            }

        chi_tiets_data    = []
        tap_the_data_list = []

        for ct in dx.chi_tiets:
            is_tap_the = ct.ten_don_vi_de_xuat is not None or ct.quan_nhan_id is None
            # Chỉ ẩn khi bị từ chối dứt điểm bởi Hội đồng/Tuyên huấn (không ẩn nếu chỉ 1 phòng ban từ chối)
            if ct.bi_loai and ct.phong_loai == "Tuyên huấn" and ct.trang_thai == TrangThaiChiTiet.TU_CHOI.value:
                continue
            if danh_hieu_filter and ct.loai_danh_hieu != danh_hieu_filter: continue

            if not is_tap_the:
                if scope_filter == 'quan_luc' and ct.doi_tuong not in BAN_QUANLUC_DOI_TUONG: continue
                if scope_filter == 'can_bo'   and ct.doi_tuong in BAN_QUANLUC_DOI_TUONG:     continue

            if search_query:
                name_val = (ct.quan_nhan.ho_ten if ct.quan_nhan else '') or (ct.ten_don_vi_de_xuat or '')
                if search_query.lower() not in name_val.lower():   continue

            # ── Build dept results (0 queries) ──────────────────
            ct_dept_results = {}
            all_dept_ok     = True

            for dept_name in tracking_dept_names:
                dept_data      = dept_lookup.get(dept_name)
                is_auto        = _is_auto_scope_approved(dept_name, ct.doi_tuong)
                # ★ Tra cứu O(1) từ ttr_gate_map thay vì gọi hàm query
                is_ttr_gate_ok = ttr_gate_map.get((ct.id, dept_name), False) \
                                 if dept_name in _TTR_GATE_COLUMNS else False

                if dept_data:
                    kq = dept_data['items'].get(ct.id)
                    if dept_data['phe_duyet'].ket_qua == KetQuaDuyet.DONG_Y.value:
                        ct_dept_results[dept_name] = {'ket_qua': KetQuaDuyet.DONG_Y.value, 'auto': is_auto}
                    elif kq:
                        ct_dept_results[dept_name] = {'ket_qua': kq.ket_qua, 'auto': is_auto}
                    elif is_auto or is_ttr_gate_ok:
                        ct_dept_results[dept_name] = {
                            'ket_qua': KetQuaDuyet.DONG_Y.value,
                            'auto': is_auto and not is_ttr_gate_ok,
                        }
                    else:
                        ct_dept_results[dept_name] = None
                else:
                    if is_auto or is_ttr_gate_ok:
                        ct_dept_results[dept_name] = {
                            'ket_qua': KetQuaDuyet.DONG_Y.value,
                            'auto': is_auto and not is_ttr_gate_ok,
                        }
                    else:
                        ct_dept_results[dept_name] = None

                if dept_name in DEPT_NAMES:
                    dept_ok = (
                        ct_dept_results[dept_name] is not None and
                        ct_dept_results[dept_name].get('ket_qua') == KetQuaDuyet.DONG_Y.value
                    )
                    if not dept_ok:
                        all_dept_ok = False

            ct_can_approve = all_dept_ok and not ct.admin_approved

            # ★ BUILD fields_dict — thay vì Jinja dùng ct|attr(field) chậm
            # getattr() ở Python nhanh hơn Jinja attr filter ~10x
            fields_dict = {}
            for field in ALL_FIELDS:
                val = getattr(ct, field, None)
                if val is not None and val != '' and val != 0:
                    fields_dict[field] = str(val)

            ct_entry = {
                'ct':               ct,
                'dept_results':     ct_dept_results,
                'can_final_approve': ct_can_approve,
                'dx_id':            dx.id,
                'fields_dict':      fields_dict,   # ★ THÊM
            }

            if is_tap_the:
                tap_the_data_list.append(ct_entry)
            else:
                chi_tiets_data.append(ct_entry)
                total_individuals += 1

        if chi_tiets_data or tap_the_data_list:
            if unit_name not in unit_groups_dict:
                unit_groups_dict[unit_name] = []
            unit_groups_dict[unit_name].append({
                'dx':               dx,
                'chi_tiets':        chi_tiets_data,
                'tap_the_chi_tiets': tap_the_data_list,
            })

    # ════════════════════════════════════════════════════════
    # 8. BUILD UNIT GROUPS + FLAT LISTS
    # ════════════════════════════════════════════════════════
    unit_groups  = []
    flat_ca_nhan = []
    flat_tap_the = []

    for unit_name in sorted(unit_groups_dict.keys()):
        nom_list = unit_groups_dict[unit_name]
        if not nom_list:
            continue

        individual_count = 0
        tap_the_count    = 0

        for nom_data in nom_list:
            dx          = nom_data['dx']
            don_vi_name = dx.don_vi.ten_don_vi if dx.don_vi else ''

            for ct_entry in nom_data['chi_tiets']:
                    flat_ca_nhan.append({
                    **ct_entry,
                    'don_vi':        don_vi_name,
                    'loai_danh_hieu': ct_entry['ct'].loai_danh_hieu or '',
                    'dx':            dx,
                })
            individual_count += 1

            for ct_entry in nom_data['tap_the_chi_tiets']:
                flat_tap_the.append({
                    **ct_entry,
                    'don_vi':        don_vi_name,
                    'loai_danh_hieu': ct_entry['ct'].loai_danh_hieu or '',
                    'dx':            dx,
                })
                tap_the_count += 1

        unit_groups.append({
            'unit_name':       unit_name,
            'nominations':     nom_list,
            'individual_count': individual_count,
            'tap_the_count':   tap_the_count,
        })

   

    # ════════════════════════════════════════════════════════
    # 9. TT CRITERIA — chỉ query khi có flat_tap_the
    # ════════════════════════════════════════════════════════
    tt_criteria_fields = []
    tt_field_labels    = {}

    if flat_tap_the:
        from app.models.nomination import TieuChi as _TieuChi, DanhHieu as _DanhHieu

        tt_all_keys = set()
        for dh in _DanhHieu.query.filter_by(pham_vi='Đơn vị', is_active=True).all():
            for ma_truong in (dh.tieu_chi or []):
                tt_all_keys.add(ma_truong)

        for item in flat_tap_the:
            tt_all_keys.update((item['ct'].tap_the_dict or {}).keys())

        if tt_all_keys:
            tt_tieu_chi_rows = _TieuChi.query.filter(
                _TieuChi.ma_truong.in_(list(tt_all_keys)),
                _TieuChi.is_active == True
            ).order_by(_TieuChi.thu_tu, _TieuChi.ten).all()

            tt_criteria_fields = [tc.ma_truong for tc in tt_tieu_chi_rows]
            tt_field_labels    = {tc.ma_truong: tc.ten for tc in tt_tieu_chi_rows}
            for k in tt_all_keys:
                if k not in tt_field_labels:
                    tt_criteria_fields.append(k)
                    tt_field_labels[k] = k

    # ════════════════════════════════════════════════════════
    # 10. MISC
    # ════════════════════════════════════════════════════════
    status_list    = [e.value for e in TrangThaiDeXuat if e != TrangThaiDeXuat.NHAP]
    danh_hieu_list = [e.value for e in LoaiDanhHieu]
    # 1. Lấy danh sách unit_names đã được sắp xếp chuẩn theo thu_tu và ten_don_vi
    unit_names = [
        u.ten_don_vi for u in
        DonVi.query
        .filter(
          
            DonVi.is_active == True
        )
        .order_by(DonVi.thu_tu.asc(), DonVi.ten_don_vi.asc())
        .all()
    ]

# 2. Tạo một dictionary map vị trí (index) của từng đơn vị
# Đơn vị nào đứng trước trong unit_names sẽ có index nhỏ hơn
    unit_order_map = {name: index for index, name in enumerate(unit_names)}

    # 3. Apply vào lambda sort
    # Nếu có đơn vị cũ (is_active=False) không nằm trong map, nó sẽ nhận giá trị 9999 và bị đẩy xuống cuối
    flat_ca_nhan.sort(key=lambda x: (
        unit_order_map.get(x['don_vi'], 9999),  # Ưu tiên 1: Sắp xếp theo vị trí thu_tu
        x['don_vi'],                            # Ưu tiên 2: Cùng thu_tu thì sắp theo Tên đơn vị
        x['loai_danh_hieu']                     # Ưu tiên 3: Sắp xếp theo Danh hiệu
    ))

    flat_tap_the.sort(key=lambda x: (
        unit_order_map.get(x['don_vi'], 9999), 
        x['don_vi'], 
        x['loai_danh_hieu']
    ))
    return render_template(
        'admin/tracking.html',
        unit_groups=unit_groups,
        flat_ca_nhan=flat_ca_nhan,
        flat_tap_the=flat_tap_the,
        status_filter=status_filter,
        unit_filter=unit_filter,
        danh_hieu_filter=danh_hieu_filter,
        search_query=search_query,
        scope_filter=scope_filter,
        view_mode=view_mode,
        status_list=status_list,
        unit_names=unit_names,
        nam_hoc_list=nam_hoc_list,
        nam_hoc_filter=nam_hoc_filter,
        danh_hieu_list=danh_hieu_list,
        stats=stats,
        dept_names=DEPT_NAMES,
        tracking_dept_columns=TRACKING_DEPT_COLUMNS,
        total_individuals=total_individuals,
        all_field_labels=ALL_FIELD_LABELS,
        all_fields=ALL_FIELDS,
        tt_criteria_fields=tt_criteria_fields,
        tt_field_labels=tt_field_labels,
    )

@admin_bp.route('/tracking/chi-tiet/<int:ct_id>')
@login_required
@admin_required
def tracking_detail(ct_id):
    """View detailed info for ONE individual (DeXuatChiTiet)."""
    ct = DeXuatChiTiet.query.get_or_404(ct_id)
    de_xuat = ct.de_xuat

    # Build per-department results for this individual
    dept_item_results = {}
    for pd in de_xuat.phe_duyets:
        kq = KetQuaDuyetChiTiet.query.filter_by(
            phe_duyet_id=pd.id, chi_tiet_id=ct.id
        ).first()
        is_auto = _is_auto_scope_approved(pd.phong_duyet, ct.doi_tuong)
        is_ttr_gate_ok = _is_thu_truong_gate_all_approved(de_xuat.id, ct, pd.phong_duyet) if pd.phong_duyet in _TTR_GATE_COLUMNS else False
        if not kq and (is_auto or is_ttr_gate_ok):
            kq = SimpleNamespace(ket_qua=KetQuaDuyet.DONG_Y.value, ly_do='Tự động duyệt theo phạm vi')
        dept_item_results[pd.phong_duyet] = {
            'phe_duyet': pd,
            'item_result': kq,
            'is_auto': is_auto or is_ttr_gate_ok,
        }

    # Check if all departments approved THIS individual
    all_dept_ok = True
    for dept_name in DEPT_NAMES:
        approved = _is_individual_dept_approved(de_xuat.id, ct, dept_name)
        if not approved:
            all_dept_ok = False
            break

    # Check if already admin_approved (in Bảng 2) or in KhenThuong
    already_approved = ct.admin_approved or KhenThuong.query.filter_by(chi_tiet_id=ct.id).first() is not None

    can_final_approve = all_dept_ok and not already_approved

    # All criteria field labels for the detail view
    all_field_labels = {
        'muc_do_hoan_thanh': 'Hoàn thành nhiệm vụ',
        'phieu_tin_nhiem': 'Phiếu tín nhiệm',
        'kiem_tra_dieu_lenh': 'Kiểm tra điều lệnh',
        'ban_sung': 'Bắn súng',
        'the_luc': 'Thể lực',
        'kiem_tra_chinh_tri': 'Kiểm tra chính trị',
        'kiem_tra_tin_hoc': 'Kỹ năng số',
        'dia_ly_quan_su': 'Địa hình quân sự',
        'danh_hieu_gv_gioi': 'Danh hiệu GV giỏi',
        'dinh_muc_giang_day': 'Định mức giảng dạy',
        'ket_qua_kiem_tra_giang': 'Kết quả kiểm tra giảng',
        'thoi_gian_lao_dong_kh': 'Thời gian LĐ khoa học',
        'tien_do_pgs': 'Tiến độ PGS',
        'danh_hieu_hv_gioi': 'Danh hiệu HV giỏi',
        'diem_tong_ket': 'Điểm tổng kết',
        'ket_qua_thuc_hanh': 'Kết quả thực hành',
        'ket_qua_ren_luyen': 'Kết quả rèn luyện',
        'hinh_thuc_tot_nghiep': 'Hình thức thi tốt nghiệp',
        'diem_tn_ctd': 'Điểm CTĐ (tốt nghiệp)',
        'diem_tn_ct': 'Điểm CT (tốt nghiệp)',
        'diem_tn_ta': 'Điểm TA (tốt nghiệp)',
        'diem_tn_mon4': 'Điểm môn thứ 4 (tốt nghiệp)',
        'diem_tn_chuyennganh': 'Điểm chuyên ngành (tốt nghiệp)',
        'diem_tn_baove': 'Điểm bảo vệ KL (tốt nghiệp)',
        'ket_qua_doan_the': 'Kết quả đoàn thể',
        'chu_tri_don_vi_danh_hieu': 'Chủ trì ĐV danh hiệu',
        'diem_nckh': 'Điểm NCKH',
        'nckh_noi_dung': 'Nội dung NCKH',
        'nckh_minh_chung': 'Minh chứng NCKH',
        'mo_ta_khoa_hoc': 'Mô tả thành tích KH',
        'diem_tot_nghiep': 'Điểm tốt nghiệp (TB)',
        'minh_chung_thanh_tich_khac': 'MC thành tích khác',
        'thanh_tich_ca_nhan_khac': 'Thành tích cá nhân khác',
    }

    # All field names in display order
    all_fields = [
        'muc_do_hoan_thanh', 'phieu_tin_nhiem',
        'kiem_tra_chinh_tri', 'kiem_tra_dieu_lenh', 'kiem_tra_tin_hoc',
        'dia_ly_quan_su', 'ban_sung', 'the_luc',
        'xep_loai_dang_vien',
        'ket_qua_doan_the', 'xep_loai_doan_vien',
        'hinh_thuc_khen_thuong_qc', 'ket_qua_phu_nu', 'hinh_thuc_khen_thuong_pn',
        'chu_tri_don_vi_danh_hieu',
        'danh_hieu_gv_gioi', 'dinh_muc_giang_day', 'ket_qua_kiem_tra_giang',
        'tien_do_pgs', 'thoi_gian_lao_dong_kh',
        'danh_hieu_hv_gioi', 'diem_tong_ket', 'ket_qua_thuc_hanh', 'ket_qua_ren_luyen',
        'hinh_thuc_tot_nghiep',
        'diem_tn_ctd', 'diem_tn_ct', 'diem_tn_ta', 'diem_tn_mon4',
        'diem_tn_chuyennganh', 'diem_tn_baove',
        'diem_nckh', 'nckh_noi_dung',
        'mo_ta_khoa_hoc', 'diem_tot_nghiep', 'minh_chung_thanh_tich_khac',
        'thanh_tich_ca_nhan_khac',
    ]

    return render_template('admin/tracking_detail.html',
                           ct=ct,
                           de_xuat=de_xuat,
                           dept_item_results=dept_item_results,
                           dept_names=DEPT_NAMES,
                           tracking_dept_columns=TRACKING_DEPT_COLUMNS,
                           can_final_approve=can_final_approve,
                           already_approved=already_approved,
                           all_field_labels=all_field_labels,
                           all_fields=all_fields,
                           hoi_dong_votes=HoiDongBieuQuyet.query.filter_by(chi_tiet_id=ct.id).all(),
                           HOI_DONG_VAI_TRO=HOI_DONG_VAI_TRO,
                           HOI_DONG_VAI_TRO_DISPLAY=HOI_DONG_VAI_TRO_DISPLAY,
                           back_url=request.referrer or url_for('admin.approval_tracking'))


@admin_bp.route('/api/chi-tiet/<int:ct_id>')
@login_required
def api_chi_tiet_detail(ct_id):
    """JSON API: return full individual info + criteria for offcanvas detail panel.
    Accessible to admin, reward_viewer, and hoi_dong roles."""
    from app.models.hoi_dong import HoiDongBieuQuyet as _BQ, HOI_DONG_VAI_TRO as _VT, HOI_DONG_VAI_TRO_DISPLAY as _VT_DISPLAY
    ct = DeXuatChiTiet.query.get_or_404(ct_id)
    qn = ct.quan_nhan
    dx = ct.de_xuat

    # Personal info
    personal = {
        'ho_ten': qn.ho_ten if qn else (ct.ten_don_vi_de_xuat or '—'),
        'can_cuoc_cong_dan': qn.can_cuoc_cong_dan if qn else None,
        'ngay_sinh': qn.ngay_sinh.strftime('%d/%m/%Y') if qn and qn.ngay_sinh else None,
        'ngay_nhap_ngu': qn.ngay_nhap_ngu if qn else None,
        'cap_bac': qn.cap_bac if qn else None,
        'chuc_vu': qn.chuc_vu if qn else None,
        'doi_tuong': ct.doi_tuong or (qn.doi_tuong if qn else None),
        'don_vi_truc_thuoc': qn.don_vi_truc_thuoc if qn else None,
        'hoc_ham': qn.hoc_ham if qn else None,
        'hoc_vi': qn.hoc_vi if qn else None,
        'trinh_do_hoc_van': qn.trinh_do_hoc_van if qn else None,
        'ngoai_ngu': qn.ngoai_ngu if qn else None,
        'la_chi_huy': qn.la_chi_huy if qn else None,
        'la_bi_thu': qn.la_bi_thu if qn else None,
        'la_doan_vien': qn.la_doan_vien if qn else None,
        'la_hoi_vien_phu_nu': qn.la_hoi_vien_phu_nu if qn else None,
    }

    # Nomination info
    nomination = {
        'don_vi': dx.don_vi.ten_don_vi if dx.don_vi else '—',
        'nam_hoc': dx.nam_hoc,
        'ngay_gui': dx.ngay_gui.strftime('%d/%m/%Y') if dx.ngay_gui else None,
        'loai_danh_hieu': ct.loai_danh_hieu,
        'trang_thai': dx.trang_thai,
    }

    # Criteria
    # When called from dept approval view, only show fields relevant to current user's dept
    dept_only = request.args.get('dept_only') == '1'

    criteria_fields = [
        ('muc_do_hoan_thanh', 'Hoàn thành nhiệm vụ', None),
        ('phieu_tin_nhiem', 'Phiếu tín nhiệm', None),
        ('xep_loai_dang_vien', 'Xếp loại đảng viên', None),
        ('ket_qua_doan_the', 'Kết quả đoàn thể', None),
        ('xep_loai_doan_vien', 'Xếp loại đoàn viên', None),
        ('hinh_thuc_khen_thuong_qc', 'KT phong trào Quần chúng', None),
        ('ket_qua_phu_nu', 'Xếp loại Phụ nữ hằng năm', None),
        ('hinh_thuc_khen_thuong_pn', 'KT hội Phụ nữ', None),
        ('kiem_tra_chinh_tri', 'Kiểm tra chính trị', 'diem_kiem_tra_chinh_tri'),
        ('kiem_tra_dieu_lenh', 'Kiểm tra điều lệnh', 'diem_kiem_tra_dieu_lenh'),
        ('kiem_tra_tin_hoc', 'Kỹ năng số', 'diem_kiem_tra_tin_hoc'),
        ('dia_ly_quan_su', 'Địa hình quân sự', 'diem_dia_ly_quan_su'),
        ('ban_sung', 'Bắn súng', 'diem_ban_sung'),
        ('the_luc', 'Thể lực', 'diem_the_luc'),
        ('danh_hieu_gv_gioi', 'Danh hiệu GV giỏi', None),
        ('dinh_muc_giang_day', 'Định mức giảng dạy', None),
        ('ket_qua_kiem_tra_giang', 'KT kiểm tra giảng', None),
        ('tien_do_pgs', 'Tiến độ PGS', None),
        ('thoi_gian_lao_dong_kh', 'Thời gian LĐ khoa học', None),
        ('danh_hieu_hv_gioi', 'Danh hiệu HV giỏi', None),
        ('diem_tong_ket', 'Điểm tổng kết', None),
        ('ket_qua_thuc_hanh', 'Kết quả thực hành', None),
        ('ket_qua_ren_luyen', 'Kết quả rèn luyện', None),
        ('hinh_thuc_tot_nghiep', 'Hình thức thi tốt nghiệp', None),
        ('diem_tn_ctd', 'Điểm CTĐ (tốt nghiệp)', None),
        ('diem_tn_ct', 'Điểm CT (tốt nghiệp)', None),
        ('diem_tn_ta', 'Điểm TA (tốt nghiệp)', None),
        ('diem_tn_mon4', 'Điểm môn thứ 4 (tốt nghiệp)', None),
        ('diem_tn_chuyennganh', 'Điểm chuyên ngành (tốt nghiệp)', None),
        ('diem_tn_baove', 'Điểm bảo vệ KL (tốt nghiệp)', None),
        ('diem_nckh', 'Điểm NCKH', None),
        ('nckh_noi_dung', 'Nội dung NCKH', None),
        ('nckh_minh_chung', 'Minh chứng NCKH (text)', None),
        ('mo_ta_khoa_hoc', 'Mô tả thành tích KH', None),
        ('diem_tot_nghiep', 'Điểm tốt nghiệp (TB)', None),
        ('minh_chung_thanh_tich_khac', 'Minh chứng thành tích khác', None),
        ('chu_tri_don_vi_danh_hieu', 'Chủ trì ĐV danh hiệu', None),
        ('thanh_tich_ca_nhan_khac', 'Thành tích cá nhân khác', None),
        ('ghi_chu', 'Ghi chú', None),
    ]

    criteria = []
    if qn is None:
        # 1. Định nghĩa danh sách các trường của tập thể
        tap_the_fields_list = [
            ('kiem_tra_giang', 'Tỷ lệ chất lượng, kết quả kiểm tra bài giảng đạt yêu cầu (%)', None),
            ('cb_xeploai_canthu_pct_hoanttot', 'Tỷ lệ HTTNV cán bộ (%)', None),
            ('cb_xeploai_canthu_pct_xuatsac', 'Tỷ lệ HTT+XS cán bộ (%)', None),
            ('tc_xeploai_dangvien_pct_hoanttot', 'Tỷ lệ HTTNV đảng viên (%)', None),
            ('tc_xeploai_dangvien_pct_xuatsac', 'Tỷ lệ HTT+XS đảng viên (%)', None),
            ('tc_xeploai_tcdcs', 'Xếp loại tổ chức Đoàn cơ sở', None),
            ('th_diem_tdtx_quy1', 'Điểm thi đua Quý I', None),
            ('th_diem_tdtx_quy2', 'Điểm thi đua Quý II', None),
            ('th_diem_tdtx_quy3', 'Điểm thi đua Quý III', None),
            ('th_diem_tdtx_quy4', 'Điểm Thi đua  Quý IV', None),
            ('th_kq_ktra_ct_pct_dyc', 'Tỷ lệ kết quả kiểm tra chính trị đạt yêu cầu trở lên', None),
            ('th_kq_ktra_ct_xeploai', 'Kết quả kiểm tra chính trị đơn vị đạt', None),
            ('ctqc_xeploai_doanvien', 'Xếp loại đoàn viên', None),
            ('ctqc_xeploai_tcd', 'Xếp loại tổ chức đoàn', None),
            ('ctqc_xeploai_hoivien_phunu', 'Tỷ lệ xếp loại hội viên phụ nữ hoàn thành nhiệm vụ trở lên (%)', None),
            ('ctqc_xeploai_hoi_phunu_coso', 'Xếp loại Hội Phụ nữ cơ sở', None),
            ('cntt_chuyen_doi_so', 'Thực hiện Kế hoạch chuyển đổi số hằng năm của Nhà trường', None),
            ('cntt_an_toan_thong_tin', 'Bảo đảm an toàn thông tin, an ninh mạng', None),
            ('tachuan_an_toan_tuyet_doi', 'Đơn vị an toàn tuyệt đối về người, VKTB, tài sản, phương tiện', None),
            ('tachuan_vmtd_mau_muc', 'Đơn vị đạt vững mạnh toàn diện "Mẫu mực, tiêu biểu"', None),
            ('tachuan_dinh_luong', 'Tỷ lệ kết quả kiểm tra điều lệnh đạt yêu cầu trở lên (%)', None),
            ('tachuan_dinh_hinh', 'Tỷ lệ kết quả kiểm tra địa hình quân sự đạt yêu cầu trở lên (%)', None),
            ('tachuan_ban_sung_pct', 'Kết quả kiểm tra bắn súng đơn vị đạt', None),
            ('tachuan_the_luc_pct', 'Xếp loại kết quả kiểm tra thể lực', None),
            ('dt_dinh_muc_ldsp_pct_vuot', 'Tỷ lệ vượt định mức Lao động sư phạm sau uy đổi (%)', None),
            ('dt_cl_bai_giang_pct_dyc', 'Tỷ lệ chất lượng, kết quả bài giảng đạt yêu cầu trở lên (%)', None),
            ('dt_cl_bai_giang_pct_kdat', 'Tỷ lệ chất lượng, kết quả bài giảng đạt khá, tốt trở lên (%)', None),
            ('dt_kq_ktra_giang', 'Tỷ lệ kết quả kiểm tra giảng của Thủ trưởng Nhà trường, cơ quan chức năng đạt khá, tốt', None),
            ('dt_gv_gioi_pct', 'Tỷ lệ quân số đạt danh hiệu "Giảng viên giỏi" (%)', None),
            ('dt_kq_hoc_tap_pct_khagioi', 'Tỷ lệ kết quả học tập đạt khá, giỏi, xuất sắc (%)', None),
            ('dt_kq_hoc_tap_pct_gioi', 'Tỷ lệ kết quả học tập đạt giỏi (%)', None),
            ('dt_luan_van_sdh', 'Tỷ lệ kết quả bảo vệ luận văn của học viên sau đại học đạt khá trở lên (%)', None),
            ('dt_tieng_anh_sdh_pct', 'Tỷ lệ kết quả phúc tra của học viên sau đại học đạt yêu cầu trở lên (%)', None),
            ('dt_ren_luyen_hv_xeploai', 'Tỷ lệ kết quả phân loại rèn luyện tốt của học viên (%)', None),
            ('kh_dinh_muc_ldkh_pct_vuot', 'Tỷ lệ vượt định mức lao động khoa học sau quy đổi (%)', None),
            ('kh_vuot_chi_tieu_nckh', 'Vượt chỉ tiêu, định mức nghiên cứu khoa học', None),
            ('kh_kq_nghiem_thu', 'Kết quả nghiệm thu nghiên cứu đề tài, sáng kiến khoa học các cấp', None),
            ('kh_giao_trinh_so_luong', 'Biên soạn giáo trình, tài liệu dạy học đúng tiến độ, nghiệm thu đạt chất lượng tốt', None),
            ('kh_nckh_ca_nhan_so_luong', 'Tỷ lệ kết quả nghiên cứu đề tài cá nhân đạt khá, xuất sắc (%)', None),
            ('kh_bai_bao_so_luong', 'Tỷ lệ cán bộ, giảng viên là ThS, TS, PGS có đủ số lượng bài báo khoa học theo quy định (%)', None),
            ('kh_sang_kien_hieu_qua', 'Tên sáng kiến, giải pháp công tác hoặc đề tài nghiên cứu mang lại hiệu quả cao và có phạm vi ảnh hưởng đối với đơn vị được cấp có tẩm quyền công nhận', None),
            ('hckt_bao_dam_tieu_chuan', 'Bảo đảm đầy đủ tiêu chuẩn, chế độ cho các đối tượng', None),
            ('hckt_tgsx_tieu_doan', 'Kết quả TGSX so với chỉ tiêu Nghị quyết lãnh đạo thực hiện nhiệm vụ năm học của Đảng uỷ Nhà trường', None),
            ('hckt_phong_chong_dich', 'Chấp hành nghiêm quy định về Phòng chống dịch', None),
            ('hckt_quan_so_khoe_pct', 'Quân số khỏe (%)', None),
            ('hckt_csvc_xeploai', 'Quản lý, sử dụng cơ sở vật chất hậu cần - kỹ thuật: Tốt, bền, an toàn, tiết kiệm', None),
            ('hckt_an_toan_gt', 'An toàn giao thông', None),
            ('kt_kq_ktra_giang_pct_khatot', 'Tỷ lệ kết quả kiểm tra giảng của Thủ trưởng Nhà trường, cơ quan chức năng đạt khá, tốt', None),
            ('kt_kq_ktra_giang_pct_tot', 'Tỷ lệ kết quả kiểm tra giảng của Thủ trưởng Nhà trường, cơ quan chức năng đạt tốt', None),
            ('muc_do_hoan_thanh', 'Mức độ hoàn thành NV', None),
            ('chu_tri_don_vi_danh_hieu', 'Chủ trì ĐV danh hiệu', None),
            ('thanh_tich_ca_nhan_khac', 'Thành tích khác', None),
            ('ghi_chu', 'Ghi chú', None),
        ]

        # 2. Tạo Dictionary dùng để tra cứu nhanh (Map từ field -> label)
        label_mapping = {f[0]: f[1] for f in tap_the_fields_list}

        # 3. Quét dữ liệu trong chuỗi JSON (tap_the_dict) và dịch key sang label
        tap_the_dict = ct.tap_the_dict or {}
        for k, v in tap_the_dict.items():
            if v is not None and v != '':
                # Lấy tên tiếng Việt từ label_mapping, nếu không có thì hiển thị lại key gốc
                mapped_label = label_mapping.get(k, k) 
                criteria.append({'label': mapped_label, 'value': str(v), 'score': None})

        # 4. Quét thêm các cột lưu trực tiếp trên bảng (tránh lấy trùng với các key đã có ở JSON)
        existing_json_keys = set(tap_the_dict.keys())
        for field, label, score_field in tap_the_fields_list:
            if field not in existing_json_keys:
                val = getattr(ct, field, None)
                if val:
                    criteria.append({'label': label, 'value': str(val), 'score': None})
    else:
        for field, label, score_field in criteria_fields:
            val = getattr(ct, field, None)
            if val is None or val == '':
                continue
            if val == 'true':
                val = 'Có'
            elif val == 'false':
                val = 'Không'
            score = getattr(ct, score_field, None) if score_field else None
            criteria.append({'field': field, 'label': label, 'value': str(val), 'score': str(score) if score is not None else None})

    # Filter criteria to dept-specific fields when requested
    if dept_only and qn is not None:
        from app.routes.approval import get_phong_fields as _get_phong_fields
        dept_fields = set(_get_phong_fields().get(current_user.role, []))
        if dept_fields:
            criteria = [c for c in criteria if c.get('field') in dept_fields]

    # Hội đồng votes
    votes = []
    for vai_tro in _VT:
        bq = _BQ.query.filter_by(chi_tiet_id=ct.id, vai_tro=vai_tro).first()
        votes.append({
            'vai_tro': _VT_DISPLAY.get(vai_tro, vai_tro),
            'ket_qua': bq.ket_qua if bq else None,
            'ghi_chu': bq.ghi_chu if bq else None,
        })

    # Dept approval results
    dept_results = []
    for dept in DEPT_NAMES:
        is_auto = _is_auto_scope_approved(dept, ct.doi_tuong)
        if is_auto:
            dept_results.append({'dept': dept, 'ket_qua': 'auto', 'ghi_chu': None, 'ngay_duyet': None})
            continue
        pd_dept = PheDuyet.query.filter_by(de_xuat_id=dx.id, phong_duyet=dept).first()
        if not pd_dept:
            dept_results.append({'dept': dept, 'ket_qua': None, 'ghi_chu': None, 'ngay_duyet': None})
            continue
        kq_ct = KetQuaDuyetChiTiet.query.filter_by(phe_duyet_id=pd_dept.id, chi_tiet_id=ct.id).first()
        if pd_dept.ket_qua == KetQuaDuyet.DONG_Y.value:
            dept_results.append({'dept': dept, 'ket_qua': 'Đồng ý', 'ghi_chu': pd_dept.ghi_chu,
                                  'ngay_duyet': pd_dept.ngay_duyet.strftime('%d/%m/%Y %H:%M') if pd_dept.ngay_duyet else None})
        elif pd_dept.ket_qua == KetQuaDuyet.TU_CHOI.value:
            dept_results.append({'dept': dept, 'ket_qua': 'Từ chối', 'ghi_chu': pd_dept.ghi_chu,
                                  'ngay_duyet': pd_dept.ngay_duyet.strftime('%d/%m/%Y %H:%M') if pd_dept.ngay_duyet else None})
        elif kq_ct:
            dept_results.append({'dept': dept, 'ket_qua': kq_ct.ket_qua, 'ghi_chu': kq_ct.ly_do,
                                  'ngay_duyet': pd_dept.ngay_duyet.strftime('%d/%m/%Y %H:%M') if pd_dept.ngay_duyet else None})
        else:
            dept_results.append({'dept': dept, 'ket_qua': None, 'ghi_chu': None, 'ngay_duyet': None})

    # Minh chứng files grouped by type — include URL for download/preview
    minh_chungs = []
    MINH_CHUNG_LABELS = {
        'nckh_minh_chung': 'Minh chứng NCKH',
        'minh_chung_chung': 'Minh chứng chung',
        'minh_chung_thanh_tich_khac': 'Minh chứng thành tích khác',
        'minh_chung_thanh_tich_ca_nhan_khac': 'Minh chứng thành tích cá nhân khác',
    }
    IMAGE_EXTS = {'.jpg', '.jpeg', '.png', '.gif', '.webp', '.bmp'}
    for mc in ct.minh_chungs:
        ext = ('.' + mc.duong_dan.rsplit('.', 1)[-1].lower()) if '.' in mc.duong_dan else ''
        minh_chungs.append({
            'loai': MINH_CHUNG_LABELS.get(mc.loai_minh_chung, mc.loai_minh_chung),
            'ten_file': mc.ten_file_goc or mc.duong_dan.split('/')[-1],
            'url': '/uploads/' + mc.duong_dan,
            'is_image': ext in IMAGE_EXTS,
        })

    return jsonify({
        'personal': personal,
        'nomination': nomination,
        'criteria': criteria,
        'votes': votes,
        'dept_results': dept_results,
        'minh_chungs': minh_chungs,
    })


@admin_bp.route('/tracking/chi-tiet/<int:ct_id>/final-approve', methods=['POST'])
@login_required
@admin_required
def final_approve_individual(ct_id):
    """Admin pre-approves a single individual (Bảng 1 → Bảng 2).
    Does NOT create KhenThuong yet; marks ct.admin_approved = True.
    When all individuals in the nomination are admin-approved → moves to PHE_DUYET_CUOI
    for Hội đồng voting (Bảng 2).
    """
    ct = DeXuatChiTiet.query.get_or_404(ct_id)
    de_xuat = ct.de_xuat

    # Check if already admin-approved
    if ct.admin_approved:
        flash('Cá nhân này đã được phê duyệt (chờ Hội đồng biểu quyết).', 'warning')
        return redirect(url_for('admin.reward_list'))

    # Verify all departments approved this individual (including auto-scope)
    for dept_name in DEPT_NAMES:
        if not _is_individual_dept_approved(de_xuat.id, ct, dept_name):
            flash(f'{dept_name} chưa đồng ý cho cá nhân này.', 'warning')
            return redirect(url_for('admin.reward_list'))

    now = datetime.utcnow()
    ghi_chu = request.form.get('ghi_chu', '').strip() or None

    # Mark this individual as admin pre-approved
    ct.admin_approved = True

    # Check if ALL individuals in this nomination are now admin-approved
    all_ct_ids = {c.id for c in de_xuat.chi_tiets}
    already_approved = {c.id for c in de_xuat.chi_tiets if c.admin_approved}
    # Include the one we just approved (not yet committed)
    already_approved.add(ct.id)

    if all_ct_ids <= already_approved:
        # All individuals admin-approved → move to PHE_DUYET_CUOI for Hội đồng voting
        de_xuat.trang_thai = TrangThaiDeXuat.PHE_DUYET_CUOI.value
        admin_pd = PheDuyet.query.filter_by(
            de_xuat_id=de_xuat.id, phong_duyet=PhongDuyet.ADMIN_TUYENHUAN.value
        ).first()
        if admin_pd:
            admin_pd.ket_qua = KetQuaDuyet.DONG_Y.value
            admin_pd.nguoi_duyet_id = current_user.id
            admin_pd.ngay_duyet = now
            admin_pd.ghi_chu = ghi_chu

    # Sync per-item trang_thai
    try:
        from app.routes.approval import _recompute_chi_tiet_status
        _recompute_chi_tiet_status(de_xuat)
    except Exception:
        pass

    db.session.commit()
    ho_ten = ct.quan_nhan.ho_ten if ct.quan_nhan else de_xuat.don_vi.ten_don_vi
    log_action('admin_pre_approve', resource_type='chi_tiet', resource_id=ct.id,
               detail=f'{ho_ten} — đề xuất #{de_xuat.id} năm học {de_xuat.nam_hoc}')
    db.session.commit()
    flash(f'Đã đồng ý cho "{ho_ten}". Khi toàn bộ đề xuất được duyệt sẽ chuyển sang Hội đồng biểu quyết.', 'success')
    return redirect(url_for('admin.reward_list', nam_hoc=de_xuat.nam_hoc, _anchor='bang2'))


@admin_bp.route('/tracking/<int:id>/final-approve', methods=['POST'])
@login_required
@admin_required
def final_approve_from_tracking(id):
    """Admin pre-approves entire nomination at once (Bảng 1 → Bảng 2).
    Marks all individuals as admin_approved and moves nomination to PHE_DUYET_CUOI
    for Hội đồng voting. Does NOT create KhenThuong yet.
    """
    de_xuat = DeXuat.query.get_or_404(id)

    if de_xuat.trang_thai != TrangThaiDeXuat.HOI_DONG.value:
        flash('Đề xuất này chưa qua giai đoạn Xét duyệt của cơ quan thường trực.', 'warning')
        return redirect(url_for('admin.approval_tracking'))

    # Verify departments approved this nomination, honoring auto-scope
    for dept_name in DEPT_NAMES:
        pd = PheDuyet.query.filter_by(
            de_xuat_id=id, phong_duyet=dept_name
        ).first()
        if not pd:
            # acceptable only if all individuals are auto-scope for this department
            has_in_scope = any(not _is_auto_scope_approved(dept_name, ct.doi_tuong) for ct in de_xuat.chi_tiets_active)
            if has_in_scope:
                flash(f'{dept_name} chưa phê duyệt xong.', 'warning')
                return redirect(url_for('admin.approval_tracking'))
            continue
        if pd.ket_qua == KetQuaDuyet.TU_CHOI.value:
            flash(f'{dept_name} chưa phê duyệt xong.', 'warning')
            return redirect(url_for('admin.approval_tracking'))

    now = datetime.utcnow()
    ghi_chu = request.form.get('ghi_chu', '').strip() or None

    # Mark all individuals as admin pre-approved
    for ct in de_xuat.chi_tiets_active:
        ct.admin_approved = True

    # Update admin PheDuyet record
    admin_pd = PheDuyet.query.filter_by(
        de_xuat_id=id, phong_duyet=PhongDuyet.ADMIN_TUYENHUAN.value
    ).first()
    if admin_pd:
        admin_pd.ket_qua = KetQuaDuyet.DONG_Y.value
        admin_pd.nguoi_duyet_id = current_user.id
        admin_pd.ngay_duyet = now
        admin_pd.ghi_chu = ghi_chu

    de_xuat.trang_thai = TrangThaiDeXuat.PHE_DUYET_CUOI.value

    # Sync per-item trang_thai
    try:
        from app.routes.approval import _recompute_chi_tiet_status
        _recompute_chi_tiet_status(de_xuat)
    except Exception:
        pass

    db.session.commit()
    flash('Đã phê duyệt đề xuất. Chuyển sang Hội đồng biểu quyết (Bảng 2).', 'success')
    return redirect(url_for('admin.reward_list', nam_hoc=de_xuat.nam_hoc, _anchor='bang2'))


@admin_bp.route('/reward-list/confirm-khen-thuong/<int:id>', methods=['POST'])
@login_required
@admin_required
def confirm_khen_thuong(id):
    """Final step: After all 6 Hội đồng roles voted DONG_Y, Admin creates KhenThuong records.
    This moves nomination from PHE_DUYET_CUOI to fully finalized (KhenThuong records exist).
    """
    de_xuat = DeXuat.query.get_or_404(id)

    if de_xuat.trang_thai != TrangThaiDeXuat.PHE_DUYET_CUOI.value:
        flash('Đề xuất này không ở giai đoạn chờ xác nhận khen thưởng.', 'warning')
        return redirect(url_for('admin.reward_list'))

    # Check that all 6 Hội đồng roles voted DONG_Y for all chi_tiets
    from app.models.hoi_dong import HOI_DONG_VAI_TRO, KET_QUA_DONG_Y
    for ct in de_xuat.chi_tiets_active:
        for vai_tro in HOI_DONG_VAI_TRO:
            bq = HoiDongBieuQuyet.query.filter_by(chi_tiet_id=ct.id, vai_tro=vai_tro).first()
            if not bq or bq.ket_qua != KET_QUA_DONG_Y:
                name = ct.quan_nhan.ho_ten if ct.quan_nhan else ct.ten_don_vi_de_xuat or 'Đơn vị'
                from app.models.hoi_dong import HOI_DONG_VAI_TRO_DISPLAY
                vt_name = HOI_DONG_VAI_TRO_DISPLAY.get(vai_tro, vai_tro)
                flash(f'{vt_name} chưa biểu quyết đồng ý cho "{name}".', 'warning')
                return redirect(url_for('admin.reward_list'))

    now = datetime.utcnow()
    ghi_chu = request.form.get('ghi_chu', '').strip() or None

    # Create KhenThuong for each individual
    for ct in de_xuat.chi_tiets_active:
        existing = KhenThuong.query.filter_by(de_xuat_id=de_xuat.id, chi_tiet_id=ct.id).first()
        if not existing:
            khen_thuong = KhenThuong(
                de_xuat_id=de_xuat.id,
                chi_tiet_id=ct.id,
                quan_nhan_id=ct.quan_nhan_id,
                don_vi_id=de_xuat.don_vi_id,
                ho_ten=ct.quan_nhan.ho_ten if ct.quan_nhan else de_xuat.don_vi.ten_don_vi,
                cap_bac=ct.quan_nhan.cap_bac if ct.quan_nhan else None,
                chuc_vu=ct.quan_nhan.chuc_vu if ct.quan_nhan else None,
                doi_tuong=ct.doi_tuong,
                loai_danh_hieu=ct.loai_danh_hieu,
                nam_hoc=de_xuat.nam_hoc,
                nguoi_duyet_id=current_user.id,
                ngay_duyet=now,
                ghi_chu=ghi_chu,
            )
            db.session.add(khen_thuong)

    db.session.commit()
    flash(f'Đã xác nhận khen thưởng cho đề xuất của {de_xuat.don_vi.ten_don_vi}.', 'success')
    return redirect(url_for('admin.reward_list'))


@admin_bp.route('/reward-list/confirm-khen-thuong-ct/<int:ct_id>', methods=['POST'])
@login_required
@admin_required
def confirm_khen_thuong_ct(ct_id):
    """Admin confirms KhenThuong for a single DeXuatChiTiet after all 7 organs have voted (Bảng 2 → Bảng 3)."""
    from app.models.hoi_dong import HOI_DONG_VAI_TRO, HOI_DONG_VAI_TRO_DISPLAY
    ct = DeXuatChiTiet.query.get_or_404(ct_id)
    de_xuat = ct.de_xuat

    if de_xuat.trang_thai != TrangThaiDeXuat.PHE_DUYET_CUOI.value:
        flash('Đề xuất không ở giai đoạn chờ xác nhận khen thưởng.', 'warning')
        return redirect(url_for('admin.reward_list', nam_hoc=de_xuat.nam_hoc))

    # Require all 7 organs to have cast a vote (any result) before Admin decides
    for vai_tro in HOI_DONG_VAI_TRO:
        bq = HoiDongBieuQuyet.query.filter_by(chi_tiet_id=ct.id, vai_tro=vai_tro).first()
        if not bq:
            vt_name = HOI_DONG_VAI_TRO_DISPLAY.get(vai_tro, vai_tro)
            flash(f'{vt_name} chưa biểu quyết cho cá nhân/tập thể này.', 'warning')
            return redirect(url_for('admin.reward_list', nam_hoc=de_xuat.nam_hoc))

    existing = KhenThuong.query.filter_by(de_xuat_id=de_xuat.id, chi_tiet_id=ct.id).first()
    if existing:
        flash('Cá nhân/tập thể này đã được xác nhận khen thưởng rồi.', 'info')
        return redirect(url_for('admin.reward_list', nam_hoc=de_xuat.nam_hoc))

    now = datetime.utcnow()
    name = ct.quan_nhan.ho_ten if ct.quan_nhan else (ct.ten_don_vi_de_xuat or de_xuat.don_vi.ten_don_vi)
    kt = KhenThuong(
        de_xuat_id=de_xuat.id,
        chi_tiet_id=ct.id,
        quan_nhan_id=ct.quan_nhan_id,
        don_vi_id=de_xuat.don_vi_id,
        ho_ten=name,
        cap_bac=ct.quan_nhan.cap_bac if ct.quan_nhan else None,
        chuc_vu=ct.quan_nhan.chuc_vu if ct.quan_nhan else None,
        doi_tuong=ct.doi_tuong,
        loai_danh_hieu=ct.loai_danh_hieu,
        nam_hoc=de_xuat.nam_hoc,
        nguoi_duyet_id=current_user.id,
        ngay_duyet=now,
    )
    db.session.add(kt)
    db.session.commit()
    flash(f'Đã xác nhận khen thưởng cho {name}.', 'success')
    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        return jsonify({'ok': True, 'message': f'Đã xác nhận khen thưởng cho {name}.'})
    return redirect(url_for('admin.reward_list', nam_hoc=de_xuat.nam_hoc))


@admin_bp.route('/reward-list/confirm-khong-dong-y-ct/<int:ct_id>', methods=['POST'])
@login_required
@admin_required
def confirm_khong_dong_y_ct(ct_id):
    """Admin xác nhận Không đồng ý cho một cá nhân/tập thể trong Bảng 2 → đưa vào Bảng 3 mục Bị từ chối."""
    from app.models.hoi_dong import (
        HoiDongBieuQuyet, HOI_DONG_VAI_TRO, HOI_DONG_VAI_TRO_DISPLAY,
        KET_QUA_KHONG_DONG_Y,
    )
    ct = DeXuatChiTiet.query.get_or_404(ct_id)
    de_xuat = ct.de_xuat

    if de_xuat.trang_thai != TrangThaiDeXuat.PHE_DUYET_CUOI.value:
        flash('Đề xuất không ở giai đoạn xét duyệt của Hội đồng.', 'warning')
        return redirect(url_for('admin.reward_list', nam_hoc=de_xuat.nam_hoc))

    # Require all 7 organs to have cast a vote before Admin decides
    for vai_tro in HOI_DONG_VAI_TRO:
        bq = HoiDongBieuQuyet.query.filter_by(chi_tiet_id=ct.id, vai_tro=vai_tro).first()
        if not bq:
            vt_name = HOI_DONG_VAI_TRO_DISPLAY.get(vai_tro, vai_tro)
            flash(f'{vt_name} chưa biểu quyết cho cá nhân/tập thể này.', 'warning')
            return redirect(url_for('admin.reward_list', nam_hoc=de_xuat.nam_hoc))

    # Check not already confirmed as KhenThuong
    if KhenThuong.query.filter_by(chi_tiet_id=ct.id).first():
        flash('Cá nhân/tập thể này đã được xác nhận khen thưởng, không thể đánh dấu không đồng ý.', 'warning')
        return redirect(url_for('admin.reward_list', nam_hoc=de_xuat.nam_hoc))

    ghi_chu = request.form.get('ghi_chu', '').strip() or None

    existing = HoiDongBieuQuyet.query.filter_by(chi_tiet_id=ct_id, vai_tro='admin_final').first()
    if existing:
        existing.ket_qua = KET_QUA_KHONG_DONG_Y
        existing.ghi_chu = ghi_chu
        existing.nguoi_bieu_quyet_id = current_user.id
    else:
        bq = HoiDongBieuQuyet(
            de_xuat_id=ct.de_xuat_id,
            chi_tiet_id=ct_id,
            nguoi_bieu_quyet_id=current_user.id,
            vai_tro='admin_final',
            ket_qua=KET_QUA_KHONG_DONG_Y,
            ghi_chu=ghi_chu,
        )
        db.session.add(bq)

    db.session.commit()
    name = ct.quan_nhan.ho_ten if ct.quan_nhan else (ct.ten_don_vi_de_xuat or de_xuat.don_vi.ten_don_vi)
    flash(f'Đã xác nhận Không đồng ý cho {name}. Sẽ hiển thị trong Bảng 3 mục Bị từ chối.', 'info')
    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        return jsonify({'ok': True, 'message': f'Đã xác nhận Không đồng ý cho {name}.'})
    return redirect(url_for('admin.reward_list', nam_hoc=de_xuat.nam_hoc))
# ── Mapping hạ danh hiệu ──────────────────────────────────────────────────
_DOWNGRADE_MAP = {
    'Chiến sĩ thi đua':   'Chiến sĩ tiên tiến',
    'Đơn vị quyết thắng': 'Đơn vị tiên tiến',
}

@admin_bp.route('/reward-list/downgrade-ct/<int:ct_id>', methods=['POST'])
@login_required
@admin_required
def downgrade_danh_hieu_ct(ct_id):
    """
    Hạ danh hiệu cho một chi tiết đề xuất đang ở Bảng 2 (PHE_DUYET_CUOI):
      CSTD  → CSTT
      DVQT  → DVTT
    Đồng thời tạo KhenThuong với danh hiệu mới (hạ cấp).
    """
    ct = DeXuatChiTiet.query.get_or_404(ct_id)
    de_xuat = ct.de_xuat

    if de_xuat.trang_thai != TrangThaiDeXuat.PHE_DUYET_CUOI.value:
        flash('Đề xuất không ở giai đoạn xét duyệt của Hội đồng.', 'warning')
        return redirect(url_for('admin.reward_list', nam_hoc=de_xuat.nam_hoc))

    new_dh = _DOWNGRADE_MAP.get(ct.loai_danh_hieu)
    if not new_dh:
        flash(f'Danh hiệu "{ct.loai_danh_hieu}" không thể hạ cấp.', 'warning')
        return redirect(url_for('admin.reward_list', nam_hoc=de_xuat.nam_hoc))

    # Kiểm tra chưa có KhenThuong cho ct này
    if KhenThuong.query.filter_by(chi_tiet_id=ct.id).first():
        flash('Cá nhân/tập thể này đã được xác nhận khen thưởng, không thể hạ danh hiệu.', 'warning')
        return redirect(url_for('admin.reward_list', nam_hoc=de_xuat.nam_hoc))

    old_dh = ct.loai_danh_hieu
    ct.loai_danh_hieu = new_dh  # Cập nhật trực tiếp trên chi tiết đề xuất

    name = ct.quan_nhan.ho_ten if ct.quan_nhan else (ct.ten_don_vi_de_xuat or de_xuat.don_vi.ten_don_vi)
    now  = datetime.utcnow()

    kt = KhenThuong(
        de_xuat_id=de_xuat.id,
        chi_tiet_id=ct.id,
        quan_nhan_id=ct.quan_nhan_id,
        don_vi_id=de_xuat.don_vi_id,
        ho_ten=name,
        cap_bac=ct.quan_nhan.cap_bac if ct.quan_nhan else None,
        chuc_vu=ct.quan_nhan.chuc_vu if ct.quan_nhan else None,
        doi_tuong=ct.doi_tuong,
        loai_danh_hieu=new_dh,   # ← danh hiệu mới (hạ cấp)
        nam_hoc=de_xuat.nam_hoc,
        nguoi_duyet_id=current_user.id,
        ngay_duyet=now,
    )
    db.session.add(kt)
    db.session.commit()

    log_action('admin_downgrade_danh_hieu', resource_type='chi_tiet', resource_id=ct.id,
               detail=f'{name}: {old_dh} → {new_dh}, năm học {de_xuat.nam_hoc}')
    db.session.commit()

    flash(f'Đã hạ danh hiệu {name}: {old_dh} → {new_dh}.', 'success')
    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        return jsonify({'ok': True, 'message': f'Đã hạ danh hiệu {name}: {old_dh} → {new_dh}.'})
    return redirect(url_for('admin.reward_list', nam_hoc=de_xuat.nam_hoc))


@admin_bp.route('/reward-list/downgrade-kt/<int:kt_id>', methods=['POST'])
@login_required
@admin_required
def downgrade_khen_thuong(kt_id):
    """
    Hạ danh hiệu cho một KhenThuong đã xác nhận (Bảng 3):
      CSTD  → CSTT
      DVQT  → DVTT
    Cập nhật cả KhenThuong lẫn DeXuatChiTiet.
    """
    kt = KhenThuong.query.get_or_404(kt_id)
    ct = DeXuatChiTiet.query.get(kt.chi_tiet_id)

    new_dh = _DOWNGRADE_MAP.get(kt.loai_danh_hieu)
    if not new_dh:
        flash(f'Danh hiệu "{kt.loai_danh_hieu}" không thể hạ cấp.', 'warning')
        return redirect(url_for('admin.reward_list', nam_hoc=kt.nam_hoc))

    old_dh = kt.loai_danh_hieu
    kt.loai_danh_hieu = new_dh
    if ct:
        ct.loai_danh_hieu = new_dh  # đồng bộ chi tiết đề xuất

    db.session.commit()

    log_action('admin_downgrade_khen_thuong', resource_type='khen_thuong', resource_id=kt.id,
               detail=f'{kt.ho_ten}: {old_dh} → {new_dh}, năm học {kt.nam_hoc}')
    db.session.commit()

    flash(f'Đã hạ danh hiệu {kt.ho_ten}: {old_dh} → {new_dh}.', 'success')
    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        return jsonify({'ok': True, 'message': f'Đã hạ danh hiệu {kt.ho_ten}: {old_dh} → {new_dh}.'})
    return redirect(url_for('admin.reward_list', nam_hoc=kt.nam_hoc))


@admin_bp.route('/reward-list/vote-ct/<int:ct_id>', methods=['POST'])
@login_required
def hoi_dong_vote_ct(ct_id):
    """Hội đồng member casts a vote (đồng ý / không đồng ý) for one chi_tiet."""
    from app.models.hoi_dong import HoiDongBieuQuyet, KET_QUA_DONG_Y, KET_QUA_KHONG_DONG_Y, HA_CAP
    vai_tro = current_user.hoi_dong_vai_tro
    # Admin can vote using the vai_tro submitted in the form (or 'admin' as fallback)
    if not vai_tro:
        if current_user.is_admin:
            vai_tro = request.form.get('vai_tro', 'admin').strip() or 'admin'
        else:
            flash('Bạn không có quyền biểu quyết.', 'danger')
            return redirect(url_for('admin.reward_list'))

    ct = DeXuatChiTiet.query.get_or_404(ct_id)
    if ct.de_xuat.trang_thai != TrangThaiDeXuat.PHE_DUYET_CUOI.value and ct.admin_approved != 1:
        flash('Đề xuất không ở giai đoạn biểu quyết của Hội đồng.', 'warning')
        return redirect(url_for('admin.reward_list'))

    ket_qua = request.form.get('ket_qua', '').strip()
    if ket_qua not in (KET_QUA_DONG_Y, KET_QUA_KHONG_DONG_Y, HA_CAP):
        flash('Kết quả biểu quyết không hợp lệ.', 'danger')
        return redirect(url_for('admin.reward_list'))

    ghi_chu = request.form.get('ghi_chu', '').strip() or None

    existing = HoiDongBieuQuyet.query.filter_by(chi_tiet_id=ct_id, vai_tro=vai_tro).first()
    if existing:
        existing.ket_qua = ket_qua
        existing.ghi_chu = ghi_chu
        existing.nguoi_bieu_quyet_id = current_user.id
    else:
        bq = HoiDongBieuQuyet(
            de_xuat_id=ct.de_xuat_id,
            chi_tiet_id=ct_id,
            nguoi_bieu_quyet_id=current_user.id,
            vai_tro=vai_tro,
            ket_qua=ket_qua,
            ghi_chu=ghi_chu,
        )
        db.session.add(bq)

    db.session.commit()
    name = ct.quan_nhan.ho_ten if ct.quan_nhan else ct.de_xuat.don_vi.ten_don_vi
    log_action('hoi_dong_vote', resource_type='chi_tiet', resource_id=ct_id,
               detail=f'{name} — kết quả: {ket_qua}, vai trò: {vai_tro}')
    db.session.commit()
    flash(f'Đã ghi nhận biểu quyết "{ket_qua}" cho {name}.', 'success')
    nam_hoc = ct.de_xuat.nam_hoc if ct.de_xuat else ''
    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        return jsonify({'ok': True, 'message': f'Đã ghi nhận biểu quyết "{ket_qua}" cho {name}.'})
    return redirect(url_for('admin.reward_list', nam_hoc=nam_hoc))



@admin_bp.route('/tracking/<int:id>/reject', methods=['POST'])
@login_required
@admin_required
def reject_from_tracking(id):
    """Admin rejects a nomination from tracking detail page."""
    de_xuat = DeXuat.query.get_or_404(id)

    ly_do = request.form.get('ly_do', '').strip()
    if not ly_do:
        flash('Vui lòng nhập lý do từ chối.', 'danger')
        return redirect(url_for('admin.approval_tracking'))

    admin_pd = PheDuyet.query.filter_by(
        de_xuat_id=id, phong_duyet=PhongDuyet.ADMIN_TUYENHUAN.value
    ).first()

    if admin_pd:
        admin_pd.ket_qua = KetQuaDuyet.TU_CHOI.value
        admin_pd.nguoi_duyet_id = current_user.id
        admin_pd.ngay_duyet = datetime.utcnow()
        admin_pd.ly_do = ly_do

    de_xuat.trang_thai = TrangThaiDeXuat.TU_CHOI.value
    db.session.commit()
    log_action('admin_reject', resource_type='de_xuat', resource_id=id,
               detail=f'Từ chối đề xuất #{id} năm học {de_xuat.nam_hoc} — {ly_do}')
    db.session.commit()

    # Notify unit account
    unit_user = User.query.filter_by(don_vi_id=de_xuat.don_vi_id, role=Role.UNIT_USER).first()
    if unit_user:
        thong_bao = ThongBao(
            user_id=unit_user.id,
            de_xuat_id=de_xuat.id,
            loai='tu_choi',
            tieu_de=f'Ban Tuyên huấn (Admin) từ chối đề xuất năm học {de_xuat.nam_hoc}',
            noi_dung=f'Lý do: {ly_do}',
        )
        db.session.add(thong_bao)
        db.session.commit()

    flash('Đã từ chối đề xuất.', 'warning')
    return redirect(url_for('admin.approval_tracking'))
# ─────────────────────────────────────────────────────────────────────────────
# HELPER: logic từ chối 1 chi tiết — dùng chung cho cả đơn lẻ lẫn hàng loạt
# ─────────────────────────────────────────────────────────────────────────────
def _reject_single_chi_tiet(ct, ly_do: str) -> str:
    """
    Từ chối 1 DeXuatChiTiet (cá nhân hoặc tập thể).
    - Tạo/cập nhật PheDuyet + KetQuaDuyetChiTiet
    - Đánh dấu ct.bi_loai, reset ct.admin_approved
    - Recompute trạng thái đề xuất
    - Gửi ThongBao về đơn vị
    - Ghi log

    Trả về: tên hiển thị (ho_ten hoặc ten_don_vi)
    KHÔNG commit — caller tự commit sau khi xử lý xong batch.
    """
    de_xuat = ct.de_xuat

    # ── 1. Đảm bảo PheDuyet admin tồn tại ───────────────────────────────────
    admin_pd = PheDuyet.query.filter_by(
        de_xuat_id=de_xuat.id,
        phong_duyet=PhongDuyet.ADMIN_TUYENHUAN.value
    ).first()
    if not admin_pd:
        admin_pd = PheDuyet(
            de_xuat_id=de_xuat.id,
            phong_duyet=PhongDuyet.ADMIN_TUYENHUAN.value,
            ket_qua=KetQuaDuyet.CHO_DUYET.value,
        )
        db.session.add(admin_pd)
        db.session.flush()  # lấy admin_pd.id

    # ── 2. Tạo/cập nhật KetQuaDuyetChiTiet ──────────────────────────────────
    kq_ct = KetQuaDuyetChiTiet.query.filter_by(
        phe_duyet_id=admin_pd.id,
        chi_tiet_id=ct.id
    ).first()
    if kq_ct:
        kq_ct.ket_qua = KetQuaDuyet.TU_CHOI.value
        kq_ct.ly_do   = ly_do
    else:
        kq_ct = KetQuaDuyetChiTiet(
            phe_duyet_id=admin_pd.id,
            chi_tiet_id=ct.id,
            ket_qua=KetQuaDuyet.TU_CHOI.value,
            ly_do=ly_do,
        )
        db.session.add(kq_ct)

    # ── 3. Reset admin_approved + đánh dấu bi_loai ───────────────────────────
    ct.admin_approved = False
    if not ct.bi_loai:
        ct.bi_loai    = True
        ct.ly_do_loai = ly_do
        ct.phong_loai = PhongDuyet.ADMIN_TUYENHUAN.value
        ct.ngay_loai  = datetime.utcnow()

    # ── 4. Recompute trạng thái đề xuất ──────────────────────────────────────
    from app.routes.approval import _recompute_de_xuat_status
    _recompute_de_xuat_status(de_xuat)

    # ── 5. Tên hiển thị ───────────────────────────────────────────────────────
    ho_ten = (
        ct.quan_nhan.ho_ten
        if ct.quan_nhan
        else (de_xuat.don_vi.ten_don_vi if de_xuat.don_vi else f'ID:{ct.id}')
    )

    # ── 6. Gửi ThongBao về đơn vị ────────────────────────────────────────────
    unit_user = User.query.filter_by(
        don_vi_id=de_xuat.don_vi_id,
        role=Role.UNIT_USER
    ).first()
    if unit_user:
        thong_bao = ThongBao(
            user_id=unit_user.id,
            de_xuat_id=de_xuat.id,
            chi_tiet_id=ct.id,
            loai='tu_choi',
            tieu_de=f'Ban Tuyên huấn (Admin) loại khỏi đề xuất: {ho_ten}',
            noi_dung=(
                f'Lý do: {ly_do}. {ho_ten} đã bị loại khỏi đề xuất '
                f'năm học {de_xuat.nam_hoc}. '
                f'Các cá nhân/tập thể còn lại vẫn tiếp tục được xét duyệt.'
            ),
        )
        db.session.add(thong_bao)

    # ── 7. Ghi log ────────────────────────────────────────────────────────────
    log_action(
        'admin_reject_individual',
        resource_type='chi_tiet',
        resource_id=ct.id,
        detail=f'{ho_ten} — lý do: {ly_do}',
    )

    return ho_ten
# ─────────────────────────────────────────────────────────────────────────────
# ROUTE: Từ chối hàng loạt (JSON API, dùng helper trong vòng lặp)
# ─────────────────────────────────────────────────────────────────────────────
@admin_bp.route('/batch-reject-individuals', methods=['POST'])
@login_required
@admin_required
def batch_reject_individuals():
    """Từ chối hàng loạt các chi tiết đề xuất đã chọn trên trang tracking."""
    data  = request.get_json(force=True) or {}
    ids   = data.get('ids') or []
    ly_do = (data.get('ly_do') or '').strip()

    if not ids:
        return jsonify(success=False, message='Không có mục nào được chọn.')
    if not ly_do:
        return jsonify(success=False, message='Vui lòng nhập lý do từ chối.')

    rejected = 0
    skipped  = []   # ids không tìm thấy
    errors   = []   # ids gặp lỗi runtime

    for ct_id in ids:
        ct = DeXuatChiTiet.query.get(ct_id)
        if not ct:
            skipped.append(ct_id)
            continue
        try:
            _reject_single_chi_tiet(ct, ly_do)
            rejected += 1
        except Exception as e:
            errors.append(f'ID {ct_id}: {str(e)}')

    # Commit toàn bộ batch 1 lần duy nhất
    try:
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        return jsonify(success=False, message=f'Lỗi lưu dữ liệu: {str(e)}')

    # Tổng hợp kết quả
    msg_parts = [f'Đã từ chối {rejected} mục.']
    if skipped:
        msg_parts.append(f'Không tìm thấy: {len(skipped)} mục (ID: {skipped}).')
    if errors:
        msg_parts.append(f'Lỗi: {"; ".join(errors)}')

    return jsonify(
        success=True,
        message=' '.join(msg_parts),
        rejected=rejected,
        skipped=len(skipped),
        errors=errors,
    )

@admin_bp.route('/tracking/ct/<int:ct_id>/reject', methods=['POST'])
@login_required
@admin_required
def reject_individual_from_tracking(ct_id):
    """Admin rejects a single individual from tracking page, notifies unit."""
    ct = DeXuatChiTiet.query.get_or_404(ct_id)
    de_xuat = ct.de_xuat

    ly_do = request.form.get('ly_do', '').strip()
    if not ly_do:
        flash('Vui lòng nhập lý do từ chối.', 'danger')
        return redirect(url_for('admin.approval_tracking'))

    # Mark this individual as rejected via admin PheDuyet chi_tiet record
    # First ensure admin PheDuyet row exists
    admin_pd = PheDuyet.query.filter_by(
        de_xuat_id=de_xuat.id, phong_duyet=PhongDuyet.ADMIN_TUYENHUAN.value
    ).first()
    if not admin_pd:
        admin_pd = PheDuyet(
            de_xuat_id=de_xuat.id,
            phong_duyet=PhongDuyet.ADMIN_TUYENHUAN.value,
            ket_qua=KetQuaDuyet.CHO_DUYET.value,
        )
        db.session.add(admin_pd)
        db.session.flush()

    # Update or create per-individual record
    kq_ct = KetQuaDuyetChiTiet.query.filter_by(
        phe_duyet_id=admin_pd.id, chi_tiet_id=ct.id
    ).first()
    if kq_ct:
        kq_ct.ket_qua = KetQuaDuyet.TU_CHOI.value
        kq_ct.ly_do = ly_do
    else:
        kq_ct = KetQuaDuyetChiTiet(
            phe_duyet_id=admin_pd.id,
            chi_tiet_id=ct.id,
            ket_qua=KetQuaDuyet.TU_CHOI.value,
            ly_do=ly_do,
        )
        db.session.add(kq_ct)

    # Reset admin_approved flag if it was set
    ct.admin_approved = False

    # Soft-remove ONLY this cá nhân/tập thể; the rest of the đề xuất continues.
    if not ct.bi_loai:
        ct.bi_loai = True
        ct.ly_do_loai = ly_do
        ct.phong_loai = PhongDuyet.ADMIN_TUYENHUAN.value
        ct.ngay_loai = datetime.utcnow()

    # Recompute status of the remaining active items
    from app.routes.approval import _recompute_de_xuat_status
    _recompute_de_xuat_status(de_xuat)
    db.session.commit()

    # Notify unit account
    unit_user = User.query.filter_by(don_vi_id=de_xuat.don_vi_id, role=Role.UNIT_USER).first()
    if unit_user:
        ho_ten = ct.quan_nhan.ho_ten if ct.quan_nhan else de_xuat.don_vi.ten_don_vi
        thong_bao = ThongBao(
            user_id=unit_user.id,
            de_xuat_id=de_xuat.id,
            chi_tiet_id=ct.id,
            loai='tu_choi',
            tieu_de=f'Ban Tuyên huấn (Admin) loại khỏi đề xuất: {ho_ten}',
            noi_dung=f'Lý do: {ly_do}. {ho_ten} đã bị loại khỏi đề xuất năm học {de_xuat.nam_hoc}. Các cá nhân/tập thể còn lại vẫn tiếp tục được xét duyệt.',
        )
        db.session.add(thong_bao)
        db.session.commit()

    ho_ten = ct.quan_nhan.ho_ten if ct.quan_nhan else de_xuat.don_vi.ten_don_vi
    log_action('admin_reject_individual', resource_type='chi_tiet', resource_id=ct.id,
               detail=f'{ho_ten} — lý do: {ly_do}')
    db.session.commit()
    flash(f'Đã loại "{ho_ten}" khỏi đề xuất và gửi thông báo về đơn vị.', 'warning')
    return redirect(url_for('admin.approval_tracking'))


@admin_bp.route('/batch-final-approve', methods=['POST'])
@login_required
@admin_required
def batch_final_approve():
    """Batch final approval for multiple nominations at once."""
    data = request.get_json()
    ids = data.get('ids', [])
    ghi_chu = data.get('ghi_chu', '').strip() or None

    if not ids:
        return jsonify({'success': False, 'message': 'Không có đề xuất nào được chọn.'}), 400

    approved_count = 0
    now = datetime.utcnow()

    for dx_id in ids:
        de_xuat = DeXuat.query.get(dx_id)
        if not de_xuat or de_xuat.trang_thai != TrangThaiDeXuat.HOI_DONG.value:
            continue

        # Verify departments approved, honoring auto-scope
        all_ok = True
        for dept_name in DEPT_NAMES:
            pd = PheDuyet.query.filter_by(
                de_xuat_id=dx_id, phong_duyet=dept_name
            ).first()
            if not pd:
                has_in_scope = any(not _is_auto_scope_approved(dept_name, ct.doi_tuong) for ct in de_xuat.chi_tiets)
                if has_in_scope:
                    all_ok = False
                    break
                continue
            if pd.ket_qua == KetQuaDuyet.TU_CHOI.value:
                all_ok = False
                break

        if not all_ok:
            continue

        # Update admin PheDuyet record
        admin_pd = PheDuyet.query.filter_by(
            de_xuat_id=dx_id, phong_duyet=PhongDuyet.ADMIN_TUYENHUAN.value
        ).first()
        if admin_pd:
            admin_pd.ket_qua = KetQuaDuyet.DONG_Y.value
            admin_pd.nguoi_duyet_id = current_user.id
            admin_pd.ngay_duyet = now
            admin_pd.ghi_chu = ghi_chu

        de_xuat.trang_thai = TrangThaiDeXuat.PHE_DUYET_CUOI.value

        # Create KhenThuong records
        for ct in de_xuat.chi_tiets:
            all_approved = True
            for dept_name in DEPT_NAMES:
                if not _is_individual_dept_approved(dx_id, ct, dept_name):
                    all_approved = False
                    break

            if all_approved:
                existing = KhenThuong.query.filter_by(
                    de_xuat_id=de_xuat.id, chi_tiet_id=ct.id
                ).first()
                if not existing:
                    khen_thuong = KhenThuong(
                        de_xuat_id=de_xuat.id,
                        chi_tiet_id=ct.id,
                        quan_nhan_id=ct.quan_nhan_id,
                        don_vi_id=de_xuat.don_vi_id,
                        ho_ten=ct.quan_nhan.ho_ten if ct.quan_nhan else de_xuat.don_vi.ten_don_vi,
                        cap_bac=ct.quan_nhan.cap_bac if ct.quan_nhan else None,
                        chuc_vu=ct.quan_nhan.chuc_vu if ct.quan_nhan else None,
                        doi_tuong=ct.doi_tuong,
                        loai_danh_hieu=ct.loai_danh_hieu,
                        nam_hoc=de_xuat.nam_hoc,
                        nguoi_duyet_id=current_user.id,
                        ngay_duyet=now,
                        ghi_chu=ghi_chu,
                    )
                    db.session.add(khen_thuong)
                ct.admin_approved = True

        approved_count += 1

    db.session.commit()
    log_action('batch_final_approve', detail=f'Phê duyệt cuối {approved_count} đề xuất (ids={ids})')
    db.session.commit()
    return jsonify({
        'success': True,
        'message': f'Đã phê duyệt cuối {approved_count} đề xuất.',
        'approved_count': approved_count,
    })


@admin_bp.route('/revoke-final/<int:id>', methods=['POST'])
@login_required
@admin_required
def revoke_final_approval(id):
    """Revoke final approval — deletes KhenThuong records, clears Hội đồng votes,
    resets admin_approved flags, and moves nomination back to HOI_DONG (Bảng 1)."""
    de_xuat = DeXuat.query.get_or_404(id)

    if de_xuat.trang_thai != TrangThaiDeXuat.PHE_DUYET_CUOI.value:
        flash('Đề xuất này chưa được phê duyệt cuối.', 'warning')
        return redirect(url_for('admin.reward_list'))

    # Delete all KhenThuong records for this nomination
    KhenThuong.query.filter_by(de_xuat_id=id).delete()

    # Delete all Hội đồng biểu quyết records for this nomination
    HoiDongBieuQuyet.query.filter_by(de_xuat_id=id).delete()

    # Reset per-individual admin_approved flags so they re-appear in Bảng 1
    for ct in de_xuat.chi_tiets:
        ct.admin_approved = False

    # Reset admin PheDuyet back to pending
    admin_pd = PheDuyet.query.filter_by(
        de_xuat_id=id, phong_duyet=PhongDuyet.ADMIN_TUYENHUAN.value
    ).first()
    if admin_pd:
        admin_pd.ket_qua = KetQuaDuyet.CHO_DUYET.value
        admin_pd.nguoi_duyet_id = None
        admin_pd.ngay_duyet = None
        admin_pd.ly_do = None
        admin_pd.ghi_chu = None

    # Revert status back to HOI_DONG (Bảng 1) — all 13 depts still approved
    de_xuat.trang_thai = TrangThaiDeXuat.HOI_DONG.value

    db.session.commit()
    flash(f'Đã thu hồi phê duyệt cuối cho đề xuất của {de_xuat.don_vi.ten_don_vi}.', 'success')
    return redirect(url_for('admin.reward_list'))


@admin_bp.route('/reward-stats')
@login_required
@admin_or_reward_viewer_required
def reward_stats():
    """Thống kê khen thưởng: tìm người đạt danh hiệu X trong N năm liên tiếp."""
    danh_hieu_list = [d[0] for d in db.session.query(KhenThuong.loai_danh_hieu)
                      .distinct().order_by(KhenThuong.loai_danh_hieu).all() if d[0]]

    selected_danh_hieu = request.args.get('danh_hieu', '').strip()
    so_nam = request.args.get('so_nam', '', type=str).strip()
    lien_tiep = request.args.get('lien_tiep', '1') == '1'
    results = []
    searched = False

    if selected_danh_hieu and so_nam:
        try:
            so_nam_int = int(so_nam)
        except ValueError:
            so_nam_int = 0

        if so_nam_int >= 1:
            searched = True

            def _nam_hoc_start(nh):
                try:
                    return int(str(nh).split('-')[0])
                except Exception:
                    return 0

            rows = db.session.query(KhenThuong.quan_nhan_id, KhenThuong.ho_ten,
                                    KhenThuong.don_vi_id, KhenThuong.nam_hoc).filter(
                KhenThuong.loai_danh_hieu == selected_danh_hieu,
                KhenThuong.quan_nhan_id.isnot(None)
            ).all()

            # Group by person
            by_person = {}
            names = {}
            don_vi_ids = {}
            for qn_id, ho_ten, dv_id, nam_hoc in rows:
                by_person.setdefault(qn_id, set()).add(nam_hoc)
                names[qn_id] = ho_ten
                don_vi_ids[qn_id] = dv_id

            for qn_id, years in by_person.items():
                if len(years) < so_nam_int:
                    continue
                years_sorted = sorted(list(years), key=_nam_hoc_start)
                starts = sorted([_nam_hoc_start(y) for y in years_sorted if _nam_hoc_start(y) > 0])

                if lien_tiep:
                    # Find max consecutive streak
                    streak = 1
                    max_streak = 1
                    for i in range(1, len(starts)):
                        if starts[i] == starts[i - 1] + 1:
                            streak += 1
                            max_streak = max(max_streak, streak)
                        else:
                            streak = 1
                    if max_streak < so_nam_int:
                        continue
                    # Find the streak years
                    streak_years = []
                    cur = [starts[0]]
                    for i in range(1, len(starts)):
                        if starts[i] == starts[i - 1] + 1:
                            cur.append(starts[i])
                        else:
                            if len(cur) >= so_nam_int:
                                streak_years.extend(cur)
                            cur = [starts[i]]
                    if len(cur) >= so_nam_int:
                        streak_years.extend(cur)
                    display_years = [y for y in years_sorted if _nam_hoc_start(y) in streak_years]
                    results.append({
                        'qn': QuanNhan.query.get(qn_id),
                        'ho_ten': names[qn_id],
                        'years': display_years,
                        'count': max_streak,
                        'don_vi': DonVi.query.get(don_vi_ids[qn_id]),
                    })
                else:
                    results.append({
                        'qn': QuanNhan.query.get(qn_id),
                        'ho_ten': names[qn_id],
                        'years': years_sorted,
                        'count': len(years_sorted),
                        'don_vi': DonVi.query.get(don_vi_ids[qn_id]),
                    })

            results.sort(key=lambda x: (x['don_vi'].ten_don_vi if x['don_vi'] else '', x['ho_ten']))

    return render_template('admin/reward_stats.html',
                           danh_hieu_list=danh_hieu_list,
                           selected_danh_hieu=selected_danh_hieu,
                           so_nam=so_nam,
                           lien_tiep=lien_tiep,
                           results=results,
                           searched=searched)


@admin_bp.route('/don-vi-stats')
@login_required
@admin_or_reward_viewer_required
def don_vi_stats():
    """Thống kê tỷ lệ đề xuất khen thưởng của từng đơn vị theo năm học."""
    # Năm học options
    nam_hoc_list = [r[0] for r in db.session.query(DeXuat.nam_hoc).filter(
        DeXuat.nam_hoc.isnot(None)
    ).distinct().order_by(DeXuat.nam_hoc.desc()).all() if r[0]]

    nam_hoc_filter = request.args.get('nam_hoc', '')
    if not nam_hoc_filter and nam_hoc_list:
        nam_hoc_filter = nam_hoc_list[0]

    stats = []
    if nam_hoc_filter:
        don_vi_list = DonVi.query.filter_by(is_active=True).order_by(DonVi.thu_tu, DonVi.ten_don_vi).all()

        # Pre-fetch DeXuatChiTiet for this nam_hoc (exclude NHAP drafts)
        submitted_statuses = [
            TrangThaiDeXuat.CHO_DUYET.value,
            TrangThaiDeXuat.DANG_DUYET.value,
            TrangThaiDeXuat.HOI_DONG.value,
            TrangThaiDeXuat.PHE_DUYET_CUOI.value,
        ]
        # Also include approved/rejected to count all submitted
        all_non_draft = submitted_statuses + [TrangThaiDeXuat.TU_CHOI.value]

        from sqlalchemy import and_
        rows = db.session.query(
            DeXuatChiTiet.quan_nhan_id,
            DeXuatChiTiet.loai_danh_hieu,
            DeXuat.don_vi_id,
        ).join(DeXuat, DeXuatChiTiet.de_xuat_id == DeXuat.id).filter(
            DeXuat.nam_hoc == nam_hoc_filter,
            DeXuat.trang_thai.in_(all_non_draft),
            DeXuatChiTiet.quan_nhan_id.isnot(None),
        ).all()

        # Build per-donvi lookup: {dv_id: {qn_id: set(loai_danh_hieu)}}
        from collections import defaultdict
        dv_qn_map = defaultdict(lambda: defaultdict(set))
        for qn_id, loai_dh, dv_id in rows:
            if loai_dh:
                dv_qn_map[dv_id][qn_id].add(loai_dh)

        for dv in don_vi_list:
            # Đếm quân nhân đang hoạt động (is_active=True, chưa bị xóa)
            total_qn = QuanNhan.query.filter_by(don_vi_id=dv.id, is_active=True).count()
            qn_map = dv_qn_map.get(dv.id, {})

            # Count CSTD (Chiến sĩ thi đua) and CSTT (Chiến sĩ tiên tiến)
            cstd_label = LoaiDanhHieu.CHIEN_SI_THI_DUA.value
            cstt_label = LoaiDanhHieu.CHIEN_SI_TIEN_TIEN.value

            # A person may be nominated for multiple danh hieu; count each separately
            cstd_count = sum(1 for danh_hieus in qn_map.values() if cstd_label in danh_hieus)
            cstt_count = sum(1 for danh_hieus in qn_map.values() if cstt_label in danh_hieus)
            # Total unique nominated (either CSTD or CSTT)
            total_nom = sum(1 for danh_hieus in qn_map.values()
                            if cstd_label in danh_hieus or cstt_label in danh_hieus)

            def pct(num, denom):
                if not denom:
                    return None
                return round(num * 100 / denom, 1)

            stats.append({
                'don_vi': dv,
                'total_qn': total_qn,
                'total_nom': total_nom,
                'cstd_count': cstd_count,
                'cstt_count': cstt_count,
                'total_pct': pct(total_nom, total_qn),
                'cstd_pct': pct(cstd_count, total_qn),
                'cstt_pct': pct(cstt_count, total_qn),
            })

    return render_template('admin/don_vi_stats.html',
                           nam_hoc_list=nam_hoc_list,
                           nam_hoc_filter=nam_hoc_filter,
                           stats=stats)


@admin_bp.route('/reward-list')
@login_required
@admin_or_reward_viewer_required
def reward_list():
    from app.models.nomination import DanhHieu as _DanhHieu, DeXuat as _DeXuat
    from app.models.hoi_dong import HoiDongBieuQuyet, KET_QUA_KHONG_DONG_Y
    from sqlalchemy import collate as _collate, func

    page           = request.args.get('page', 1, type=int)
    nam_hoc_filter = request.args.get('nam_hoc', '')
    unit_filter    = request.args.get('unit', '')
    danh_hieu_filter = request.args.get('danh_hieu', '')
    search_query   = request.args.get('q', '').strip()

    # ── Bảng 3 chính: KhenThuong paginated ──────────────────────────────────
    _thu_tu_subq = (
        db.session.query(_DanhHieu.thu_tu)
        .filter(
            _collate(_DanhHieu.ten_danh_hieu, 'utf8mb4_unicode_ci') ==
            _collate(KhenThuong.loai_danh_hieu, 'utf8mb4_unicode_ci')
        )
        .correlate(KhenThuong)
        .scalar_subquery()
    )
    query = KhenThuong.query.join(DonVi, KhenThuong.don_vi_id == DonVi.id)
    if nam_hoc_filter:
        query = query.filter(KhenThuong.nam_hoc == nam_hoc_filter)
    if unit_filter:
        query = query.filter(DonVi.ten_don_vi == unit_filter)
    if danh_hieu_filter:
        query = query.filter(KhenThuong.loai_danh_hieu == danh_hieu_filter)
    if search_query:
        query = query.filter(KhenThuong.ho_ten.ilike(f'%{search_query}%'))
    rewards = query.order_by(
        KhenThuong.nam_hoc.desc(), _thu_tu_subq,
        DonVi.ten_don_vi, KhenThuong.ho_ten.asc(),
    ).paginate(page=page, per_page=20, error_out=False)

    # ── Filter options ───────────────────────────────────────────────────────
    _nh_kt = {n[0] for n in db.session.query(KhenThuong.nam_hoc).distinct().all() if n[0]}
    _nh_dx = {n[0] for n in db.session.query(_DeXuat.nam_hoc).distinct().all() if n[0]}
    nam_hoc_list = sorted(_nh_kt | _nh_dx, reverse=True)

    unit_names = [u[0] for u in db.session.query(DonVi.ten_don_vi)
                  .join(KhenThuong, KhenThuong.don_vi_id == DonVi.id)
                  .distinct().order_by(DonVi.ten_don_vi).all()]

    danh_hieu_list = [d[0] for d in db.session.query(KhenThuong.loai_danh_hieu).distinct().all()]

    # ── Stats: 2 queries thay vì N+1 ────────────────────────────────────────
    total_rewards = KhenThuong.query.count()
    stats_by_danh_hieu = {
        dh: cnt for dh, cnt in
        db.session.query(KhenThuong.loai_danh_hieu, func.count(KhenThuong.id))
        .group_by(KhenThuong.loai_danh_hieu).all()
    }

    # ── Bảng 2 ──────────────────────────────────────────────────────────────
    # Bảng 1 (pending_final_nominations) đã xóa — không cần thiết, gây load chậm
    phe_duyet_cuoi_items      = _get_phe_duyet_cuoi_items(nam_hoc=nam_hoc_filter)

    # ── Bảng 3b: Bị từ chối ─────────────────────────────────────────────────
    # ★ 1 query lấy tất cả admin_final rejected votes
    rej_votes = HoiDongBieuQuyet.query.filter_by(
        vai_tro='admin_final', ket_qua=KET_QUA_KHONG_DONG_Y
    ).all()

    if rej_votes:
        rej_ct_ids = [rv.chi_tiet_id for rv in rej_votes]

        # ★ Batch load chi_tiet + eager load de_xuat
        from sqlalchemy.orm import joinedload
        rej_ct_map = {
            ct.id: ct for ct in
            DeXuatChiTiet.query
            .filter(DeXuatChiTiet.id.in_(rej_ct_ids))
            .options(joinedload(DeXuatChiTiet.de_xuat))
            .all()
        }

        # ★ Batch check KhenThuong đã confirm
        confirmed_ct_ids = {
            r[0] for r in
            db.session.query(KhenThuong.chi_tiet_id)
            .filter(KhenThuong.chi_tiet_id.in_(rej_ct_ids))
            .all()
        }

        rejected_items = []
        for rv in rej_votes:
            ct_r = rej_ct_map.get(rv.chi_tiet_id)
            if not ct_r:
                continue
            dx_r = ct_r.de_xuat
            if not dx_r:
                continue
            if nam_hoc_filter and dx_r.nam_hoc != nam_hoc_filter:
                continue
            if rv.chi_tiet_id in confirmed_ct_ids:
                continue
            rejected_items.append({'ct': ct_r, 'dx': dx_r, 'admin_vote': rv})
    else:
        rejected_items = []

    _rej_ct_ids_set = {r['ct'].id for r in rejected_items}
    khong_dong_y_items = [
        r for r in phe_duyet_cuoi_items
        if r['admin_final_vote'] is not None
        and r['admin_final_vote'].ket_qua == KET_QUA_KHONG_DONG_Y
        and r['ct'].id not in _rej_ct_ids_set
    ]

    # ── CSTD >= 3 năm ────────────────────────────────────────────────────────
    cstd_rows = db.session.query(KhenThuong.quan_nhan_id, KhenThuong.nam_hoc).filter(
        KhenThuong.loai_danh_hieu == LoaiDanhHieu.CHIEN_SI_THI_DUA.value,
        KhenThuong.quan_nhan_id.isnot(None)
    ).all()

    by_person = {}
    for qn_id, nam_hoc in cstd_rows:
        by_person.setdefault(qn_id, set()).add(nam_hoc)

    # ★ Batch load QuanNhan — 1 query thay vì N queries
    eligible_ids = [qn_id for qn_id, years in by_person.items() if len(years) >= 3]
    qn_map = {}
    if eligible_ids:
        qn_map = {
            qn.id: qn for qn in
            QuanNhan.query.filter(QuanNhan.id.in_(eligible_ids)).all()
        }

    def _nam_hoc_start(nh):
        try:
            return int(str(nh).split('-')[0])
        except Exception:
            return 0

    cstd_non_consecutive = []
    cstd_consecutive     = []

    for qn_id, years in by_person.items():
        if len(years) < 3:
            continue
        qn = qn_map.get(qn_id)
        if not qn:
            continue
        years_sorted = sorted(list(years), key=_nam_hoc_start)
        cstd_non_consecutive.append({
            'qn': qn, 'years': years_sorted, 'count': len(years_sorted)
        })
        starts = sorted([s for s in (_nam_hoc_start(y) for y in years_sorted) if s > 0])
        streak = max_streak = 1
        for i in range(1, len(starts)):
            streak = streak + 1 if starts[i] == starts[i - 1] + 1 else 1
            max_streak = max(max_streak, streak)
        if max_streak >= 3:
            cstd_consecutive.append({
                'qn': qn, 'years': years_sorted, 'max_streak': max_streak
            })

    return render_template('admin/reward_list.html',
                           rewards=rewards,
                           nam_hoc_filter=nam_hoc_filter,
                           unit_filter=unit_filter,
                           danh_hieu_filter=danh_hieu_filter,
                           search_query=search_query,
                           nam_hoc_list=nam_hoc_list,
                           unit_names=unit_names,
                           danh_hieu_list=danh_hieu_list,
                           total_rewards=total_rewards,
                           stats_by_danh_hieu=stats_by_danh_hieu,
                           can_admin_action=current_user.is_admin,
                           current_user_vai_tro=current_user.hoi_dong_vai_tro,
                           can_view_pending_final=(current_user.is_admin or current_user.is_reward_viewer),
                           phe_duyet_cuoi_items=phe_duyet_cuoi_items,
                           khong_dong_y_items=khong_dong_y_items,
                           rejected_items=rejected_items,
                           HOI_DONG_VAI_TRO=HOI_DONG_VAI_TRO,
                           HOI_DONG_VAI_TRO_DISPLAY=HOI_DONG_VAI_TRO_DISPLAY,
                           cstd_non_consecutive=cstd_non_consecutive,
                           cstd_consecutive=cstd_consecutive)


# ─────────────────────────────────────────────────────────────────────────────

def _get_pending_final_individuals(nam_hoc=None):
    """Bảng 1 — Batch load toàn bộ, tránh N×M queries."""
    from sqlalchemy.orm import joinedload, subqueryload

    q = (DeXuat.query
         .filter(DeXuat.trang_thai != TrangThaiDeXuat.NHAP.value)
         .options(
             subqueryload(DeXuat.chi_tiets)
         ))
    if nam_hoc:
        q = q.filter(DeXuat.nam_hoc == nam_hoc)
    nominations = q.order_by(DeXuat.ngay_gui.desc()).all()

    if not nominations:
        return []

    # ★ Batch load tất cả KetQuaDuyetChiTiet liên quan
    dx_ids = [dx.id for dx in nominations]
    ct_ids = [ct.id for dx in nominations for ct in dx.chi_tiets]

    if not ct_ids:
        return []

    # Lấy tất cả phe_duyet của các dx này
    phe_duyet_list = PheDuyet.query.filter(PheDuyet.de_xuat_id.in_(dx_ids)).all()
    pd_map = {}  # {de_xuat_id: {phong_duyet: PheDuyet}}
    for pd in phe_duyet_list:
        pd_map.setdefault(pd.de_xuat_id, {})[pd.phong_duyet] = pd

    pd_ids = [pd.id for pd in phe_duyet_list]

    # Batch load KetQuaDuyetChiTiet
    kq_list = KetQuaDuyetChiTiet.query.filter(
        KetQuaDuyetChiTiet.phe_duyet_id.in_(pd_ids),
        KetQuaDuyetChiTiet.chi_tiet_id.in_(ct_ids),
    ).all() if pd_ids else []

    # Build lookup: {(phe_duyet_id, chi_tiet_id): KetQuaDuyetChiTiet}
    kq_map = {(kq.phe_duyet_id, kq.chi_tiet_id): kq for kq in kq_list}

    pending = []
    for dx in nominations:
        for ct in dx.chi_tiets:
            all_approved = True
            for dept_name in DEPT_NAMES:
                if not _is_individual_dept_approved(dx.id, ct, dept_name):
                    all_approved = False
                    break
            if all_approved == False:
                continue
            if ct.admin_approved or ct.bi_loai:
                continue
            pending.append({'dx': dx, 'ct': ct})
    return pending

  


def _get_phe_duyet_cuoi_items(nam_hoc=None):
    """Bảng 2 — Batch load HoiDongBieuQuyet và KhenThuong."""
    from app.models.hoi_dong import HOI_DONG_VAI_TRO, KET_QUA_DONG_Y
    from app.models.nomination import DanhHieu
    from sqlalchemy.orm import subqueryload, joinedload

    q = (DeXuat.query
         .filter(DeXuat.trang_thai != TrangThaiDeXuat.NHAP.value)
         .join(DonVi, DeXuat.don_vi_id == DonVi.id)
         .options(
             joinedload(DeXuat.don_vi),
             subqueryload(DeXuat.chi_tiets).joinedload(DeXuatChiTiet.quan_nhan),
         ))
    if nam_hoc:
        q = q.filter(DeXuat.nam_hoc == nam_hoc)
    nominations = q.order_by(DonVi.ten_don_vi).all()

    # Lọc chi_tiets hợp lệ (admin_approved hoặc PHE_DUYET_CUOI)
    eligible_cts = [
        ct for dx in nominations for ct in dx.chi_tiets
        if ct.admin_approved or ct.trang_thai == TrangThaiChiTiet.PHE_DUYET_CUOI.value
    ]
    if not eligible_cts:
        return []

    ct_ids = [ct.id for ct in eligible_cts]

    # ★ Batch load HoiDongBieuQuyet — 1 query
    bq_list = HoiDongBieuQuyet.query.filter(
        HoiDongBieuQuyet.chi_tiet_id.in_(ct_ids)
    ).all()
    # {chi_tiet_id: {vai_tro: HoiDongBieuQuyet}}
    bq_map = {}
    for bq in bq_list:
        bq_map.setdefault(bq.chi_tiet_id, {})[bq.vai_tro] = bq

    # ★ Batch check KhenThuong đã confirm — 1 query
    confirmed_ct_ids = {
        r[0] for r in
        db.session.query(KhenThuong.chi_tiet_id)
        .filter(KhenThuong.chi_tiet_id.in_(ct_ids))
        .all()
    }

    # Award order map
    danh_hieu_order = {dh.ten_danh_hieu: dh.thu_tu for dh in DanhHieu.query.all()}

    # Build ct → dx map
    ct_dx_map = {ct.id: dx for dx in nominations for ct in dx.chi_tiets}

    rows = []
    for ct in eligible_cts:
        dx        = ct_dx_map[ct.id]
        ct_votes  = bq_map.get(ct.id, {})

        votes         = {vai_tro: ct_votes.get(vai_tro) for vai_tro in HOI_DONG_VAI_TRO}
        ct_all_dong_y = all(bq and bq.ket_qua == KET_QUA_DONG_Y for bq in votes.values())
        ct_all_voted  = all(bq is not None for bq in votes.values())

        rows.append({
            'dx':             dx,
            'ct':             ct,
            'votes':          votes,
            'voted_count':    sum(1 for bq in votes.values() if bq is not None),
            'all_voted_dong_y': ct_all_dong_y,
            'all_voted':      ct_all_voted,
            'is_confirmed':   ct.id in confirmed_ct_ids,
            'admin_final_vote': ct_votes.get('admin_final'),
            '_sort_award':    danh_hieu_order.get(ct.loai_danh_hieu, 999),
            '_sort_unit':     dx.don_vi.ten_don_vi if dx.don_vi else '',
        })

    rows.sort(key=lambda r: (r['_sort_award'], r['_sort_unit']))
    return rows





@admin_bp.route('/tracking/export-excel')
@login_required
@admin_required
def export_tracking_excel():
    """Export bảng theo dõi quy trình phê duyệt ra Excel chi tiết."""
    import datetime as _dt
    status_filter = request.args.get('status', '')
    unit_filter = request.args.get('unit', '')
    danh_hieu_filter = request.args.get('danh_hieu', '')
    search_query = request.args.get('q', '').strip()
    scope_filter = request.args.get('scope', '')
    nam_hoc_filter = request.args.get('nam_hoc', '')
    view_mode = request.args.get('view', 'detail')

    query = DeXuat.query.filter(DeXuat.trang_thai != TrangThaiDeXuat.NHAP.value)
    if nam_hoc_filter:
        query = query.filter(DeXuat.nam_hoc == nam_hoc_filter)
    if status_filter:
        query = query.filter(DeXuat.trang_thai == status_filter)
    if unit_filter:
        query = query.join(DonVi).filter(DonVi.ten_don_vi == unit_filter)
    else:
        # Join DonVi even if not filtering, for sorting
        query = query.join(DonVi, DeXuat.don_vi_id == DonVi.id)
    
    # Sort by unit hierarchy (Phòng > Khoa > Đơn vị), then by date
    nominations = query.order_by(DonVi.thu_tu.asc(), DeXuat.nam_hoc.desc(), DeXuat.ngay_gui.desc()).all()

    approved_ct_ids = set(row[0] for row in db.session.query(KhenThuong.chi_tiet_id).all())

    wb = Workbook()
    ws = wb.active
    ws.title = 'Theo doi phe duyet'

    from openpyxl.styles import Font as _Font, Alignment as _Align, PatternFill as _Fill, Border as _Border, Side as _Side
    bold = _Font(name='Times New Roman', bold=True, size=11)
    normal = _Font(name='Times New Roman', size=10)
    center = _Align(horizontal='center', vertical='center', wrap_text=True)
    left = _Align(horizontal='left', vertical='center', wrap_text=True)
    thin = _Border(left=_Side(style='thin'), right=_Side(style='thin'),
                   top=_Side(style='thin'), bottom=_Side(style='thin'))
    navy_fill = _Fill(start_color='1F4E79', end_color='1F4E79', fill_type='solid')
    white_bold = _Font(name='Times New Roman', bold=True, size=10, color='FFFFFF')
    unit_fill = _Fill(start_color='D9E1F2', end_color='D9E1F2', fill_type='solid')

    # Title
    col_count = 12 if view_mode == 'detail' else 8
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=col_count)
    ws.cell(1, 1).value = 'THEO DÕI QUY TRÌNH PHÊ DUYỆT'
    ws.cell(1, 1).font = _Font(name='Times New Roman', bold=True, size=13)
    ws.cell(1, 1).alignment = center

    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=col_count)
    filter_info = []
    if nam_hoc_filter: filter_info.append(f'Năm học: {nam_hoc_filter}')
    if status_filter: filter_info.append(f'Trạng thái: {status_filter}')
    if unit_filter: filter_info.append(f'Đơn vị: {unit_filter}')
    ws.cell(2, 1).value = ('  |  '.join(filter_info) if filter_info else 'Tất cả dữ liệu') + \
        f'  —  Xuất ngày: {_dt.datetime.now().strftime("%d/%m/%Y %H:%M")}'
    ws.cell(2, 1).font = _Font(name='Times New Roman', italic=True, size=10)
    ws.cell(2, 1).alignment = center

    # Headers
    if view_mode == 'detail':
        headers = ['STT', 'Đơn vị', 'Họ và tên', 'CCCD', 'Cấp bậc', 'Chức vụ',
                   'Đối tượng', 'Danh hiệu', 'Năm học', 'Trạng thái', 'Ngày gửi', 'Ghi chú']
        widths =  [5,    28,     22,       14,      12,      20,
                   16,        22,        12,      16,        12,       30]
    else:
        headers = ['STT', 'Đơn vị', 'Họ và tên', 'Danh hiệu', 'Năm học', 'Trạng thái', 'Ngày gửi', 'Ghi chú']
        widths  = [5,     28,       22,           22,           12,        16,             12,          30]

    for col_idx, (h, w) in enumerate(zip(headers, widths), 1):
        c = ws.cell(row=3, column=col_idx, value=h)
        c.font = white_bold
        c.fill = navy_fill
        c.alignment = center
        c.border = thin
        ws.column_dimensions[get_column_letter(col_idx)].width = w

    row_num = 4
    stt = 0
    for dx in nominations:
        unit_name = dx.don_vi.ten_don_vi if dx.don_vi else '—'
        for ct in dx.chi_tiets:
            if danh_hieu_filter and ct.loai_danh_hieu != danh_hieu_filter:
                continue
            if scope_filter == 'quan_luc' and ct.doi_tuong not in BAN_QUANLUC_DOI_TUONG:
                continue
            if scope_filter == 'can_bo' and ct.doi_tuong in BAN_QUANLUC_DOI_TUONG:
                continue
            qn = ct.quan_nhan
            ho_ten = qn.ho_ten if qn else (ct.ten_don_vi_de_xuat or '—')
            if search_query and search_query.lower() not in ho_ten.lower():
                continue

            stt += 1
            trang_thai = dx.trang_thai
            if ct.id in approved_ct_ids:
                trang_thai = 'Đã khen thưởng'

            if view_mode == 'detail':
                row_data = [
                    stt, unit_name, ho_ten,
                    (qn.can_cuoc_cong_dan if qn else ''),
                    (qn.cap_bac if qn else ''),
                    (qn.chuc_vu if qn else ''),
                    (ct.doi_tuong or ''),
                    ct.loai_danh_hieu or '',
                    dx.nam_hoc or '',
                    trang_thai,
                    dx.ngay_gui.strftime('%d/%m/%Y') if dx.ngay_gui else '',
                    ct.ghi_chu or '',
                ]
            else:
                row_data = [
                    stt, unit_name, ho_ten,
                    ct.loai_danh_hieu or '',
                    dx.nam_hoc or '',
                    trang_thai,
                    dx.ngay_gui.strftime('%d/%m/%Y') if dx.ngay_gui else '',
                    ct.ghi_chu or '',
                ]

            for col_idx, val in enumerate(row_data, 1):
                c = ws.cell(row=row_num, column=col_idx, value=val)
                c.font = normal
                c.border = thin
                c.alignment = center if col_idx == 1 else left
            row_num += 1

    # Summary
    ws.merge_cells(start_row=row_num, start_column=1, end_row=row_num, end_column=col_count)
    c = ws.cell(row=row_num, column=1, value=f'Tổng cộng: {stt} cá nhân')
    c.font = bold
    c.alignment = left

    ws.page_setup.paperSize = 9
    ws.page_setup.orientation = 'landscape'
    ws.page_setup.fitToPage = True
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    ws.freeze_panes = 'A4'

    ws.protection.sheet = True
    ws.protection.password = 'bth123'

    output = BytesIO()
    wb.save(output)
    output.seek(0)
    fname_parts = ['TheoDoiPheduyet']
    if nam_hoc_filter:
        fname_parts.append(nam_hoc_filter.replace('-', '_'))
    fname_parts.append(_dt.datetime.now().strftime('%d%m%Y'))
    return send_file(output, as_attachment=True,
                     download_name='_'.join(fname_parts) + '.xlsx',
                     mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')

@admin_bp.route('/tracking/export-word')
@login_required
@admin_required
def export_tracking_word():
    """Export bảng theo dõi quy trình phê duyệt ra Word (theo mẫu khen thưởng)."""
    import datetime as _dt
    from io import BytesIO
    from datetime import date
    from docx import Document
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.text import WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # Lưu ý: KHÔNG dùng db.session.expire_all() ở đây — buộc query lại DB cho MỌI
    # thuộc tính truy cập sau đó (ct.quan_nhan, dx.don_vi, ...), có thể gây ra
    # hàng chục/trăm round-trip DB dư thừa thay vì eager-load 1 lần bên dưới.

    # ── Tham số lọc (giống export_tracking_excel) ────────────────────────────
    status_filter    = request.args.get('status', '')
    unit_filter      = request.args.get('unit', '')
    danh_hieu_filter = request.args.get('danh_hieu', '')
    search_query     = request.args.get('q', '').strip()
    scope_filter     = request.args.get('scope', '')
    nam_hoc_filter   = request.args.get('nam_hoc', '')

    query = DeXuat.query.filter(DeXuat.trang_thai != TrangThaiDeXuat.NHAP.value).options(
        joinedload(DeXuat.don_vi),
        subqueryload(DeXuat.chi_tiets).joinedload(DeXuatChiTiet.quan_nhan),
    )
    if nam_hoc_filter:
        query = query.filter(DeXuat.nam_hoc == nam_hoc_filter)
    if status_filter:
        query = query.filter(DeXuat.trang_thai == status_filter)
    if unit_filter:
        query = query.join(DonVi).filter(DonVi.ten_don_vi == unit_filter)
    else:
        query = query.join(DonVi, DeXuat.don_vi_id == DonVi.id)

    nominations = query.order_by(
        DonVi.thu_tu.asc(), DeXuat.nam_hoc.desc(), DeXuat.ngay_gui.desc()
    ).all()

    approved_ct_ids = set(
        row[0] for row in db.session.query(KhenThuong.chi_tiet_id).all()
    )

    # ── Gom chi tiết theo danh hiệu ──────────────────────────────────────────
    ds_quyet_thang   = []
    ds_tien_tien_dv  = []
    ds_chien_si_tdcs = []
    ds_chien_si_tt   = []
    ds_khac          = {}
    seen_ids         = set()
    units_by_type = {}
    for loai in LoaiDonVi:
        units_by_type[loai] = DonVi.query.filter_by(loai_don_vi=loai) \
            .order_by(DonVi.thu_tu).all()

    _sort_map = {}
    for loai_idx, (loai, units) in enumerate(units_by_type.items()):
        for unit in units:
            _sort_map[unit.id] = (loai_idx, unit.thu_tu or 0, unit.ten_don_vi or '')

    # ── Hàm sort key ────────────────────────────────────────────────────────────
    def _unit_sort_key(ct_dx_tuple):
        ct, dx = ct_dx_tuple
        don_vi_id = dx.don_vi_id if dx else None
        base = _sort_map.get(don_vi_id, (999, 999, 'zzz'))
        ho_ten = ''
        if ct:
            ho_ten = (ct.quan_nhan.ho_ten if ct.quan_nhan else ct.ten_don_vi_de_xuat) or ''
        return (*base, ho_ten.lower())
    for dx in nominations:
        for ct in dx.chi_tiets:
            if ct.id in seen_ids:
                continue
            seen_ids.add(ct.id)

            if danh_hieu_filter and ct.loai_danh_hieu != danh_hieu_filter:
                continue
            if scope_filter == 'quan_luc' and ct.doi_tuong not in BAN_QUANLUC_DOI_TUONG:
                continue
            if scope_filter == 'can_bo' and ct.doi_tuong in BAN_QUANLUC_DOI_TUONG:
                continue

            ho_ten = (ct.quan_nhan.ho_ten if ct.quan_nhan
                      else (ct.ten_don_vi_de_xuat or ''))
            if search_query and search_query.lower() not in ho_ten.lower():
                continue

            dh = (ct.loai_danh_hieu or '').strip()
            if   dh == 'Đơn vị quyết thắng':  ds_quyet_thang.append((ct, dx))
            elif dh == 'Đơn vị tiên tiến':     ds_tien_tien_dv.append((ct, dx))
            elif dh == 'Chiến sĩ thi đua':     ds_chien_si_tdcs.append((ct, dx))
            elif dh == 'Chiến sĩ tiên tiến':   ds_chien_si_tt.append((ct, dx))
            else: ds_khac.setdefault(dh, []).append((ct, dx))

            # Sort từng nhóm
   

    # ── Sort từng nhóm ──────────────────────────────────────────────────────────
    ds_quyet_thang   = sorted(ds_quyet_thang,   key=_unit_sort_key)
    ds_tien_tien_dv  = sorted(ds_tien_tien_dv,  key=_unit_sort_key)
    ds_chien_si_tdcs = sorted(ds_chien_si_tdcs, key=_unit_sort_key)
    ds_chien_si_tt   = sorted(ds_chien_si_tt,   key=_unit_sort_key)
    ds_khac          = {dh: sorted(lst, key=_unit_sort_key) for dh, lst in ds_khac.items()}

    title_nam_hoc = nam_hoc_filter or 'TẤT CẢ NĂM HỌC'
    





    # ── Tạo docx nhanh bằng XML template (thay python-docx ~700x nhanh hơn) ──────
    from app.utils.docx_fast import (
        cm_to_twips, _para, _build_table, _data_row,
        _build_document_xml, build_docx,
    )
    import datetime as _dt

    SP = 'xml:space="preserve"'

    def _section_heading(text):
        return _para(text, bold=True, size_pt=12, space_before=120, space_after=40)

    def _empty_notice():
        return _para('(Không có)', italic=True, size_pt=10, space_before=40, space_after=40)

    # ── Chiều rộng cột cá nhân: STT|Họ tên|Cấp bậc|Chức vụ|Đơn vị|Năm học|Tóm tắt ──
    CN_WIDTHS = [
        cm_to_twips(0.8),   # STT
        cm_to_twips(3.5),   # Họ tên
        cm_to_twips(2.0),   # Cấp bậc
        cm_to_twips(2.5),   # Chức vụ
        cm_to_twips(2.5),   # Đơn vị
       
        cm_to_twips(5.5),   # Tóm tắt
    ]
    CN_HEADERS = ['STT', 'Họ và tên', 'Cấp bậc', 'Chức vụ', 'Đơn vị', 'Năm học', 'Tóm tắt thành tích']

    # ── Chiều rộng cột tập thể: STT|Tên đơn vị|Đề xuất đơn vị|Năm học ──
    TT_WIDTHS = [
        cm_to_twips(0.8),
        cm_to_twips(4.5),
        cm_to_twips(11.5),
       
    ]
    TT_HEADERS = ['STT', 'Tên đơn vị', 'Đề xuất của đơn vị']

    # ── TieuChi map (batch load 1 lần) ───────────────────────────────────────────
    from app.models.nomination import TieuChi as _TieuChi
    _tieu_chi_map_all = {tc.ma_truong: tc.ten for tc in _TieuChi.query.all()}

    def _build_tomtat(ct):
        """Tóm tắt thành tích ngắn gọn cho cột Word."""
        parts = []
        if ct.muc_do_hoan_thanh:
            parts.append(ct.muc_do_hoan_thanh)
        if ct.diem_tong_ket:
            parts.append(f'ĐTK: {ct.diem_tong_ket}')
        if ct.ket_qua_ren_luyen:
            parts.append(f'RL: {ct.ket_qua_ren_luyen}')
        if ct.xep_loai_dang_vien:
            parts.append(f'ĐV: {ct.xep_loai_dang_vien}')
        if ct.thanh_tich_ca_nhan_khac:
            parts.append(ct.thanh_tich_ca_nhan_khac[:80])
        return '; '.join(parts)

    def _build_tt_criteria(ct):
        """Tiêu chí tập thể dạng text cho cột Word."""
        td = ct.tap_the_dict or {}
        lines = []
        for k, v in td.items():
            if v and str(v).strip() not in ('', '0', 'None'):
                label = _tieu_chi_map_all.get(k, k)
                lines.append(f'- {label}: {v}')
        if ct.muc_do_hoan_thanh:
            lines.insert(0, f'- Mức độ HT: {ct.muc_do_hoan_thanh}')
        return '\n'.join(lines)

    def _cn_rows(items):
        rows_xml = []
        for i, (ct, dx) in enumerate(items, 1):
            qn = ct.quan_nhan
            row_cells = [
                (str(i), False, 'center'),
                (qn.ho_ten if qn else '', True, 'left'),
                (qn.cap_bac if qn else '', False, 'left'),
                (qn.chuc_vu if qn else '', False, 'left'),
                (dx.don_vi.ten_don_vi if dx.don_vi else '', False, 'left'),
                (_build_tomtat(ct), False, 'left'),
            ]
            shade = 'F8F9FA' if i % 2 == 0 else None
            rows_xml.append(_data_row(row_cells, CN_WIDTHS, size_pt=9, shade=shade))
        return rows_xml

    def _tt_rows(items):
        rows_xml = []
        for i, (ct, dx) in enumerate(items, 1):
            criteria_text = _build_tt_criteria(ct)
            row_cells = [
                (str(i), False, 'center'),
                (ct.ten_don_vi_de_xuat or (dx.don_vi.ten_don_vi if dx.don_vi else ''), True, 'left'),
                (criteria_text, False, 'left'),
                (dx.nam_hoc or '', False, 'center'),
            ]
            shade = 'F8F9FA' if i % 2 == 0 else None
            rows_xml.append(_data_row(row_cells, TT_WIDTHS, size_pt=9, shade=shade))
        return rows_xml

    # ── Xây dựng nội dung tài liệu ───────────────────────────────────────────────
    body = []

    # Tiêu đề
    today_str = _dt.date.today().strftime('%d/%m/%Y')
    body.append(_para('TRƯỜNG SĨ QUAN CHÍNH TRỊ', bold=True, size_pt=12, align='center', space_before=0, space_after=20))
    body.append(_para(f'THEO DÕI PHÊ DUYỆT KHEN THƯỞNG — {title_nam_hoc}', bold=True, size_pt=14, align='center', space_before=60, space_after=20))
    body.append(_para(f'(Xuất lúc {_dt.datetime.now().strftime("%H:%M")} ngày {today_str})', italic=True, size_pt=10, align='center', space_before=0, space_after=120))

    def _add_section(label, items, is_tap_the=False):
        if not items:
            return
        body.append(_section_heading(label))
        if is_tap_the:
            rows_xml = _tt_rows(items)
            total = f'Tổng cộng: {len(items)} đơn vị'
            body.append(_build_table(TT_HEADERS, rows_xml, TT_WIDTHS, total_label=total, size_pt=9))
        else:
            rows_xml = _cn_rows(items)
            total = f'Tổng cộng: {len(items)} người'
            body.append(_build_table(CN_HEADERS, rows_xml, CN_WIDTHS, total_label=total, size_pt=9))
        body.append(_para('', space_before=60, space_after=0))

    _add_section('I. DANH HIỆU ĐƠN VỊ QUYẾT THẮNG', ds_quyet_thang, is_tap_the=True)
    _add_section('II. DANH HIỆU ĐƠN VỊ TIÊN TIẾN', ds_tien_tien_dv, is_tap_the=True)
    _add_section('III. CHIẾN SĨ THI ĐUA', ds_chien_si_tdcs)
    _add_section('IV. CHIẾN SĨ TIÊN TIẾN', ds_chien_si_tt)
    for extra_dh, extra_items in ds_khac.items():
        _add_section(extra_dh, extra_items)

    # Footer
    body.append(_para(f'(Xuất lúc {_dt.datetime.now().strftime("%H:%M ngày %d/%m/%Y")})',
                      italic=True, size_pt=9, align='right', space_before=120, space_after=0))

    doc_xml = _build_document_xml(body, margin_left=2016, margin_right=720, margin_top=1440, margin_bottom=1440)
    buf = build_docx(doc_xml)

    fname_parts = ['TheoDoiPheduyet']
    if nam_hoc_filter:
        fname_parts.append(nam_hoc_filter.replace('-', '_'))
    fname_parts.append(_dt.datetime.now().strftime('%d%m%Y'))
    filename = '_'.join(fname_parts) + '.docx'

    response = send_file(
        buf, as_attachment=True, download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response.set_cookie(
        'export_done', '1',
        max_age=600, httponly=False, samesite='Lax', path='/'
    )
    return response


@admin_bp.route('/tracking/export-word-less')
@login_required
@admin_required
def export_tracking_word_less():
    """Export bảng theo dõi quy trình phê duyệt ra Word (theo mẫu khen thưởng)."""
    import datetime as _dt
    import zipfile
    import os
    import binascii
    import hashlib
    from io import BytesIO
    from datetime import date
    
    # --- Tham số lọc (giống export_tracking_excel) ---
    status_filter    = request.args.get('status', '')
    unit_filter      = request.args.get('unit', '')
    danh_hieu_filter = request.args.get('danh_hieu', '')
    search_query     = request.args.get('q', '').strip()
    scope_filter     = request.args.get('scope', '')
    nam_hoc_filter   = request.args.get('nam_hoc', '')

    query = DeXuat.query.filter(DeXuat.trang_thai != TrangThaiDeXuat.NHAP.value).options(
        joinedload(DeXuat.don_vi),
        subqueryload(DeXuat.chi_tiets).joinedload(DeXuatChiTiet.quan_nhan),
    )
    if nam_hoc_filter:
        query = query.filter(DeXuat.nam_hoc == nam_hoc_filter)
    if status_filter:
        query = query.filter(DeXuat.trang_thai == status_filter)
    if unit_filter:
        query = query.join(DonVi).filter(DonVi.ten_don_vi == unit_filter)
    else:
        query = query.join(DonVi, DeXuat.don_vi_id == DonVi.id)

    nominations = query.order_by(
        DonVi.thu_tu.asc(), DeXuat.nam_hoc.desc(), DeXuat.ngay_gui.desc()
    ).all()

    # --- Gom chi tiết theo danh hiệu ---
    ds_quyet_thang   = []
    ds_tien_tien_dv  = []
    ds_chien_si_tdcs = []
    ds_chien_si_tt   = []
    ds_khac          = {}
    seen_ids         = set()
    units_by_type = {}
    for loai in LoaiDonVi:
        units_by_type[loai] = DonVi.query.filter_by(loai_don_vi=loai) \
            .order_by(DonVi.thu_tu).all()

    _sort_map = {}
    for loai_idx, (loai, units) in enumerate(units_by_type.items()):
        for unit in units:
            _sort_map[unit.id] = (loai_idx, unit.thu_tu or 0, unit.ten_don_vi or '')

    # --- Hàm sort key ---
    def _unit_sort_key(ct_dx_tuple):
        ct, dx = ct_dx_tuple
        don_vi_id = dx.don_vi_id if dx else None
        base = _sort_map.get(don_vi_id, (999, 999, 'zzz'))
        ho_ten = ''
        if ct:
            ho_ten = (ct.quan_nhan.ho_ten if ct.quan_nhan else ct.ten_don_vi_de_xuat) or ''
        return (*base, ho_ten.lower())
        
    for dx in nominations:
        for ct in dx.chi_tiets:
            if ct.id in seen_ids: continue
            seen_ids.add(ct.id)

            if danh_hieu_filter and ct.loai_danh_hieu != danh_hieu_filter: continue
            if scope_filter == 'quan_luc' and ct.doi_tuong not in BAN_QUANLUC_DOI_TUONG: continue
            if scope_filter == 'can_bo' and ct.doi_tuong in BAN_QUANLUC_DOI_TUONG: continue

            ho_ten = (ct.quan_nhan.ho_ten if ct.quan_nhan else (ct.ten_don_vi_de_xuat or ''))
            if search_query and search_query.lower() not in ho_ten.lower(): continue

            dh = (ct.loai_danh_hieu or '').strip()
            if   dh == 'Đơn vị quyết thắng':  ds_quyet_thang.append((ct, dx))
            elif dh == 'Đơn vị tiên tiến':     ds_tien_tien_dv.append((ct, dx))
            elif dh == 'Chiến sĩ thi đua':     ds_chien_si_tdcs.append((ct, dx))
            elif dh == 'Chiến sĩ tiên tiến':   ds_chien_si_tt.append((ct, dx))
            else: ds_khac.setdefault(dh, []).append((ct, dx))

    # --- Sort từng nhóm ---
    ds_quyet_thang   = sorted(ds_quyet_thang,   key=_unit_sort_key)
    ds_tien_tien_dv  = sorted(ds_tien_tien_dv,  key=_unit_sort_key)
    ds_chien_si_tdcs = sorted(ds_chien_si_tdcs, key=_unit_sort_key)
    ds_chien_si_tt   = sorted(ds_chien_si_tt,   key=_unit_sort_key)
    ds_khac          = {dh: sorted(lst, key=_unit_sort_key) for dh, lst in ds_khac.items()}

    title_nam_hoc = nam_hoc_filter or 'TẤT CẢ NĂM HỌC'

    # --- Tạo docx nhanh bằng XML template ---
    from app.utils.docx_fast import (
        cm_to_twips, _para, _build_table, _data_row,
        _build_document_xml, build_docx,
    )

    # --- Chiều rộng cột cá nhân ---
    CN_WIDTHS = [
        cm_to_twips(0.8), cm_to_twips(3.5), cm_to_twips(2.0),
        cm_to_twips(2.5), cm_to_twips(2.5), cm_to_twips(5.5),
    ]
    CN_HEADERS = ['STT', 'Họ và tên', 'Cấp bậc', 'Chức vụ', 'Đơn vị', 'Ghi chú']

    from app.models.nomination import TieuChi as _TieuChi
    _tieu_chi_map_all = {tc.ma_truong: tc.ten for tc in _TieuChi.query.all()}
    
    def _build_tomtat(ct):
        parts = []
        if ct.muc_do_hoan_thanh: parts.append(ct.muc_do_hoan_thanh)
        if ct.diem_tong_ket: parts.append(f'HT: {ct.diem_tong_ket}')
        if ct.ket_qua_ren_luyen: parts.append(f'RL: {ct.ket_qua_ren_luyen}')
        return '; '.join(parts)

    def _cn_rows(items):
        rows_xml = []
        for i, (ct, dx) in enumerate(items, 1):
            qn_obj = ct.quan_nhan
            row_cells = [
                (str(i), False, 'center'),
                (qn_obj.ho_ten if qn_obj else '', True, 'left'),
                (qn_obj.cap_bac if qn_obj else '', False, 'left'),
                (qn_obj.chuc_vu if qn_obj else '', False, 'left'),
                (dx.don_vi.ten_don_vi if dx.don_vi else '', False, 'left'),
                (_build_tomtat(ct) or '', False, 'left'),
            ]
            rows_xml.append(_data_row(row_cells, CN_WIDTHS, size_pt=9, shade=None)) 
        return rows_xml

    now = _dt.datetime.now()
    today_str = now.strftime('%d/%m/%Y')
    
    # Thiết lập Tên đơn vị hiển thị động ở góc trái tiêu đề (Nếu có filter đơn vị cụ thể thì lấy tên đó)
    ten_don_vi_header = unit_filter.upper() if unit_filter else "KHOA SƯ PHẠM QUÂN SỰ"

    body = []
    
    # ─────────────────────────────────────────────────────────────────────────
    # [TÍCH HỢP] BẢNG QUỐC HIỆU TIÊU NGỮ & TIÊU ĐỀ CHUẨN ĐÚNG THEO XML MẪU
    # ─────────────────────────────────────────────────────────────────────────
    header_table_xml = f"""<w:tbl xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
    <w:tblPr>
    <w:tblW w:type="dxa" w:w="9070"/>
    <w:jc w:val="left"/>
    <w:tblLayout w:type="fixed"/>
    <w:tblLook w:firstColumn="1" w:firstRow="1" w:lastColumn="0" w:lastRow="0" w:noHBand="0" w:noVBand="1" w:val="04A0"/>
    </w:tblPr>
    <w:tblGrid>
    <w:gridCol w:w="4535"/>
    <w:gridCol w:w="4535"/>
    </w:tblGrid>
    <w:tr>
    <w:tc>
    <w:tcPr>
    <w:tcW w:type="dxa" w:w="4535"/>
    <w:tcBorders>
    <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>
    <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>
    <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>
    <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>
    <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>
    <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>
    </w:tcBorders>
    </w:tcPr>
    <w:p>
    <w:pPr>
    <w:spacing w:before="0" w:after="0" w:line="240" w:lineRule="auto"/>
    <w:jc w:val="center"/>
    </w:pPr>
    <w:r>
    <w:rPr>
    <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>
    <w:b w:val="0"/>
    <w:i w:val="0"/>
    <w:sz w:val="24"/>
    </w:rPr>
    <w:t>TRƯỜNG SĨ QUAN CHÍNH TRỊ</w:t>
    </w:r>
    </w:p>
    <w:p>
    <w:pPr>
    <w:spacing w:before="0" w:after="0" w:line="240" w:lineRule="auto"/>
    <w:jc w:val="center"/>
    </w:pPr>
    <w:r>
    <w:rPr>
    <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>
    <w:b/>
    <w:i w:val="0"/>
    <w:sz w:val="24"/>
    <w:u w:val="single"/>
    </w:rPr>
    <w:t>{ten_don_vi_header}</w:t>
    </w:r>
    </w:p>
    <w:p>
    <w:pPr>
    <w:spacing w:before="0" w:after="0" w:line="240" w:lineRule="auto"/>
    </w:pPr>
    <w:r>
    <w:rPr>
    <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>
    <w:b w:val="0"/>
    <w:i w:val="0"/>
    <w:sz w:val="22"/>
    </w:rPr>
    </w:r>
    </w:p>
    </w:tc>
    <w:tc>
    <w:tcPr>
    <w:tcW w:type="dxa" w:w="4535"/>
    <w:tcBorders>
    <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>
    <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>
    <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>
    <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>
    <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>
    <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>
    </w:tcBorders>
    </w:tcPr>
    <w:p>
    <w:pPr>
    <w:spacing w:before="0" w:after="0" w:line="240" w:lineRule="auto"/>
    <w:jc w:val="center"/>
    </w:pPr>
    <w:r>
    <w:rPr>
    <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>
    <w:b/>
    <w:i w:val="0"/>
    <w:sz w:val="24"/>
    </w:rPr>
    <w:t>CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM</w:t>
    </w:r>
    </w:p>
    <w:p>
    <w:pPr>
    <w:spacing w:before="0" w:after="0" w:line="240" w:lineRule="auto"/>
    <w:jc w:val="center"/>
    </w:pPr>
    <w:r>
    <w:rPr>
    <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>
    <w:b/>
    <w:i w:val="0"/>
    <w:sz w:val="24"/>
    <w:u w:val="single"/>
    </w:rPr>
    <w:t>Độc lập - Tự do - Hạnh phúc</w:t>
    </w:r>
    </w:p>
    <w:p>
    <w:pPr>
    <w:spacing w:before="0" w:after="0" w:line="240" w:lineRule="auto"/>
    <w:jc w:val="center"/>
    </w:pPr>
    <w:r>
    <w:rPr>
    <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>
    <w:b w:val="0"/>
    <w:i/>
    <w:sz w:val="22"/>
    </w:rPr>
    <w:t>Hà Nội, ngày {now.day} tháng {now.month} năm {now.year}</w:t>
    </w:r>
    </w:p>
    </w:tc>
    </w:tr>
    </w:tbl>"""

    title_xml = f"""<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
    <w:pPr>
    <w:jc w:val="center"/>
    <w:spacing w:before="180" w:after="40"/>
    </w:pPr>
    <w:r>
    <w:rPr>
    <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>
    <w:b/>
    <w:i w:val="0"/>
    <w:sz w:val="26"/>
    </w:rPr>
    <w:t>DANH SÁCH ĐỀ NGHỊ KHEN THƯỞNG NĂM HỌC {title_nam_hoc.upper()}</w:t>
    </w:r>
    </w:p>"""

    body.append(header_table_xml)
    body.append("<w:p/>")
    body.append(title_xml)
    body.append(_para(f'(Xuất lúc {now.strftime("%H:%M")} ngày {today_str})', italic=True, size_pt=10, align='center', space_before=0, space_after=120))

    def _add_section(label, items, is_tap_the=False):
        if not items:
            return
        body.append(_para(label, bold=True, size_pt=12, space_before=120, space_after=40))
        
        if is_tap_the:
            # HIỂN THỊ DẠNG DANH SÁCH TỪNG DÒNG (BỎ BẢNG)
            for i, (ct, dx) in enumerate(items, 1):
                ten_dv = ct.ten_don_vi_de_xuat or (dx.don_vi.ten_don_vi if dx.don_vi else '')
                body.append(_para(f"{i}. {ten_dv}", size_pt=11, align='left', space_before=40, space_after=0))
                
            body.append(_para(f'Tổng cộng: {len(items)} đơn vị', italic=True, size_pt=10, align='right', space_before=60, space_after=60))
        else:
            # CÁ NHÂN GIỮ NGUYÊN DẠNG BẢNG NỀN TRẮNG CHỮ ĐEN
            rows_xml = _cn_rows(items)
            body.append(_build_table(CN_HEADERS, rows_xml, CN_WIDTHS, total_label=f'Tổng cộng: {len(items)} người', size_pt=9))
            body.append(_para('', space_before=60, space_after=0))

    _add_section('I. DANH HIỆU ĐƠN VỊ QUYẾT THẮNG', ds_quyet_thang, is_tap_the=True)
    _add_section('II. DANH HIỆU ĐƠN VỊ TIÊN TIẾN', ds_tien_tien_dv, is_tap_the=True)
    _add_section('III. CHIẾN SĨ THI ĐUA', ds_chien_si_tdcs)
    _add_section('IV. CHIẾN SĨ TIÊN TIẾN', ds_chien_si_tt)
    for extra_dh, extra_items in ds_khac.items():
        _add_section(extra_dh, extra_items)

    body.append(_para(f'(Xuất lúc {now.strftime("%H:%M ngày %d/%m/%Y")})',
                      italic=True, size_pt=9, align='right', space_before=120, space_after=0))
    
    doc_xml = _build_document_xml(body, margin_left=2016, margin_right=720, margin_top=1440, margin_bottom=1440)
    
    # --- Sinh file ban đầu bằng XML Fast ---
    buf = build_docx(doc_xml)
    
    # =========================================================================
    # BẢO VỆ CHỐNG CHỈNH SỬA (CÓ MẬT KHẨU BĂM) BẰNG ZIP/XML TRỰC TIẾP
    # =========================================================================
    buf.seek(0)
    final_buf = BytesIO()

    # Thuật toán hash mật khẩu Office 2010+ (Agile Encryption)
    password = "123" # <--- THAY ĐỔI MẬT KHẨU TẠI ĐÂY
    salt = os.urandom(16)
    salt_b64 = binascii.b2a_base64(salt).strip().decode()

    key = hashlib.sha512(salt + password.encode('utf-16le')).digest()
    spin_count = 10000
    for i in range(spin_count):
        iterator = i.to_bytes(4, byteorder='little')
        key = hashlib.sha512(key + iterator).digest()
        
    hash_b64 = binascii.b2a_base64(key).strip().decode()

    # Tạo chuỗi XML bảo vệ theo chuẩn Agile
    protection_tag = (
        f'<w:documentProtection w:edit="readOnly" w:enforcement="1" '
        f'w:algorithmName="SHA-512" w:spinCount="{spin_count}" '
        f'w:hashValue="{hash_b64}" w:saltValue="{salt_b64}"/>'
    ).encode('utf-8')

    with zipfile.ZipFile(buf, 'r') as zin:
        with zipfile.ZipFile(final_buf, 'w') as zout:
            for item in zin.infolist():
                file_content = zin.read(item.filename)
                
                # Tìm file cấu hình settings.xml
                if item.filename == 'word/settings.xml':
                    # Tiêm mã bảo vệ vào trước khi kết thúc thẻ w:settings
                    if b'</w:settings>' in file_content:
                        # Tránh trùng lặp nếu thẻ đã tồn tại
                        if b'w:documentProtection' not in file_content:
                            file_content = file_content.replace(b'</w:settings>', protection_tag + b'</w:settings>')
                
                # Copy toàn bộ nội dung sang file zip mới
                zout.writestr(item, file_content)
    
    final_buf.seek(0)
    # =========================================================================

    fname_parts = ['DanhSachKhenThuong']
    if nam_hoc_filter:
        fname_parts.append(nam_hoc_filter.replace('-', '_'))
    fname_parts.append(now.strftime('%d%m%Y'))
    filename = '_'.join(fname_parts) + '.docx'

    response = send_file(
        final_buf, as_attachment=True, download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response.set_cookie(
        'export_done', '1',
        max_age=600, httponly=False, samesite='Lax', path='/'
    )
    return response

def add_corner_logo(doc):
    """Thêm logo nhỏ ở góc phải trên cùng của trang (sau header table hiện tại)."""
    import os
    from flask import current_app

    # ★ Ưu tiên logo nhỏ (19 KB) để giảm kích thước file docx
    logo_path = os.path.join(current_app.root_path, 'static', 'img', 'logo-Si-quan.png')
    if not os.path.exists(logo_path):
        logo_path = os.path.join(current_app.root_path, 'static', 'img', 'watermark.png')
        if not os.path.exists(logo_path):
            return
    
    try:
        for section in doc.sections:
            header = section.header
            
            # Thêm paragraph mới vào cuối header (sau table header hiện tại)
            para = header.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Set paragraph spacing để logo sát lề trên
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            
            # Thêm logo nhỏ căn phải (1.5cm)
            run = para.add_run()
            run.add_picture(logo_path, width=Cm(1.5))
            
    except Exception as e:
        print(f"Warning: Could not add corner logo: {e}")


def add_logo_watermark(doc):
    """Thêm watermark logo vào tất cả các trang - DEPRECATED, use specific functions above."""
    import os
    from flask import current_app
    
    # Đường dẫn logo
    logo_path = os.path.join(current_app.root_path, 'static', 'img', 'logo-Si-quan.png')
    
    if not os.path.exists(logo_path):
        return  # Skip if logo not found
    
    try:
        for section in doc.sections:
            # Thêm watermark vào header của section
            header = section.header
            
            # Tạo paragraph cho watermark - đặt ở cuối header
            para = header.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Thêm hình ảnh watermark với opacity thấp
            run = para.add_run()
            picture = run.add_picture(logo_path, width=Cm(6))
            
            # Note: Đơn giản hóa - chỉ thêm logo vào header
            # Không cần chuyển thành anchor phức tạp, vì có thể gây lỗi
            # Logo sẽ xuất hiện ở đầu mỗi trang trong header
            
    except Exception as e:
        # Nếu watermark fails, không làm crash toàn bộ export
        print(f"Warning: Could not add watermark: {e}")
        pass


def protect_document_formatting_only(doc, password: str):
    """
    Khóa tài liệu: chỉ đọc nội dung (readOnly).
    Mật khẩu được hash theo chuẩn Office 2010+ (Agile Encryption).
    """
    # 1. Tạo salt ngẫu nhiên (16 bytes)
    salt = os.urandom(16)
    salt_b64 = binascii.b2a_base64(salt).strip().decode()

    # 2. Hash lần đầu: SHA-512(salt + password)
    # Lưu ý: password bắt buộc encode sang chuẩn UTF-16 Little Endian
    key = hashlib.sha512(salt + password.encode('utf-16le')).digest()
    
    # 3. Lặp N vòng để chống brute-force.
    # 10.000 vẫn tuân thủ chuẩn Agile Encryption, nhanh hơn ~10x so với 100.000 vòng.
    spin_count = 10000
    for i in range(spin_count):
        iterator = i.to_bytes(4, byteorder='little')
        key = hashlib.sha512(key + iterator).digest()
        
    hash_b64 = binascii.b2a_base64(key).strip().decode()

    # 4. Lấy cấu hình settings của docx
    settings = doc.settings.element

    # Xóa thẻ documentProtection cũ nếu có
    for old in settings.findall(qn('w:documentProtection')):
        settings.remove(old)

    # 5. Tạo thẻ <w:documentProtection> theo chuẩn Office đời mới
    doc_prot = OxmlElement('w:documentProtection')
    doc_prot.set(qn('w:edit'),          'readOnly')
    doc_prot.set(qn('w:enforcement'),   '1')
    
    # BỎ CÁC THẺ CŨ (cryptProviderType, v.v.). SỬ DỤNG CHUẨN AGILE MỚI:
    doc_prot.set(qn('w:algorithmName'), 'SHA-512')
    doc_prot.set(qn('w:spinCount'),     str(spin_count))
    doc_prot.set(qn('w:hashValue'),     hash_b64)     # Đã đổi từ w:hash thành w:hashValue
    doc_prot.set(qn('w:saltValue'),     salt_b64)     # Đã đổi từ w:salt thành w:saltValue

    # Chèn vào đầu <w:settings>
    settings.insert(0, doc_prot)

@admin_bp.route('/reward-list/export-hoi-dong-excel')
@login_required
@admin_or_reward_viewer_required
def export_hoi_dong_excel():
    """Export danh sách chờ (HOI_DONG + PHE_DUYET_CUOI) kèm điểm tiêu chí + biểu quyết hội đồng."""
    import datetime as _dt
    from openpyxl.styles import Font as _Font, Alignment as _Align, PatternFill as _Fill, Border as _Border, Side as _Side
    from app.models.hoi_dong import HoiDongBieuQuyet as _BQ, HOI_DONG_VAI_TRO as _VT, HOI_DONG_VAI_TRO_DISPLAY as _VTD

    nam_hoc_filter = request.args.get('nam_hoc', '')
    unit_filter = request.args.get('unit', '')
    danh_hieu_filter = request.args.get('danh_hieu', '')
    trang_thai_filter = request.args.get('trang_thai', '')  # 'hoi_dong' | 'phe_duyet_cuoi' | '' (cả hai)

    # Build query
    statuses = []
    if trang_thai_filter == 'hoi_dong':
        statuses = [TrangThaiDeXuat.HOI_DONG.value]
    elif trang_thai_filter == 'phe_duyet_cuoi':
        statuses = [TrangThaiDeXuat.PHE_DUYET_CUOI.value]
    else:
        statuses = [TrangThaiDeXuat.HOI_DONG.value, TrangThaiDeXuat.PHE_DUYET_CUOI.value]

    query = DeXuat.query.filter(DeXuat.trang_thai.in_(statuses))
    if nam_hoc_filter:
        query = query.filter(DeXuat.nam_hoc == nam_hoc_filter)
    if unit_filter:
        query = query.join(DonVi).filter(DonVi.ten_don_vi == unit_filter)
    nominations = query.order_by(DeXuat.nam_hoc.desc(), DeXuat.ngay_gui.desc()).all()

    # Pre-fetch all votes for these de_xuat ids
    dx_ids = [dx.id for dx in nominations]
    all_votes = {}  # chi_tiet_id -> {vai_tro: ket_qua}
    if dx_ids:
        votes = _BQ.query.filter(_BQ.de_xuat_id.in_(dx_ids)).all()
        for v in votes:
            all_votes.setdefault(v.chi_tiet_id, {})[v.vai_tro] = v.ket_qua

    # Styles
    bold = _Font(name='Times New Roman', bold=True, size=11)
    normal = _Font(name='Times New Roman', size=10)
    center = _Align(horizontal='center', vertical='center', wrap_text=True)
    left = _Align(horizontal='left', vertical='center', wrap_text=True)
    thin = _Border(left=_Side(style='thin'), right=_Side(style='thin'),
                   top=_Side(style='thin'), bottom=_Side(style='thin'))
    navy_fill = _Fill(start_color='1F4E79', end_color='1F4E79', fill_type='solid')
    green_fill = _Fill(start_color='C6EFCE', end_color='C6EFCE', fill_type='solid')
    red_fill = _Fill(start_color='FFC7CE', end_color='FFC7CE', fill_type='solid')
    yellow_fill = _Fill(start_color='FFEB9C', end_color='FFEB9C', fill_type='solid')
    white_bold = _Font(name='Times New Roman', bold=True, size=10, color='FFFFFF')

    wb = Workbook()
    ws = wb.active
    ws.title = 'DS Hoi dong Bieu quyet'

    # Title row 1
    vote_col_count = len(_VT)
    total_cols = 13 + vote_col_count  # info cols + vote cols + tong + ket_luan
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=total_cols)
    ws.cell(1, 1).value = 'DANH SÁCH XÉT DUYỆT HỘI ĐỒNG THI ĐUA KHEN THƯỞNG'
    ws.cell(1, 1).font = _Font(name='Times New Roman', bold=True, size=13)
    ws.cell(1, 1).alignment = center

    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=total_cols)
    filter_info = []
    if nam_hoc_filter: filter_info.append(f'Năm học: {nam_hoc_filter}')
    if unit_filter: filter_info.append(f'Đơn vị: {unit_filter}')
    if danh_hieu_filter: filter_info.append(f'Danh hiệu: {danh_hieu_filter}')
    ws.cell(2, 1).value = ('  |  '.join(filter_info) if filter_info else 'Tất cả') + \
        f'  —  Xuất ngày: {_dt.datetime.now().strftime("%d/%m/%Y %H:%M")}'
    ws.cell(2, 1).font = _Font(name='Times New Roman', italic=True, size=10)
    ws.cell(2, 1).alignment = center

    # Header row 3: fixed info columns
    info_headers = [
        'STT', 'Đơn vị', 'Họ và tên', 'CCCD', 'Cấp bậc', 'Chức vụ',
        'Đối tượng', 'Danh hiệu', 'Năm học', 'Trạng thái',
        'Mức độ HT nhiệm vụ', 'Phiếu tín nhiệm', 'Ghi chú',
    ]
    info_widths = [5, 26, 22, 14, 12, 20, 16, 20, 12, 18, 22, 18, 20]

    for col_idx, (h, w) in enumerate(zip(info_headers, info_widths), 1):
        c = ws.cell(row=3, column=col_idx, value=h)
        c.font = white_bold
        c.fill = navy_fill
        c.alignment = center
        c.border = thin
        ws.column_dimensions[get_column_letter(col_idx)].width = w

    # Vote columns
    vote_start_col = len(info_headers) + 1
    for i, vt in enumerate(_VT):
        col = vote_start_col + i
        label = _VTD.get(vt, vt)
        c = ws.cell(row=3, column=col, value=label)
        c.font = white_bold
        c.fill = _Fill(start_color='375623', end_color='375623', fill_type='solid')
        c.alignment = center
        c.border = thin
        ws.column_dimensions[get_column_letter(col)].width = 18

    # Tổng đồng ý + Kết luận
    tong_col = vote_start_col + len(_VT)
    kl_col = tong_col + 1
    c = ws.cell(row=3, column=tong_col, value='Tổng ĐY')
    c.font = white_bold; c.fill = navy_fill; c.alignment = center; c.border = thin
    ws.column_dimensions[get_column_letter(tong_col)].width = 10
    c = ws.cell(row=3, column=kl_col, value='Kết luận')
    c.font = white_bold; c.fill = navy_fill; c.alignment = center; c.border = thin
    ws.column_dimensions[get_column_letter(kl_col)].width = 16

    # Data rows
    row_num = 4
    stt = 0
    for dx in nominations:
        unit_name = dx.don_vi.ten_don_vi if dx.don_vi else '—'
        for ct in dx.chi_tiets:
            if danh_hieu_filter and ct.loai_danh_hieu != danh_hieu_filter:
                continue
            qn = ct.quan_nhan
            ho_ten = qn.ho_ten if qn else (ct.ten_don_vi_de_xuat or '—')
            stt += 1

            row_data = [
                stt,
                unit_name,
                ho_ten,
                (qn.can_cuoc_cong_dan if qn else ''),
                (qn.cap_bac if qn else ''),
                (qn.chuc_vu if qn else ''),
                (ct.doi_tuong or ''),
                ct.loai_danh_hieu or '',
                dx.nam_hoc or '',
                dx.trang_thai,
                ct.muc_do_hoan_thanh or '',
                ct.phieu_tin_nhiem or '',
                ct.ghi_chu or '',
            ]
            for col_idx, val in enumerate(row_data, 1):
                c = ws.cell(row=row_num, column=col_idx, value=val)
                c.font = normal
                c.border = thin
                c.alignment = center if col_idx == 1 else left

            # Vote cells
            ct_votes = all_votes.get(ct.id, {})
            dong_y_count = 0
            for i, vt in enumerate(_VT):
                col = vote_start_col + i
                ket_qua = ct_votes.get(vt, '')
                c = ws.cell(row=row_num, column=col, value=ket_qua)
                c.font = normal
                c.border = thin
                c.alignment = center
                if ket_qua == 'Đồng ý':
                    c.fill = green_fill
                    dong_y_count += 1
                elif ket_qua == 'Không đồng ý':
                    c.fill = red_fill

            # Tổng ĐY
            c = ws.cell(row=row_num, column=tong_col, value=dong_y_count)
            c.font = _Font(name='Times New Roman', bold=True, size=10)
            c.border = thin
            c.alignment = center

            # Kết luận
            voted_count = len(ct_votes)
            if voted_count == 0:
                ket_luan = 'Chưa biểu quyết'
                kl_fill = None
            elif dong_y_count > len(_VT) / 2:
                ket_luan = 'Đề nghị khen thưởng'
                kl_fill = green_fill
            else:
                ket_luan = 'Không đạt'
                kl_fill = red_fill
            c = ws.cell(row=row_num, column=kl_col, value=ket_luan)
            c.font = _Font(name='Times New Roman', bold=True, size=10)
            c.border = thin
            c.alignment = center
            if kl_fill:
                c.fill = kl_fill

            row_num += 1

    # Summary
    ws.merge_cells(start_row=row_num, start_column=1, end_row=row_num, end_column=total_cols)
    c = ws.cell(row=row_num, column=1, value=f'Tổng cộng: {stt} người')
    c.font = bold
    c.alignment = left

    ws.page_setup.paperSize = 9
    ws.page_setup.orientation = 'landscape'
    ws.page_setup.fitToPage = True
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    ws.freeze_panes = 'A4'

    ws.protection.sheet = True
    ws.protection.password = 'bth123'

    output = BytesIO()
    wb.save(output)
    output.seek(0)
    fname_parts = ['DanhSach_HoiDong_BieuQuyet']
    if nam_hoc_filter:
        fname_parts.append(nam_hoc_filter.replace('-', '_'))
    fname_parts.append(_dt.datetime.now().strftime('%d%m%Y'))
    return send_file(output, as_attachment=True,
                     download_name='_'.join(fname_parts) + '.xlsx',
                     mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')


@admin_bp.route('/reward-list/pending-final/export-excel')
@login_required
@admin_or_reward_viewer_required
def export_pending_final_excel():
    items = _get_pending_final_individuals()

    wb = Workbook()
    ws = wb.active
    ws.title = 'Cho phe duyet cuoi'

    headers = [
        'STT', 'Họ và tên', 'CCCD', 'Cấp bậc', 'Chức vụ', 'Đối tượng',
        'Đơn vị', 'Danh hiệu', 'Năm học', 'Ngày gửi'
    ]
    for i, h in enumerate(headers, 1):
        c = ws.cell(row=1, column=i, value=h)
        c.font = Font(bold=True)

    for idx, item in enumerate(items, 1):
        dx = item['dx']
        ct = item['ct']
        qn = ct.quan_nhan
        ws.append([
            idx,
            qn.ho_ten if qn else dx.don_vi.ten_don_vi,
            qn.can_cuoc_cong_dan if qn else '',
            qn.cap_bac if qn else '',
            qn.chuc_vu if qn else '',
            ct.doi_tuong or '',
            dx.don_vi.ten_don_vi,
            ct.loai_danh_hieu,
            dx.nam_hoc,
            dx.ngay_gui.strftime('%d/%m/%Y') if dx.ngay_gui else '',
        ])

    widths = [6, 24, 16, 12, 20, 18, 28, 20, 12, 12]
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w

    ws.page_setup.paperSize = 9
    ws.page_setup.orientation = 'landscape'
    ws.page_setup.fitToPage = True
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0
    ws.sheet_properties.pageSetUpPr.fitToPage = True

    ws.protection.sheet = True
    ws.protection.password = 'bth123'

    out = BytesIO()
    wb.save(out)
    out.seek(0)
    return send_file(
        out,
        as_attachment=True,
        download_name='danh_sach_cho_phe_duyet_cuoi.xlsx',
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )


@admin_bp.route('/reward-list/pending-final/export-word')
@login_required
@admin_or_reward_viewer_required
def export_pending_final_word():
    items = _get_pending_final_individuals()

    rows_html = []
    for idx, item in enumerate(items, 1):
        dx = item['dx']
        ct = item['ct']
        qn = ct.quan_nhan
        rows_html.append(
            '<tr>'
            f'<td>{idx}</td>'
            f'<td>{escape(qn.ho_ten if qn else dx.don_vi.ten_don_vi)}</td>'
            f'<td>{escape(qn.can_cuoc_cong_dan if qn and qn.can_cuoc_cong_dan else "")}</td>'
            f'<td>{escape(qn.cap_bac if qn and qn.cap_bac else "")}</td>'
            f'<td>{escape(qn.chuc_vu if qn and qn.chuc_vu else "")}</td>'
            f'<td>{escape(ct.doi_tuong or "")}</td>'
            f'<td>{escape(dx.don_vi.ten_don_vi)}</td>'
            f'<td>{escape(ct.loai_danh_hieu or "")}</td>'
            f'<td>{escape(dx.nam_hoc or "")}</td>'
            f'<td>{escape(dx.ngay_gui.strftime("%d/%m/%Y") if dx.ngay_gui else "")}</td>'
            '</tr>'
        )

    html = (
        '<html><head><meta charset="utf-8"></head><body>'
        '<h3>Danh sách chờ phê duyệt cuối</h3>'
        '<table border="1" cellspacing="0" cellpadding="4" style="border-collapse:collapse; font-family:Times New Roman; font-size:12pt;">'
        '<tr><th>STT</th><th>Họ và tên</th><th>CCCD</th><th>Cấp bậc</th><th>Chức vụ</th><th>Đối tượng</th><th>Đơn vị</th><th>Danh hiệu</th><th>Năm học</th><th>Ngày gửi</th></tr>'
        + ''.join(rows_html) +
        '</table></body></html>'
    )

    return Response(
        html,
        mimetype='application/msword',
        headers={'Content-Disposition': 'attachment; filename=danh_sach_cho_phe_duyet_cuoi.doc'}
    )

@admin_bp.route('/reward-list/bang2/export-excel')
@login_required
@admin_or_reward_viewer_required
def export_bang2_word():
    # Implementation for exporting to Word
    pass

@admin_bp.route('/reward-list/bang2/export-excel')
@login_required
@admin_or_reward_viewer_required
def export_bang2_excel():
    """Export Bảng 2 (PHE_DUYET_CUOI – Hội đồng biểu quyết) ra Excel đầy đủ thông tin."""
    from app.models.hoi_dong import HOI_DONG_VAI_TRO, HOI_DONG_VAI_TRO_DISPLAY
    import datetime as _dt
    items = _get_phe_duyet_cuoi_items()

    wb = Workbook()
    ws = wb.active
    ws.title = 'Bang 2 - Xet duyet HD'

    from openpyxl.styles import Font as _Font, Alignment as _Align, PatternFill as _Fill, Border as _Border, Side as _Side
    normal = _Font(name='Times New Roman', size=10)
    center = _Align(horizontal='center', vertical='center', wrap_text=True)
    left = _Align(horizontal='left', vertical='center', wrap_text=True)
    thin = _Border(left=_Side(style='thin'), right=_Side(style='thin'),
                   top=_Side(style='thin'), bottom=_Side(style='thin'))
    navy_fill = _Fill(start_color='1F4E79', end_color='1F4E79', fill_type='solid')
    white_bold = _Font(name='Times New Roman', bold=True, size=10, color='FFFFFF')

    # Columns: STT + thông tin cá nhân đầy đủ + 6 ban + Trạng thái
    personal_headers = [
        'STT', 'Họ và tên', 'CCCD', 'Cấp bậc', 'Chức vụ', 'Đối tượng',
        'Đơn vị', 'Danh hiệu', 'Năm học',
    ]
    personal_widths = [5, 22, 14, 12, 20, 16, 28, 20, 12]
    ban_headers = [HOI_DONG_VAI_TRO_DISPLAY[vt] for vt in HOI_DONG_VAI_TRO]
    all_headers = personal_headers + ban_headers + ['Trạng thái']
    all_widths = personal_widths + [16] * len(HOI_DONG_VAI_TRO) + [16]
    col_count = len(all_headers)

    # Title rows
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=col_count)
    ws.cell(1, 1).value = 'BẢNG 2: XÉT DUYỆT CỦA CƠ QUAN THƯỜNG TRỰC – HỘI ĐỒNG BIỂU QUYẾT'
    ws.cell(1, 1).font = _Font(name='Times New Roman', bold=True, size=13)
    ws.cell(1, 1).alignment = center

    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=col_count)
    ws.cell(2, 1).value = f'Xuất ngày: {_dt.datetime.now().strftime("%d/%m/%Y %H:%M")}'
    ws.cell(2, 1).font = _Font(name='Times New Roman', italic=True, size=10)
    ws.cell(2, 1).alignment = center

    # Header row
    for col_idx, (h, w) in enumerate(zip(all_headers, all_widths), 1):
        c = ws.cell(row=3, column=col_idx, value=h)
        c.font = white_bold
        c.fill = navy_fill
        c.alignment = center
        c.border = thin
        ws.column_dimensions[get_column_letter(col_idx)].width = w

    # Data rows — items is list of dict
    for row_num, row in enumerate(items, 4):
        ct = row['ct']
        dx = row['dx']
        votes = row['votes']
        qn = ct.quan_nhan
        name = qn.ho_ten if qn else (ct.ten_don_vi_de_xuat or '—')
        don_vi = dx.don_vi.ten_don_vi if dx.don_vi else '—'

        row_data = [
            row_num - 3,
            name,
            (qn.can_cuoc_cong_dan if qn else '') or '',
            (qn.cap_bac if qn else '') or '',
            (qn.chuc_vu if qn else '') or '',
            (ct.doi_tuong or ''),
            don_vi,
            ct.loai_danh_hieu or '',
            dx.nam_hoc or '',
        ]
        for vt in HOI_DONG_VAI_TRO:
            bq = votes.get(vt)
            row_data.append(bq.ket_qua if bq else '—')
        trang_thai = ('Đã xác nhận' if row['is_confirmed']
                      else ('Đủ phiếu – chờ xác nhận' if row['all_voted_dong_y']
                            else 'Chờ biểu quyết'))
        row_data.append(trang_thai)

        for col_idx, val in enumerate(row_data, 1):
            c = ws.cell(row=row_num, column=col_idx, value=val)
            c.font = normal
            c.alignment = center if col_idx == 1 else left
            c.border = thin

    ws.page_setup.paperSize = 9
    ws.page_setup.orientation = 'landscape'
    ws.page_setup.fitToPage = True
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    ws.freeze_panes = 'A4'

    ws.protection.sheet = True
    ws.protection.password = 'bth123'

    output = BytesIO()
    wb.save(output)
    output.seek(0)
    return send_file(output, as_attachment=True,
                     download_name='bang2_xet_duyet_hoi_dong.xlsx',
                     mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')


@admin_bp.route('/reward-list/bang2/export-detail-excel')
@login_required
@admin_or_reward_viewer_required
def export_bang2_detail_excel():
    """Export Bảng 2 chi tiết: Sheet 1 = danh sách tổng; Sheet per người = đầy đủ tiêu chí."""
    from app.models.hoi_dong import HOI_DONG_VAI_TRO, HOI_DONG_VAI_TRO_DISPLAY
    import datetime as _dt
    from openpyxl.styles import Font as _Font, Alignment as _Align, PatternFill as _Fill, Border as _Border, Side as _Side

    nam_hoc_filter = request.args.get('nam_hoc', '')
    items = _get_phe_duyet_cuoi_items(nam_hoc=nam_hoc_filter or None)

    wb = Workbook()
    # ── Sheet 1: Danh sách tổng ──────────────────────────────────────────────
    ws = wb.active
    ws.title = 'Tong hop'

    normal = _Font(name='Times New Roman', size=10)
    center = _Align(horizontal='center', vertical='center', wrap_text=True)
    left = _Align(horizontal='left', vertical='center', wrap_text=True)
    thin = _Border(left=_Side(style='thin'), right=_Side(style='thin'),
                   top=_Side(style='thin'), bottom=_Side(style='thin'))
    navy_fill = _Fill(start_color='1F4E79', end_color='1F4E79', fill_type='solid')
    white_bold = _Font(name='Times New Roman', bold=True, size=10, color='FFFFFF')
    green_fill = _Fill(start_color='C6EFCE', end_color='C6EFCE', fill_type='solid')

    personal_headers = ['STT', 'Họ và tên / Tên tập thể', 'CCCD', 'Cấp bậc', 'Chức vụ',
                        'Đối tượng', 'Đơn vị', 'Danh hiệu', 'Năm học']
    personal_widths = [5, 24, 14, 14, 22, 16, 28, 22, 12]
    ban_headers = [HOI_DONG_VAI_TRO_DISPLAY[vt] for vt in HOI_DONG_VAI_TRO]
    all_headers = personal_headers + ban_headers + ['Tổng ĐY', 'Trạng thái']
    all_widths = personal_widths + [16] * len(HOI_DONG_VAI_TRO) + [10, 18]
    col_count = len(all_headers)

    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=col_count)
    ws.cell(1, 1).value = 'BẢNG 2: XÉT DUYỆT CỦA CƠ QUAN THƯỜNG TRỰC – DANH SÁCH CHI TIẾT'
    ws.cell(1, 1).font = _Font(name='Times New Roman', bold=True, size=13)
    ws.cell(1, 1).alignment = center

    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=col_count)
    filter_txt = f'Năm học: {nam_hoc_filter}  ' if nam_hoc_filter else ''
    ws.cell(2, 1).value = f'{filter_txt}Xuất ngày: {_dt.datetime.now().strftime("%d/%m/%Y %H:%M")}'
    ws.cell(2, 1).font = _Font(name='Times New Roman', italic=True, size=10)
    ws.cell(2, 1).alignment = center

    for col_idx, (h, w) in enumerate(zip(all_headers, all_widths), 1):
        c = ws.cell(row=3, column=col_idx, value=h)
        c.font = white_bold; c.fill = navy_fill; c.alignment = center; c.border = thin
        ws.column_dimensions[get_column_letter(col_idx)].width = w

    for row_num, row in enumerate(items, 4):
        ct = row['ct']; dx = row['dx']; votes = row['votes']
        qn = ct.quan_nhan
        name = qn.ho_ten if qn else (ct.ten_don_vi_de_xuat or '—')
        don_vi = dx.don_vi.ten_don_vi if dx.don_vi else '—'
        dong_y_count = sum(1 for vt in HOI_DONG_VAI_TRO if votes.get(vt) and votes[vt].ket_qua == 'Đồng ý')
        row_data = [
            row_num - 3, name,
            (qn.can_cuoc_cong_dan if qn else '') or '',
            (qn.cap_bac if qn else '') or '',
            (qn.chuc_vu if qn else '') or '',
            ct.doi_tuong or '', don_vi,
            ct.loai_danh_hieu or '', dx.nam_hoc or '',
        ]
        for vt in HOI_DONG_VAI_TRO:
            bq = votes.get(vt)
            row_data.append(bq.ket_qua if bq else '—')
        row_data.append(dong_y_count)
        trang_thai = ('Đã xác nhận' if row['is_confirmed']
                      else ('Đủ phiếu – chờ XN' if row.get('all_voted_dong_y') else 'Chờ biểu quyết'))
        row_data.append(trang_thai)

        for col_idx, val in enumerate(row_data, 1):
            c = ws.cell(row=row_num, column=col_idx, value=val)
            c.font = normal
            c.alignment = center if col_idx in (1, col_count - 1) else left
            c.border = thin
            if row['is_confirmed']:
                c.fill = green_fill

    ws.freeze_panes = 'A4'
    ws.page_setup.paperSize = 9
    ws.page_setup.orientation = 'landscape'
    ws.page_setup.fitToPage = True; ws.page_setup.fitToWidth = 1; ws.page_setup.fitToHeight = 0
    ws.sheet_properties.pageSetUpPr.fitToPage = True

    # ── Sheet per người: đầy đủ tiêu chí ────────────────────────────────────
    TIEU_CHI_LABELS = [
        ('Mức độ hoàn thành NV', 'muc_do_hoan_thanh'),
        ('Phiếu tín nhiệm', 'phieu_tin_nhiem'),
        ('Xếp loại đảng viên', 'xep_loai_dang_vien'),
        ('KQ đoàn thể', 'ket_qua_doan_the'),
        ('KT điều lệnh', 'kiem_tra_dieu_lenh'),
        ('Điểm KT điều lệnh', 'diem_kiem_tra_dieu_lenh'),
        ('Bắn súng', 'ban_sung'),
        ('Điểm bắn súng', 'diem_ban_sung'),
        ('Thể lực', 'the_luc'),
        ('Điểm thể lực', 'diem_the_luc'),
        ('KT chính trị', 'kiem_tra_chinh_tri'),
        ('Điểm KT chính trị', 'diem_kiem_tra_chinh_tri'),
        ('Kỹ năng số', 'kiem_tra_tin_hoc'),
        ('Điểm kỹ năng số', 'diem_kiem_tra_tin_hoc'),
        ('Địa hình QS', 'dia_ly_quan_su'),
        ('Điểm địa hình QS', 'diem_dia_ly_quan_su'),
        ('GV giỏi', 'danh_hieu_gv_gioi'),
        ('Định mức giảng dạy', 'dinh_muc_giang_day'),
        ('KT giảng', 'ket_qua_kiem_tra_giang'),
        ('Tiến độ PGS', 'tien_do_pgs'),
        ('LĐ khoa học', 'thoi_gian_lao_dong_kh'),
        ('HV giỏi', 'danh_hieu_hv_gioi'),
        ('Điểm tổng kết', 'diem_tong_ket'),
        ('KQ thực hành', 'ket_qua_thuc_hanh'),
        ('KQ rèn luyện', 'ket_qua_ren_luyen'),
        ('HT thi tốt nghiệp', 'hinh_thuc_tot_nghiep'),
        ('Điểm CTĐ (TN)', 'diem_tn_ctd'),
        ('Điểm CT (TN)', 'diem_tn_ct'),
        ('Điểm TA (TN)', 'diem_tn_ta'),
        ('Điểm môn 4 (TN)', 'diem_tn_mon4'),
        ('Điểm CN (TN)', 'diem_tn_chuyennganh'),
        ('Điểm BV KL (TN)', 'diem_tn_baove'),
        ('Chủ trì ĐV danh hiệu', 'chu_tri_don_vi_danh_hieu'),
        ('Điểm NCKH', 'diem_nckh'),
        ('Nội dung NCKH', 'nckh_noi_dung'),
        ('Thành tích khác', 'thanh_tich_ca_nhan_khac'),
        ('Ghi chú', 'ghi_chu'),
    ]

    label_font = _Font(name='Times New Roman', bold=True, size=10)
    value_font = _Font(name='Times New Roman', size=10)
    label_align = _Align(horizontal='left', vertical='center', wrap_text=True)
    val_align = _Align(horizontal='left', vertical='center', wrap_text=True)
    section_fill = _Fill(start_color='D6E4F0', end_color='D6E4F0', fill_type='solid')
    section_font = _Font(name='Times New Roman', bold=True, size=11, color='1F4E79')

    for idx, row in enumerate(items, 1):
        ct = row['ct']; dx = row['dx']; votes = row['votes']
        qn = ct.quan_nhan
        is_tap_the = qn is None
        name = qn.ho_ten if qn else (ct.ten_don_vi_de_xuat or '—')
        # Sheet title: truncate to 31 chars (Excel limit)
        sheet_name = f'{idx:02d}_{name}'[:31]
        ws2 = wb.create_sheet(title=sheet_name)
        ws2.column_dimensions['A'].width = 30
        ws2.column_dimensions['B'].width = 40

        r = 1
        # Section: Thông tin chung
        ws2.merge_cells(start_row=r, start_column=1, end_row=r, end_column=2)
        ws2.cell(r, 1).value = 'THÔNG TIN CHUNG'
        ws2.cell(r, 1).font = section_font; ws2.cell(r, 1).fill = section_fill
        ws2.cell(r, 1).alignment = center; r += 1

        info_fields = [
            ('Họ và tên / Tên tập thể', name),
            ('Đơn vị', dx.don_vi.ten_don_vi if dx.don_vi else '—'),
            ('Danh hiệu đề xuất', ct.loai_danh_hieu or '—'),
            ('Năm học', dx.nam_hoc or '—'),
            ('Đối tượng', ct.doi_tuong or '—'),
        ]
        if qn:
            info_fields += [
                ('CCCD', qn.can_cuoc_cong_dan or '—'),
                ('Ngày sinh', qn.ngay_sinh.strftime('%d/%m/%Y') if qn.ngay_sinh else '—'),
                ('Ngày nhập ngũ', qn.ngay_nhap_ngu or '—'),
                ('Cấp bậc', qn.cap_bac or '—'),
                ('Chức vụ', qn.chuc_vu or '—'),
                ('Đơn vị trực thuộc', qn.don_vi_truc_thuoc or '—'),
                ('Học hàm / Học vị', f"{qn.hoc_ham or '—'} / {qn.hoc_vi or '—'}"),
                ('Trình độ học vấn', qn.trinh_do_hoc_van or '—'),
                ('Ngoại ngữ', qn.ngoai_ngu or '—'),
            ]

        for lbl, val in info_fields:
            ws2.cell(r, 1).value = lbl; ws2.cell(r, 1).font = label_font
            ws2.cell(r, 1).alignment = label_align; ws2.cell(r, 1).border = thin
            ws2.cell(r, 2).value = val; ws2.cell(r, 2).font = value_font
            ws2.cell(r, 2).alignment = val_align; ws2.cell(r, 2).border = thin
            r += 1

        if not is_tap_the:
            # Section: Tiêu chí xét duyệt
            ws2.merge_cells(start_row=r, start_column=1, end_row=r, end_column=2)
            ws2.cell(r, 1).value = 'TIÊU CHÍ XÉT DUYỆT'
            ws2.cell(r, 1).font = section_font; ws2.cell(r, 1).fill = section_fill
            ws2.cell(r, 1).alignment = center; r += 1
            for lbl, attr in TIEU_CHI_LABELS:
                val = getattr(ct, attr, None)
                if val is None:
                    val = '—'
                elif attr == 'danh_hieu_hv_gioi':
                    val = 'Có' if str(val) == 'true' else ('Không' if str(val) == 'false' else val)
                ws2.cell(r, 1).value = lbl; ws2.cell(r, 1).font = label_font
                ws2.cell(r, 1).alignment = label_align; ws2.cell(r, 1).border = thin
                ws2.cell(r, 2).value = str(val) if val is not None else '—'
                ws2.cell(r, 2).font = value_font
                ws2.cell(r, 2).alignment = val_align; ws2.cell(r, 2).border = thin
                r += 1
        else:
            # Tập thể: chỉ mục tiêu chí đặc biệt
            ws2.merge_cells(start_row=r, start_column=1, end_row=r, end_column=2)
            ws2.cell(r, 1).value = 'TIÊU CHÍ TẬP THỂ'
            ws2.cell(r, 1).font = section_font; ws2.cell(r, 1).fill = section_fill
            ws2.cell(r, 1).alignment = center; r += 1
            for lbl, attr in [('Mức độ HT nhiệm vụ', 'muc_do_hoan_thanh'),
                               ('Chủ trì ĐV danh hiệu', 'chu_tri_don_vi_danh_hieu'),
                               ('Ghi chú', 'ghi_chu')]:
                val = getattr(ct, attr, None) or '—'
                ws2.cell(r, 1).value = lbl; ws2.cell(r, 1).font = label_font
                ws2.cell(r, 1).alignment = label_align; ws2.cell(r, 1).border = thin
                ws2.cell(r, 2).value = val; ws2.cell(r, 2).font = value_font
                ws2.cell(r, 2).alignment = val_align; ws2.cell(r, 2).border = thin
                r += 1

        # Section: Kết quả biểu quyết Hội đồng
        ws2.merge_cells(start_row=r, start_column=1, end_row=r, end_column=2)
        ws2.cell(r, 1).value = 'KẾT QUẢ BIỂU QUYẾT HỘI ĐỒNG'
        ws2.cell(r, 1).font = section_font; ws2.cell(r, 1).fill = section_fill
        ws2.cell(r, 1).alignment = center; r += 1
        dong_y = 0
        for vt in HOI_DONG_VAI_TRO:
            bq = votes.get(vt)
            ket_qua = bq.ket_qua if bq else 'Chưa bỏ phiếu'
            if ket_qua == 'Đồng ý':
                dong_y += 1
            ws2.cell(r, 1).value = HOI_DONG_VAI_TRO_DISPLAY[vt]
            ws2.cell(r, 1).font = label_font; ws2.cell(r, 1).alignment = label_align; ws2.cell(r, 1).border = thin
            ws2.cell(r, 2).value = ket_qua; ws2.cell(r, 2).font = value_font
            ws2.cell(r, 2).alignment = val_align; ws2.cell(r, 2).border = thin
            if ket_qua == 'Đồng ý':
                ws2.cell(r, 2).fill = green_fill
            elif ket_qua == 'Không đồng ý':
                ws2.cell(r, 2).fill = _Fill(start_color='FFC7CE', end_color='FFC7CE', fill_type='solid')
            r += 1
        trang_thai_txt = ('Đã xác nhận khen thưởng' if row['is_confirmed']
                          else f'Đã bỏ phiếu: {dong_y}/{len(HOI_DONG_VAI_TRO)} Đồng ý')
        ws2.merge_cells(start_row=r, start_column=1, end_row=r, end_column=2)
        ws2.cell(r, 1).value = f'→ Tổng kết: {trang_thai_txt}'
        ws2.cell(r, 1).font = _Font(name='Times New Roman', bold=True, size=10)
        ws2.cell(r, 1).alignment = center
        if row['is_confirmed']:
            ws2.cell(r, 1).fill = green_fill

        ws2.page_setup.paperSize = 9  # A4
        ws2.page_setup.orientation = 'portrait'
        ws2.page_setup.fitToPage = True; ws2.page_setup.fitToWidth = 1; ws2.page_setup.fitToHeight = 0
        ws2.sheet_properties.pageSetUpPr.fitToPage = True

    wb['Tong hop'].protection.sheet = True
    wb['Tong hop'].protection.password = 'hktd@2025'

    output = BytesIO()
    wb.save(output)
    output.seek(0)
    fname = f'bang2_chitiet{"_" + nam_hoc_filter if nam_hoc_filter else ""}.xlsx'
    return send_file(output, as_attachment=True, download_name=fname,
                     mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')


@admin_bp.route('/reward-list/export')
@login_required
@admin_or_reward_viewer_required
def export_reward_list():
    """Export danh sách khen thưởng ra file Excel với đầy đủ thông tin."""
    nam_hoc_filter = request.args.get('nam_hoc', '')
    unit_filter = request.args.get('unit', '')
    danh_hieu_filter = request.args.get('danh_hieu', '')
    search_query = request.args.get('q', '').strip()

    query = KhenThuong.query

    if nam_hoc_filter:
        query = query.filter(KhenThuong.nam_hoc == nam_hoc_filter)
    if unit_filter:
        query = query.join(DonVi).filter(DonVi.ten_don_vi == unit_filter)
    if danh_hieu_filter:
        query = query.filter(KhenThuong.loai_danh_hieu == danh_hieu_filter)
    if search_query:
        query = query.filter(KhenThuong.ho_ten.ilike(f'%{search_query}%'))

    rewards = query.order_by(KhenThuong.ngay_duyet.desc()).all()

    # Create Excel workbook
    wb = Workbook()
    ws = wb.active
    ws.title = 'Danh sách khen thưởng'

    # Styles
    header_font = Font(name='Times New Roman', bold=True, size=13)
    sub_header_font = Font(name='Times New Roman', bold=True, size=11)
    col_header_font = Font(name='Times New Roman', bold=True, size=10)
    col_header_fill = PatternFill(start_color='1F4E79', end_color='1F4E79', fill_type='solid')
    col_header_font_white = Font(name='Times New Roman', bold=True, size=10, color='FFFFFF')
    data_font = Font(name='Times New Roman', size=10)
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin'),
    )
    center_align = Alignment(horizontal='center', vertical='center', wrap_text=True)
    left_align = Alignment(horizontal='left', vertical='center', wrap_text=True)

    # Title rows
    ws.merge_cells('A1:Y1')
    cell_title = ws['A1']
    cell_title.value = 'TRƯỜNG SĨ QUAN CHÍNH TRỊ'
    cell_title.font = Font(name='Times New Roman', bold=True, size=13)
    cell_title.alignment = center_align

    ws.merge_cells('A2:Y2')
    cell_sub = ws['A2']
    cell_sub.value = 'DANH SÁCH KHEN THƯỞNG'
    cell_sub.font = Font(name='Times New Roman', bold=True, size=14)
    cell_sub.alignment = center_align

    # Filter info row
    filter_parts = []
    if nam_hoc_filter:
        filter_parts.append(f'Năm học: {nam_hoc_filter}')
    if unit_filter:
        filter_parts.append(f'Đơn vị: {unit_filter}')
    if danh_hieu_filter:
        filter_parts.append(f'Danh hiệu: {danh_hieu_filter}')
    if search_query:
        filter_parts.append(f'Tìm kiếm: {search_query}')

    ws.merge_cells('A3:Y3')
    cell_filter = ws['A3']
    cell_filter.value = ' | '.join(filter_parts) if filter_parts else 'Tất cả dữ liệu'
    cell_filter.font = Font(name='Times New Roman', italic=True, size=10)
    cell_filter.alignment = center_align

    # Column headers (row 5)
    headers = [
        ('STT', 6),
        ('Họ và tên', 22),
        ('Cấp bậc', 14),
        ('Chức vụ', 20),
        ('Đối tượng', 18),
        ('Đơn vị', 30),
        ('Danh hiệu thi đua', 22),
        ('Năm học', 12),
        ('Hoàn thành NV', 16),
        ('Phiếu tín nhiệm', 14),
        ('KT Chính trị', 14),
        ('KT Điều lệnh', 14),
        ('Kỹ năng số', 14),
        ('ĐHQS', 14),
        ('Bắn súng', 12),
        ('Thể lực', 12),
        ('Điểm NCKH', 12),
        ('Nội dung NCKH', 30),
        ('XL Đảng viên', 22),
        ('XL Cán bộ', 22),
        ('XL Đoàn viên', 22),
        ('XL Phụ nữ', 22),
        ('KQ Đoàn thể', 22),
        ('KQ Phụ nữ', 22),
        ('Ngày phê duyệt', 16),
    ]

    for col_idx, (header_text, width) in enumerate(headers, 1):
        cell = ws.cell(row=5, column=col_idx, value=header_text)
        cell.font = col_header_font_white
        cell.fill = col_header_fill
        cell.alignment = center_align
        cell.border = thin_border
        ws.column_dimensions[get_column_letter(col_idx)].width = width

    # Data rows
    for row_idx, kt in enumerate(rewards, 1):
        row_num = row_idx + 5  # data starts at row 6

        # Load original DeXuatChiTiet for criteria fields
        ct = DeXuatChiTiet.query.get(kt.chi_tiet_id) if kt.chi_tiet_id else None

        # Load DanhGiaHangNam for annual evaluation fields
        dg = None
        if kt.quan_nhan_id and kt.nam_hoc:
            dg = DanhGiaHangNam.query.filter_by(
                quan_nhan_id=kt.quan_nhan_id,
                nam_hoc=kt.nam_hoc
            ).first()

        row_data = [
            row_idx,                                                    # STT
            kt.ho_ten,                                                  # Họ tên
            kt.cap_bac or '',                                           # Cấp bậc
            kt.chuc_vu or '',                                           # Chức vụ
            kt.doi_tuong or '',                                         # Đối tượng
            kt.don_vi.ten_don_vi if kt.don_vi else '',                 # Đơn vị
            kt.loai_danh_hieu or '',                                    # Danh hiệu
            kt.nam_hoc or '',                                           # Năm học
            (ct.muc_do_hoan_thanh or '') if ct else '',                # HTNV
            (ct.phieu_tin_nhiem or '') if ct else '',                   # Phiếu tín nhiệm
            (ct.kiem_tra_chinh_tri or '') if ct else '',                # KT CT
            (ct.kiem_tra_dieu_lenh or '') if ct else '',                # KT ĐL
            (ct.kiem_tra_tin_hoc or '') if ct else '',                  # Kỹ năng số
            (ct.dia_ly_quan_su or '') if ct else '',                    # ĐHQS
            (ct.ban_sung or '') if ct else '',                          # Bắn súng
            (ct.the_luc or '') if ct else '',                           # Thể lực
            (ct.diem_nckh or '') if ct else '',                         # Điểm NCKH
            (ct.nckh_noi_dung or '') if ct else '',                    # Nội dung NCKH
            (dg.xep_loai_dang_vien or '') if dg else '',               # XL Đảng viên
            (dg.xep_loai_can_bo or '') if dg else '',                  # XL Cán bộ
            (dg.xep_loai_doan_vien or '') if dg else '',               # XL Đoàn viên
            (dg.xep_loai_phu_nu or '') if dg else '',                  # XL Phụ nữ
            (ct.ket_qua_doan_the or '') if ct else '',                 # KQ Đoàn thể
            (ct.ket_qua_phu_nu or '') if ct else '',                   # KQ Phụ nữ
            kt.ngay_duyet.strftime('%d/%m/%Y') if kt.ngay_duyet else '',  # Ngày duyệt
        ]

        for col_idx, value in enumerate(row_data, 1):
            cell = ws.cell(row=row_num, column=col_idx, value=value)
            cell.font = data_font
            cell.border = thin_border
            if col_idx == 1:  # STT
                cell.alignment = center_align
            elif col_idx in (8, 10, 11, 12, 13, 14, 15, 16, 17, 19):  # center-aligned cols
                cell.alignment = center_align
            else:
                cell.alignment = left_align

    # Summary row
    summary_row = len(rewards) + 6
    ws.merge_cells(f'A{summary_row}:F{summary_row}')
    cell_sum = ws.cell(row=summary_row, column=1, value=f'Tổng cộng: {len(rewards)} cá nhân')
    cell_sum.font = Font(name='Times New Roman', bold=True, size=10)
    cell_sum.alignment = left_align

    # Signature section
    sig_row = summary_row + 2
    ws.merge_cells(f'T{sig_row}:Y{sig_row}')
    cell_date = ws.cell(row=sig_row, column=20,
                        value=f'Ngày {datetime.now().day} tháng {datetime.now().month} năm {datetime.now().year}')
    cell_date.font = Font(name='Times New Roman', italic=True, size=10)
    cell_date.alignment = center_align

    ws.merge_cells(f'T{sig_row+1}:Y{sig_row+1}')
    cell_signer = ws.cell(row=sig_row + 1, column=20, value='BAN THƯ KÝ HỘI ĐỒNG THI ĐUA KHEN THƯỞNG')
    cell_signer.font = Font(name='Times New Roman', bold=True, size=11)
    cell_signer.alignment = center_align

    # Set print area and page setup
    ws.page_setup.paperSize = 9
    ws.page_setup.orientation = 'landscape'
    ws.page_setup.fitToPage = True
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0
    ws.sheet_properties.pageSetUpPr.fitToPage = True

    # Freeze panes
    ws.freeze_panes = 'A6'

    ws.protection.sheet = True
    ws.protection.password = 'bth123'

    # Write to BytesIO
    output = BytesIO()
    wb.save(output)
    output.seek(0)

    # Generate filename
    filename_parts = ['DanhSachKhenThuong']
    if nam_hoc_filter:
        filename_parts.append(nam_hoc_filter.replace('-', '_'))
    filename_parts.append(datetime.now().strftime('%d%m%Y'))
    filename = '_'.join(filename_parts) + '.xlsx'

    return send_file(
        output,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        as_attachment=True,
        download_name=filename,
    )


@admin_bp.route('/reward-list/export-b3')
@login_required
@admin_or_reward_viewer_required
def export_b3_excel():
    """Export Bảng 3 – Danh sách khen thưởng đã xác nhận (gọn, có tóm tắt thành tích)."""
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    from openpyxl import Workbook

    nam_hoc_filter = request.args.get('nam_hoc', '')
    unit_filter    = request.args.get('unit', '')
    danh_hieu_filter = request.args.get('danh_hieu', '')
    search_query   = request.args.get('q', '').strip()

    query = KhenThuong.query
    if nam_hoc_filter:
        query = query.filter(KhenThuong.nam_hoc == nam_hoc_filter)
    if unit_filter:
        query = query.join(DonVi).filter(DonVi.ten_don_vi == unit_filter)
    if danh_hieu_filter:
        query = query.filter(KhenThuong.loai_danh_hieu == danh_hieu_filter)
    if search_query:
        query = query.filter(KhenThuong.ho_ten.ilike(f'%{search_query}%'))

    # Tách cá nhân và tập thể
    all_rewards = query.order_by(KhenThuong.don_vi_id, KhenThuong.loai_danh_hieu, KhenThuong.ho_ten).all()
    ca_nhan  = [kt for kt in all_rewards if kt.quan_nhan_id]
    tap_the  = [kt for kt in all_rewards if not kt.quan_nhan_id]

    # ── Styles ──────────────────────────────────────────────────────────────
    wb = Workbook()
    ws = wb.active
    ws.title = 'Danh sách KT'

    hdr_font   = Font(name='Times New Roman', bold=True, size=11)
    title_font = Font(name='Times New Roman', bold=True, size=13)
    sub_font   = Font(name='Times New Roman', bold=True, size=14)
    col_font   = Font(name='Times New Roman', bold=True, size=10, color='FFFFFF')
    data_font  = Font(name='Times New Roman', size=10)
    grp_font   = Font(name='Times New Roman', bold=True, italic=True, size=10)
    col_fill   = PatternFill(start_color='1F4E79', end_color='1F4E79', fill_type='solid')
    grp_fill   = PatternFill(start_color='D9E1F2', end_color='D9E1F2', fill_type='solid')
    thin       = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'),  bottom=Side(style='thin'),
    )
    c_align = Alignment(horizontal='center', vertical='center', wrap_text=True)
    l_align = Alignment(horizontal='left',   vertical='center', wrap_text=True)

    NUM_COLS = 7
    col_letters = [get_column_letter(i) for i in range(1, NUM_COLS + 1)]

    def _merge_title(row, text, font, fill=None):
        ws.merge_cells(f'A{row}:{col_letters[-1]}{row}')
        c = ws['A' + str(row)]
        c.value = text; c.font = font; c.alignment = c_align
        if fill: c.fill = fill

    # ── Tiêu đề ─────────────────────────────────────────────────────────────
    _merge_title(1, 'TRƯỜNG SĨ QUAN CHÍNH TRỊ', title_font)
    _merge_title(2, 'DANH SÁCH KHEN THƯỞNG ĐÃ XÁC NHẬN', sub_font)

    filter_txt = ' | '.join(filter(None, [
        f'Năm học: {nam_hoc_filter}' if nam_hoc_filter else '',
        f'Đơn vị: {unit_filter}' if unit_filter else '',
        f'Danh hiệu: {danh_hieu_filter}' if danh_hieu_filter else '',
        f'Tìm kiếm: {search_query}' if search_query else '',
    ])) or 'Tất cả dữ liệu'
    _merge_title(3, filter_txt, Font(name='Times New Roman', italic=True, size=10))

    # ── Header cột ──────────────────────────────────────────────────────────
    COL_WIDTHS = [6, 28, 14, 22, 30, 12, 48]
    
    COL_HEADERS = ['TT', 'Họ và tên', 'Cấp bậc', 'Chức vụ', 'Đơn vị', 'Danh hiệu', 'Tóm tắt thành tích']
    for ci, (hdr, w) in enumerate(zip(COL_HEADERS, COL_WIDTHS), 1):
        c = ws.cell(row=5, column=ci, value=hdr)
        c.font = col_font; c.fill = col_fill
        c.alignment = c_align; c.border = thin
        ws.column_dimensions[get_column_letter(ci)].width = w
    ws.row_dimensions[5].height = 28

    # ── Helper: tóm tắt thành tích ──────────────────────────────────────────
    def _tom_tat(kt):
        ct = DeXuatChiTiet.query.get(kt.chi_tiet_id) if kt.chi_tiet_id else None
        doi_tuong = (kt.doi_tuong or '').lower()
        is_hv = 'học viên' in doi_tuong

        if is_hv:
            parts = []
            if ct and ct.diem_tong_ket:
                parts.append(f'Điểm TK: {ct.diem_tong_ket}')
            if ct and ct.ket_qua_ren_luyen:
                parts.append(f'KQ rèn luyện: {ct.ket_qua_ren_luyen}')
            return '; '.join(parts) if parts else ''
        else:
            # Cán bộ, giảng viên, QNCN, CNV...
            if ct and ct.muc_do_hoan_thanh:
                return f'Hoàn thành NV: {ct.muc_do_hoan_thanh}'
            return ''

    # ── Hàm ghi một nhóm danh sách ──────────────────────────────────────────
    def _write_group(items, start_row, group_label):
        r = start_row
        # Nhóm header
        ws.merge_cells(f'A{r}:{col_letters[-1]}{r}')
        c = ws['A' + str(r)]
        c.value = group_label; c.font = grp_font
        c.fill = grp_fill; c.alignment = l_align; c.border = thin
        r += 1

        prev_dh = None
        stt = 0
        for kt in items:
            # Separator theo danh hiệu
            if kt.loai_danh_hieu != prev_dh:
                prev_dh = kt.loai_danh_hieu
                ws.merge_cells(f'A{r}:{col_letters[-1]}{r}')
                dh_c = ws['A' + str(r)]
                dh_c.value = f'── {kt.loai_danh_hieu} ──'
                dh_c.font = Font(name='Times New Roman', bold=True, italic=True, size=9, color='1F4E79')
                dh_c.alignment = c_align
                for ci in range(1, NUM_COLS + 1):
                    ws.cell(row=r, column=ci).border = thin
                r += 1

            stt += 1
            tom_tat = _tom_tat(kt)
            row_vals = [
                stt,
                kt.ho_ten,
                kt.cap_bac or '',
                kt.chuc_vu or '',
                kt.don_vi.ten_don_vi if kt.don_vi else '',
                kt.loai_danh_hieu or '',
                tom_tat,
            ]
            for ci, val in enumerate(row_vals, 1):
                c = ws.cell(row=r, column=ci, value=val)
                c.font = data_font; c.border = thin
                c.alignment = c_align if ci in (1, 6) else l_align
            ws.row_dimensions[r].height = 18
            r += 1
        return r

    # ── Ghi dữ liệu ─────────────────────────────────────────────────────────
    cur_row = 6
    if ca_nhan:
        cur_row = _write_group(ca_nhan, cur_row, f'I. DANH HIỆU CÁ NHÂN ({len(ca_nhan)} người)')
    if tap_the:
        if ca_nhan:
            cur_row += 1  # dòng trống giữa 2 nhóm
        cur_row = _write_group(tap_the, cur_row, f'II. DANH HIỆU TẬP THỂ ({len(tap_the)} đơn vị)')

    # Tổng kết
    ws.merge_cells(f'A{cur_row}:{col_letters[-1]}{cur_row}')
    sum_c = ws['A' + str(cur_row)]
    sum_c.value = f'Tổng cộng: {len(ca_nhan)} cá nhân, {len(tap_the)} tập thể'
    sum_c.font = Font(name='Times New Roman', bold=True, size=10)
    sum_c.alignment = l_align

    # Ký tên
    sig_row = cur_row + 2
    ws.merge_cells(f'E{sig_row}:{col_letters[-1]}{sig_row}')
    ws.cell(row=sig_row, column=5,
            value=f'Ngày {datetime.now().day} tháng {datetime.now().month} năm {datetime.now().year}'
            ).font = Font(name='Times New Roman', italic=True, size=10)
    ws.cell(row=sig_row, column=5).alignment = c_align
    ws.merge_cells(f'E{sig_row+1}:{col_letters[-1]}{sig_row+1}')
    ws.cell(row=sig_row + 1, column=5,
            value='BAN THƯ KÝ HỘI ĐỒNG THI ĐUA KHEN THƯỞNG'
            ).font = Font(name='Times New Roman', bold=True, size=11)
    ws.cell(row=sig_row + 1, column=5).alignment = c_align

    # Page setup
    ws.page_setup.paperSize = 9
    ws.page_setup.orientation = 'portrait'
    ws.page_setup.fitToPage = True
    ws.page_setup.fitToWidth = 1
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    ws.freeze_panes = 'A6'
    ws.protection.sheet = True
    ws.protection.password = 'bth123'

    output = BytesIO()
    wb.save(output)
    output.seek(0)

    parts = ['DanhSachKhenThuong_B3']
    if nam_hoc_filter:
        parts.append(nam_hoc_filter.replace('-', '_'))
    parts.append(datetime.now().strftime('%d%m%Y'))
    fname = '_'.join(parts) + '.xlsx'

    return send_file(output, as_attachment=True, download_name=fname,
                     mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')


@admin_bp.route('/reward-detail/<int:kt_id>')
@login_required
@admin_or_reward_viewer_required
def reward_detail(kt_id):
    """View detailed info for one KhenThuong record (individual award)."""
    kt = KhenThuong.query.get_or_404(kt_id)

    # Load the original DeXuatChiTiet for full criteria
    ct = DeXuatChiTiet.query.get(kt.chi_tiet_id) if kt.chi_tiet_id else None
    de_xuat = DeXuat.query.get(kt.de_xuat_id) if kt.de_xuat_id else None

    # Build per-department results for this individual
    dept_item_results = {}
    if de_xuat:
        for pd in de_xuat.phe_duyets:
            kq = None
            if ct:
                kq = KetQuaDuyetChiTiet.query.filter_by(
                    phe_duyet_id=pd.id, chi_tiet_id=ct.id
                ).first()
            dept_item_results[pd.phong_duyet] = {
                'phe_duyet': pd,
                'item_result': kq,
            }

    # All criteria field labels
    all_field_labels = {
        'muc_do_hoan_thanh': 'Hoàn thành nhiệm vụ',
        'phieu_tin_nhiem': 'Phiếu tín nhiệm',
        'kiem_tra_dieu_lenh': 'Kiểm tra điều lệnh',
        'ban_sung': 'Bắn súng',
        'the_luc': 'Thể lực',
        'kiem_tra_chinh_tri': 'Kiểm tra chính trị',
        'kiem_tra_tin_hoc': 'Kỹ năng số',
        'dia_ly_quan_su': 'Địa hình quân sự',
        'danh_hieu_gv_gioi': 'Danh hiệu GV giỏi',
        'dinh_muc_giang_day': 'Định mức giảng dạy',
        'ket_qua_kiem_tra_giang': 'Kết quả kiểm tra giảng',
        'thoi_gian_lao_dong_kh': 'Thời gian LĐ khoa học',
        'tien_do_pgs': 'Tiến độ PGS',
        'danh_hieu_hv_gioi': 'Danh hiệu HV giỏi',
        'diem_tong_ket': 'Điểm tổng kết',
        'ket_qua_thuc_hanh': 'Kết quả thực hành',
        'ket_qua_ren_luyen': 'Kết quả rèn luyện',
        'hinh_thuc_tot_nghiep': 'Hình thức thi tốt nghiệp',
        'diem_tn_ctd': 'Điểm CTĐ (tốt nghiệp)',
        'diem_tn_ct': 'Điểm CT (tốt nghiệp)',
        'diem_tn_ta': 'Điểm TA (tốt nghiệp)',
        'diem_tn_mon4': 'Điểm môn thứ 4 (tốt nghiệp)',
        'diem_tn_chuyennganh': 'Điểm chuyên ngành (tốt nghiệp)',
        'diem_tn_baove': 'Điểm bảo vệ KL (tốt nghiệp)',
        'ket_qua_doan_the': 'Kết quả đoàn thể',
        'chu_tri_don_vi_danh_hieu': 'Chủ trì ĐV danh hiệu',
        'diem_nckh': 'Điểm NCKH',
        'nckh_noi_dung': 'Nội dung NCKH',
        'nckh_minh_chung': 'Minh chứng NCKH',
        'mo_ta_khoa_hoc': 'Mô tả thành tích KH',
        'diem_tot_nghiep': 'Điểm tốt nghiệp (TB)',
        'minh_chung_thanh_tich_khac': 'MC thành tích khác',
        'thanh_tich_ca_nhan_khac': 'Thành tích cá nhân khác',
    }

    all_fields = [
        'muc_do_hoan_thanh', 'phieu_tin_nhiem',
        'kiem_tra_chinh_tri', 'kiem_tra_dieu_lenh', 'kiem_tra_tin_hoc',
        'dia_ly_quan_su', 'ban_sung', 'the_luc',
        'ket_qua_doan_the', 'chu_tri_don_vi_danh_hieu',
        'danh_hieu_gv_gioi', 'dinh_muc_giang_day', 'ket_qua_kiem_tra_giang',
        'tien_do_pgs', 'thoi_gian_lao_dong_kh',
        'danh_hieu_hv_gioi', 'diem_tong_ket', 'ket_qua_thuc_hanh', 'ket_qua_ren_luyen',
        'hinh_thuc_tot_nghiep',
        'diem_tn_ctd', 'diem_tn_ct', 'diem_tn_ta', 'diem_tn_mon4',
        'diem_tn_chuyennganh', 'diem_tn_baove',
        'diem_nckh', 'nckh_noi_dung', 'nckh_minh_chung',
        'mo_ta_khoa_hoc', 'diem_tot_nghiep', 'minh_chung_thanh_tich_khac',
        'thanh_tich_ca_nhan_khac',
    ]

    return render_template('admin/reward_detail.html',
                           kt=kt,
                           ct=ct,
                           de_xuat=de_xuat,
                           dept_item_results=dept_item_results,
                           dept_names=DEPT_NAMES,
                           all_field_labels=all_field_labels,
                           all_fields=all_fields)


# ===== Legacy routes - redirect to new locations =====

@admin_bp.route('/final-review')
@login_required
@admin_required
def final_review_list():
    """Redirect to tracking page filtered for awaiting final approval."""
    return redirect(url_for('admin.approval_tracking', status='Đã duyệt'))


@admin_bp.route('/final-review/<int:id>')
@login_required
@admin_required
def final_review_detail(id):
    """Redirect to tracking page (legacy route)."""
    return redirect(url_for('admin.approval_tracking'))


@admin_bp.route('/final-review/<int:id>/approve', methods=['POST'])
@login_required
@admin_required
def final_approve(id):
    """Legacy final approve - redirect to new endpoint."""
    return redirect(url_for('admin.final_approve_from_tracking', id=id), code=307)


@admin_bp.route('/final-review/<int:id>/reject', methods=['POST'])
@login_required
@admin_required
def final_reject(id):
    """Legacy final reject - redirect to new endpoint."""
    return redirect(url_for('admin.reject_from_tracking', id=id), code=307)


# ===== User/Unit management (unchanged) =====

@admin_bp.route('/users')
@login_required
@admin_required
def manage_users():
    users = User.query.order_by(User.role, User.username).all()
    units = DonVi.query.filter_by(is_active=True).order_by(DonVi.thu_tu).all()
    roles = [(r.value, ROLE_DISPLAY.get(r, r.name)) for r in Role]
    return render_template('admin/manage_users.html',
                           users=users, units=units, roles=roles)


@admin_bp.route('/users/create', methods=['POST'])
@login_required
@admin_required
def create_user():
    username = request.form.get('username', '').strip()
    password = request.form.get('password', '').strip()
    ho_ten = request.form.get('ho_ten', '').strip()
    role_val = request.form.get('role', '').strip()
    don_vi_id = request.form.get('don_vi_id', type=int)

    if not username or not password or not ho_ten or not role_val:
        flash('Vui lòng điền đầy đủ thông tin.', 'danger')
        return redirect(url_for('admin.manage_users'))

    if User.query.filter_by(username=username).first():
        flash('Tên đăng nhập đã tồn tại.', 'danger')
        return redirect(url_for('admin.manage_users'))

    try:
        role = Role(role_val)
    except ValueError:
        flash('Vai trò không hợp lệ.', 'danger')
        return redirect(url_for('admin.manage_users'))

    user = User(
        username=username,
        ho_ten=ho_ten,
        role=role,
        don_vi_id=don_vi_id if role == Role.UNIT_USER else None,
    )
    user.set_password(password)
    db.session.add(user)
    db.session.commit()
    flash(f'Đã tạo tài khoản: {username}', 'success')
    return redirect(url_for('admin.manage_users'))


@admin_bp.route('/users/<int:id>/toggle', methods=['POST'])
@login_required
@admin_required
def toggle_user(id):
    user = User.query.get_or_404(id)
    if user.id == current_user.id:
        flash('Không thể tự vô hiệu hóa tài khoản của mình.', 'danger')
        return redirect(url_for('admin.manage_users'))

    user.is_active_account = not user.is_active_account
    db.session.commit()
    status = 'kích hoạt' if user.is_active_account else 'vô hiệu hóa'
    flash(f'Đã {status} tài khoản: {user.username}', 'success')
    return redirect(url_for('admin.manage_users'))


@admin_bp.route('/users/<int:id>/reset-password', methods=['POST'])
@login_required
@admin_required
def reset_password(id):
    user = User.query.get_or_404(id)
    new_pw = request.form.get('new_password', '').strip()
    if not new_pw:
        flash('Mật khẩu mới không được để trống.', 'danger')
        return redirect(url_for('admin.manage_users'))

    user.set_password(new_pw)
    db.session.commit()
    flash(f'Đã đặt lại mật khẩu cho: {user.username}', 'success')
    return redirect(url_for('admin.manage_users'))


@admin_bp.route('/units')
@login_required
@admin_required
def manage_units():
    units_by_type = {}
    for loai in LoaiDonVi:
        units_by_type[loai] = DonVi.query.filter_by(loai_don_vi=loai)\
            .order_by(DonVi.thu_tu).all()

    loai_don_vi_list = [(l.value, l.name) for l in LoaiDonVi]
    return render_template('admin/manage_units.html',
                           units_by_type=units_by_type,
                           loai_don_vi_list=loai_don_vi_list)


@admin_bp.route('/units/create', methods=['POST'])
@login_required
@admin_required
def create_unit():
    ma_don_vi = request.form.get('ma_don_vi', '').strip()
    ten_don_vi = request.form.get('ten_don_vi', '').strip()
    loai_don_vi_val = request.form.get('loai_don_vi', '').strip()
    thu_tu = request.form.get('thu_tu', 0, type=int)

    if not ma_don_vi or not ten_don_vi or not loai_don_vi_val:
        flash('Vui lòng điền đầy đủ thông tin.', 'danger')
        return redirect(url_for('admin.manage_units'))

    # Check unique ma_don_vi
    if DonVi.query.filter_by(ma_don_vi=ma_don_vi).first():
        flash(f'Mã đơn vị "{ma_don_vi}" đã tồn tại.', 'danger')
        return redirect(url_for('admin.manage_units'))

    try:
        loai = LoaiDonVi(loai_don_vi_val)
    except ValueError:
        flash('Loại đơn vị không hợp lệ.', 'danger')
        return redirect(url_for('admin.manage_units'))

    unit = DonVi(
        ma_don_vi=ma_don_vi,
        ten_don_vi=ten_don_vi,
        loai_don_vi=loai,
        thu_tu=thu_tu,
        is_active=True,
    )
    db.session.add(unit)
    db.session.commit()
    flash(f'Đã tạo đơn vị: {ten_don_vi}', 'success')
    return redirect(url_for('admin.manage_units'))


@admin_bp.route('/units/<int:id>/edit', methods=['POST'])
@login_required
@admin_required
def edit_unit(id):
    unit = DonVi.query.get_or_404(id)
    ma_don_vi = request.form.get('ma_don_vi', '').strip()
    ten_don_vi = request.form.get('ten_don_vi', '').strip()
    loai_don_vi_val = request.form.get('loai_don_vi', '').strip()
    thu_tu = request.form.get('thu_tu', 0, type=int)

    if not ma_don_vi or not ten_don_vi or not loai_don_vi_val:
        flash('Vui lòng điền đầy đủ thông tin.', 'danger')
        return redirect(url_for('admin.manage_units'))

    # Check unique ma_don_vi (exclude current unit)
    existing = DonVi.query.filter_by(ma_don_vi=ma_don_vi).first()
    if existing and existing.id != unit.id:
        flash(f'Mã đơn vị "{ma_don_vi}" đã được sử dụng.', 'danger')
        return redirect(url_for('admin.manage_units'))

    try:
        loai = LoaiDonVi(loai_don_vi_val)
    except ValueError:
        flash('Loại đơn vị không hợp lệ.', 'danger')
        return redirect(url_for('admin.manage_units'))

    unit.ma_don_vi = ma_don_vi
    unit.ten_don_vi = ten_don_vi
    unit.loai_don_vi = loai
    unit.thu_tu = thu_tu
    db.session.commit()
    flash(f'Đã cập nhật đơn vị: {ten_don_vi}', 'success')
    return redirect(url_for('admin.manage_units'))


@admin_bp.route('/units/<int:id>/toggle', methods=['POST'])
@login_required
@admin_required
def toggle_unit(id):
    unit = DonVi.query.get_or_404(id)
    unit.is_active = not unit.is_active
    db.session.commit()
    status = 'kích hoạt' if unit.is_active else 'ngừng hoạt động'
    flash(f'Đã {status} đơn vị: {unit.ten_don_vi}', 'success')
    return redirect(url_for('admin.manage_units'))


@admin_bp.route('/units/<int:id>/update-thu-tu', methods=['POST'])
@login_required
@admin_required
def update_unit_thu_tu(id):
    """AJAX endpoint to update thu_tu inline."""
    unit = DonVi.query.get_or_404(id)
    try:
        thu_tu = request.json.get('thu_tu', 0)
        unit.thu_tu = int(thu_tu)
        db.session.commit()
        return jsonify({'success': True, 'message': 'Đã cập nhật thứ tự'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)}), 400


@admin_bp.route('/report')
@login_required
@admin_required
def report_summary():
    nam_hoc_filter = request.args.get('nam_hoc', '')

    # Get available nam_hoc values
    nam_hoc_list = db.session.query(DeXuat.nam_hoc).distinct()\
        .order_by(DeXuat.nam_hoc.desc()).all()
    nam_hoc_list = [n[0] for n in nam_hoc_list if n[0]]

    # Base queries (optionally filtered by nam_hoc)
    dx_query = DeXuat.query
    kt_query = KhenThuong.query
    if nam_hoc_filter:
        dx_query = dx_query.filter(DeXuat.nam_hoc == nam_hoc_filter)
        kt_query = kt_query.filter(KhenThuong.nam_hoc == nam_hoc_filter)

    # Overall stats
    total_personnel = QuanNhan.query.filter_by(is_active=True).count()
    total_nominations = dx_query.count()
    total_approved = dx_query.filter(
        DeXuat.trang_thai == TrangThaiDeXuat.PHE_DUYET_CUOI.value
    ).count()
    total_rewards = kt_query.count()

    # Stats by nomination status
    status_stats = {}
    for status in TrangThaiDeXuat:
        if status == TrangThaiDeXuat.NHAP:
            continue
        count = dx_query.filter(DeXuat.trang_thai == status.value).count()
        status_stats[status.value] = count

    # Stats by danh_hieu (from KhenThuong)
    danh_hieu_stats = {}
    for dh in LoaiDanhHieu:
        count = kt_query.filter(KhenThuong.loai_danh_hieu == dh.value).count()
        danh_hieu_stats[dh.value] = count

    # Stats by unit type (from KhenThuong)
    loai_dv_stats = {}
    for ldv in LoaiDonVi:
        count = kt_query.join(DonVi, KhenThuong.don_vi_id == DonVi.id)\
            .filter(DonVi.loai_don_vi == ldv).count()
        loai_dv_stats[ldv.value] = count

    # Per-unit stats
    units = DonVi.query.filter_by(is_active=True).order_by(DonVi.thu_tu).all()
    unit_stats = []
    for unit in units:
        personnel_count = QuanNhan.query.filter_by(don_vi_id=unit.id, is_active=True).count()

        nom_q = DeXuat.query.filter_by(don_vi_id=unit.id)
        kt_q = KhenThuong.query.filter_by(don_vi_id=unit.id)
        if nam_hoc_filter:
            nom_q = nom_q.filter(DeXuat.nam_hoc == nam_hoc_filter)
            kt_q = kt_q.filter(KhenThuong.nam_hoc == nam_hoc_filter)

        nomination_count = nom_q.count()
        approved_count = nom_q.filter(
            DeXuat.trang_thai == TrangThaiDeXuat.PHE_DUYET_CUOI.value
        ).count()
        reward_count = kt_q.count()

        # Per-danh-hieu breakdown for this unit
        unit_dh = {}
        for dh in LoaiDanhHieu:
            unit_dh[dh.value] = kt_q.filter(KhenThuong.loai_danh_hieu == dh.value).count()

        unit_stats.append({
            'unit': unit,
            'personnel': personnel_count,
            'nominations': nomination_count,
            'approved': approved_count,
            'rewards': reward_count,
            'by_danh_hieu': unit_dh,
        })

    # Department approval stats (how many individuals each dept has reviewed)
    dept_stats = {}
    for dept_name in DEPT_NAMES:
        total_reviewed = db.session.query(KetQuaDuyetChiTiet).join(
            PheDuyet, KetQuaDuyetChiTiet.phe_duyet_id == PheDuyet.id
        ).filter(PheDuyet.phong_duyet == dept_name)

        if nam_hoc_filter:
            total_reviewed = total_reviewed.join(
                DeXuat, PheDuyet.de_xuat_id == DeXuat.id
            ).filter(DeXuat.nam_hoc == nam_hoc_filter)

        total_count = total_reviewed.count()
        approved_count = total_reviewed.filter(
            KetQuaDuyetChiTiet.ket_qua == KetQuaDuyet.DONG_Y.value
        ).count()
        rejected_count = total_reviewed.filter(
            KetQuaDuyetChiTiet.ket_qua == KetQuaDuyet.TU_CHOI.value
        ).count()

        dept_stats[dept_name] = {
            'total': total_count,
            'approved': approved_count,
            'rejected': rejected_count,
        }

    # Chart data: rewards per unit (top units with rewards)
    chart_unit_names = []
    chart_unit_rewards = []
    for s in unit_stats:
        if s['rewards'] > 0:
            chart_unit_names.append(s['unit'].ma_don_vi)
            chart_unit_rewards.append(s['rewards'])

    return render_template('admin/report.html',
                           unit_stats=unit_stats,
                           total_personnel=total_personnel,
                           total_nominations=total_nominations,
                           total_approved=total_approved,
                           total_rewards=total_rewards,
                           status_stats=status_stats,
                           danh_hieu_stats=danh_hieu_stats,
                           loai_dv_stats=loai_dv_stats,
                           dept_stats=dept_stats,
                           dept_names=DEPT_NAMES,
                           nam_hoc_filter=nam_hoc_filter,
                           nam_hoc_list=nam_hoc_list,
                            chart_unit_names=chart_unit_names,
                           chart_unit_rewards=chart_unit_rewards)


# ------------------------------------------------------------------
# Admin: View all personnel across all units
# ------------------------------------------------------------------
# Admin: Danh sách xóa
# ------------------------------------------------------------------
@admin_bp.route('/personnel/deleted')
@login_required
@admin_required
def admin_deleted_personnel():
    search = request.args.get('search', '').strip()
    don_vi_id = request.args.get('don_vi_id', '', type=str)
    try:
        query = QuanNhan.query.filter(QuanNhan.is_deleted == True).join(DonVi)
    except Exception:
        query = QuanNhan.query.filter(QuanNhan.id == -1)  # empty safe fallback
    if search:
        query = query.filter(QuanNhan.ho_ten.ilike(f'%{search}%'))
    if don_vi_id:
        query = query.filter(QuanNhan.don_vi_id == int(don_vi_id))
    deleted_list = query.order_by(QuanNhan.deleted_at.desc()).all()
    units = DonVi.query.filter_by(is_active=True).order_by(DonVi.ten_don_vi).all()
    return render_template('admin/deleted_personnel.html',
                           deleted_list=deleted_list,
                           search=search,
                           don_vi_id=don_vi_id,
                           units=units)


@admin_bp.route('/personnel/deleted/bulk-action', methods=['POST'])
@login_required
@admin_required
def admin_bulk_deleted_action():
    """Bulk restore hoặc hard-delete nhiều quân nhân từ trang deleted."""
    action = request.form.get('action', '').strip()
    ids = request.form.getlist('ids')
    if not ids:
        flash('Chưa chọn quân nhân nào.', 'warning')
        return redirect(url_for('admin.admin_deleted_personnel'))
    ids = [int(i) for i in ids if i.isdigit()]
    personnel = QuanNhan.query.filter(QuanNhan.id.in_(ids), QuanNhan.is_deleted == True).all()
    if action == 'restore':
        for qn in personnel:
            qn.is_deleted = False
            qn.is_active = True
            qn.deleted_at = None
            qn.deleted_by_id = None
        db.session.commit()
        flash(f'Đã khôi phục {len(personnel)} quân nhân.', 'success')
    elif action == 'hard_delete':
        names = [qn.ho_ten for qn in personnel]
        for qn in personnel:
            db.session.delete(qn)
        db.session.commit()
        flash(f'Đã xóa vĩnh viễn {len(names)} quân nhân.', 'danger')
    else:
        flash('Hành động không hợp lệ.', 'danger')
    return redirect(url_for('admin.admin_deleted_personnel'))


# ------------------------------------------------------------------
@admin_bp.route('/personnel')
@login_required
@admin_required
def all_personnel():
    search = request.args.get('search', '').strip()
    don_vi_id = request.args.get('don_vi_id', '', type=str)
    doi_tuong = request.args.get('doi_tuong', '').strip()
    cap_bac = request.args.get('cap_bac', '').strip()
    chuc_vu = request.args.get('chuc_vu', '').strip()
    sort_by = request.args.get('sort_by', 'chuc_vu').strip()  # Changed default from 'don_vi' to 'chuc_vu'
    page = request.args.get('page', 1, type=int)
    per_page = request.args.get('per_page', 30, type=int)
    if per_page not in (20, 30, 50, 100):
        per_page = 30

    from sqlalchemy.sql.expression import collate as _collate
    from sqlalchemy import case as _case
    from app.models.catalog import ChucVuOption as _ChucVuOption
    _chuc_vu_alias = db.aliased(_ChucVuOption)
    base_q = QuanNhan.query.filter(
    QuanNhan.is_active == True,
    QuanNhan.is_deleted == False,
    db.or_(QuanNhan.is_chuyen_vung.is_(None), QuanNhan.is_chuyen_vung == False)  # ★
)
    # try:
    #     base_q = base_q.filter(QuanNhan.is_deleted == False,QuanNhan.is_chuyen_vung==False)
    # except Exception:
    #     pass
    query = base_q.join(DonVi)\
        .outerjoin(_chuc_vu_alias,
                   _collate(_chuc_vu_alias.ten, 'utf8mb4_unicode_ci') ==
                   _collate(QuanNhan.chuc_vu, 'utf8mb4_unicode_ci'))

    if search:
        query = query.filter(QuanNhan.ho_ten.ilike(f'%{search}%'))
    if don_vi_id:
        query = query.filter(QuanNhan.don_vi_id == int(don_vi_id))
    if doi_tuong:
        query = query.filter(QuanNhan.doi_tuong == doi_tuong)
    if cap_bac:
        query = query.filter(QuanNhan.cap_bac.ilike(f'%{cap_bac}%'))
    if chuc_vu:
        query = query.filter(QuanNhan.chuc_vu.ilike(f'%{chuc_vu}%'))

    if sort_by == 'ho_ten':
        query = query.order_by(QuanNhan.ho_ten.asc())
    elif sort_by == 'cap_bac':
        query = query.order_by(QuanNhan.cap_bac.asc(), QuanNhan.ho_ten.asc())
    elif sort_by == 'doi_tuong':
        query = query.order_by(QuanNhan.doi_tuong.asc(), QuanNhan.ho_ten.asc())
    elif sort_by == 'chuc_vu':  # New sort option
        query = query.order_by(
            _case((_chuc_vu_alias.thu_tu.is_(None), 1), else_=0),
            _chuc_vu_alias.thu_tu.asc(),
            QuanNhan.chuc_vu.asc(),
            QuanNhan.ho_ten.asc()
        )
    else:  # don_vi
        query = query.order_by(
            DonVi.thu_tu,
            DonVi.ten_don_vi,
            _case((_chuc_vu_alias.thu_tu.is_(None), 1), else_=0),
            _chuc_vu_alias.thu_tu.asc(),
            QuanNhan.chuc_vu.asc(),
            QuanNhan.ho_ten.asc()
        )
    personnel = query.paginate(page=page, per_page=per_page, error_out=False)

    units = DonVi.query.filter_by(is_active=True).order_by(DonVi.thu_tu, DonVi.ten_don_vi).all()
    doi_tuong_list = _get_doi_tuong_option_list()

    # cap_bac and chuc_vu distinct lists for dropdowns
    from sqlalchemy import distinct as _distinct
    cap_bac_list = [r[0] for r in db.session.query(_distinct(QuanNhan.cap_bac)).filter(
        QuanNhan.is_active == True, QuanNhan.cap_bac != None, QuanNhan.cap_bac != ''
    ).order_by(QuanNhan.cap_bac).all()]
    chuc_vu_list = [r[0] for r in db.session.query(_distinct(QuanNhan.chuc_vu)).filter(
        QuanNhan.is_active == True, QuanNhan.chuc_vu != None, QuanNhan.chuc_vu != ''
    ).order_by(QuanNhan.chuc_vu).all()]

    return render_template('admin/all_personnel.html',
                           personnel=personnel,
                           search=search,
                           don_vi_id=don_vi_id,
                           doi_tuong_filter=doi_tuong,
                           cap_bac_filter=cap_bac,
                           chuc_vu_filter=chuc_vu,
                           sort_by=sort_by,
                           per_page=per_page,
                           units=units,
                           doi_tuong_list=doi_tuong_list,
                           cap_bac_list=cap_bac_list,
                           chuc_vu_list=chuc_vu_list)


# ------------------------------------------------------------------
# Admin: View personnel detail (any unit)
# ------------------------------------------------------------------
@admin_bp.route('/personnel/<int:id>')
@login_required
@admin_required
def admin_personnel_detail(id):
    qn = QuanNhan.query.get_or_404(id)

    # Nomination history: all DeXuatChiTiet rows for this person
    nominations_history = DeXuatChiTiet.query.filter_by(quan_nhan_id=qn.id)\
        .join(DeXuat, DeXuatChiTiet.de_xuat_id == DeXuat.id)\
        .order_by(DeXuat.nam_hoc.desc()).all()

    # Reward history: all KhenThuong rows for this person
    rewards_history = KhenThuong.query.filter_by(quan_nhan_id=qn.id)\
        .order_by(KhenThuong.nam_hoc.desc()).all()

    return render_template('admin/personnel_detail.html', qn=qn,
                           loai_chung_chi_list=[e.value for e in LoaiChungChi],
                           nominations_history=nominations_history,
                           rewards_history=rewards_history)


# ------------------------------------------------------------------
# Admin: Bulk action on personnel (across all units)
# ------------------------------------------------------------------
@admin_bp.route('/personnel/bulk-action', methods=['POST'])
@login_required
@admin_required
def admin_bulk_action():
    data = request.get_json()
    if not data:
        return jsonify({'success': False, 'message': 'Dữ liệu không hợp lệ'}), 400
    action = data.get('action')
    ids = data.get('ids', [])
    if not ids:
        return jsonify({'success': False, 'message': 'Không có quân nhân nào được chọn'}), 400

    qn_list = QuanNhan.query.filter(QuanNhan.id.in_(ids), QuanNhan.is_active == True).all()

    if action == 'doi_tuong':
        doi_tuong = data.get('doi_tuong', '').strip()
        if not doi_tuong:
            return jsonify({'success': False, 'message': 'Chưa chọn đối tượng'}), 400
        for qn in qn_list:
            qn.doi_tuong = doi_tuong
        db.session.commit()
        return jsonify({'success': True, 'message': f'Đã cập nhật đối tượng cho {len(qn_list)} quân nhân'})

    elif action == 'chuyen_vung':
        from datetime import datetime as _dt
        for qn in qn_list:
            qn.is_chuyen_vung = True
            qn.ngay_chuyen_vung = _dt.utcnow()
        db.session.commit()
        return jsonify({'success': True, 'message': f'Đã chuyển vùng {len(qn_list)} quân nhân'})

    elif action == 'chuyen_don_vi':
        don_vi_id = data.get('don_vi_id')
        ly_do = data.get('ly_do', '').strip()
        if not don_vi_id:
            return jsonify({'success': False, 'message': 'Chưa chọn đơn vị đích'}), 400
        target_unit = DonVi.query.get(don_vi_id)
        if not target_unit:
            return jsonify({'success': False, 'message': 'Đơn vị đích không tồn tại'}), 404
        from app.models.transfer import ChuyenDonVi as _ChuyenDonVi, TrangThaiChuyen as _TrangThaiChuyen
        for qn in qn_list:
            chuyen = _ChuyenDonVi(
                quan_nhan_id=qn.id,
                don_vi_nguon_id=qn.don_vi_id,
                don_vi_dich_id=don_vi_id,
                ly_do=ly_do or None,
                trang_thai=_TrangThaiChuyen.PENDING,
                nguoi_tao_id=current_user.id,
            )
            db.session.add(chuyen)
        db.session.commit()
        return jsonify({'success': True, 'message': f'Đã tạo yêu cầu chuyển đơn vị cho {len(qn_list)} quân nhân'})

    elif action == 'delete':
        from datetime import datetime as _dt
        for qn in qn_list:
            qn.is_active = False
            try:
                qn.is_deleted = True
                qn.deleted_at = _dt.utcnow()
                qn.deleted_by_id = current_user.id
            except Exception:
                pass
        db.session.commit()
        return jsonify({'success': True, 'message': f'Đã xóa {len(qn_list)} quân nhân'})

    else:
        return jsonify({'success': False, 'message': 'Hành động không hợp lệ'}), 400


# ------------------------------------------------------------------
# Admin: Edit personnel (any unit)
# ------------------------------------------------------------------
@admin_bp.route('/personnel/<int:id>/edit', methods=['GET', 'POST'])
@login_required
@admin_required
def admin_personnel_edit(id):
    qn = QuanNhan.query.get_or_404(id)

    if request.method == 'POST':
        qn.ho_ten = request.form.get('ho_ten', '').strip()
        qn.cap_bac = request.form.get('cap_bac', '').strip() or None
        qn.chuc_danh = None
        qn.chuc_vu = request.form.get('chuc_vu', '').strip() or None
        qn.don_vi_truc_thuoc = request.form.get('don_vi_truc_thuoc', '').strip() or None
        qn.can_cuoc_cong_dan = request.form.get('can_cuoc_cong_dan', '').strip() or None
        qn.doi_tuong = request.form.get('doi_tuong', '').strip() or None
        qn.hoc_ham = request.form.get('hoc_ham', 'Không').strip()
        qn.hoc_vi = request.form.get('hoc_vi', 'Không').strip()
        qn.trinh_do_hoc_van = request.form.get('trinh_do_hoc_van', '').strip() or None
        qn.ngoai_ngu = request.form.get('ngoai_ngu', '').strip() or None
        qn.ngay_nhap_ngu = request.form.get('ngay_nhap_ngu', '').strip() or None
        qn.la_chi_huy = 'la_chi_huy' in request.form
        qn.la_bi_thu = 'la_bi_thu' in request.form

        ns_str = request.form.get('ngay_sinh', '').strip()
        if ns_str:
            try:
                qn.ngay_sinh = datetime.strptime(ns_str, '%Y-%m-%d').date()
            except ValueError:
                pass
        else:
            qn.ngay_sinh = None

        if not qn.ho_ten:
            flash('Họ tên không được để trống.', 'danger')
        else:
            db.session.commit()
            flash('Đã cập nhật thông tin.', 'success')
            return redirect(url_for('admin.admin_personnel_detail', id=qn.id))

    return render_template('admin/personnel_edit.html', qn=qn,
                           cap_bac_list=[x.ten for x in CapBacOption.query.filter_by(is_active=True).order_by(CapBacOption.thu_tu, CapBacOption.ten).all()] or [e.value for e in CapBac],
                           hoc_ham_list=[e.value for e in HocHam],
                           hoc_vi_list=[e.value for e in HocVi],
                           doi_tuong_list=_get_doi_tuong_option_list(),
                           chuc_vu_options=ChucVuOption.query.filter_by(is_active=True).order_by(ChucVuOption.thu_tu, ChucVuOption.ten).all())


@admin_bp.route('/doi-tuong')
@login_required
@admin_required
def manage_doi_tuong():
    if DoiTuongOption.query.count() == 0:
        defaults = [e.value for e in DoiTuong]
        for idx, ten in enumerate(defaults, start=1):
            db.session.add(DoiTuongOption(ten=ten, thu_tu=idx, is_active=True))
        db.session.commit()
    items = DoiTuongOption.query.order_by(DoiTuongOption.thu_tu, DoiTuongOption.ten).all()
    return render_template('admin/manage_doi_tuong.html', items=items)


@admin_bp.route('/doi-tuong/create', methods=['POST'])
@login_required
@admin_required
def create_doi_tuong():
    ten = request.form.get('ten', '').strip()
    thu_tu = request.form.get('thu_tu', 0, type=int)
    if not ten:
        flash('Tên đối tượng không được để trống.', 'danger')
        return redirect(url_for('admin.manage_doi_tuong'))
    if DoiTuongOption.query.filter_by(ten=ten).first():
        flash('Đối tượng đã tồn tại.', 'warning')
        return redirect(url_for('admin.manage_doi_tuong'))
    db.session.add(DoiTuongOption(ten=ten, thu_tu=thu_tu, is_active=True))
    db.session.commit()
    flash('Đã thêm đối tượng.', 'success')
    return redirect(url_for('admin.manage_doi_tuong'))


@admin_bp.route('/doi-tuong/<int:id>/edit', methods=['POST'])
@login_required
@admin_required
def edit_doi_tuong(id):
    item = DoiTuongOption.query.get_or_404(id)
    ten = request.form.get('ten', '').strip()
    thu_tu = request.form.get('thu_tu', 0, type=int)
    if not ten:
        flash('Tên đối tượng không được để trống.', 'danger')
        return redirect(url_for('admin.manage_doi_tuong'))
    dup = DoiTuongOption.query.filter(DoiTuongOption.ten == ten, DoiTuongOption.id != id).first()
    if dup:
        flash('Tên đối tượng đã tồn tại.', 'warning')
        return redirect(url_for('admin.manage_doi_tuong'))
    item.ten = ten
    item.thu_tu = thu_tu
    db.session.commit()
    flash('Đã cập nhật đối tượng.', 'success')
    return redirect(url_for('admin.manage_doi_tuong'))


@admin_bp.route('/doi-tuong/<int:id>/toggle', methods=['POST'])
@login_required
@admin_required
def toggle_doi_tuong(id):
    item = DoiTuongOption.query.get_or_404(id)
    item.is_active = not item.is_active
    db.session.commit()
    flash('Đã cập nhật trạng thái đối tượng.', 'success')
    return redirect(url_for('admin.manage_doi_tuong'))


@admin_bp.route('/doi-tuong/<int:id>/delete', methods=['POST'])
@login_required
@admin_required
def delete_doi_tuong(id):
    item = DoiTuongOption.query.get_or_404(id)
    db.session.delete(item)
    db.session.commit()
    flash('Đã xóa đối tượng.', 'success')
    return redirect(url_for('admin.manage_doi_tuong'))


@admin_bp.route('/doi-tuong/<int:id>/update-thu-tu', methods=['POST'])
@login_required
@admin_required
def update_doi_tuong_thu_tu(id):
    """AJAX endpoint to update thu_tu inline."""
    item = DoiTuongOption.query.get_or_404(id)
    try:
        thu_tu = request.json.get('thu_tu', 0)
        item.thu_tu = int(thu_tu)
        db.session.commit()
        return jsonify({'success': True, 'message': 'Đã cập nhật thứ tự'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)}), 400


@admin_bp.route('/diem-quy-dinh')
@login_required
@admin_required
def manage_diem_quy_dinh():
    rules = DiemQuyDinhDanhHieu.query.order_by(
        DiemQuyDinhDanhHieu.loai_danh_hieu,
        DiemQuyDinhDanhHieu.tieu_chi_field,
    ).all()
    danh_hieus = DanhHieu.query.filter_by(is_active=True).order_by(DanhHieu.thu_tu, DanhHieu.ten_danh_hieu).all()
    return render_template('admin/manage_diem_quy_dinh.html',
                           rules=rules,
                           danh_hieus=danh_hieus,
                           diem_fields=DIEM_FIELDS,
                           diem_field_labels=DIEM_FIELD_LABELS)


@admin_bp.route('/diem-quy-dinh/create', methods=['POST'])
@login_required
@admin_required
def create_diem_quy_dinh():
    loai_danh_hieu = request.form.get('loai_danh_hieu', '').strip()
    tieu_chi_field = request.form.get('tieu_chi_field', '').strip()
    min_diem = request.form.get('min_diem', '').strip()

    if not loai_danh_hieu or not tieu_chi_field or not min_diem:
        flash('Vui lòng nhập đầy đủ thông tin quy định điểm.', 'danger')
        return redirect(url_for('admin.manage_diem_quy_dinh'))

    if tieu_chi_field not in DIEM_FIELDS:
        flash('Tiêu chí điểm không hợp lệ.', 'danger')
        return redirect(url_for('admin.manage_diem_quy_dinh'))

    dup = DiemQuyDinhDanhHieu.query.filter_by(
        loai_danh_hieu=loai_danh_hieu,
        tieu_chi_field=tieu_chi_field
    ).first()
    if dup:
        flash('Quy định cho danh hiệu + tiêu chí này đã tồn tại.', 'warning')
        return redirect(url_for('admin.manage_diem_quy_dinh'))

    db.session.add(DiemQuyDinhDanhHieu(
        loai_danh_hieu=loai_danh_hieu,
        tieu_chi_field=tieu_chi_field,
        min_diem=min_diem,
        is_active=True,
    ))
    db.session.commit()
    flash('Đã thêm quy định điểm.', 'success')
    return redirect(url_for('admin.manage_diem_quy_dinh'))


@admin_bp.route('/diem-quy-dinh/<int:id>/edit', methods=['POST'])
@login_required
@admin_required
def edit_diem_quy_dinh(id):
    row = DiemQuyDinhDanhHieu.query.get_or_404(id)
    loai_danh_hieu = request.form.get('loai_danh_hieu', '').strip()
    tieu_chi_field = request.form.get('tieu_chi_field', '').strip()
    min_diem = request.form.get('min_diem', '').strip()

    if not loai_danh_hieu or not tieu_chi_field or not min_diem:
        flash('Vui lòng nhập đầy đủ thông tin quy định điểm.', 'danger')
        return redirect(url_for('admin.manage_diem_quy_dinh'))
    if tieu_chi_field not in DIEM_FIELDS:
        flash('Tiêu chí điểm không hợp lệ.', 'danger')
        return redirect(url_for('admin.manage_diem_quy_dinh'))

    dup = DiemQuyDinhDanhHieu.query.filter(
        DiemQuyDinhDanhHieu.loai_danh_hieu == loai_danh_hieu,
        DiemQuyDinhDanhHieu.tieu_chi_field == tieu_chi_field,
        DiemQuyDinhDanhHieu.id != id,
    ).first()
    if dup:
        flash('Quy định cho danh hiệu + tiêu chí này đã tồn tại.', 'warning')
        return redirect(url_for('admin.manage_diem_quy_dinh'))

    row.loai_danh_hieu = loai_danh_hieu
    row.tieu_chi_field = tieu_chi_field
    row.min_diem = min_diem
    db.session.commit()
    flash('Đã cập nhật quy định điểm.', 'success')
    return redirect(url_for('admin.manage_diem_quy_dinh'))


@admin_bp.route('/diem-quy-dinh/<int:id>/toggle', methods=['POST'])
@login_required
@admin_required
def toggle_diem_quy_dinh(id):
    row = DiemQuyDinhDanhHieu.query.get_or_404(id)
    row.is_active = not row.is_active
    db.session.commit()
    flash('Đã cập nhật trạng thái quy định điểm.', 'success')
    return redirect(url_for('admin.manage_diem_quy_dinh'))


@admin_bp.route('/diem-quy-dinh/<int:id>/delete', methods=['POST'])
@login_required
@admin_required
def delete_diem_quy_dinh(id):
    row = DiemQuyDinhDanhHieu.query.get_or_404(id)
    db.session.delete(row)
    db.session.commit()
    flash('Đã xóa quy định điểm.', 'success')
    return redirect(url_for('admin.manage_diem_quy_dinh'))


@admin_bp.route('/cap-bac')
@login_required
@admin_required
def manage_cap_bac():
    items = CapBacOption.query.order_by(CapBacOption.thu_tu, CapBacOption.ten).all()
    return render_template('admin/manage_cap_bac.html', items=items)


@admin_bp.route('/cap-bac/create', methods=['POST'])
@login_required
@admin_required
def create_cap_bac():
    ten = request.form.get('ten', '').strip()
    thu_tu = request.form.get('thu_tu', 0, type=int)
    if not ten:
        flash('Tên cấp bậc không được để trống.', 'danger')
        return redirect(url_for('admin.manage_cap_bac'))
    if CapBacOption.query.filter_by(ten=ten).first():
        flash('Cấp bậc đã tồn tại.', 'warning')
        return redirect(url_for('admin.manage_cap_bac'))
    db.session.add(CapBacOption(ten=ten, thu_tu=thu_tu, is_active=True))
    db.session.commit()
    flash('Đã thêm cấp bậc.', 'success')
    return redirect(url_for('admin.manage_cap_bac'))


@admin_bp.route('/cap-bac/<int:id>/edit', methods=['POST'])
@login_required
@admin_required
def edit_cap_bac(id):
    item = CapBacOption.query.get_or_404(id)
    ten = request.form.get('ten', '').strip()
    thu_tu = request.form.get('thu_tu', 0, type=int)
    if not ten:
        flash('Tên cấp bậc không được để trống.', 'danger')
        return redirect(url_for('admin.manage_cap_bac'))
    dup = CapBacOption.query.filter(CapBacOption.ten == ten, CapBacOption.id != id).first()
    if dup:
        flash('Tên cấp bậc đã tồn tại.', 'warning')
        return redirect(url_for('admin.manage_cap_bac'))
    item.ten = ten
    item.thu_tu = thu_tu
    db.session.commit()
    flash('Đã cập nhật cấp bậc.', 'success')
    return redirect(url_for('admin.manage_cap_bac'))


@admin_bp.route('/cap-bac/<int:id>/toggle', methods=['POST'])
@login_required
@admin_required
def toggle_cap_bac(id):
    item = CapBacOption.query.get_or_404(id)
    item.is_active = not item.is_active
    db.session.commit()
    flash('Đã cập nhật trạng thái cấp bậc.', 'success')
    return redirect(url_for('admin.manage_cap_bac'))


@admin_bp.route('/cap-bac/<int:id>/delete', methods=['POST'])
@login_required
@admin_required
def delete_cap_bac(id):
    item = CapBacOption.query.get_or_404(id)
    db.session.delete(item)
    db.session.commit()
    flash('Đã xóa cấp bậc.', 'success')
    return redirect(url_for('admin.manage_cap_bac'))


@admin_bp.route('/cap-bac/<int:id>/update-thu-tu', methods=['POST'])
@login_required
@admin_required
def update_cap_bac_thu_tu(id):
    """AJAX endpoint to update thu_tu inline."""
    item = CapBacOption.query.get_or_404(id)
    try:
        thu_tu = request.json.get('thu_tu', 0)
        item.thu_tu = int(thu_tu)
        db.session.commit()
        return jsonify({'success': True, 'message': 'Đã cập nhật thứ tự'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)}), 400


@admin_bp.route('/chuc-vu')
@login_required
@admin_required
def manage_chuc_vu():
    items = ChucVuOption.query.order_by(ChucVuOption.thu_tu, ChucVuOption.ten).all()
    return render_template('admin/manage_chuc_vu.html', items=items)


@admin_bp.route('/chuc-vu/create', methods=['POST'])
@login_required
@admin_required
def create_chuc_vu():
    ten = request.form.get('ten', '').strip()
    thu_tu = request.form.get('thu_tu', 0, type=int)
    if not ten:
        flash('Tên chức vụ không được để trống.', 'danger')
        return redirect(url_for('admin.manage_chuc_vu'))
    if ChucVuOption.query.filter_by(ten=ten).first():
        flash('Chức vụ đã tồn tại.', 'warning')
        return redirect(url_for('admin.manage_chuc_vu'))
    db.session.add(ChucVuOption(ten=ten, thu_tu=thu_tu, is_active=True))
    db.session.commit()
    flash('Đã thêm chức vụ.', 'success')
    return redirect(url_for('admin.manage_chuc_vu'))


@admin_bp.route('/chuc-vu/<int:id>/edit', methods=['POST'])
@login_required
@admin_required
def edit_chuc_vu(id):
    item = ChucVuOption.query.get_or_404(id)
    ten = request.form.get('ten', '').strip()
    thu_tu = request.form.get('thu_tu', 0, type=int)
    if not ten:
        flash('Tên chức vụ không được để trống.', 'danger')
        return redirect(url_for('admin.manage_chuc_vu'))
    dup = ChucVuOption.query.filter(ChucVuOption.ten == ten, ChucVuOption.id != id).first()
    if dup:
        flash('Tên chức vụ đã tồn tại.', 'warning')
        return redirect(url_for('admin.manage_chuc_vu'))
    item.ten = ten
    item.thu_tu = thu_tu
    db.session.commit()
    flash('Đã cập nhật chức vụ.', 'success')
    return redirect(url_for('admin.manage_chuc_vu'))


@admin_bp.route('/chuc-vu/<int:id>/toggle', methods=['POST'])
@login_required
@admin_required
def toggle_chuc_vu(id):
    item = ChucVuOption.query.get_or_404(id)
    item.is_active = not item.is_active
    db.session.commit()
    flash('Đã cập nhật trạng thái chức vụ.', 'success')
    return redirect(url_for('admin.manage_chuc_vu'))


@admin_bp.route('/chuc-vu/<int:id>/delete', methods=['POST'])
@login_required
@admin_required
def delete_chuc_vu(id):
    item = ChucVuOption.query.get_or_404(id)
    db.session.delete(item)
    db.session.commit()
    flash('Đã xóa chức vụ.', 'success')
    return redirect(url_for('admin.manage_chuc_vu'))


@admin_bp.route('/chuc-vu/<int:id>/update-thu-tu', methods=['POST'])
@login_required
@admin_required
def update_chuc_vu_thu_tu(id):
    """AJAX endpoint to update thu_tu inline."""
    item = ChucVuOption.query.get_or_404(id)
    try:
        thu_tu = request.json.get('thu_tu', 0)
        item.thu_tu = int(thu_tu)
        db.session.commit()
        return jsonify({'success': True, 'message': 'Đã cập nhật thứ tự'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)}), 400


# ------------------------------------------------------------------
# Admin: Add certificate for any personnel
# ------------------------------------------------------------------
@admin_bp.route('/personnel/<int:id>/certificate', methods=['POST'])
@login_required
@admin_required
def admin_add_certificate(id):
    qn = QuanNhan.query.get_or_404(id)

    ten = request.form.get('ten_chung_chi', '').strip()
    if not ten:
        flash('Tên chứng chỉ không được để trống.', 'danger')
        return redirect(url_for('admin.admin_personnel_detail', id=qn.id))

    duong_dan_anh = None
    file = request.files.get('anh_minh_chung')
    if file and file.filename:
        duong_dan_anh = save_upload(file, 'certificates')

    ngay_cap = None
    nc_str = request.form.get('ngay_cap', '').strip()
    if nc_str:
        try:
            ngay_cap = datetime.strptime(nc_str, '%Y-%m-%d').date()
        except ValueError:
            pass

    cc = ChungChi(
        quan_nhan_id=qn.id,
        loai=request.form.get('loai', LoaiChungChi.THANH_TICH_KHAC.value),
        ten_chung_chi=ten,
        so_hieu=request.form.get('so_hieu', '').strip() or None,
        ngay_cap=ngay_cap,
        co_quan_cap=request.form.get('co_quan_cap', '').strip() or None,
        duong_dan_anh=duong_dan_anh,
        ghi_chu=request.form.get('ghi_chu', '').strip() or None,
    )
    db.session.add(cc)
    db.session.commit()
    flash(f'Đã thêm: {ten}', 'success')
    return redirect(url_for('admin.admin_personnel_detail', id=qn.id))


# ------------------------------------------------------------------
# Admin: Delete certificate for any personnel
# ------------------------------------------------------------------
@admin_bp.route('/certificate/<int:id>/delete', methods=['POST'])
@login_required
@admin_required
def admin_delete_certificate(id):
    cc = ChungChi.query.get_or_404(id)
    qn_id = cc.quan_nhan_id

    if cc.duong_dan_anh:
        delete_upload(cc.duong_dan_anh)

    db.session.delete(cc)
    db.session.commit()
    flash('Đã xóa chứng chỉ.', 'success')
    return redirect(url_for('admin.admin_personnel_detail', id=qn_id))


# ------------------------------------------------------------------
# Admin: Manage DanhHieu (Award Titles) - List
# ------------------------------------------------------------------
# Master list of all possible criteria fields (column names on DeXuatChiTiet)
TIEU_CHI_OPTIONS = [
    ('muc_do_hoan_thanh', 'Mức độ hoàn thành nhiệm vụ'),
    ('phieu_tin_nhiem', 'Phiếu tín nhiệm'),
    ('kiem_tra_chinh_tri', 'Kiểm tra chính trị'),
    ('diem_kiem_tra_chinh_tri', 'Điểm kiểm tra chính trị'),
    ('kiem_tra_dieu_lenh', 'Kiểm tra điều lệnh'),
    ('diem_kiem_tra_dieu_lenh', 'Điểm kiểm tra điều lệnh'),
    ('kiem_tra_tin_hoc', 'Kỹ năng số'),
    ('diem_kiem_tra_tin_hoc', 'Điểm kỹ năng số'),
    ('dia_ly_quan_su', 'Địa hình quân sự'),
    ('diem_dia_ly_quan_su', 'Điểm địa hình quân sự'),
    ('ban_sung', 'Bắn súng'),
    ('diem_ban_sung', 'Điểm bắn súng'),
    ('the_luc', 'Thể lực'),
    ('diem_the_luc', 'Điểm thể lực'),
    ('xep_loai_dang_vien', 'Xếp loại đảng viên'),
    ('ket_qua_doan_the', 'Kết quả đoàn thể'),
    ('chu_tri_don_vi_danh_hieu', 'Chủ trì đơn vị đạt danh hiệu'),
    ('danh_hieu_gv_gioi', 'Danh hiệu GV giỏi'),
    ('dinh_muc_giang_day', 'Định mức giảng dạy'),
    ('ket_qua_kiem_tra_giang', 'Kết quả kiểm tra giảng'),
    ('tien_do_pgs', 'Tiến độ PGS'),
    ('thoi_gian_lao_dong_kh', 'Thời gian lao động khoa học'),
    ('danh_hieu_hv_gioi', 'Danh hiệu HV giỏi (true/false)'),
    ('diem_tong_ket', 'Điểm tổng kết'),
    ('ket_qua_thuc_hanh', 'Kết quả thực hành'),
    ('ket_qua_ren_luyen', 'Kết quả rèn luyện'),
    ('hinh_thuc_tot_nghiep', 'Hình thức thi tốt nghiệp'),
    ('diem_tn_ctd', 'Điểm CTĐ (tốt nghiệp)'),
    ('diem_tn_ct', 'Điểm CT (tốt nghiệp)'),
    ('diem_tn_ta', 'Điểm TA (tốt nghiệp)'),
    ('diem_tn_mon4', 'Điểm môn thứ 4 (tốt nghiệp)'),
    ('diem_tn_chuyennganh', 'Điểm chuyên ngành (tốt nghiệp)'),
    ('diem_tn_baove', 'Điểm bảo vệ KL (tốt nghiệp)'),
    ('diem_nckh', 'Điểm NCKH'),
    ('nckh_noi_dung', 'Nội dung NCKH'),
    ('nckh_minh_chung', 'Minh chứng NCKH'),
    ('mo_ta_khoa_hoc', 'Mô tả thành tích khoa học'),
    ('diem_tot_nghiep', 'Điểm tốt nghiệp (TB)'),
    ('minh_chung_thanh_tich_khac', 'Minh chứng thành tích khác'),
    ('thanh_tich_ca_nhan_khac', 'Thành tích cá nhân khác'),
]


@admin_bp.route('/danh-hieu')
@login_required
@admin_required
def manage_danh_hieu():
    danh_hieus = DanhHieu.query.order_by(DanhHieu.thu_tu, DanhHieu.ten_danh_hieu).all()
    # Build tieu_chi_options from DB (fallback to hardcoded TIEU_CHI_OPTIONS for safety)
    db_tieu_chi = TieuChi.query.filter_by(is_active=True).order_by(TieuChi.thu_tu).all()
    if db_tieu_chi:
        tc_options = [(tc.ma_truong, tc.ten) for tc in db_tieu_chi]
    else:
        tc_options = TIEU_CHI_OPTIONS
    return render_template('admin/manage_danh_hieu.html',
                           danh_hieus=danh_hieus,
                           tieu_chi_options=tc_options)


# ------------------------------------------------------------------
# Admin: Create DanhHieu
# ------------------------------------------------------------------
@admin_bp.route('/danh-hieu/create', methods=['POST'])
@login_required
@admin_required
def create_danh_hieu():
    ten = request.form.get('ten_danh_hieu', '').strip()
    ma = request.form.get('ma_danh_hieu', '').strip()
    pham_vi = request.form.get('pham_vi', 'Cá nhân').strip()
    thu_tu = request.form.get('thu_tu', 0, type=int)
    tieu_chi = request.form.getlist('tieu_chi')

    if not ten or not ma:
        flash('Tên và mã danh hiệu không được để trống.', 'danger')
        return redirect(url_for('admin.manage_danh_hieu'))

    # Check uniqueness
    if DanhHieu.query.filter_by(ten_danh_hieu=ten).first():
        flash(f'Danh hiệu "{ten}" đã tồn tại.', 'danger')
        return redirect(url_for('admin.manage_danh_hieu'))
    if DanhHieu.query.filter_by(ma_danh_hieu=ma).first():
        flash(f'Mã "{ma}" đã tồn tại.', 'danger')
        return redirect(url_for('admin.manage_danh_hieu'))

    dh = DanhHieu(
        ten_danh_hieu=ten,
        ma_danh_hieu=ma,
        pham_vi=pham_vi,
        thu_tu=thu_tu,
        tieu_chi=tieu_chi,
    )
    db.session.add(dh)
    db.session.commit()
    flash(f'Đã thêm danh hiệu: {ten}', 'success')
    return redirect(url_for('admin.manage_danh_hieu'))


# ------------------------------------------------------------------
# Admin: Edit DanhHieu
# ------------------------------------------------------------------
@admin_bp.route('/danh-hieu/<int:id>/edit', methods=['GET', 'POST'])
@login_required
@admin_required
def edit_danh_hieu(id):
    dh = DanhHieu.query.get_or_404(id)

    if request.method == 'POST':
        ten = request.form.get('ten_danh_hieu', '').strip()
        ma = request.form.get('ma_danh_hieu', '').strip()
        pham_vi = request.form.get('pham_vi', 'Cá nhân').strip()
        thu_tu = request.form.get('thu_tu', 0, type=int)
        tieu_chi = request.form.getlist('tieu_chi')

        if not ten or not ma:
            flash('Tên và mã danh hiệu không được để trống.', 'danger')
            return redirect(url_for('admin.edit_danh_hieu', id=id))

        # Check uniqueness (excluding current)
        dup_ten = DanhHieu.query.filter(DanhHieu.ten_danh_hieu == ten, DanhHieu.id != id).first()
        if dup_ten:
            flash(f'Danh hiệu "{ten}" đã tồn tại.', 'danger')
            return redirect(url_for('admin.edit_danh_hieu', id=id))
        dup_ma = DanhHieu.query.filter(DanhHieu.ma_danh_hieu == ma, DanhHieu.id != id).first()
        if dup_ma:
            flash(f'Mã "{ma}" đã tồn tại.', 'danger')
            return redirect(url_for('admin.edit_danh_hieu', id=id))

        dh.ten_danh_hieu = ten
        dh.ma_danh_hieu = ma
        dh.pham_vi = pham_vi
        dh.thu_tu = thu_tu
        dh.tieu_chi = tieu_chi
        db.session.commit()
        flash(f'Đã cập nhật danh hiệu: {ten}', 'success')
        return redirect(url_for('admin.manage_danh_hieu'))

    # ── Build pre-grouped criteria dict in Python (avoids Jinja2 selectattr bugs
    #    with None / '' nhom values) ──────────────────────────────────────────
    from collections import OrderedDict as _OD
    nhom_labels = _get_nhom_tieu_chi_choices()
    all_tcs = TieuChi.query.filter_by(is_active=True).order_by(TieuChi.thu_tu, TieuChi.ten).all()

    tieu_chi_by_nhom = _OD()
    tc_ma_set = set()
    for tc in all_tcs:
        nhom = (tc.nhom or '').strip() or 'khac'
        if nhom not in tieu_chi_by_nhom:
            tieu_chi_by_nhom[nhom] = {
                'label': nhom_labels.get(nhom, nhom),
                'items': [],
            }
        tieu_chi_by_nhom[nhom]['items'].append(tc)
        tc_ma_set.add(tc.ma_truong)

    # Orphan fields: saved in dh.tieu_chi but no longer in any active TieuChi record.
    # Still show them pre-checked so the admin doesn't accidentally delete them.
    orphan_fields = [ma for ma in (dh.tieu_chi or []) if ma not in tc_ma_set]

    return render_template('admin/edit_danh_hieu.html',
                           dh=dh,
                           tieu_chi_by_nhom=tieu_chi_by_nhom,
                           orphan_fields=orphan_fields,
                           nhom_labels=nhom_labels,
                           tieu_chi_db=all_tcs,
                           tieu_chi_options=[(tc.ma_truong, tc.ten) for tc in all_tcs] or TIEU_CHI_OPTIONS)


# ------------------------------------------------------------------
# Admin: Toggle DanhHieu active/inactive
# ------------------------------------------------------------------
@admin_bp.route('/danh-hieu/<int:id>/toggle', methods=['POST'])
@login_required
@admin_required
def toggle_danh_hieu(id):
    dh = DanhHieu.query.get_or_404(id)
    dh.is_active = not dh.is_active
    db.session.commit()
    status = 'kích hoạt' if dh.is_active else 'vô hiệu hóa'
    flash(f'Đã {status} danh hiệu: {dh.ten_danh_hieu}', 'success')
    return redirect(url_for('admin.manage_danh_hieu'))


# ------------------------------------------------------------------
# Admin: Delete DanhHieu
# ------------------------------------------------------------------
@admin_bp.route('/danh-hieu/<int:id>/delete', methods=['POST'])
@login_required
@admin_required
def delete_danh_hieu(id):
    dh = DanhHieu.query.get_or_404(id)

    # Check if any DeXuatChiTiet uses this danh_hieu
    usage_count = DeXuatChiTiet.query.filter_by(loai_danh_hieu=dh.ten_danh_hieu).count()
    if usage_count > 0:
        flash(f'Không thể xóa: Danh hiệu đang được sử dụng bởi {usage_count} đề xuất.', 'danger')
        return redirect(url_for('admin.manage_danh_hieu'))

    db.session.delete(dh)
    db.session.commit()
    flash(f'Đã xóa danh hiệu: {dh.ten_danh_hieu}', 'success')
    return redirect(url_for('admin.manage_danh_hieu'))


# ------------------------------------------------------------------
# Admin: Manage TieuChi (Criteria) - List
# ------------------------------------------------------------------
PHONG_DUYET_OPTIONS = [
    ('Phòng Khoa học', 'Phòng Khoa học'),
    ('Phòng Đào tạo', 'Phòng Đào tạo'),
    ('Thủ trưởng Phòng TM-HC', 'Thủ trưởng Phòng TM-HC'),
    ('Ban Cán bộ', 'Ban Cán bộ'),
    ('Ban Tổ chức', 'Ban Tổ chức'),
    ('Ban Tuyên huấn', 'Ban Tuyên huấn'),
    ('Ban Công tác quần chúng', 'Ban Công tác quần chúng'),
    ('Ban Công nghệ thông tin', 'Ban Công nghệ thông tin'),
    ('Ban Tác huấn', 'Ban Tác huấn'),
    ('Ban Khảo thí', 'Ban Khảo thí'),
    ('Ban Bảo vệ an ninh', 'Ban Bảo vệ an ninh'),
    ('Ủy ban Kiểm tra', 'Ủy ban Kiểm tra'),
    ('Ban Quân lực', 'Ban Quân lực'),
    ('Phòng Hậu cần - Kỹ thuật', 'Phòng Hậu cần - Kỹ thuật'),
    ('Ban Sau đại học', 'Ban Sau đại học'),
]


def _get_nhom_tieu_chi_choices(include_inactive=False):
    try:
        _ensure_default_nhom_tieu_chi()
        query = NhomTieuChi.query
        if not include_inactive:
            query = query.filter_by(is_active=True)
        rows = query.order_by(NhomTieuChi.thu_tu, NhomTieuChi.ten_nhom).all()
        if rows:
            return {row.ma_nhom: row.ten_nhom for row in rows}
    except (ProgrammingError, OperationalError):
        db.session.rollback()
    return dict(TieuChi.NHOM_CHOICES)


DEFAULT_NHOM_TIEU_CHI = [
    {'ma_nhom': 'chung', 'ten_nhom': 'Tiêu chí chung', 'thu_tu': 1},
    {'ma_nhom': 'giang_vien', 'ten_nhom': 'Tiêu chí giảng viên', 'thu_tu': 2},
    {'ma_nhom': 'hoc_vien', 'ten_nhom': 'Tiêu chí học viên', 'thu_tu': 3},
    {'ma_nhom': 'nckh', 'ten_nhom': 'Tiêu chí NCKH', 'thu_tu': 4},
    {'ma_nhom': 'khac', 'ten_nhom': 'Khác', 'thu_tu': 5},
]


def _ensure_default_nhom_tieu_chi():
    existing_codes = {row.ma_nhom for row in NhomTieuChi.query.with_entities(NhomTieuChi.ma_nhom).all()}
    created = 0
    for item in DEFAULT_NHOM_TIEU_CHI:
        if item['ma_nhom'] in existing_codes:
            continue
        row = NhomTieuChi(
            ma_nhom=item['ma_nhom'],
            ten_nhom=item['ten_nhom'],
            mo_ta=None,
            thu_tu=item['thu_tu'],
            is_active=True,
        )
        row.doi_tuong_ap_dung = []
        db.session.add(row)
        created += 1
    if created:
        db.session.commit()


@admin_bp.route('/nhom-tieu-chi')
@login_required
@admin_required
def manage_nhom_tieu_chi():
    try:
        _ensure_default_nhom_tieu_chi()
        groups = NhomTieuChi.query.order_by(NhomTieuChi.thu_tu, NhomTieuChi.ten_nhom).all()
        doi_tuong_list = _get_doi_tuong_option_list()
        return render_template('admin/manage_nhom_tieu_chi.html',
                               groups=groups,
                               doi_tuong_list=doi_tuong_list,
                               migration_missing=False)
    except (ProgrammingError, OperationalError):
        db.session.rollback()
        flash('Thiếu bảng nhóm tiêu chí. Vui lòng chạy: flask db upgrade', 'danger')
        return render_template('admin/manage_nhom_tieu_chi.html',
                               groups=[],
                               doi_tuong_list=_get_doi_tuong_option_list(),
                               migration_missing=True)


@admin_bp.route('/nhom-tieu-chi/create', methods=['POST'])
@login_required
@admin_required
def create_nhom_tieu_chi():
    ma_nhom = request.form.get('ma_nhom', '').strip()
    ten_nhom = request.form.get('ten_nhom', '').strip()
    mo_ta = request.form.get('mo_ta', '').strip() or None
    thu_tu = request.form.get('thu_tu', 0, type=int)
    doi_tuong_ap_dung = request.form.getlist('doi_tuong_ap_dung')

    if not ma_nhom or not ten_nhom:
        flash('Mã nhóm và tên nhóm không được để trống.', 'danger')
        return redirect(url_for('admin.manage_nhom_tieu_chi'))

    if NhomTieuChi.query.filter_by(ma_nhom=ma_nhom).first():
        flash(f'Mã nhóm "{ma_nhom}" đã tồn tại.', 'danger')
        return redirect(url_for('admin.manage_nhom_tieu_chi'))

    row = NhomTieuChi(
        ma_nhom=ma_nhom,
        ten_nhom=ten_nhom,
        mo_ta=mo_ta,
        thu_tu=thu_tu,
        is_active=True,
    )
    row.doi_tuong_ap_dung = doi_tuong_ap_dung
    db.session.add(row)
    db.session.commit()
    flash(f'Đã thêm nhóm tiêu chí: {ten_nhom}', 'success')
    return redirect(url_for('admin.manage_nhom_tieu_chi'))


@admin_bp.route('/nhom-tieu-chi/<int:id>/edit', methods=['POST'])
@login_required
@admin_required
def edit_nhom_tieu_chi(id):
    row = NhomTieuChi.query.get_or_404(id)
    ma_nhom = request.form.get('ma_nhom', '').strip()
    ten_nhom = request.form.get('ten_nhom', '').strip()
    mo_ta = request.form.get('mo_ta', '').strip() or None
    thu_tu = request.form.get('thu_tu', 0, type=int)
    doi_tuong_ap_dung = request.form.getlist('doi_tuong_ap_dung')

    if not ma_nhom or not ten_nhom:
        flash('Mã nhóm và tên nhóm không được để trống.', 'danger')
        return redirect(url_for('admin.manage_nhom_tieu_chi'))

    dup = NhomTieuChi.query.filter(NhomTieuChi.ma_nhom == ma_nhom, NhomTieuChi.id != id).first()
    if dup:
        flash(f'Mã nhóm "{ma_nhom}" đã tồn tại.', 'danger')
        return redirect(url_for('admin.manage_nhom_tieu_chi'))

    row.ma_nhom = ma_nhom
    row.ten_nhom = ten_nhom
    row.mo_ta = mo_ta
    row.thu_tu = thu_tu
    row.doi_tuong_ap_dung = doi_tuong_ap_dung
    db.session.commit()
    flash('Đã cập nhật nhóm tiêu chí.', 'success')
    return redirect(url_for('admin.manage_nhom_tieu_chi'))


@admin_bp.route('/nhom-tieu-chi/<int:id>/toggle', methods=['POST'])
@login_required
@admin_required
def toggle_nhom_tieu_chi(id):
    row = NhomTieuChi.query.get_or_404(id)
    row.is_active = not row.is_active
    db.session.commit()
    flash('Đã cập nhật trạng thái nhóm tiêu chí.', 'success')
    return redirect(url_for('admin.manage_nhom_tieu_chi'))


@admin_bp.route('/nhom-tieu-chi/<int:id>/delete', methods=['POST'])
@login_required
@admin_required
def delete_nhom_tieu_chi(id):
    row = NhomTieuChi.query.get_or_404(id)
    used = TieuChi.query.filter_by(nhom=row.ma_nhom).count()
    if used > 0:
        flash(f'Không thể xóa: nhóm đang được dùng bởi {used} tiêu chí.', 'danger')
        return redirect(url_for('admin.manage_nhom_tieu_chi'))

    db.session.delete(row)
    db.session.commit()
    flash('Đã xóa nhóm tiêu chí.', 'success')
    return redirect(url_for('admin.manage_nhom_tieu_chi'))


@admin_bp.route('/tieu-chi')
@login_required
@admin_required
def manage_tieu_chi():
    tieu_chis = TieuChi.query.order_by(TieuChi.thu_tu, TieuChi.ten).all()

    model_fields = set(col.name for col in DeXuatChiTiet.__table__.columns)
    _SKIP = {'id', 'de_xuat_id', 'quan_nhan_id', 'loai_danh_hieu', 'doi_tuong', 'nam_hoc',
             'tap_the_data', 'ten_don_vi_de_xuat', 'admin_approved', 'ghi_chu', 'created_at', 'updated_at'}
    db_fields = set(tc.ma_truong for tc in tieu_chis)
    missing_fields = sorted([f for f in model_fields if f not in db_fields and f not in _SKIP])
    # TieuChi records whose ma_truong does NOT exist as a real model column
    invalid_ma_truong = set(tc.ma_truong for tc in tieu_chis if tc.ma_truong not in model_fields)

    return render_template('admin/manage_tieu_chi.html',
                           tieu_chis=tieu_chis,
                           nhom_choices=_get_nhom_tieu_chi_choices(include_inactive=True),
                           phong_duyet_options=PHONG_DUYET_OPTIONS,
                           missing_fields=missing_fields,
                           invalid_ma_truong=invalid_ma_truong)


# ------------------------------------------------------------------
# Admin: Create TieuChi
# ------------------------------------------------------------------
@admin_bp.route('/tieu-chi/create', methods=['POST'])
@login_required
@admin_required
def create_tieu_chi():
    ma = request.form.get('ma_truong', '').strip()
    ten = request.form.get('ten', '').strip()
    huong_dan = request.form.get('huong_dan', '').strip() or None
    nhom = request.form.get('nhom', 'chung').strip()
    co_minh_chung = request.form.get('co_minh_chung') == '1'
    loai_input = request.form.get('loai_input', 'textbox').strip()
    gia_tri_chon_raw = request.form.get('gia_tri_chon', '').strip()
    phong_duyet = request.form.getlist('phong_duyet')
    thu_tu = request.form.get('thu_tu', 0, type=int)

    if not ma or not ten:
        flash('Mã trường và tên tiêu chí không được để trống.', 'danger')
        return redirect(url_for('admin.manage_tieu_chi'))

    if TieuChi.query.filter_by(ma_truong=ma).first():
        flash(f'Mã trường "{ma}" đã tồn tại.', 'danger')
        return redirect(url_for('admin.manage_tieu_chi'))

    tc = TieuChi(
        ma_truong=ma,
        ten=ten,
        huong_dan=huong_dan,
        nhom=nhom,
        co_minh_chung=co_minh_chung,
        loai_input=loai_input,
        thu_tu=thu_tu,
    )
    tc.phong_duyet = phong_duyet
    if loai_input == 'combobox':
        tc.gia_tri_chon = gia_tri_chon_raw
    db.session.add(tc)
    db.session.commit()
    flash(f'Đã thêm tiêu chí: {ten}', 'success')
    return redirect(url_for('admin.manage_tieu_chi'))


# ------------------------------------------------------------------
# Admin: Edit TieuChi
# ------------------------------------------------------------------
@admin_bp.route('/tieu-chi/<int:id>/edit', methods=['GET', 'POST'])
@login_required
@admin_required
def edit_tieu_chi(id):
    tc = TieuChi.query.get_or_404(id)

    if request.method == 'POST':
        ma = request.form.get('ma_truong', '').strip()
        ten = request.form.get('ten', '').strip()
        huong_dan = request.form.get('huong_dan', '').strip() or None
        nhom = request.form.get('nhom', 'chung').strip()
        co_minh_chung = request.form.get('co_minh_chung') == '1'
        loai_input = request.form.get('loai_input', 'textbox').strip()
        gia_tri_chon_raw = request.form.get('gia_tri_chon', '').strip()
        phong_duyet = request.form.getlist('phong_duyet')
        thu_tu = request.form.get('thu_tu', 0, type=int)

        if not ma or not ten:
            flash('Mã trường và tên tiêu chí không được để trống.', 'danger')
            return redirect(url_for('admin.edit_tieu_chi', id=id))

        dup = TieuChi.query.filter(TieuChi.ma_truong == ma, TieuChi.id != id).first()
        if dup:
            flash(f'Mã trường "{ma}" đã tồn tại.', 'danger')
            return redirect(url_for('admin.edit_tieu_chi', id=id))

        tc.ma_truong = ma
        tc.ten = ten
        tc.huong_dan = huong_dan
        tc.nhom = nhom
        tc.co_minh_chung = co_minh_chung
        tc.loai_input = loai_input
        tc.phong_duyet = phong_duyet
        tc.thu_tu = thu_tu
        if loai_input == 'combobox':
            tc.gia_tri_chon = gia_tri_chon_raw
        else:
            tc._gia_tri_chon = None
        db.session.commit()
        flash(f'Đã cập nhật tiêu chí: {ten}', 'success')
        return redirect(url_for('admin.manage_tieu_chi'))

    return render_template('admin/edit_tieu_chi.html',
                           tc=tc,
                           nhom_choices=_get_nhom_tieu_chi_choices(include_inactive=True),
                           phong_duyet_options=PHONG_DUYET_OPTIONS)


@admin_bp.route('/evaluations')
@login_required
@admin_required
def list_annual_evaluations():
    nam_hoc = request.args.get('nam_hoc', '').strip()
    don_vi_id = request.args.get('don_vi_id', type=int)
    search = request.args.get('search', '').strip()
    page = request.args.get('page', 1, type=int)

    migration_missing = False
    evaluations = None
    nam_hoc_list = []
    try:
        query = DanhGiaHangNam.query.join(QuanNhan).join(DonVi)
        if nam_hoc:
            query = query.filter(DanhGiaHangNam.nam_hoc == nam_hoc)
        if don_vi_id:
            query = query.filter(DanhGiaHangNam.don_vi_id == don_vi_id)
        if search:
            query = query.filter(QuanNhan.ho_ten.ilike(f'%{search}%'))

        evaluations = query.order_by(
            DanhGiaHangNam.nam_hoc.desc(),
            DonVi.thu_tu,
            DonVi.ten_don_vi,
            QuanNhan.ho_ten,
        ).paginate(page=page, per_page=30, error_out=False)

        nam_hoc_list = [row[0] for row in db.session.query(DanhGiaHangNam.nam_hoc).distinct().order_by(DanhGiaHangNam.nam_hoc.desc()).all()]
    except (ProgrammingError, OperationalError):
        db.session.rollback()
        migration_missing = True
        flash('Thiếu bảng đánh giá hằng năm. Vui lòng chạy: flask db upgrade', 'danger')

    units = DonVi.query.filter_by(is_active=True).order_by(DonVi.thu_tu, DonVi.ten_don_vi).all()

    return render_template('admin/evaluation_list.html',
                           evaluations=evaluations,
                           nam_hoc=nam_hoc,
                           don_vi_id=don_vi_id,
                           search=search,
                           nam_hoc_list=nam_hoc_list,
                           units=units,
                           migration_missing=migration_missing)


# ------------------------------------------------------------------
# Admin: Toggle TieuChi active/inactive
# ------------------------------------------------------------------
@admin_bp.route('/tieu-chi/<int:id>/toggle', methods=['POST'])
@login_required
@admin_required
def toggle_tieu_chi(id):
    tc = TieuChi.query.get_or_404(id)
    tc.is_active = not tc.is_active
    db.session.commit()
    status = 'kích hoạt' if tc.is_active else 'vô hiệu hóa'
    flash(f'Đã {status} tiêu chí: {tc.ten}', 'success')
    return redirect(url_for('admin.manage_tieu_chi'))


# ------------------------------------------------------------------
# Admin: Delete TieuChi
# ------------------------------------------------------------------
@admin_bp.route('/tieu-chi/<int:id>/delete', methods=['POST'])
@login_required
@admin_required
def delete_tieu_chi(id):
    tc = TieuChi.query.get_or_404(id)
    db.session.delete(tc)
    db.session.commit()
    flash(f'Đã xóa tiêu chí: {tc.ten}', 'success')
    return redirect(url_for('admin.manage_tieu_chi'))


# ------------------------------------------------------------------
# Admin: Clear all nomination + approval data (reset for testing)
# ------------------------------------------------------------------
@admin_bp.route('/clear-data', methods=['GET', 'POST'])
@login_required
@admin_required
def clear_data():
    if request.method == 'GET':
        counts = {
            'de_xuat': DeXuat.query.count(),
            'chi_tiet': DeXuatChiTiet.query.count(),
            'phe_duyet': PheDuyet.query.count(),
            'ket_qua_ct': KetQuaDuyetChiTiet.query.count(),
            'khen_thuong': KhenThuong.query.count(),
            'hoi_dong': HoiDongBieuQuyet.query.count(),
            'thong_bao': ThongBao.query.count(),
        }
        return render_template('admin/clear_data.html', counts=counts)

    # POST: confirm=yes required
    if request.form.get('confirm') != 'XAC_NHAN_XOA':
        flash('Mã xác nhận không đúng. Thao tác bị hủy.', 'danger')
        return redirect(url_for('admin.clear_data'))

    try:
        KetQuaDuyetChiTiet.query.delete()
        KhenThuong.query.delete()
        HoiDongBieuQuyet.query.delete()
        ThongBao.query.delete()
        PheDuyet.query.delete()
        DeXuatChiTiet.query.delete()
        DeXuat.query.delete()
        db.session.commit()
        flash('Đã xóa toàn bộ dữ liệu đề xuất và phê duyệt.', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Lỗi khi xóa dữ liệu: {str(e)}', 'danger')

    return redirect(url_for('admin.approval_tracking'))


# ---------------------------------------------------------------------------
# Activity Log viewer
# ---------------------------------------------------------------------------

@admin_bp.route('/logs')
@login_required
@admin_required
def activity_log():
    from app.models.activity_log import ActivityLog, ACTION_LABELS
    from sqlalchemy.exc import ProgrammingError

    try:
        page       = request.args.get('page', 1, type=int)
        per_page   = 50
        action_f   = request.args.get('action', '').strip()
        user_f     = request.args.get('user', '').strip()
        role_f     = request.args.get('role', '').strip()
        date_from  = request.args.get('date_from', '').strip()
        date_to    = request.args.get('date_to', '').strip()

        q = ActivityLog.query

        if action_f:
            q = q.filter(ActivityLog.action == action_f)
        if user_f:
            q = q.filter(
                db.or_(
                    ActivityLog.username.ilike(f'%{user_f}%'),
                    ActivityLog.ho_ten.ilike(f'%{user_f}%'),
                )
            )
        if role_f:
            q = q.filter(ActivityLog.role == role_f)
        if date_from:
            try:
                from datetime import datetime as _dt
                q = q.filter(ActivityLog.created_at >= _dt.strptime(date_from, '%Y-%m-%d'))
            except ValueError:
                pass
        if date_to:
            try:
                from datetime import datetime as _dt
                q = q.filter(ActivityLog.created_at < _dt.strptime(date_to, '%Y-%m-%d').replace(hour=23, minute=59, second=59))
            except ValueError:
                pass

        logs = q.order_by(ActivityLog.created_at.desc()).paginate(page=page, per_page=per_page, error_out=False)

        # Distinct roles & actions for filter dropdowns
        all_roles   = [r[0] for r in db.session.query(ActivityLog.role).distinct().order_by(ActivityLog.role).all() if r[0]]
        all_actions = [a[0] for a in db.session.query(ActivityLog.action).distinct().order_by(ActivityLog.action).all() if a[0]]

        return render_template(
            'admin/activity_log.html',
            logs=logs,
            action_labels=ACTION_LABELS,
            all_roles=all_roles,
            all_actions=all_actions,
            filter_action=action_f,
            filter_user=user_f,
            filter_role=role_f,
            filter_date_from=date_from,
            filter_date_to=date_to,
        )
    except ProgrammingError:
        # Table doesn't exist yet — migration not applied
        flash('Bảng nhật ký hoạt động chưa được tạo. Vui lòng chạy: flask db upgrade', 'warning')
        return redirect(url_for('admin.approval_tracking'))
