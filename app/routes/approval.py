from flask import Blueprint, render_template, redirect, url_for, flash, request, jsonify, send_file
from flask_login import login_required, current_user
from app.extensions import db
from app.models.user import User, Role
from app.models.unit import DonVi
from app.models.nomination import DeXuat, DeXuatChiTiet, TrangThaiDeXuat, TrangThaiChiTiet, TieuChi
from app.models.approval import PheDuyet, PhongDuyet, KetQuaDuyet, KetQuaDuyetChiTiet
from app.models.notification import ThongBao
from app.models.edit_request import YeuCauChinhSua, TrangThaiYeuCauSua
from app.utils.decorators import department_required
from app.utils.activity_logger import log_action
from datetime import datetime
from io import BytesIO
from sqlalchemy.orm import joinedload, subqueryload, selectinload
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

#thaydoi import theo nhu cau
approval_bp = Blueprint('approval', __name__)

# Scope-limited depts that need auto-finalize when all their items are out-of-scope
_SCOPE_LIMITED_PHONGS = [
    PhongDuyet.BAN_QUANLUC.value,
    PhongDuyet.BAN_CANBO.value,
    PhongDuyet.PHONG_HAUCANKYTHUAT.value,
    PhongDuyet.BAN_SAUDAIHOC.value,
]

def _auto_finalize_scope_dept(de_xuat_id):
    """For BAN_QUANLUC and BAN_CANBO: if a PheDuyet has ALL chi_tiets out-of-scope,
    auto-create KetQuaDuyetChiTiet = DONG_Y for each and finalize PheDuyet.ket_qua = DONG_Y.
    Safe to call multiple times (idempotent).
    Returns list of auto-finalized phong_duyet names.
    """
    from app.models.nomination import DeXuat as _DX
    de_xuat = _DX.query.get(de_xuat_id)
    if not de_xuat:
        return []
    
    finalized = []
    for phong_val in _SCOPE_LIMITED_PHONGS:
        pd = PheDuyet.query.filter_by(
            de_xuat_id=de_xuat_id,
            phong_duyet=phong_val,
        ).first()
        
        if not pd or pd.ket_qua == KetQuaDuyet.DONG_Y.value:
            continue  # already done or not created yet
            
        # Determine scope role for this phong
        scope_role = _PHONG_TO_ROLE.get(phong_val)
        if scope_role is None:
            continue
            
        # Ensure KetQuaDuyetChiTiet records exist for all chi_tiets
        existing = {kq.chi_tiet_id for kq in pd.chi_tiet_duyet}
        
        # 1. Lấy toàn bộ ID chi tiết để truy vấn một lần duy nhất (Tránh N+1 Query)
        chi_tiet_ids = [ct.id for ct in de_xuat.chi_tiets]

        # 2. Lấy ra tất cả các kết quả duyệt hiện có của phe_duyet_id này dưới dạng Dictionary
        existing_records_list = KetQuaDuyetChiTiet.query.filter(
            KetQuaDuyetChiTiet.phe_duyet_id == pd.id,
            KetQuaDuyetChiTiet.chi_tiet_id.in_(chi_tiet_ids)
        ).all()
        existing_dict = {kq.chi_tiet_id: kq for kq in existing_records_list}

        # 3. Xử lý logic
        for ct in de_xuat.chi_tiets:
            if ct.bi_loai:
                continue
                
            in_scope = _is_in_dept_scope(scope_role, ct.doi_tuong)
            target_ket_qua = None
            skip_record = False

            # --- XÁC ĐỊNH TRẠNG THÁI (KẾT QUẢ) MONG MUỐN ---
            if ct.doi_tuong is None or ct.quan_nhan_id is None:
                if phong_val == PhongDuyet.BAN_SAUDAIHOC.value:
                    target_ket_qua = KetQuaDuyet.CHO_DUYET.value if in_scope else KetQuaDuyet.DONG_Y.value
                elif phong_val == PhongDuyet.PHONG_HAUCANKYTHUAT.value:
                    # Gán chờ duyệt và KHÔNG bỏ qua (skip_record vẫn là False)
                    target_ket_qua = KetQuaDuyet.CHO_DUYET.value
                else:
                    # Chỉ bỏ qua (skip) tập thể cho các phòng khác
                    skip_record = True
                    
            else:
                # Nếu không phải Phòng HCKT và không phải Ban SĐH
                if phong_val not in (PhongDuyet.PHONG_HAUCANKYTHUAT.value, PhongDuyet.BAN_SAUDAIHOC.value):
                    target_ket_qua = KetQuaDuyet.CHO_DUYET.value if in_scope else KetQuaDuyet.DONG_Y.value
                else:
                    # Nếu là Phòng HCKT hoặc Ban SĐH
                    if ct.doi_tuong == 'Học viên sau đại học':
                        target_ket_qua = KetQuaDuyet.CHO_DUYET.value
                    else:
                        target_ket_qua = KetQuaDuyet.DONG_Y.value

            # Nếu thuộc diện bỏ qua (tập thể của phòng khác) thì chuyển sang chi tiết tiếp theo
            if skip_record:
                continue

            # --- ÁP DỤNG THAY ĐỔI VÀO DATABASE (KHÔNG GỌI QUERY NỮA) ---
            if ct.id in existing_dict:
                # Đã tồn tại -> Cập nhật nếu giá trị khác với kết quả mong muốn
                ket_qua_hien_tai = existing_dict[ct.id]
                if ket_qua_hien_tai.ket_qua != target_ket_qua:
                    ket_qua_hien_tai.ket_qua = target_ket_qua
                    # KHÔNG commit ở đây, chờ làm xong hết mới flush
            else:
                # Chưa tồn tại -> Thêm mới
                ket_qua_moi = KetQuaDuyetChiTiet(
                    phe_duyet_id=pd.id,
                    chi_tiet_id=ct.id,
                    ket_qua=target_ket_qua
                )
                db.session.add(ket_qua_moi)

        # 4. Lưu tất cả thay đổi xuống database trong 1 lần duy nhất
        db.session.flush() 
        
        # Re-check: if no CHO_DUYET remains among ACTIVE items → auto-finalize
        pending = KetQuaDuyetChiTiet.query.filter_by(
            phe_duyet_id=pd.id,
            ket_qua=KetQuaDuyet.CHO_DUYET.value,
        ).join(DeXuatChiTiet, KetQuaDuyetChiTiet.chi_tiet_id == DeXuatChiTiet.id).filter(
            DeXuatChiTiet.bi_loai == False
        ).count()
        
        if pending == 0:
            pd.ket_qua = KetQuaDuyet.DONG_Y.value
            pd.ngay_duyet = datetime.utcnow()
            pd.ghi_chu = 'Tự động duyệt (không có đối tượng thuộc phạm vi)'
            finalized.append(phong_val)
            
    db.session.commit()
    return finalized


def _auto_finalize_scope_dept_batch(de_xuat_ids):
    """Batch version of _auto_finalize_scope_dept() cho nhiều de_xuat_id cùng lúc.

    Hàm gốc thực hiện nhiều query RIÊNG LẺ cho MỖI de_xuat_id (PheDuyet.query.filter_by
    theo từng phòng trong _SCOPE_LIMITED_PHONGS, rồi KetQuaDuyetChiTiet.query...count()).
    Khi gọi trong vòng lặp cho nhiều de_xuat_id (VD: pending_list() với vai trò
    THU_TRUONG_PHONG_TMHC — có ~70 đề xuất), điều này gây ra hàng trăm round-trip DB,
    cực kỳ chậm với kết nối DB có độ trễ cao (đã đo thực tế ~146s cho 1 trang).

    Hàm này thực hiện đúng logic tương tự nhưng chỉ dùng SỐ QUERY CỐ ĐỊNH (~3 query)
    bất kể số lượng de_xuat_id truyền vào.

    Returns {de_xuat_id: [finalized_phong_names]} — cùng dữ liệu như gọi hàm gốc nhiều lần.
    """
    from app.models.nomination import DeXuat as _DX

    de_xuat_ids = list({d for d in de_xuat_ids if d})
    result = {dxid: [] for dxid in de_xuat_ids}
    if not de_xuat_ids:
        return result

    # 1. Load tất cả DeXuat + chi_tiets liên quan (hit identity map nếu đã load trước đó)
    de_xuats = _DX.query.filter(_DX.id.in_(de_xuat_ids)).options(
        selectinload(_DX.chi_tiets)
    ).all()
    de_xuat_map = {dx.id: dx for dx in de_xuats}

    # 2. Load TẤT CẢ PheDuyet liên quan (mọi phòng scope-limited, mọi de_xuat_id) trong 1 query
    all_pds = PheDuyet.query.filter(
        PheDuyet.de_xuat_id.in_(de_xuat_ids),
        PheDuyet.phong_duyet.in_(_SCOPE_LIMITED_PHONGS),
    ).all()
    pds_to_process = [pd for pd in all_pds if pd.ket_qua != KetQuaDuyet.DONG_Y.value]
    if not pds_to_process:
        return result

    pd_ids = [pd.id for pd in pds_to_process]

    # 3. Load TẤT CẢ KetQuaDuyetChiTiet hiện có cho các PheDuyet này trong 1 query
    existing_rows = KetQuaDuyetChiTiet.query.filter(
        KetQuaDuyetChiTiet.phe_duyet_id.in_(pd_ids)
    ).all()
    existing_by_pd = {}
    for kq in existing_rows:
        existing_by_pd.setdefault(kq.phe_duyet_id, {})[kq.chi_tiet_id] = kq

    # 4. Xử lý logic — hoàn toàn trong bộ nhớ, không query thêm
    for pd in pds_to_process:
        de_xuat = de_xuat_map.get(pd.de_xuat_id)
        if not de_xuat:
            continue
        phong_val = pd.phong_duyet
        scope_role = _PHONG_TO_ROLE.get(phong_val)
        if scope_role is None:
            continue

        existing_dict = existing_by_pd.setdefault(pd.id, {})

        for ct in de_xuat.chi_tiets:
            if ct.bi_loai:
                continue

            in_scope = _is_in_dept_scope(scope_role, ct.doi_tuong)
            target_ket_qua = None
            skip_record = False

            if ct.doi_tuong is None or ct.quan_nhan_id is None:
                if phong_val == PhongDuyet.BAN_SAUDAIHOC.value:
                    target_ket_qua = KetQuaDuyet.CHO_DUYET.value if in_scope else KetQuaDuyet.DONG_Y.value
                elif phong_val == PhongDuyet.PHONG_HAUCANKYTHUAT.value:
                    target_ket_qua = KetQuaDuyet.CHO_DUYET.value
                else:
                    skip_record = True
            else:
                if phong_val not in (PhongDuyet.PHONG_HAUCANKYTHUAT.value, PhongDuyet.BAN_SAUDAIHOC.value):
                    target_ket_qua = KetQuaDuyet.CHO_DUYET.value if in_scope else KetQuaDuyet.DONG_Y.value
                else:
                    if ct.doi_tuong == 'Học viên sau đại học':
                        target_ket_qua = KetQuaDuyet.CHO_DUYET.value
                    else:
                        target_ket_qua = KetQuaDuyet.DONG_Y.value

            if skip_record:
                continue

            if ct.id in existing_dict:
                ket_qua_hien_tai = existing_dict[ct.id]
                if ket_qua_hien_tai.ket_qua != target_ket_qua:
                    ket_qua_hien_tai.ket_qua = target_ket_qua
            else:
                ket_qua_moi = KetQuaDuyetChiTiet(
                    phe_duyet_id=pd.id,
                    chi_tiet_id=ct.id,
                    ket_qua=target_ket_qua,
                )
                db.session.add(ket_qua_moi)
                existing_dict[ct.id] = ket_qua_moi

    db.session.flush()

    # 5. Re-check hoàn tất — dùng dữ liệu in-memory đã đầy đủ, KHÔNG query .count() nữa
    for pd in pds_to_process:
        de_xuat = de_xuat_map.get(pd.de_xuat_id)
        active_ct_ids = {ct.id for ct in de_xuat.chi_tiets if not ct.bi_loai} if de_xuat else set()
        existing_dict = existing_by_pd.get(pd.id, {})
        pending = sum(
            1 for ct_id, kq in existing_dict.items()
            if ct_id in active_ct_ids and kq.ket_qua == KetQuaDuyet.CHO_DUYET.value
        )
        if pending == 0:
            pd.ket_qua = KetQuaDuyet.DONG_Y.value
            pd.ngay_duyet = datetime.utcnow()
            pd.ghi_chu = 'Tự động duyệt (không có đối tượng thuộc phạm vi)'
            result.setdefault(pd.de_xuat_id, []).append(pd.phong_duyet)

    db.session.commit()
    return result

ROLE_TO_PHONG = {
    Role.PHONG_KHOAHOC: PhongDuyet.PHONG_KHOAHOC.value,
    Role.PHONG_DAOTAO: PhongDuyet.PHONG_DAOTAO.value,
    Role.THU_TRUONG_PHONG_TMHC: PhongDuyet.THU_TRUONG_PHONG_TMHC.value,
    Role.BAN_CANBO: PhongDuyet.BAN_CANBO.value,
    Role.BAN_TOCHUC: PhongDuyet.BAN_TOCHUC.value,
    Role.BAN_TUYENHUAN: PhongDuyet.BAN_TUYENHUAN.value,
    Role.BAN_CTCQ: PhongDuyet.BAN_CTCQ.value,
    Role.BAN_CNTT: PhongDuyet.BAN_CNTT.value,
    Role.BAN_TAC_HUAN: PhongDuyet.BAN_TAC_HUAN.value,
    Role.BAN_KHAOTHI: PhongDuyet.BAN_KHAOTHI.value,
    Role.BAN_BAOVE_ANNINH: PhongDuyet.BAN_BAOVE_ANNINH.value,
    Role.UY_BAN_KIEMTRA: PhongDuyet.UY_BAN_KIEMTRA.value,
    Role.BAN_QUANLUC: PhongDuyet.BAN_QUANLUC.value,
    Role.PHONG_HAUCANKYTHUAT: PhongDuyet.PHONG_HAUCANKYTHUAT.value,
    Role.BAN_SAUDAIHOC: PhongDuyet.BAN_SAUDAIHOC.value,
}


def _managed_gate_columns(role):
    if role == Role.THU_TRUONG_PHONG_TMHC:
        return [
            PhongDuyet.BAN_CNTT.value,
            PhongDuyet.BAN_TAC_HUAN.value,
            PhongDuyet.BAN_QUANLUC.value,
        ]
    return []

_GROUP_CONFIRMATION = {
    Role.THU_TRUONG_PHONG_TMHC: {
        PhongDuyet.BAN_CNTT.value,
        PhongDuyet.BAN_TAC_HUAN.value,
        PhongDuyet.BAN_QUANLUC.value,
    },
}


def _get_group_gate_for_pd(role, de_xuat_id):
    """Return gate status for group-confirmation roles.
    Output: {
      can_review: bool,
      required: [dept_names],
      approved: [dept_names],
      pending: [dept_names],
      rejected: [dept_names],
      results: {dept_name: ket_qua}
    }
    """
    if role not in _GROUP_CONFIRMATION:
        return {
            'can_review': True,
            'required': [],
            'approved': [],
            'pending': [],
            'rejected': [],
            'results': {},
        }

    required_groups = sorted(list(_GROUP_CONFIRMATION[role]))
    group_reviews = PheDuyet.query.filter(
        PheDuyet.de_xuat_id == de_xuat_id,
        PheDuyet.phong_duyet.in_(required_groups)
    ).all()
    result_map = {g.phong_duyet: g.ket_qua for g in group_reviews}

    approved = [d for d in required_groups if result_map.get(d) == KetQuaDuyet.DONG_Y.value]
    rejected = [d for d in required_groups if result_map.get(d) == KetQuaDuyet.TU_CHOI.value]
    pending = [d for d in required_groups if result_map.get(d) != KetQuaDuyet.DONG_Y.value]

    return {
        'can_review': len(pending) == 0,
        'required': required_groups,
        'approved': approved,
        'pending': pending,
        'rejected': rejected,
        'results': result_map,
    }


def _get_group_gate_for_ct(role, de_xuat_id, ct_id):
    """Return per-individual gate status for group-confirmation roles."""
    if role not in _GROUP_CONFIRMATION:
        return {
            'can_review': True,
            'required': [],
            'approved': [],
            'pending': [],
            'rejected': [],
            'results': {},
        }

    required_groups = sorted(list(_GROUP_CONFIRMATION[role]))

    # Query both the dept-level ket_qua and per-ct ket_qua in one shot
    rows = db.session.query(
        PheDuyet.phong_duyet,
        PheDuyet.ket_qua.label('pd_ket_qua'),
        KetQuaDuyetChiTiet.ket_qua.label('ct_ket_qua')
    ).outerjoin(
        KetQuaDuyetChiTiet,
        (KetQuaDuyetChiTiet.phe_duyet_id == PheDuyet.id) & (KetQuaDuyetChiTiet.chi_tiet_id == ct_id)
    ).filter(
        PheDuyet.de_xuat_id == de_xuat_id,
        PheDuyet.phong_duyet.in_(required_groups)
    ).all()

    # A dept is "approved" for this ct if:
    # 1. Per-ct KetQuaDuyetChiTiet.ket_qua == DONG_Y, OR
    # 2. PheDuyet.ket_qua == DONG_Y (dept fully finalized — all cts processed/auto-approved)
    result_map = {}
    for phong, pd_ket_qua, ct_ket_qua in rows:
        if ct_ket_qua == KetQuaDuyet.DONG_Y.value:
            result_map[phong] = KetQuaDuyet.DONG_Y.value
        elif pd_ket_qua == KetQuaDuyet.DONG_Y.value:
            # Dept fully approved → treat this ct as approved by that dept
            result_map[phong] = KetQuaDuyet.DONG_Y.value
        elif ct_ket_qua == KetQuaDuyet.TU_CHOI.value or pd_ket_qua == KetQuaDuyet.TU_CHOI.value:
            result_map[phong] = KetQuaDuyet.TU_CHOI.value
        else:
            result_map[phong] = ct_ket_qua  # None or CHO_DUYET

    approved = [d for d in required_groups if result_map.get(d) == KetQuaDuyet.DONG_Y.value]
    rejected = [d for d in required_groups if result_map.get(d) == KetQuaDuyet.TU_CHOI.value]
    pending = [d for d in required_groups if result_map.get(d) != KetQuaDuyet.DONG_Y.value]

    return {
        'can_review': len(pending) == 0,
        'required': required_groups,
        'approved': approved,
        'pending': pending,
        'rejected': rejected,
        'results': result_map,
    }


def _get_group_gate_for_pd_ct_batch(role, pd_de_xuat_ct_ids):
    """Batch version of _get_group_gate_for_pd + _get_group_gate_for_ct combined,
    to avoid N+1 queries when computing gate status for many PheDuyet/chi_tiet rows
    at once (used by pending_list()).

    pd_de_xuat_ct_ids: list of (pd_id, de_xuat_id, [ct_id, ...]) tuples.

    Returns (group_gate_by_pd, group_gate_by_ct) matching the same shapes produced by
    calling _get_group_gate_for_pd/_get_group_gate_for_ct individually.
    """
    empty_gate = {
        'can_review': True, 'required': [], 'approved': [],
        'pending': [], 'rejected': [], 'results': {},
    }

    group_gate_by_pd = {}
    group_gate_by_ct = {}

    if role not in _GROUP_CONFIRMATION:
        for pd_id, de_xuat_id, ct_ids in pd_de_xuat_ct_ids:
            group_gate_by_pd[pd_id] = dict(empty_gate)
            group_gate_by_ct[pd_id] = {ct_id: dict(empty_gate) for ct_id in ct_ids}
        return group_gate_by_pd, group_gate_by_ct

    required_groups = sorted(list(_GROUP_CONFIRMATION[role]))
    de_xuat_ids = list({dxid for _, dxid, _ in pd_de_xuat_ct_ids})
    all_ct_ids = [ct_id for _, _, ct_ids in pd_de_xuat_ct_ids for ct_id in ct_ids]

    # 1 query: dept-level ket_qua for ALL relevant de_xuat_ids at once
    pd_rows = PheDuyet.query.filter(
        PheDuyet.de_xuat_id.in_(de_xuat_ids),
        PheDuyet.phong_duyet.in_(required_groups)
    ).all() if de_xuat_ids else []
    pd_ket_qua_map = {}  # (de_xuat_id, phong_duyet) -> ket_qua
    for pd in pd_rows:
        pd_ket_qua_map[(pd.de_xuat_id, pd.phong_duyet)] = pd.ket_qua

    # 1 query: per-ct ket_qua for ALL relevant ct_ids at once
    ct_ket_qua_map = {}  # (ct_id, phong_duyet) -> ket_qua
    if all_ct_ids:
        ct_rows = db.session.query(
            KetQuaDuyetChiTiet.chi_tiet_id,
            PheDuyet.de_xuat_id,
            PheDuyet.phong_duyet,
            KetQuaDuyetChiTiet.ket_qua,
        ).join(PheDuyet, KetQuaDuyetChiTiet.phe_duyet_id == PheDuyet.id).filter(
            KetQuaDuyetChiTiet.chi_tiet_id.in_(all_ct_ids),
            PheDuyet.de_xuat_id.in_(de_xuat_ids),
            PheDuyet.phong_duyet.in_(required_groups),
        ).all()
        for chi_tiet_id, dxid, phong, ket_qua in ct_rows:
            ct_ket_qua_map[(chi_tiet_id, phong)] = ket_qua

    def _build_gate(result_map):
        approved = [d for d in required_groups if result_map.get(d) == KetQuaDuyet.DONG_Y.value]
        rejected = [d for d in required_groups if result_map.get(d) == KetQuaDuyet.TU_CHOI.value]
        pending = [d for d in required_groups if result_map.get(d) != KetQuaDuyet.DONG_Y.value]
        return {
            'can_review': len(pending) == 0,
            'required': required_groups,
            'approved': approved,
            'pending': pending,
            'rejected': rejected,
            'results': result_map,
        }

    for pd_id, de_xuat_id, ct_ids in pd_de_xuat_ct_ids:
        # group_gate_by_pd: dept-level only
        pd_result_map = {d: pd_ket_qua_map.get((de_xuat_id, d)) for d in required_groups}
        group_gate_by_pd[pd_id] = _build_gate(pd_result_map)

        # group_gate_by_ct: per-ct, falling back to dept-level if dept fully approved
        ct_gates = {}
        for ct_id in ct_ids:
            result_map = {}
            for phong in required_groups:
                ct_kq = ct_ket_qua_map.get((ct_id, phong))
                pd_kq = pd_ket_qua_map.get((de_xuat_id, phong))
                if ct_kq == KetQuaDuyet.DONG_Y.value:
                    result_map[phong] = KetQuaDuyet.DONG_Y.value
                elif pd_kq == KetQuaDuyet.DONG_Y.value:
                    result_map[phong] = KetQuaDuyet.DONG_Y.value
                elif ct_kq == KetQuaDuyet.TU_CHOI.value or pd_kq == KetQuaDuyet.TU_CHOI.value:
                    result_map[phong] = KetQuaDuyet.TU_CHOI.value
                else:
                    result_map[phong] = ct_kq  # None or CHO_DUYET
            ct_gates[ct_id] = _build_gate(result_map)
        group_gate_by_ct[pd_id] = ct_gates

    return group_gate_by_pd, group_gate_by_ct


# Reverse map: PhongDuyet display name -> Role
_PHONG_TO_ROLE = {v: k for k, v in ROLE_TO_PHONG.items()}

# --- Hardcoded fallbacks (used when DB has no tieu_chi rows) ---
_FALLBACK_PHONG_FIELDS = {
    Role.PHONG_DAOTAO: [
        'danh_hieu_gv_gioi', 'tien_do_pgs', 'dinh_muc_giang_day',
        'ket_qua_kiem_tra_giang', 'danh_hieu_hv_gioi', 'diem_tong_ket', 'ket_qua_thuc_hanh',
    ],
    Role.PHONG_KHOAHOC: ['thoi_gian_lao_dong_kh', 'diem_nckh', 'nckh_noi_dung', 'nckh_minh_chung'],
    Role.PHONG_THAMMUU: [
        'kiem_tra_tin_hoc', 'kiem_tra_dieu_lenh', 'dia_ly_quan_su', 'ban_sung', 'the_luc', 'ket_qua_doan_the',
    ],
    Role.PHONG_CHINHTRI: ['kiem_tra_chinh_tri', 'ket_qua_doan_the', 'xep_loai_dang_vien'],
    Role.BAN_CANBO: ['muc_do_hoan_thanh'],
    Role.BAN_QUANLUC: ['muc_do_hoan_thanh'],
    Role.PHONG_HAUCANKYTHUAT: ['muc_do_hoan_thanh'],
    Role.BAN_SAUDAIHOC: ['muc_do_hoan_thanh'],
}

_FALLBACK_FIELD_LABELS = {
    'muc_do_hoan_thanh': 'Hoàn thành NV', 'phieu_tin_nhiem': 'Tín nhiệm',
    'kiem_tra_dieu_lenh': 'Điều lệnh', 'ban_sung': 'Bắn súng', 'the_luc': 'Thể lực',
    'kiem_tra_chinh_tri': 'Chính trị', 'kiem_tra_tin_hoc': 'Kỹ năng số',
    'dia_ly_quan_su': 'Địa hình QS', 'danh_hieu_gv_gioi': 'GV giỏi',
    'xep_loai_dang_vien': 'Xếp loại ĐV', 'xep_loai_doan_vien': 'Xếp loại đoàn viên',
    'hinh_thuc_khen_thuong_qc': 'KT quần chúng', 'ket_qua_phu_nu': 'XL phụ nữ',
    'hinh_thuc_khen_thuong_pn': 'KT phụ nữ',
    'dinh_muc_giang_day': 'Định mức GD', 'ket_qua_kiem_tra_giang': 'KT giảng',
    'thoi_gian_lao_dong_kh': 'LĐ KH', 'tien_do_pgs': 'Tiến độ PGS',
    'danh_hieu_hv_gioi': 'HV giỏi', 'diem_tong_ket': 'Điểm TK',
    'ket_qua_thuc_hanh': 'Thực hành', 'ket_qua_ren_luyen': 'KQ rèn luyện',
    'ket_qua_doan_the': 'Đoàn thể', 'hinh_thuc_tot_nghiep': 'HT thi TN',
    'diem_tn_ctd': 'Điểm CTĐ (TN)', 'diem_tn_ct': 'Điểm CT (TN)',
    'diem_tn_ta': 'Điểm TA (TN)', 'diem_tn_mon4': 'Điểm môn 4 (TN)',
    'diem_tn_chuyennganh': 'Điểm CN (TN)', 'diem_tn_baove': 'Điểm BV KL (TN)',
    'chu_tri_don_vi_danh_hieu': 'Chủ trì ĐV', 'diem_nckh': 'Điểm KH',
    'nckh_noi_dung': 'NCKH', 'nckh_minh_chung': 'MC NCKH',
    'mo_ta_khoa_hoc': 'Mô tả TT KH',
    'diem_tot_nghiep': 'Điểm TN (TB)',
    'minh_chung_thanh_tich_khac': 'MC TT khác',
    'thanh_tich_ca_nhan_khac': 'Thành tích khác',
}

# Long text / file fields excluded from table columns
_LONG_TEXT_FIELDS = {'nckh_noi_dung', 'nckh_minh_chung', 'tien_do_pgs', 'thanh_tich_ca_nhan_khac'}

# Roles that view ALL criteria (read-only oversight), regardless of their assigned phong_duyet mapping.
# They can still only approve/reject the items themselves; this only affects which columns are shown.
_VIEW_ALL_CRITERIA_ROLES = {
    Role.BAN_CANBO,
    Role.BAN_CTCQ,
    Role.BAN_BAOVE_ANNINH,
    Role.BAN_TOCHUC,
    Role.BAN_TUYENHUAN,
}


def _all_criteria_columns():
    """Return list of all ca_nhan criteria fields — lấy động từ bảng TieuChi,
    chỉ giữ các field thực sự tồn tại trên DeXuatChiTiet.
    Bao gồm cả long-text và file fields (không loại trừ nữa)."""
    from app.models.nomination import DeXuatChiTiet as _DX, TieuChi

    # Tập hợp tất cả cột thực tế trên bảng DeXuatChiTiet
    _all_cols = {c.name for c in _DX.__table__.columns}

    # Các field hệ thống — không phải tiêu chí, luôn loại trừ
    _SYSTEM_FIELDS = {
        'id', 'de_xuat_id', 'quan_nhan_id', 'loai_danh_hieu',
        'doi_tuong', 'ten_don_vi_de_xuat', 'ghi_chu',
        'bi_loai', 'trang_thai', 'ly_do_tu_choi',
        'created_at', 'updated_at', 'tap_the_data','ly_do_loai','diem_nckh','xep_loai_tong_ket','diem_tot_nghiep','xep_loai_doan_vien','diem_the_luc','ngay_loai','admin_approved','phong_loai','diem_kiem_tra_tin_hoc','nam_hoc','ly_do_loai',
    }

   

    # Fallback nếu TieuChi chưa có dữ liệu
   

    # Chỉ loại trừ field hệ thống, giữ lại tất cả tiêu chí kể cả long-text/file
    return [f for f  in _all_cols if f not in _SYSTEM_FIELDS]







def _load_phong_fields_from_db():
    """Build PHONG_FIELDS dict from TieuChi table. Returns None if table is empty."""
    rows = TieuChi.query.filter_by(is_active=True).all()
    if not rows:
        return None
    result = {}
    for tc in rows:
        for pd_name in tc.phong_duyet:
            role = _PHONG_TO_ROLE.get(pd_name)
            if role:
                result.setdefault(role, [])
                if tc.ma_truong not in result[role]:
                    result[role].append(tc.ma_truong)
    return result


def _load_field_labels_from_db():
    """Build FIELD_LABELS dict from TieuChi table. Returns None if table is empty."""
    rows = TieuChi.query.filter_by(is_active=True).all()
    if not rows:
        return None
    return {tc.ma_truong: tc.ten for tc in rows}


def get_phong_fields():
    """Get department -> fields mapping, preferring DB over hardcoded fallback."""
    result = _load_phong_fields_from_db()
    return result if result else _FALLBACK_PHONG_FIELDS


def get_field_labels():
    """Get field -> label mapping, preferring DB over hardcoded fallback."""
    result = _load_field_labels_from_db()
    return result if result else _FALLBACK_FIELD_LABELS


def get_phong_table_columns():
    """Get department -> table column fields (excluding long text/file fields and collective-only fields)."""
    from app.models.nomination import DeXuatChiTiet as _DX
    _ca_nhan_cols = {c.key for c in _DX.__table__.columns}
    phong_fields = get_phong_fields()
    return {role: [f for f in fields if f not in _LONG_TEXT_FIELDS and f in _ca_nhan_cols]
            for role, fields in phong_fields.items()}


# Conditional field visibility by doi_tuong (remains hardcoded — specific to business logic)
PHONG_FIELD_CONDITIONS = {
    Role.BAN_CANBO: {
        'muc_do_hoan_thanh': ['Giảng viên', 'Cán bộ'],
    },
    Role.BAN_QUANLUC: {
        'muc_do_hoan_thanh': ['Công nhân viên', 'Quân nhân chuyên nghiệp','Hạ sĩ quan chiến sĩ'],
    },
}

# doi_tuong scope: which doi_tuong values each department approves
# Departments not listed approve ALL doi_tuong values
BAN_QUANLUC_DOI_TUONG = ['Công nhân viên', 'Quân nhân chuyên nghiệp','Hạ sĩ quan chiến sĩ']

DEPT_DOI_TUONG_SCOPE = {
    Role.BAN_QUANLUC: BAN_QUANLUC_DOI_TUONG,
    # BAN_CANBO approves all EXCEPT BAN_QUANLUC's scope
}


def _is_in_dept_scope(role, doi_tuong):
    """Check if an individual's doi_tuong falls within the department's approval scope."""
    # Tập thể (doi_tuong = None/'') → tất cả phòng ban đều xét duyệt
    if not doi_tuong:
        return True
    if role == Role.BAN_QUANLUC:
        return doi_tuong in BAN_QUANLUC_DOI_TUONG
    elif role == Role.BAN_CANBO:
        return doi_tuong not in BAN_QUANLUC_DOI_TUONG
    return True  # All other departments approve all doi_tuong


def _notify_rejections(phe_duyet):
    """Create ThongBao notifications for unit account when individuals are rejected."""
    de_xuat = phe_duyet.de_xuat
    # Find the unit user account for this don_vi
    unit_user = User.query.filter_by(
        don_vi_id=de_xuat.don_vi_id, role=Role.UNIT_USER
    ).first()
    if not unit_user:
        return

    phong_name = phe_duyet.phong_duyet
    for kq in phe_duyet.chi_tiet_duyet:
        if kq.ket_qua == KetQuaDuyet.TU_CHOI.value:
            ct = kq.chi_tiet
            name = ct.quan_nhan.ho_ten if ct.quan_nhan else de_xuat.don_vi.ten_don_vi
            thong_bao = ThongBao(
                user_id=unit_user.id,
                de_xuat_id=de_xuat.id,
                chi_tiet_id=ct.id,
                loai='tu_choi',
                tieu_de=f'{phong_name} không nhất trí: {name}',
                noi_dung=f'Lý do: {kq.ly_do or "Không rõ"}. Đề xuất năm học {de_xuat.nam_hoc}.',
            )
            db.session.add(thong_bao)


def _remove_chi_tiet_on_reject(phe_duyet, ct_id, ly_do=''):
    """When a department rejects a single cá nhân/tập thể, remove ONLY that item from the
    active approval process (soft-remove). The rest of the đề xuất continues unaffected.
    The unit is notified so it can see who was removed and why."""
    de_xuat = phe_duyet.de_xuat
    ct = DeXuatChiTiet.query.get(ct_id)
    if ct and not ct.bi_loai:
        ct.bi_loai = True
        ct.ly_do_loai = ly_do or 'Không đạt yêu cầu'
        ct.phong_loai = phe_duyet.phong_duyet
        ct.ngay_loai = datetime.utcnow()

    # Notify the unit account
    unit_user = User.query.filter_by(
        don_vi_id=de_xuat.don_vi_id, role=Role.UNIT_USER
    ).first()
    if unit_user and ct:
        name = (ct.quan_nhan.ho_ten if ct.quan_nhan else
                (ct.ten_don_vi_de_xuat or de_xuat.don_vi.ten_don_vi))
        thong_bao = ThongBao(
            user_id=unit_user.id,
            de_xuat_id=de_xuat.id,
            chi_tiet_id=ct_id,
            loai='tu_choi',
            tieu_de=f'{phe_duyet.phong_duyet} loại khỏi đề xuất: {name}',
            noi_dung=(f'Lý do: {ly_do or "Không đạt yêu cầu"}. '
                      f'{name} đã bị loại khỏi đề xuất năm học {de_xuat.nam_hoc} '
                      f'của {de_xuat.don_vi.ten_don_vi}. Các cá nhân/tập thể còn lại vẫn tiếp tục được xét duyệt.'),
        )
        db.session.add(thong_bao)

    # Recompute đề xuất status now that this item no longer participates
    _recompute_de_xuat_status(de_xuat)


def _recompute_de_xuat_status(de_xuat):
    """Recompute each department's finalization and the overall đề xuất status,
    counting ONLY active (non-removed) chi_tiets. Removed items never block or reject.
    A department finalizes as DONG_Y once none of its active in-scope items are still pending.
    """
    active_ct_ids = {ct.id for ct in de_xuat.chi_tiets if not ct.bi_loai}

    dept_pds = PheDuyet.query.filter_by(de_xuat_id=de_xuat.id).filter(
        PheDuyet.phong_duyet != PhongDuyet.ADMIN_TUYENHUAN.value
    ).all()

    # If there are no active items left at all, reject the whole đề xuất.
    if not active_ct_ids:
        de_xuat.trang_thai = TrangThaiDeXuat.TU_CHOI.value
        return

    for pd in dept_pds:
        if pd.ket_qua == KetQuaDuyet.DONG_Y.value:
            continue
        # KetQuaDuyetChiTiet records are created lazily (when a dept first visits review_nomination).
        # If there are no records at all, this dept has never started reviewing → do NOT auto-promote.
        all_records = [kq for kq in pd.chi_tiet_duyet if kq.chi_tiet_id in active_ct_ids]
        if not all_records:
            continue  # dept hasn't reviewed yet; skip
        # Pending = active items in this dept still waiting for a decision
        pending = [kq for kq in all_records if kq.ket_qua == KetQuaDuyet.CHO_DUYET.value]
        if not pending:
            pd.ket_qua = KetQuaDuyet.DONG_Y.value
            if pd.ngay_duyet is None:
                pd.ngay_duyet = datetime.utcnow()

    db.session.flush()

    # All departments approved → advance to Hội đồng
    all_done = all(p.ket_qua == KetQuaDuyet.DONG_Y.value for p in dept_pds)
    if all_done and dept_pds:
        if de_xuat.trang_thai not in (TrangThaiDeXuat.HOI_DONG.value,
                                      TrangThaiDeXuat.PHE_DUYET_CUOI.value):
            de_xuat.trang_thai = TrangThaiDeXuat.HOI_DONG.value
        existing_admin = PheDuyet.query.filter_by(
            de_xuat_id=de_xuat.id, phong_duyet=PhongDuyet.ADMIN_TUYENHUAN.value
        ).first()
        if not existing_admin:
            db.session.add(PheDuyet(
                de_xuat_id=de_xuat.id,
                phong_duyet=PhongDuyet.ADMIN_TUYENHUAN.value,
                ket_qua=KetQuaDuyet.CHO_DUYET.value,
            ))
    else:
        if de_xuat.trang_thai == TrangThaiDeXuat.TU_CHOI.value:
            de_xuat.trang_thai = TrangThaiDeXuat.DANG_DUYET.value

    # Update per-item trang_thai to match parent de_xuat stage
    _recompute_chi_tiet_status(de_xuat)


def _recompute_chi_tiet_status(de_xuat):
    """Sync each active DeXuatChiTiet.trang_thai with the parent đề xuất stage.

    Mapping:
      de_xuat NHAP/CHO_DUYET/DANG_DUYET → chi_tiet DANG_DUYET  (submitted, under review)
      de_xuat HOI_DONG                   → chi_tiet DA_DUYET    (all depts approved, Bảng 1)
      de_xuat PHE_DUYET_CUOI             → chi_tiet HOI_DONG    (admin_approved, Bảng 2/3)
      bi_loai = True                     → chi_tiet TU_CHOI
    Each individual item's admin_approved flag further promotes it to PHE_DUYET_CUOI.
    """
    dx_tt = de_xuat.trang_thai
    for ct in de_xuat.chi_tiets:
        if ct.bi_loai:
            ct.trang_thai = TrangThaiChiTiet.TU_CHOI.value
         
            continue
        if dx_tt in (TrangThaiDeXuat.NHAP.value, TrangThaiDeXuat.CHO_DUYET.value,
                     TrangThaiDeXuat.DANG_DUYET.value, TrangThaiDeXuat.TU_CHOI.value):
            ct.trang_thai = TrangThaiChiTiet.DANG_DUYET.value
            ct.ly_do_loai = None
            ct.phong_loai = None
            ct.ngay_loai = None
            ct.bi_loai = False
        elif dx_tt == TrangThaiDeXuat.HOI_DONG.value:
            if ct.admin_approved:
                ct.trang_thai = TrangThaiChiTiet.HOI_DONG.value
            else:
                ct.trang_thai = TrangThaiChiTiet.DA_DUYET.value
        elif dx_tt == TrangThaiDeXuat.PHE_DUYET_CUOI.value:
            ct.trang_thai = TrangThaiChiTiet.HOI_DONG.value

def _auto_prepare_pending(phong_name, nam_hoc_filter):
    """Tạo KetQuaDuyetChiTiet còn thiếu và auto-finalize nếu đủ điều kiện.
    Commit 1 lần duy nhất. Không trả về gì."""
    from app.models.nomination import DeXuat as _DeXuat

    pds = PheDuyet.query.filter_by(
        phong_duyet=phong_name,
        ket_qua=KetQuaDuyet.CHO_DUYET.value
    ).options(
        subqueryload(PheDuyet.chi_tiet_duyet),
        joinedload(PheDuyet.de_xuat).subqueryload(_DeXuat.chi_tiets),
    ).join(_DeXuat, PheDuyet.de_xuat_id == _DeXuat.id)\
     .filter(_DeXuat.nam_hoc == nam_hoc_filter).all()

    new_kq_list = []
    is_auto_dept = phong_name in (
        PhongDuyet.PHONG_HAUCANKYTHUAT.value,
        PhongDuyet.BAN_SAUDAIHOC.value,
    )

    for pd in pds:
        existing_ct_ids = {kq.chi_tiet_id for kq in pd.chi_tiet_duyet}
        for ct in pd.de_xuat.chi_tiets:
            if ct.bi_loai or ct.id in existing_ct_ids:
                continue
            in_scope = _is_in_dept_scope(
                next(r for r, p in ROLE_TO_PHONG.items() if p == phong_name),
                ct.doi_tuong
            )
            ket_qua_val = (
                KetQuaDuyet.DONG_Y.value
                if (is_auto_dept or not in_scope)
                else KetQuaDuyet.CHO_DUYET.value
            )
            new_kq_list.append(KetQuaDuyetChiTiet(
                phe_duyet_id=pd.id,
                chi_tiet_id=ct.id,
                ket_qua=ket_qua_val,
            ))

    if new_kq_list:
        db.session.add_all(new_kq_list)
        db.session.flush()

    need_commit = bool(new_kq_list)

    for pd in pds:
        if pd.ket_qua != KetQuaDuyet.CHO_DUYET.value:
            continue
        active_ct_ids = {ct.id for ct in pd.de_xuat.chi_tiets if not ct.bi_loai}
        if not active_ct_ids:
            continue
        all_kq = list(pd.chi_tiet_duyet) + [
            kq for kq in new_kq_list if kq.phe_duyet_id == pd.id
        ]
        kq_map = {kq.chi_tiet_id: kq.ket_qua for kq in all_kq}
        if any(kq_map.get(ct_id) == KetQuaDuyet.CHO_DUYET.value for ct_id in active_ct_ids):
            continue
        pd.ket_qua        = KetQuaDuyet.DONG_Y.value
        pd.nguoi_duyet_id = None
        pd.ngay_duyet     = datetime.utcnow()
        pd.ghi_chu        = 'Tự động duyệt (không có đối tượng thuộc phạm vi)'
        need_commit       = True
        de_xuat  = pd.de_xuat
        all_dept = PheDuyet.query.filter_by(de_xuat_id=de_xuat.id).filter(
            PheDuyet.phong_duyet != PhongDuyet.ADMIN_TUYENHUAN.value
        ).all()
        if all(a.ket_qua == KetQuaDuyet.DONG_Y.value for a in all_dept):
            de_xuat.trang_thai = TrangThaiDeXuat.HOI_DONG.value
            if not PheDuyet.query.filter_by(
                de_xuat_id=de_xuat.id,
                phong_duyet=PhongDuyet.ADMIN_TUYENHUAN.value
            ).first():
                db.session.add(PheDuyet(
                    de_xuat_id=de_xuat.id,
                    phong_duyet=PhongDuyet.ADMIN_TUYENHUAN.value,
                    ket_qua=KetQuaDuyet.CHO_DUYET.value,
                ))

    if need_commit:
        db.session.commit()
from sqlalchemy.orm import joinedload, selectinload, contains_eager
from sqlalchemy import exists


# ── Patch _get_group_gate_for_pd_ct_batch: chunk IN() ──────────────────────
def _get_group_gate_for_pd_ct_batch_v2(role, pd_de_xuat_ct_ids):
    """Giống bản gốc nhưng chunk IN() thành batch 200 để tránh MySQL chậm."""
    empty_gate = {
        'can_review': True, 'required': [], 'approved': [],
        'pending': [], 'rejected': [], 'results': {},
    }
    group_gate_by_pd = {}
    group_gate_by_ct = {}

    if role not in _GROUP_CONFIRMATION:
        for pd_id, de_xuat_id, ct_ids in pd_de_xuat_ct_ids:
            group_gate_by_pd[pd_id] = dict(empty_gate)
            group_gate_by_ct[pd_id] = {ct_id: dict(empty_gate) for ct_id in ct_ids}
        return group_gate_by_pd, group_gate_by_ct

    required_groups = sorted(list(_GROUP_CONFIRMATION[role]))
    de_xuat_ids     = list({dxid for _, dxid, _ in pd_de_xuat_ct_ids})
    all_ct_ids      = [ct_id for _, _, ct_ids in pd_de_xuat_ct_ids for ct_id in ct_ids]

    # ── 1 query: dept-level ket_qua ──────────────────────────────────────
    pd_rows = PheDuyet.query.filter(
        PheDuyet.de_xuat_id.in_(de_xuat_ids),
        PheDuyet.phong_duyet.in_(required_groups)
    ).all() if de_xuat_ids else []
    pd_ket_qua_map = {(pd.de_xuat_id, pd.phong_duyet): pd.ket_qua for pd in pd_rows}

    # ── FIX-D: chunk IN() thành batch 200 ────────────────────────────────
    CHUNK = 200
    ct_ket_qua_map = {}
    for i in range(0, len(all_ct_ids), CHUNK):
        chunk = all_ct_ids[i:i + CHUNK]
        rows = db.session.query(
            KetQuaDuyetChiTiet.chi_tiet_id,
            PheDuyet.de_xuat_id,
            PheDuyet.phong_duyet,
            KetQuaDuyetChiTiet.ket_qua,
        ).join(PheDuyet, KetQuaDuyetChiTiet.phe_duyet_id == PheDuyet.id).filter(
            KetQuaDuyetChiTiet.chi_tiet_id.in_(chunk),
            PheDuyet.de_xuat_id.in_(de_xuat_ids),
            PheDuyet.phong_duyet.in_(required_groups),
        ).all()
        for chi_tiet_id, dxid, phong, ket_qua in rows:
            ct_ket_qua_map[(chi_tiet_id, phong)] = ket_qua

    def _build_gate(result_map):
        approved = [d for d in required_groups if result_map.get(d) == KetQuaDuyet.DONG_Y.value]
        rejected = [d for d in required_groups if result_map.get(d) == KetQuaDuyet.TU_CHOI.value]
        pending  = [d for d in required_groups if result_map.get(d) != KetQuaDuyet.DONG_Y.value]
        return {
            'can_review': len(pending) == 0,
            'required': required_groups, 'approved': approved,
            'pending': pending, 'rejected': rejected, 'results': result_map,
        }

    for pd_id, de_xuat_id, ct_ids in pd_de_xuat_ct_ids:
        pd_result_map = {d: pd_ket_qua_map.get((de_xuat_id, d)) for d in required_groups}
        group_gate_by_pd[pd_id] = _build_gate(pd_result_map)
        ct_gates = {}
        for ct_id in ct_ids:
            result_map = {}
            for phong in required_groups:
                ct_kq = ct_ket_qua_map.get((ct_id, phong))
                pd_kq = pd_ket_qua_map.get((de_xuat_id, phong))
                if ct_kq == KetQuaDuyet.DONG_Y.value:
                    result_map[phong] = KetQuaDuyet.DONG_Y.value
                elif pd_kq == KetQuaDuyet.DONG_Y.value:
                    result_map[phong] = KetQuaDuyet.DONG_Y.value
                elif ct_kq == KetQuaDuyet.TU_CHOI.value or pd_kq == KetQuaDuyet.TU_CHOI.value:
                    result_map[phong] = KetQuaDuyet.TU_CHOI.value
                else:
                    result_map[phong] = ct_kq
            ct_gates[ct_id] = _build_gate(result_map)
        group_gate_by_ct[pd_id] = ct_gates

    return group_gate_by_pd, group_gate_by_ct


# ── FIX-C: _auto_finalize_scope_dept_batch nhận de_xuat_map từ ngoài ───────
def _auto_finalize_scope_dept_batch_v2(de_xuat_ids, de_xuat_map_hint=None):
    """Batch auto-finalize — nhận de_xuat_map_hint (đã eager-load) để tránh
    query lại DeXuat + chi_tiets. Nếu không truyền thì tự query như bản gốc.

    Returns {de_xuat_id: [finalized_phong_names]}
    """
    from app.models.nomination import DeXuat as _DX

    de_xuat_ids = list({d for d in de_xuat_ids if d})
    result = {dxid: [] for dxid in de_xuat_ids}
    if not de_xuat_ids:
        return result

    # ── FIX-C: Tái sử dụng de_xuat_map đã có, không query lại ────────────
    if de_xuat_map_hint:
        de_xuat_map = de_xuat_map_hint
    else:
        de_xuats = _DX.query.filter(_DX.id.in_(de_xuat_ids)).options(
            selectinload(_DX.chi_tiets)
        ).all()
        de_xuat_map = {dx.id: dx for dx in de_xuats}

    # Load PheDuyet scope-limited chưa hoàn tất
    all_pds = PheDuyet.query.filter(
        PheDuyet.de_xuat_id.in_(de_xuat_ids),
        PheDuyet.phong_duyet.in_(_SCOPE_LIMITED_PHONGS),
    ).all()
    pds_to_process = [pd for pd in all_pds if pd.ket_qua != KetQuaDuyet.DONG_Y.value]
    if not pds_to_process:
        return result

    pd_ids = [pd.id for pd in pds_to_process]

    # Load KetQuaDuyetChiTiet hiện có
    existing_rows = KetQuaDuyetChiTiet.query.filter(
        KetQuaDuyetChiTiet.phe_duyet_id.in_(pd_ids)
    ).all()
    existing_by_pd = {}
    for kq in existing_rows:
        existing_by_pd.setdefault(kq.phe_duyet_id, {})[kq.chi_tiet_id] = kq

    # Xử lý logic (giống bản gốc)
    for pd in pds_to_process:
        de_xuat = de_xuat_map.get(pd.de_xuat_id)
        if not de_xuat:
            continue
        phong_val  = pd.phong_duyet
        scope_role = _PHONG_TO_ROLE.get(phong_val)
        if scope_role is None:
            continue
        existing_dict = existing_by_pd.setdefault(pd.id, {})
        for ct in de_xuat.chi_tiets:
            if ct.bi_loai:
                continue
            in_scope       = _is_in_dept_scope(scope_role, ct.doi_tuong)
            target_ket_qua = None
            skip_record    = False
            if ct.doi_tuong is None or ct.quan_nhan_id is None:
                if phong_val == PhongDuyet.BAN_SAUDAIHOC.value:
                    target_ket_qua = KetQuaDuyet.CHO_DUYET.value if in_scope else KetQuaDuyet.DONG_Y.value
                elif phong_val == PhongDuyet.PHONG_HAUCANKYTHUAT.value:
                    target_ket_qua = KetQuaDuyet.CHO_DUYET.value
                else:
                    skip_record = True
            else:
                if phong_val not in (PhongDuyet.PHONG_HAUCANKYTHUAT.value, PhongDuyet.BAN_SAUDAIHOC.value):
                    target_ket_qua = KetQuaDuyet.CHO_DUYET.value if in_scope else KetQuaDuyet.DONG_Y.value
                else:
                    target_ket_qua = (KetQuaDuyet.CHO_DUYET.value
                                      if ct.doi_tuong == 'Học viên sau đại học'
                                      else KetQuaDuyet.DONG_Y.value)
            if skip_record:
                continue
            if ct.id in existing_dict:
                if existing_dict[ct.id].ket_qua != target_ket_qua:
                    existing_dict[ct.id].ket_qua = target_ket_qua
            else:
                kq_moi = KetQuaDuyetChiTiet(
                    phe_duyet_id=pd.id, chi_tiet_id=ct.id, ket_qua=target_ket_qua,
                )
                db.session.add(kq_moi)
                existing_dict[ct.id] = kq_moi

    db.session.flush()

    # Re-check hoàn tất — in-memory, không query .count()
    for pd in pds_to_process:
        de_xuat       = de_xuat_map.get(pd.de_xuat_id)
        active_ct_ids = {ct.id for ct in de_xuat.chi_tiets if not ct.bi_loai} if de_xuat else set()
        existing_dict = existing_by_pd.get(pd.id, {})
        pending = sum(
            1 for ct_id, kq in existing_dict.items()
            if ct_id in active_ct_ids and kq.ket_qua == KetQuaDuyet.CHO_DUYET.value
        )
        if pending == 0:
            pd.ket_qua   = KetQuaDuyet.DONG_Y.value
            pd.ngay_duyet = datetime.utcnow()
            pd.ghi_chu   = 'Tự động duyệt (không có đối tượng thuộc phạm vi)'
            result.setdefault(pd.de_xuat_id, []).append(pd.phong_duyet)

    db.session.commit()
    return result


# ═══════════════════════════════════════════════════════════════════════════
# pending_list() — VIẾT LẠI HOÀN TOÀN
# Từ ~930 queries → ~8 queries cố định (không phụ thuộc số đề xuất)
#
# ROOT CAUSES đã fix:
# [#1] Step 5 reload thiếu eager-load de_xuat → lazy load 70+ lần
# [#2] _auto_finalize_scope_dept_batch() query lại DeXuat đã có
# [#3] _recompute_de_xuat_status() gọi trong loop → N×2 queries
# [#4] get_field_labels()/get_phong_fields() không cache → 6 queries
# [#5] IN(700+ ids) trong group_gate batch
# [#6] quan_nhan KHÔNG được eager-load → 700 lazy loads
# ═══════════════════════════════════════════════════════════════════════════

from sqlalchemy.orm import joinedload, selectinload
from sqlalchemy import exists as sa_exists
from functools import lru_cache


# ── Module-level cache cho TieuChi (FIX-#4) ─────────────────────────────
# Dùng simple dict thay vì lru_cache để có thể invalidate khi cần
_tieu_chi_cache: dict = {}

def _get_tieu_chi_cached():
    """Load TieuChi 1 lần per process, cache vào module-level dict."""
    if _tieu_chi_cache:
        return _tieu_chi_cache
    rows = TieuChi.query.filter_by(is_active=True).all()
    _tieu_chi_cache['labels']      = {tc.ma_truong: tc.ten for tc in rows}
    _tieu_chi_cache['phong_fields'] = {}
    for tc in rows:
        for pd_name in (tc.phong_duyet or []):
            role = _PHONG_TO_ROLE.get(pd_name)
            if role:
                _tieu_chi_cache['phong_fields'].setdefault(role, [])
                if tc.ma_truong not in _tieu_chi_cache['phong_fields'][role]:
                    _tieu_chi_cache['phong_fields'][role].append(tc.ma_truong)
    return _tieu_chi_cache

def get_field_labels_cached():
    c = _get_tieu_chi_cached()
    return c.get('labels') or _FALLBACK_FIELD_LABELS

def get_phong_fields_cached():
    c = _get_tieu_chi_cached()
    return c.get('phong_fields') or _FALLBACK_PHONG_FIELDS

def invalidate_tieu_chi_cache():
    """Gọi sau khi admin thay đổi TieuChi."""
    _tieu_chi_cache.clear()


@approval_bp.route('/pending')
@login_required
@department_required
def pending_list():
    # ── Tắt expire_on_commit — an toàn trong 1 request ──────────────────
    db.session().expire_on_commit = False

    from app.models.nomination import DeXuat as _DeXuat, DanhHieu as _DanhHieu
    from app.models.unit import DonVi

    phong_name     = ROLE_TO_PHONG.get(current_user.role, '')
    nam_hoc_filter = request.args.get('nam_hoc', '')

    # ════════════════════════════════════════════════════════════════════
    # QUERY 1: Nam học dropdown
    # ════════════════════════════════════════════════════════════════════
    nam_hoc_list = [n[0] for n in db.session.query(_DeXuat.nam_hoc).join(
        PheDuyet, PheDuyet.de_xuat_id == _DeXuat.id
    ).filter(PheDuyet.phong_duyet == phong_name
    ).distinct().order_by(_DeXuat.nam_hoc.desc()).all()]

    if not nam_hoc_filter:
        # Không có filter → chỉ render dropdown, không load data
        return render_template(
            'approval/pending_list.html',
            pending_reviews=[], all_item_results={},
            phong_name=phong_name, allowed_fields=[], table_columns=[],
            field_labels=get_field_labels_cached(), field_conditions={},
            unit_names=[], out_of_scope_ct_ids=set(),
            group_gate_by_pd={}, group_gate_by_ct={},
            managed_dept_columns=[], gate_dept_fields=[],
            nam_hoc_filter='', nam_hoc_list=nam_hoc_list,
            tt_criteria_fields=[], tt_field_labels={},
            edit_requests_by_ct={}, ct_fields_map={},
            stats={'total': 0, 'approved': 0, 'rejected': 0},
        )

    # ════════════════════════════════════════════════════════════════════
    # QUERY 2+3+4: Query chính — EAGER LOAD ĐẦY ĐỦ bao gồm quan_nhan
    # FIX-#6: thêm joinedload(chi_tiets.quan_nhan) → tránh 700 lazy loads
    # FIX-#1: KHÔNG reload lại sau này → giữ nguyên object trong identity map
    # ════════════════════════════════════════════════════════════════════
    has_active_ct = (
        db.session.query(DeXuatChiTiet.id)
        .filter(
            DeXuatChiTiet.de_xuat_id == _DeXuat.id,
            DeXuatChiTiet.bi_loai == False,
        ).correlate(_DeXuat).exists()
    )

    pending_reviews = (
        db.session.query(PheDuyet)
        .join(_DeXuat, PheDuyet.de_xuat_id == _DeXuat.id)
        .join(DonVi, _DeXuat.don_vi_id == DonVi.id)
        .filter(
            PheDuyet.phong_duyet == phong_name,
            _DeXuat.nam_hoc == nam_hoc_filter,
            _DeXuat.id != None,
            has_active_ct,
        )
        .options(
            # ★ KEY: eager-load TOÀN BỘ quan hệ cần thiết trong 1 lần
            selectinload(PheDuyet.de_xuat).options(
                joinedload(_DeXuat.don_vi),
                selectinload(_DeXuat.chi_tiets).options(
                    joinedload(DeXuatChiTiet.quan_nhan),  # FIX-#6
                ),
            ),
            selectinload(PheDuyet.chi_tiet_duyet),
        )
        .order_by(DonVi.thu_tu.asc(), _DeXuat.ngay_gui.desc())
        .all()
    )

    if not pending_reviews:
        return render_template(
            'approval/pending_list.html',
            pending_reviews=[], all_item_results={},
            phong_name=phong_name, allowed_fields=[], table_columns=[],
            field_labels=get_field_labels_cached(), field_conditions={},
            unit_names=[], out_of_scope_ct_ids=set(),
            group_gate_by_pd={}, group_gate_by_ct={},
            managed_dept_columns=[], gate_dept_fields=[],
            nam_hoc_filter=nam_hoc_filter, nam_hoc_list=nam_hoc_list,
            tt_criteria_fields=[], tt_field_labels={},
            edit_requests_by_ct={}, ct_fields_map={},
            stats={'total': 0, 'approved': 0, 'rejected': 0},
        )

    # ════════════════════════════════════════════════════════════════════
    # BUILD IN-MEMORY MAPS — không query thêm
    # ════════════════════════════════════════════════════════════════════
    de_xuat_map   = {pd.de_xuat_id: pd.de_xuat for pd in pending_reviews if pd.de_xuat}
    existing_map  = {pd.id: {kq.chi_tiet_id for kq in pd.chi_tiet_duyet}
                     for pd in pending_reviews}

    # ════════════════════════════════════════════════════════════════════
    # QUERY 5 (optional): Tạo KetQuaDuyetChiTiet còn thiếu — bulk insert
    # FIX-#1: KHÔNG reload lại sau bulk_save → dùng in-memory update
    # ════════════════════════════════════════════════════════════════════
    new_records = []
    # Map pd_id → list KetQuaDuyetChiTiet mới để merge vào chi_tiet_duyet
    new_kq_by_pd: dict[int, list] = {}

    is_special = phong_name in (
        PhongDuyet.PHONG_HAUCANKYTHUAT.value,
        PhongDuyet.BAN_SAUDAIHOC.value,
    )
    for pd in pending_reviews:
        existing_ct_ids = existing_map[pd.id]
        for ct in pd.de_xuat.chi_tiets:
            if ct.phong_loai == 'Tuyên huấn' or ct.id in existing_ct_ids:
                continue
            in_scope    = _is_in_dept_scope(current_user.role, ct.doi_tuong)
            ket_qua_val = (
                KetQuaDuyet.DONG_Y.value
                if is_special or not in_scope
                else KetQuaDuyet.CHO_DUYET.value
            )
            kq_new = KetQuaDuyetChiTiet(
                phe_duyet_id=pd.id,
                chi_tiet_id=ct.id,
                ket_qua=ket_qua_val,
            )
            new_records.append(kq_new)
            new_kq_by_pd.setdefault(pd.id, []).append(kq_new)

    if new_records:
        db.session.add_all(new_records)   # add_all để có id sau flush
        db.session.flush()
        # FIX-#1: Cập nhật chi_tiet_duyet IN-MEMORY, không reload từ DB
        # SQLAlchemy identity map đã có các object mới sau add_all + flush
        for pd in pending_reviews:
            if pd.id in new_kq_by_pd:
                # Append trực tiếp vào collection đã loaded
                for kq_new in new_kq_by_pd[pd.id]:
                    if kq_new not in pd.chi_tiet_duyet:
                        pd.chi_tiet_duyet.append(kq_new)

    # ════════════════════════════════════════════════════════════════════
    # AUTO-FINALIZE — FIX-#3: inline logic, KHÔNG gọi _recompute_de_xuat_status
    # Dùng all_related_pd_map đã load 1 lần
    # ════════════════════════════════════════════════════════════════════
    de_xuat_ids_all = list(de_xuat_map.keys())
    auto_finalized_ids: set[int] = set()
    need_commit = bool(new_records)

    # QUERY 6: Load tất cả PheDuyet liên quan (1 query)
    all_related_pd_map: dict[int, list] = {}
    existing_admin_set: set[int] = set()

    if de_xuat_ids_all:
        for rpd in PheDuyet.query.filter(
            PheDuyet.de_xuat_id.in_(de_xuat_ids_all),
            PheDuyet.phong_duyet != PhongDuyet.ADMIN_TUYENHUAN.value,
        ).all():
            all_related_pd_map.setdefault(rpd.de_xuat_id, []).append(rpd)

        existing_admin_set = {
            r[0] for r in db.session.query(PheDuyet.de_xuat_id).filter(
                PheDuyet.de_xuat_id.in_(de_xuat_ids_all),
                PheDuyet.phong_duyet == PhongDuyet.ADMIN_TUYENHUAN.value,
            ).all()
        }

    for pd in pending_reviews:
        if pd.ket_qua != KetQuaDuyet.CHO_DUYET.value:
            continue
        active_ct_ids = {ct.id for ct in pd.de_xuat.chi_tiets if not ct.bi_loai}
        if not active_ct_ids:
            continue
        # Kiểm tra còn item nào CHO_DUYET không (dùng in-memory)
        still_pending = any(
            kq.chi_tiet_id in active_ct_ids and kq.ket_qua == KetQuaDuyet.CHO_DUYET.value
            for kq in pd.chi_tiet_duyet
        )
        if still_pending:
            continue

        # Tất cả active items đã có kết quả → auto-finalize
        pd.ket_qua    = KetQuaDuyet.DONG_Y.value
        pd.ngay_duyet = datetime.utcnow()
        pd.ghi_chu    = 'Tự động duyệt (không có đối tượng thuộc phạm vi)'
        auto_finalized_ids.add(pd.id)
        need_commit = True

        # FIX-#3: Kiểm tra toàn dept inline — không gọi _recompute_de_xuat_status
        de_xuat  = pd.de_xuat
        all_dept = all_related_pd_map.get(de_xuat.id, [])
        all_done = all(a.ket_qua == KetQuaDuyet.DONG_Y.value for a in all_dept)
        if all_done and all_dept:
            de_xuat.trang_thai = TrangThaiDeXuat.HOI_DONG.value
            if de_xuat.id not in existing_admin_set:
                db.session.add(PheDuyet(
                    de_xuat_id=de_xuat.id,
                    phong_duyet=PhongDuyet.ADMIN_TUYENHUAN.value,
                    ket_qua=KetQuaDuyet.CHO_DUYET.value,
                ))
                existing_admin_set.add(de_xuat.id)

    if need_commit:
        db.session.commit()

    # Loại auto-finalized khỏi danh sách hiển thị
    pending_reviews = [pd for pd in pending_reviews if pd.id not in auto_finalized_ids]

    # ════════════════════════════════════════════════════════════════════
    # BUILD LOOKUP MAPS — hoàn toàn in-memory, 0 queries
    # ════════════════════════════════════════════════════════════════════
    all_item_results = {
        pd.id: {kq.chi_tiet_id: kq for kq in pd.chi_tiet_duyet}
        for pd in pending_reviews
    }

    out_of_scope_ct_ids: set[int] = set()
    if current_user.role in (Role.BAN_QUANLUC, Role.BAN_CANBO):
        for pd in pending_reviews:
            for ct in pd.de_xuat.chi_tiets:
                if not _is_in_dept_scope(current_user.role, ct.doi_tuong):
                    out_of_scope_ct_ids.add(ct.id)

    # ════════════════════════════════════════════════════════════════════
    # FIX-#4: Dùng cached field labels/fields — 0 queries thêm
    # ════════════════════════════════════════════════════════════════════
    _field_labels    = get_field_labels_cached()
    _phong_fields    = get_phong_fields_cached()
    allowed_fields   = _phong_fields.get(current_user.role, [])
    _all_cols        = {c.name for c in DeXuatChiTiet.__table__.columns}
    _SYSTEM_FIELDS   = {
        'id','de_xuat_id','quan_nhan_id','loai_danh_hieu','doi_tuong',
        'ten_don_vi_de_xuat','ghi_chu','bi_loai','trang_thai','ly_do_tu_choi',
        'created_at','updated_at','tap_the_data','ly_do_loai','diem_nckh',
        'xep_loai_tong_ket','diem_tot_nghiep','xep_loai_doan_vien','diem_the_luc',
        'ngay_loai','admin_approved','phong_loai','diem_kiem_tra_tin_hoc',
        'nam_hoc','ly_do_loai',
    }
    table_columns = [
        f for f in (_phong_fields.get(current_user.role, []))
        if f not in _LONG_TEXT_FIELDS and f in _all_cols
    ]
    if current_user.role in _VIEW_ALL_CRITERIA_ROLES or not table_columns:
        table_columns = [f for f in _all_cols if f not in _SYSTEM_FIELDS]

    field_conditions     = PHONG_FIELD_CONDITIONS.get(current_user.role, {})
    managed_dept_columns = _managed_gate_columns(current_user.role)

    # ════════════════════════════════════════════════════════════════════
    # GATE DEPT FIELDS + AUTO-FINALIZE SCOPE (THU_TRUONG_PHONG_TMHC)
    # FIX-#2: Truyền de_xuat_map đã có → không query lại DeXuat
    # ════════════════════════════════════════════════════════════════════
    gate_dept_fields: list = []
    if current_user.role in _GROUP_CONFIRMATION:
        _auto_finalize_scope_dept_batch_v2(
            list(de_xuat_map.keys()),
            de_xuat_map_hint=de_xuat_map,   # FIX-#2
        )
        for gate_dept_name in managed_dept_columns:
            gate_role = _PHONG_TO_ROLE.get(gate_dept_name)
            if gate_role:
                fields = [f for f in _phong_fields.get(gate_role, [])
                          if f not in _LONG_TEXT_FIELDS]
                if fields:
                    gate_dept_fields.append({'dept': gate_dept_name, 'fields': fields})

    # ════════════════════════════════════════════════════════════════════
    # QUERY 7: Group gate batch — FIX-#5: chunk IN() 200
    # ════════════════════════════════════════════════════════════════════
    group_gate_by_pd: dict = {}
    group_gate_by_ct: dict = {}
    if current_user.role in _GROUP_CONFIRMATION:
        pd_de_xuat_ct_ids = [
            (pd.id, pd.de_xuat_id, [ct.id for ct in pd.de_xuat.chi_tiets])
            for pd in pending_reviews
        ]
        group_gate_by_pd, group_gate_by_ct = _get_group_gate_for_pd_ct_batch_v2(
            current_user.role, pd_de_xuat_ct_ids
        )

    # ════════════════════════════════════════════════════════════════════
    # Unit names — in-memory
    # ════════════════════════════════════════════════════════════════════
    seen: set[str] = set()
    unit_names: list[str] = []
    for pd in pending_reviews:
        name = pd.de_xuat.don_vi.ten_don_vi if pd.de_xuat and pd.de_xuat.don_vi else ''
        if name and name not in seen:
            seen.add(name)
            unit_names.append(name)

    # ════════════════════════════════════════════════════════════════════
    # QUERY 8: Tap the criteria columns
    # ════════════════════════════════════════════════════════════════════
    tt_all_keys: set[str] = set()
    for dh in _DanhHieu.query.filter_by(pham_vi='Đơn vị', is_active=True).all():
        for ma_truong in (dh.tieu_chi or []):
            tt_all_keys.add(ma_truong)
    for pd in pending_reviews:
        for ct in pd.de_xuat.chi_tiets:
            if ct.quan_nhan_id is None:
                tt_all_keys.update((ct.tap_the_dict or {}).keys())

    tt_criteria_fields: list[str] = []
    tt_field_labels_map: dict[str, str] = {}
    if tt_all_keys:
        tt_rows = TieuChi.query.filter(
            TieuChi.ma_truong.in_(list(tt_all_keys)),
            TieuChi.is_active == True,
        ).order_by(TieuChi.thu_tu, TieuChi.ten).all()
        tt_criteria_fields  = [tc.ma_truong for tc in tt_rows]
        tt_field_labels_map = {tc.ma_truong: tc.ten for tc in tt_rows}
        for k in tt_all_keys:
            if k not in tt_field_labels_map:
                tt_criteria_fields.append(k)
                tt_field_labels_map[k] = k

    # ════════════════════════════════════════════════════════════════════
    # QUERY 9: Edit requests — 1 query với IN()
    # ════════════════════════════════════════════════════════════════════
    edit_requests_by_ct: dict = {}
    all_ct_ids = [
        ct.id for pd in pending_reviews if pd.de_xuat
        for ct in pd.de_xuat.chi_tiets if not ct.bi_loai
    ]
    if all_ct_ids:
        for req in YeuCauChinhSua.query.filter(
            YeuCauChinhSua.chi_tiet_id.in_(all_ct_ids),
            YeuCauChinhSua.trang_thai == TrangThaiYeuCauSua.CHO_SUA.value,
            YeuCauChinhSua.phong_yeu_cau == phong_name,
        ).all():
            edit_requests_by_ct[req.chi_tiet_id] = req

    # ════════════════════════════════════════════════════════════════════
    # ct_fields_map — Python getattr(), 0 queries
    # ════════════════════════════════════════════════════════════════════
    all_display_fields = set(table_columns) | {
        f for gd in gate_dept_fields for f in gd['fields']
    }
    ct_fields_map: dict = {}
    if all_display_fields:
        for pd in pending_reviews:
            if not pd.de_xuat:
                continue
            for ct in pd.de_xuat.chi_tiets:
                if ct.id in ct_fields_map or ct.quan_nhan_id is None:
                    continue
                fd = {}
                for field in all_display_fields:
                    val = getattr(ct, field, None)
                    if val is not None and val != '':
                        fd[field] = val
                ct_fields_map[ct.id] = fd

    # ════════════════════════════════════════════════════════════════════
    # Stats — in-memory, 0 queries
    # ════════════════════════════════════════════════════════════════════
    ns_total = ns_approved = ns_rejected = 0
    for pd in pending_reviews:
        results_pd = all_item_results.get(pd.id, {})
        for ct in pd.de_xuat.chi_tiets:
            if ct.id in out_of_scope_ct_ids or ct.bi_loai:
                continue
            ns_total += 1
            kq = results_pd.get(ct.id)
            if kq:
                if kq.ket_qua == KetQuaDuyet.DONG_Y.value:
                    ns_approved += 1
                elif kq.ket_qua == KetQuaDuyet.TU_CHOI.value:
                    ns_rejected += 1

    return render_template(
        'approval/pending_list.html',
        pending_reviews=pending_reviews,
        all_item_results=all_item_results,
        phong_name=phong_name,
        allowed_fields=allowed_fields,
        table_columns=table_columns,
        field_labels=_field_labels,          # dùng cached
        field_conditions=field_conditions,
        unit_names=unit_names,
        out_of_scope_ct_ids=out_of_scope_ct_ids,
        group_gate_by_pd=group_gate_by_pd,
        group_gate_by_ct=group_gate_by_ct,
        managed_dept_columns=managed_dept_columns,
        gate_dept_fields=gate_dept_fields,
        nam_hoc_filter=nam_hoc_filter,
        nam_hoc_list=nam_hoc_list,
        tt_criteria_fields=tt_criteria_fields,
        tt_field_labels=tt_field_labels_map,
        edit_requests_by_ct=edit_requests_by_ct,
        ct_fields_map=ct_fields_map,
        stats={
            'total':    ns_total,
            'approved': ns_approved,
            'rejected': ns_rejected,
        },
    )

@approval_bp.route('/review/<int:id>', methods=['GET'])
@login_required
@department_required
def review_nomination(id):
    de_xuat = DeXuat.query.get_or_404(id)
    phong_name = ROLE_TO_PHONG.get(current_user.role, '')

    phe_duyet = PheDuyet.query.filter_by(
        de_xuat_id=id, phong_duyet=phong_name
    ).first_or_404()
   
    group_gate = _get_group_gate_for_pd(current_user.role, id)
    group_gate_by_ct = {}
    if current_user.role in _GROUP_CONFIRMATION:
        for ct in de_xuat.chi_tiets:
            group_gate_by_ct[ct.id] = _get_group_gate_for_ct(current_user.role, id, ct.id)

    # Ensure per-item records exist for all chi_tiets
    # For BAN_QUANLUC/BAN_CANBO: auto-approve out-of-scope items
    existing_ct_ids = {kq.chi_tiet_id for kq in phe_duyet.chi_tiet_duyet}
    for ct in de_xuat.chi_tiets:
        if ct.bi_loai:
            continue
        if ct.doi_tuong is None:
            if phong_name == PhongDuyet.PHONG_HAUCANKYTHUAT.value and phong_name == PhongDuyet.BAN_SAUDAIHOC.value:

                ket_qua_1 = KetQuaDuyetChiTiet(
                            phe_duyet_id=phe_duyet.id,
                            chi_tiet_id=ct.id,
                            ket_qua=KetQuaDuyet.DONG_Y.value,
                        )
        else:
            if ct.id not in existing_ct_ids:
                in_scope = _is_in_dept_scope(current_user.role, ct.doi_tuong)
                if phong_name != PhongDuyet.PHONG_HAUCANKYTHUAT.value and phong_name != PhongDuyet.BAN_SAUDAIHOC.value:

                    ket_qua_1 = KetQuaDuyetChiTiet(
                            phe_duyet_id=phe_duyet.id,
                            chi_tiet_id=ct.id,
                            ket_qua=KetQuaDuyet.CHO_DUYET.value if in_scope else KetQuaDuyet.DONG_Y.value,
                        )
            else:
                if ct.doi_tuong in ['Học viên sau đại học']:
                    ket_qua_1 = KetQuaDuyetChiTiet(
                        phe_duyet_id=phe_duyet.id,
                        chi_tiet_id=ct.id,
                        ket_qua=KetQuaDuyet.CHO_DUYET.value,
                    )
                else:
                    ket_qua_1 = KetQuaDuyetChiTiet(
                        phe_duyet_id=phe_duyet.id,
                        chi_tiet_id=ct.id,
                        ket_qua=KetQuaDuyet.DONG_Y.value,
                    )
        db.session.add(ket_qua_1)                                                                           
    db.session.commit()

    # Reload
    phe_duyet = PheDuyet.query.get(phe_duyet.id)

    # Build lookup: chi_tiet_id -> KetQuaDuyetChiTiet
    item_results = {kq.chi_tiet_id: kq for kq in phe_duyet.chi_tiet_duyet}

    # Build out-of-scope set for template
    out_of_scope_ct_ids = set()
    if current_user.role in (Role.BAN_QUANLUC, Role.BAN_CANBO):
        for ct in de_xuat.chi_tiets:
            if not _is_in_dept_scope(current_user.role, ct.doi_tuong):
                out_of_scope_ct_ids.add(ct.id)

    allowed_fields = get_phong_fields().get(current_user.role, [])
    table_columns = get_phong_table_columns().get(current_user.role, [])
    field_conditions = PHONG_FIELD_CONDITIONS.get(current_user.role, {})

    if current_user.role in _VIEW_ALL_CRITERIA_ROLES or not table_columns:
        table_columns = _all_criteria_columns()

    return render_template('approval/review.html',
                           de_xuat=de_xuat, phe_duyet=phe_duyet,
                           phong_name=phong_name, item_results=item_results,
                           allowed_fields=allowed_fields,
                           table_columns=table_columns,
                           field_labels=get_field_labels(),
                           field_conditions=field_conditions,
                           out_of_scope_ct_ids=out_of_scope_ct_ids,
                           group_gate=group_gate,
                           group_gate_by_ct=group_gate_by_ct)


@approval_bp.route('/review/<int:id>/item/<int:ct_id>/approve', methods=['POST'])
@login_required
@department_required
def approve_item(id, ct_id):
    phong_name = ROLE_TO_PHONG.get(current_user.role, '')
    phe_duyet = PheDuyet.query.filter_by(
        de_xuat_id=id, phong_duyet=phong_name
    ).first_or_404()

    if current_user.role in _GROUP_CONFIRMATION:
        group_gate = _get_group_gate_for_ct(current_user.role, id, ct_id)
        if not group_gate['can_review']:
            flash('Chưa đủ điều kiện phê duyệt của nhóm ban liên quan.', 'warning')
            return redirect(url_for('approval.review_nomination', id=id))

    # Block if out of scope for BAN_QUANLUC/BAN_CANBO
    ct = DeXuatChiTiet.query.get_or_404(ct_id)
    if not _is_in_dept_scope(current_user.role, ct.doi_tuong):
        flash('Cá nhân này không thuộc phạm vi duyệt của bạn.', 'warning')
        return redirect(url_for('approval.review_nomination', id=id))

    kq = KetQuaDuyetChiTiet.query.filter_by(
        phe_duyet_id=phe_duyet.id, chi_tiet_id=ct_id
    ).first_or_404()

    kq.ket_qua = KetQuaDuyet.DONG_Y.value
    kq.ly_do = None
    db.session.commit()

    name = ct.quan_nhan.ho_ten if ct.quan_nhan else 'Đơn vị'
    log_action('dept_approve_item', resource_type='chi_tiet', resource_id=ct_id,
               detail=f'{name} — {phong_name} nhất trí')
    db.session.commit()
    flash(f'Đã nhất trí: {name}', 'success')
    return redirect(url_for('approval.review_nomination', id=id))


@approval_bp.route('/review/<int:id>/item/<int:ct_id>/reject', methods=['POST'])
@login_required
@department_required
def reject_item(id, ct_id):
    phong_name = ROLE_TO_PHONG.get(current_user.role, '')
    phe_duyet = PheDuyet.query.filter_by(
        de_xuat_id=id, phong_duyet=phong_name
    ).first_or_404()

    if current_user.role in _GROUP_CONFIRMATION:
        group_gate = _get_group_gate_for_ct(current_user.role, id, ct_id)
        if not group_gate['can_review']:
            flash('Chưa đủ điều kiện phê duyệt của nhóm ban liên quan.', 'warning')
            return redirect(url_for('approval.review_nomination', id=id))

    # Block if out of scope for BAN_QUANLUC/BAN_CANBO
    ct_obj = DeXuatChiTiet.query.get_or_404(ct_id)
    if not _is_in_dept_scope(current_user.role, ct_obj.doi_tuong):
        flash('Cá nhân này không thuộc phạm vi duyệt của bạn.', 'warning')
        return redirect(url_for('approval.review_nomination', id=id))

    kq = KetQuaDuyetChiTiet.query.filter_by(
        phe_duyet_id=phe_duyet.id, chi_tiet_id=ct_id
    ).first_or_404()

    ly_do = request.form.get('ly_do', '').strip()
    if not ly_do:
        flash('Vui lòng nhập lý do không nhất trí.', 'danger')
        return redirect(url_for('approval.review_nomination', id=id))

    kq.ket_qua = KetQuaDuyet.TU_CHOI.value
    kq.ly_do = ly_do
    _remove_chi_tiet_on_reject(phe_duyet, ct_id, ly_do)
    db.session.commit()

    ct = DeXuatChiTiet.query.get(ct_id)
    name = ct.quan_nhan.ho_ten if ct and ct.quan_nhan else 'Tập thể'
    log_action('dept_reject_item', resource_type='chi_tiet', resource_id=ct_id,
               detail=f'{name} — {phong_name} không nhất trí: {ly_do}')
    db.session.commit()
    flash(f'Đã loại khỏi đề xuất: {name}. Các cá nhân/tập thể còn lại vẫn tiếp tục được xét duyệt. Đã gửi thông báo cho đơn vị.', 'warning')
    return redirect(url_for('approval.review_nomination', id=id))


@approval_bp.route('/review/<int:id>/submit', methods=['POST'])
@login_required
@department_required
def submit_review(id):
    phong_name = ROLE_TO_PHONG.get(current_user.role, '')
    phe_duyet = PheDuyet.query.filter_by(
        de_xuat_id=id, phong_duyet=phong_name
    ).first_or_404()

    if current_user.role in _GROUP_CONFIRMATION:
        blocked = []
        for kq in phe_duyet.chi_tiet_duyet:
            if kq.chi_tiet.bi_loai:
                continue
            ct_gate = _get_group_gate_for_ct(current_user.role, id, kq.chi_tiet_id)
            if not ct_gate['can_review']:
                blocked.append(kq.chi_tiet_id)
        if blocked:
            flash(f'Có {len(blocked)} cá nhân chưa đủ điều kiện theo nhóm ban liên quan.', 'warning')
            return redirect(url_for('approval.review_nomination', id=id))

    # Check all ACTIVE items have been reviewed (removed items are excluded)
    pending_items = [kq for kq in phe_duyet.chi_tiet_duyet
                     if not kq.chi_tiet.bi_loai and kq.ket_qua == KetQuaDuyet.CHO_DUYET.value]
    if pending_items:
        flash(f'Còn {len(pending_items)} cá nhân chưa được duyệt. Vui lòng duyệt tất cả trước khi hoàn tất.', 'danger')
        return redirect(url_for('approval.review_nomination', id=id))

    de_xuat = DeXuat.query.get(id)

    # Rejecting an item already removed it from the process, so finalize as DONG_Y
    # based on the remaining active items.
    phe_duyet.ket_qua = KetQuaDuyet.DONG_Y.value
    phe_duyet.nguoi_duyet_id = current_user.id
    phe_duyet.ngay_duyet = datetime.utcnow()
    phe_duyet.ghi_chu = request.form.get('ghi_chu', '').strip() or None

    _recompute_de_xuat_status(de_xuat)

    db.session.commit()
    log_action('dept_submit_review', resource_type='de_xuat', resource_id=id,
               detail=f'{phong_name} hoàn tất duyệt đề xuất #{id}')
    db.session.commit()
    flash(f'{phong_name} đã hoàn tất duyệt đề xuất.', 'success')
    return redirect(url_for('approval.pending_list'))


@approval_bp.route('/toggle/<int:pd_id>/<int:ct_id>', methods=['POST'])
@login_required
@department_required
def toggle_item(pd_id, ct_id):
    phong_name = ROLE_TO_PHONG.get(current_user.role, '')
    phe_duyet = PheDuyet.query.filter_by(
        id=pd_id, phong_duyet=phong_name
    ).first_or_404()

    if current_user.role in _GROUP_CONFIRMATION:
        group_gate = _get_group_gate_for_ct(current_user.role, phe_duyet.de_xuat_id, ct_id)
        if not group_gate['can_review']:
            return jsonify({'success': False, 'message': 'Chưa đủ điều kiện phê duyệt của nhóm ban liên quan.'}), 403

    # Block if out of scope for BAN_QUANLUC/BAN_CANBO
    ct_obj = DeXuatChiTiet.query.get_or_404(ct_id)
    if not _is_in_dept_scope(current_user.role, ct_obj.doi_tuong):
        return jsonify({'success': False, 'message': 'Cá nhân này không thuộc phạm vi duyệt của bạn.'}), 403

    kq = KetQuaDuyetChiTiet.query.filter_by(
        phe_duyet_id=phe_duyet.id, chi_tiet_id=ct_id
    ).first_or_404()

    data = request.get_json()
    approved = data.get('approved', True)
    ly_do = data.get('ly_do', '').strip()

    if approved:
        kq.ket_qua = KetQuaDuyet.DONG_Y.value
        kq.ly_do = None
        db.session.commit()
    else:
        if not ly_do:
            return jsonify({'success': False, 'message': 'Vui lòng nhập lý do'}), 400
        kq.ket_qua = KetQuaDuyet.TU_CHOI.value
        kq.ly_do = ly_do
        # Reject = remove ONLY this cá nhân/tập thể; the rest of the đề xuất continues.
        _remove_chi_tiet_on_reject(phe_duyet, ct_id, ly_do)
        db.session.commit()

    de_xuat = phe_duyet.de_xuat

    # Auto-finalize this department once none of its ACTIVE (non-removed) in-scope
    # items are still pending. Removed items never block or reject.
    active_ct_ids = {ct.id for ct in de_xuat.chi_tiets if not ct.bi_loai}
    pending_count = KetQuaDuyetChiTiet.query.filter_by(
        phe_duyet_id=phe_duyet.id,
        ket_qua=KetQuaDuyet.CHO_DUYET.value
    ).join(DeXuatChiTiet, KetQuaDuyetChiTiet.chi_tiet_id == DeXuatChiTiet.id).filter(
        DeXuatChiTiet.bi_loai == False
    ).count()

    auto_finalized = False
    if pending_count == 0 and active_ct_ids:
        if phe_duyet.ket_qua != KetQuaDuyet.DONG_Y.value:
            phe_duyet.ket_qua = KetQuaDuyet.DONG_Y.value
            phe_duyet.nguoi_duyet_id = current_user.id
            phe_duyet.ngay_duyet = datetime.utcnow()
        # Advance the whole đề xuất if every department has now approved.
        _recompute_de_xuat_status(de_xuat)
        db.session.commit()
        auto_finalized = True

    # Build stats over ACTIVE items only
    all_kq = [k for k in phe_duyet.chi_tiet_duyet if not k.chi_tiet.bi_loai]
    total = len(all_kq)
    approved_count = sum(1 for k in all_kq if k.ket_qua == KetQuaDuyet.DONG_Y.value)
    rejected_count = sum(1 for k in all_kq if k.ket_qua == KetQuaDuyet.TU_CHOI.value)

    return jsonify({
        'success': True,
        'ket_qua': kq.ket_qua,
        'auto_finalized': auto_finalized,
        'stats': {
            'total': total,
            'reviewed': approved_count + rejected_count,
            'approved': approved_count,
            'rejected': rejected_count,
        }
    })


def _reviewable_fields_for_role(role, ct):
    """Return the set of ma_truong this department may flag for editing on the
    given chi_tiet (cá nhân or tập thể)."""
    if ct.quan_nhan_id is None:
        # ★ FIX: Tập thể — trả về TẤT CẢ field tiêu chí tập thể theo config,
        # KHÔNG chỉ những field đã có giá trị trong tap_the_dict.
        # Điều này cho phép yêu cầu điền mới field đang rỗng.
        
        # ★ Tập thể: lấy từ config động + union với dict hiện có
        config_fields   = set(_all_tap_the_columns())
        existing_fields = set((ct.tap_the_dict or {}).keys())
        return config_fields | existing_fields

    # Cá nhân — giữ nguyên logic cũ
    if role in _VIEW_ALL_CRITERIA_ROLES:
        fields = set(_all_criteria_columns())
    else:
        fields = set(get_phong_table_columns().get(role, []))
        if not fields:
            fields = set(_all_criteria_columns())
    return fields
# ── Cache module-level để tránh query DB nhiều lần ──────────────────────────

# Python 3.9 trở xuống — dùng Optional từ typing
from typing import Optional

_criteria_cache: Optional[dict] = None

def _get_criteria_by_type() -> dict:
    """Query bảng TieuChi một lần, phân loại theo nhóm:
    - nhom bắt đầu bằng 'ban_' hoặc 'phong_' → tập thể
    - còn lại → cá nhân
    Trả về dict: {'ca_nhan': [...], 'tap_the': [...]}
    """
    global _criteria_cache
    if _criteria_cache is not None:
        return _criteria_cache

    from app.models.nomination import TieuChi

    all_tc = TieuChi.query.order_by(TieuChi.thu_tu.asc()).all()

    ca_nhan = []
    tap_the = []
    for tc in all_tc:
        if not tc.ma_truong:
            continue
        nhom = (tc.nhom or '').strip().lower()
        if nhom.startswith('ban_') or nhom.startswith('phong_'):
            tap_the.append(tc.ma_truong)
        else:
            ca_nhan.append(tc.ma_truong)

    _criteria_cache = {'ca_nhan': ca_nhan, 'tap_the': tap_the}
    return _criteria_cache

def _all_tap_the_columns():
    """Trả về list ma_truong tiêu chí tập thể — lấy động từ TieuChi."""
    fields = _get_criteria_by_type()['tap_the']
    return fields


@approval_bp.route('/request-edit/<int:pd_id>/<int:ct_id>', methods=['POST'])
@login_required
@department_required
def request_edit(pd_id, ct_id):
    """Approver flags one or more criteria of a single cá nhân/tập thể and asks the
    unit to fix them. Only the flagged criteria become editable by the unit; all other
    data stays locked. The flagging department's result for this item is reset to
    CHO_DUYET so it must re-review after the unit resubmits."""
    phong_name = ROLE_TO_PHONG.get(current_user.role, '')
    phe_duyet = PheDuyet.query.filter_by(
        id=pd_id, phong_duyet=phong_name
    ).first_or_404()

    ct = DeXuatChiTiet.query.get_or_404(ct_id)
    if ct.de_xuat_id != phe_duyet.de_xuat_id or ct.bi_loai:
        return jsonify({'success': False, 'message': 'Cá nhân/tập thể không hợp lệ.'}), 400

    if not _is_in_dept_scope(current_user.role, ct.doi_tuong):
        return jsonify({'success': False, 'message': 'Cá nhân này không thuộc phạm vi duyệt của bạn.'}), 403

    data = request.get_json(silent=True) or {}
    fields = data.get('fields') or []
    ly_do = (data.get('ly_do') or '').strip()
    if not isinstance(fields, list) or not fields:
        return jsonify({'success': False, 'message': 'Vui lòng chọn ít nhất một tiêu chí cần chỉnh sửa.'}), 400

    allowed = _reviewable_fields_for_role(current_user.role, ct)
    fields = [f for f in fields if f in allowed]
    if not fields and not phong_name == PhongDuyet.BAN_TUYENHUAN.value:
        return jsonify({'success': False, 'message': 'Các tiêu chí được chọn không thuộc phạm vi duyệt của bạn.'}), 400

    # Reuse an existing open request for the same item from the same department.
    yc = YeuCauChinhSua.query.filter_by(
        chi_tiet_id=ct_id,
        phong_yeu_cau=phong_name,
        trang_thai=TrangThaiYeuCauSua.CHO_SUA.value,
    ).first()
    if yc:
        merged = list(dict.fromkeys((yc.cac_truong or []) + fields))
        yc.cac_truong = merged
        yc.ly_do = ly_do or yc.ly_do
        yc.nguoi_yeu_cau_id = current_user.id
    else:
        yc = YeuCauChinhSua(
            de_xuat_id=phe_duyet.de_xuat_id,
            chi_tiet_id=ct_id,
            phong_yeu_cau=phong_name,
            nguoi_yeu_cau_id=current_user.id,
            ly_do=ly_do,
            trang_thai=TrangThaiYeuCauSua.CHO_SUA.value,
        )
        yc.cac_truong = fields
        db.session.add(yc)

    # Reset this department's result for the item so it must re-review after edit.
    kq = KetQuaDuyetChiTiet.query.filter_by(
        phe_duyet_id=phe_duyet.id, chi_tiet_id=ct_id
    ).first()
    if kq:
        kq.ket_qua = KetQuaDuyet.CHO_DUYET.value
        kq.ly_do = None
    # Keep this department open (not finalized) while the edit is pending.
    if phe_duyet.ket_qua == KetQuaDuyet.DONG_Y.value:
        phe_duyet.ket_qua = KetQuaDuyet.CHO_DUYET.value
        phe_duyet.ngay_duyet = None

    # If de_xuat had already advanced to HOI_DONG, revert it back to DANG_DUYET
    # so it clearly shows as "in departmental review" again, and the admin/hội đồng
    # PheDuyet (which was created prematurely) is removed.
    de_xuat = phe_duyet.de_xuat
    if de_xuat.trang_thai == TrangThaiDeXuat.HOI_DONG.value:
        de_xuat.trang_thai = TrangThaiDeXuat.DANG_DUYET.value
        admin_pd = PheDuyet.query.filter_by(
            de_xuat_id=de_xuat.id,
            phong_duyet=PhongDuyet.ADMIN_TUYENHUAN.value,
            ket_qua=KetQuaDuyet.CHO_DUYET.value,
        ).first()
        if admin_pd:
            db.session.delete(admin_pd)
    unit_user = User.query.filter_by(
        don_vi_id=de_xuat.don_vi_id, role=Role.UNIT_USER
    ).first()
    if unit_user:
        name = (ct.quan_nhan.ho_ten if ct.quan_nhan else
                (ct.ten_don_vi_de_xuat or de_xuat.don_vi.ten_don_vi))
        labels = get_field_labels()
        field_names = ', '.join(labels.get(f, f) for f in fields)
        db.session.add(ThongBao(
            user_id=unit_user.id,
            de_xuat_id=de_xuat.id,
            chi_tiet_id=ct_id,
            loai='yeu_cau_sua',
            tieu_de=f'{phong_name} yêu cầu chỉnh sửa: {name}',
            noi_dung=(f'Tiêu chí cần chỉnh sửa: {field_names}. '
                      f'Lý do: {ly_do or "Không rõ"}. '
                      f'Đề xuất năm học {de_xuat.nam_hoc} của {de_xuat.don_vi.ten_don_vi}.'),
        ))

    db.session.commit()
    return jsonify({'success': True, 'message': f'Đã gửi yêu cầu chỉnh sửa cho đơn vị ({len(fields)} tiêu chí).'})


@approval_bp.route('/revoke-item/<int:pd_id>/<int:ct_id>', methods=['POST'])
@login_required
@department_required
def revoke_item(pd_id, ct_id):
    """Thu hồi kết quả duyệt của 1 cá nhân/tập thể, đưa về CHO_DUYET."""
    phong_name = ROLE_TO_PHONG.get(current_user.role, '')
    phe_duyet = PheDuyet.query.filter_by(id=pd_id, phong_duyet=phong_name).first_or_404()
    de_xuat = phe_duyet.de_xuat

    if de_xuat.trang_thai in (TrangThaiDeXuat.PHE_DUYET_CUOI.value, TrangThaiDeXuat.HOI_DONG.value):
        return jsonify({'success': False, 'message': 'Không thể thu hồi - đề xuất đã qua giai đoạn duyệt của bộ phận.'}), 403

    kq = KetQuaDuyetChiTiet.query.filter_by(
        phe_duyet_id=phe_duyet.id, chi_tiet_id=ct_id
    ).first_or_404()

    if kq.ket_qua == KetQuaDuyet.CHO_DUYET.value:
        return jsonify({'success': False, 'message': 'Mục này chưa được duyệt.'}), 400

    # Reset this item
    kq.ket_qua = KetQuaDuyet.CHO_DUYET.value
    kq.ly_do = None

    # If PheDuyet was finalized, revert it back to pending
    if phe_duyet.ket_qua != KetQuaDuyet.CHO_DUYET.value:
        phe_duyet.ket_qua = KetQuaDuyet.CHO_DUYET.value
        phe_duyet.nguoi_duyet_id = None
        phe_duyet.ngay_duyet = None
        phe_duyet.ly_do = None
        # Revert DeXuat status to DANG_DUYET if it was HOI_DONG
        if de_xuat.trang_thai == TrangThaiDeXuat.HOI_DONG.value:
            de_xuat.trang_thai = TrangThaiDeXuat.DANG_DUYET.value

    db.session.commit()
    return jsonify({'success': True, 'message': 'Đã thu hồi kết quả duyệt cho cá nhân này.'})


@approval_bp.route('/batch-approve', methods=['POST'])
@login_required
@department_required
def batch_approve():
    phong_name = ROLE_TO_PHONG.get(current_user.role, '')
    data = request.get_json(silent=True)
    if not data:
        return jsonify({'success': False, 'message': 'Không parse được JSON'}), 400

    pd_id  = data.get('pd_id')
    ct_ids = data.get('ct_ids', [])

    if not pd_id or not ct_ids:
        return jsonify({'success': False, 'message': f'Thiếu dữ liệu — pd_id={pd_id!r}, ct_ids={ct_ids!r}'}), 400

    phe_duyet = PheDuyet.query.filter_by(
        id=pd_id, phong_duyet=phong_name
    ).first_or_404()

    # ── FIX 1: Cho phép duyệt tiếp dù đã auto-finalize trước đó ──────────
    # Bỏ check cứng, chỉ skip nếu đã TỪ CHỐI hoàn toàn
    if phe_duyet.ket_qua == KetQuaDuyet.TU_CHOI.value:
        return jsonify({'success': False, 'message': 'Đề xuất đã bị từ chối, không thể duyệt.'}), 400

    if current_user.role in _GROUP_CONFIRMATION:
        blocked_ids = []
        for ct_id in ct_ids:
            ct_gate = _get_group_gate_for_ct(current_user.role, phe_duyet.de_xuat_id, ct_id)
            if not ct_gate['can_review']:
                blocked_ids.append(ct_id)
        if blocked_ids:
            return jsonify({'success': False, 'message': f'Có {len(blocked_ids)} cá nhân chưa đủ điều kiện phê duyệt.'}), 403

    # ── FIX 2: Batch load ct và kq thay vì query từng cái trong loop ──────
    ct_map = {
        ct.id: ct for ct in
        DeXuatChiTiet.query.filter(DeXuatChiTiet.id.in_(ct_ids)).all()
    }
    kq_map = {
        kq.chi_tiet_id: kq for kq in
        KetQuaDuyetChiTiet.query.filter_by(phe_duyet_id=phe_duyet.id)
        .filter(KetQuaDuyetChiTiet.chi_tiet_id.in_(ct_ids)).all()
    }

    approved_names = []
    for ct_id in ct_ids:
        chi_tiet = ct_map.get(ct_id)
        if not chi_tiet:
            continue
        if not _is_in_dept_scope(current_user.role, chi_tiet.doi_tuong):
            continue
        kq = kq_map.get(ct_id)
        if kq and kq.ket_qua == KetQuaDuyet.CHO_DUYET.value:
            kq.ket_qua = KetQuaDuyet.DONG_Y.value
            kq.ly_do   = None
            name = chi_tiet.quan_nhan.ho_ten if chi_tiet.quan_nhan else 'Đơn vị'
            approved_names.append(name)

    db.session.commit()

    de_xuat = phe_duyet.de_xuat

    # ── FIX 3: Đổi tên biến tránh shadow 'ct' ────────────────────────────
    active_ct_ids = {item.id for item in de_xuat.chi_tiets if not item.bi_loai}

    pending_count = (
        KetQuaDuyetChiTiet.query
        .filter_by(phe_duyet_id=phe_duyet.id, ket_qua=KetQuaDuyet.CHO_DUYET.value)
        .join(DeXuatChiTiet, KetQuaDuyetChiTiet.chi_tiet_id == DeXuatChiTiet.id)
        .filter(DeXuatChiTiet.bi_loai == False)
        .count()
    )

    auto_finalized = False
    if pending_count == 0 and active_ct_ids:
        phe_duyet.ket_qua        = KetQuaDuyet.DONG_Y.value
        phe_duyet.nguoi_duyet_id = current_user.id
        phe_duyet.ngay_duyet     = datetime.utcnow()
        _recompute_de_xuat_status(de_xuat)
        db.session.commit()
        auto_finalized = True

    # Build stats
    all_kq         = [k for k in phe_duyet.chi_tiet_duyet if not k.chi_tiet.bi_loai]
    total          = len(all_kq)
    approved_count = sum(1 for k in all_kq if k.ket_qua == KetQuaDuyet.DONG_Y.value)
    rejected_count = sum(1 for k in all_kq if k.ket_qua == KetQuaDuyet.TU_CHOI.value)

    return jsonify({
        'success':        True,
        'approved_count': len(approved_names),
        'auto_finalized': auto_finalized,
        'stats': {
            'total':    total,
            'reviewed': approved_count + rejected_count,
            'approved': approved_count,
            'rejected': rejected_count,
        }
    })



@approval_bp.route('/history')
@login_required
@department_required
def history():
    phong_name = ROLE_TO_PHONG.get(current_user.role, '')

    page = request.args.get('page', 1, type=int)
    unit_filter = request.args.get('unit', '')
    danh_hieu_filter = request.args.get('danh_hieu', '')
    ket_qua_filter = request.args.get('ket_qua', '')

    # Query per-individual results (KetQuaDuyetChiTiet) for this department
    query = db.session.query(KetQuaDuyetChiTiet).join(
        PheDuyet, KetQuaDuyetChiTiet.phe_duyet_id == PheDuyet.id
    ).filter(
        PheDuyet.phong_duyet == phong_name,
        KetQuaDuyetChiTiet.ket_qua != KetQuaDuyet.CHO_DUYET.value,
    )

    # Apply filters
    if ket_qua_filter:
        query = query.filter(KetQuaDuyetChiTiet.ket_qua == ket_qua_filter)

    if unit_filter:
        query = query.join(DeXuat, PheDuyet.de_xuat_id == DeXuat.id).join(
            DonVi, DeXuat.don_vi_id == DonVi.id
        ).filter(DonVi.ten_don_vi == unit_filter)
    else:
        query = query.join(DeXuat, PheDuyet.de_xuat_id == DeXuat.id)

    if danh_hieu_filter:
        query = query.join(
            DeXuatChiTiet, KetQuaDuyetChiTiet.chi_tiet_id == DeXuatChiTiet.id
        ).filter(DeXuatChiTiet.loai_danh_hieu == danh_hieu_filter)

    individual_results = query.order_by(
        PheDuyet.ngay_duyet.desc(), KetQuaDuyetChiTiet.id.desc()
    ).paginate(page=page, per_page=20, error_out=False)

    # Get filter options
    unit_names_q = db.session.query(DonVi.ten_don_vi).join(
        DeXuat, DeXuat.don_vi_id == DonVi.id
    ).join(PheDuyet, PheDuyet.de_xuat_id == DeXuat.id).join(
        KetQuaDuyetChiTiet, KetQuaDuyetChiTiet.phe_duyet_id == PheDuyet.id
    ).filter(
        PheDuyet.phong_duyet == phong_name,
        KetQuaDuyetChiTiet.ket_qua != KetQuaDuyet.CHO_DUYET.value,
    ).distinct().order_by(DonVi.ten_don_vi).all()
    unit_names = [u[0] for u in unit_names_q]

    # Summary stats (unfiltered)
    base_q = db.session.query(KetQuaDuyetChiTiet).join(
        PheDuyet, KetQuaDuyetChiTiet.phe_duyet_id == PheDuyet.id
    ).filter(
        PheDuyet.phong_duyet == phong_name,
        KetQuaDuyetChiTiet.ket_qua != KetQuaDuyet.CHO_DUYET.value,
    )
    stats = {
        'total': base_q.count(),
        'approved': base_q.filter(KetQuaDuyetChiTiet.ket_qua == KetQuaDuyet.DONG_Y.value).count(),
        'rejected': base_q.filter(KetQuaDuyetChiTiet.ket_qua == KetQuaDuyet.TU_CHOI.value).count(),
    }

    allowed_fields = get_phong_fields().get(current_user.role, [])
    table_columns = get_phong_table_columns().get(current_user.role, [])

    if current_user.role in _VIEW_ALL_CRITERIA_ROLES or not table_columns:
        table_columns = _all_criteria_columns()

    # Build tt_criteria_fields from ALL active collective DanhHieu definitions,
    # not from the current page items (which is paginated and would miss criteria
    # from items on other pages or in items not yet loaded).
    from app.models.nomination import DanhHieu
    tt_all_keys = set()
    collective_danh_hieus = DanhHieu.query.filter_by(pham_vi='Đơn vị', is_active=True).all()
    for dh in collective_danh_hieus:
        for ma_truong in (dh.tieu_chi or []):
            tt_all_keys.add(ma_truong)

    # Also collect any keys actually present in current-page items that may not
    # be in DanhHieu definitions (legacy data).
    for kq_item in individual_results.items:
        ct = kq_item.chi_tiet
        if ct and ct.quan_nhan_id is None:
            td = ct.tap_the_dict or {}
            tt_all_keys.update(td.keys())

    if tt_all_keys:
        tt_tieu_chi_rows = TieuChi.query.filter(
            TieuChi.ma_truong.in_(list(tt_all_keys)), TieuChi.is_active == True
        ).order_by(TieuChi.thu_tu, TieuChi.ten).all()
        tt_history_fields = [tc.ma_truong for tc in tt_tieu_chi_rows]
        tt_field_labels_h = {tc.ma_truong: tc.ten for tc in tt_tieu_chi_rows}
        for k in tt_all_keys:
            if k not in tt_field_labels_h:
                tt_history_fields.append(k)
                tt_field_labels_h[k] = k
    else:
        tt_history_fields = []
        tt_field_labels_h = {}

    return render_template('approval/history.html',
                           individual_results=individual_results,
                           phong_name=phong_name,
                           unit_filter=unit_filter,
                           danh_hieu_filter=danh_hieu_filter,
                           ket_qua_filter=ket_qua_filter,
                           unit_names=unit_names,
                           stats=stats,
                           allowed_fields=allowed_fields,
                           table_columns=table_columns,
                           field_labels=get_field_labels(),
                           tt_history_fields=tt_history_fields,
                           tt_field_labels=tt_field_labels_h)


@approval_bp.route('/history/chi-tiet/<int:ct_id>')
@login_required
@department_required
def history_detail(ct_id):
    """View detailed info for one individual in history, scoped to current department."""
    phong_name = ROLE_TO_PHONG.get(current_user.role, '')
    ct = DeXuatChiTiet.query.get_or_404(ct_id)
    de_xuat = ct.de_xuat

    # Get this department's PheDuyet and individual result
    phe_duyet = PheDuyet.query.filter_by(
        de_xuat_id=de_xuat.id, phong_duyet=phong_name
    ).first_or_404()

    kq = KetQuaDuyetChiTiet.query.filter_by(
        phe_duyet_id=phe_duyet.id, chi_tiet_id=ct.id
    ).first()

    allowed_fields = get_phong_fields().get(current_user.role, [])

    return render_template('approval/history_detail.html',
                           ct=ct,
                           de_xuat=de_xuat,
                           phe_duyet=phe_duyet,
                           kq=kq,
                           phong_name=phong_name,
                           allowed_fields=allowed_fields,
                           field_labels=get_field_labels())


@approval_bp.route('/revoke/<int:pd_id>', methods=['POST'])
@login_required
@department_required
def revoke_review(pd_id):
    """Revoke (thu hồi) a completed department review, resetting it to pending."""
    phong_name = ROLE_TO_PHONG.get(current_user.role, '')

    phe_duyet = PheDuyet.query.filter_by(
        id=pd_id, phong_duyet=phong_name
    ).first_or_404()

    de_xuat = phe_duyet.de_xuat

    # Cannot revoke after admin final approval
    if de_xuat.trang_thai == TrangThaiDeXuat.PHE_DUYET_CUOI.value:
        flash('Không thể thu hồi - đề xuất đã được phê duyệt cuối cùng.', 'danger')
        return redirect(url_for('approval.history'))

    # Can only revoke a completed review (not one that's still pending)
    if phe_duyet.ket_qua == KetQuaDuyet.CHO_DUYET.value:
        flash('Kết quả duyệt này vẫn đang chờ, không cần thu hồi.', 'warning')
        return redirect(url_for('approval.history'))

    # For paired scope logic, disallow revoke if this department is auto-approved-by-scope
    if current_user.role in (Role.BAN_CANBO, Role.BAN_QUANLUC):
        in_scope_items = [
            kq for kq in phe_duyet.chi_tiet_duyet
            if _is_in_dept_scope(current_user.role, kq.chi_tiet.doi_tuong)
        ]
        if not in_scope_items:
            flash('Kết quả tự động theo phạm vi, không thể thu hồi.', 'warning')
            return redirect(url_for('approval.history'))

    # 1. Reset the PheDuyet record
    phe_duyet.ket_qua = KetQuaDuyet.CHO_DUYET.value
    phe_duyet.nguoi_duyet_id = None
    phe_duyet.ngay_duyet = None
    phe_duyet.ly_do = None
    phe_duyet.ghi_chu = None

    # 2. Reset KetQuaDuyetChiTiet records back to pending
    #    For BAN_QUANLUC/BAN_CANBO: only reset in-scope items, keep out-of-scope as DONG_Y
    for kq in phe_duyet.chi_tiet_duyet:
        ct = kq.chi_tiet
        if ct and not _is_in_dept_scope(current_user.role, ct.doi_tuong):
            # Out-of-scope: keep auto-approved
            continue
        kq.ket_qua = KetQuaDuyet.CHO_DUYET.value
        kq.ly_do = None

    # 3. Handle DeXuat status changes
    old_status = de_xuat.trang_thai

    # If admin Tuyên huấn PheDuyet was already created (status was 'Đã duyệt'),
    # delete it since not all 6 depts approve anymore
    if old_status == TrangThaiDeXuat.HOI_DONG.value:
        admin_pd = PheDuyet.query.filter_by(
            de_xuat_id=de_xuat.id,
            phong_duyet=PhongDuyet.ADMIN_TUYENHUAN.value
        ).first()
        if admin_pd:
            # Delete admin's chi_tiet_duyet records if any
            KetQuaDuyetChiTiet.query.filter_by(phe_duyet_id=admin_pd.id).delete()
            db.session.delete(admin_pd)

    # Determine new status: check other departments
    other_depts = PheDuyet.query.filter_by(de_xuat_id=de_xuat.id).filter(
        PheDuyet.phong_duyet != PhongDuyet.ADMIN_TUYENHUAN.value,
        PheDuyet.id != phe_duyet.id  # exclude the one we just revoked
    ).all()

    has_other_completed = any(
        pd.ket_qua != KetQuaDuyet.CHO_DUYET.value for pd in other_depts
    )

    if has_other_completed:
        # At least one other dept has completed their review
        de_xuat.trang_thai = TrangThaiDeXuat.DANG_DUYET.value
    else:
        # No dept has completed review - back to waiting
        de_xuat.trang_thai = TrangThaiDeXuat.CHO_DUYET.value

    db.session.commit()
    flash(f'{phong_name} đã thu hồi kết quả duyệt cho đề xuất của {de_xuat.don_vi.ten_don_vi}.', 'success')
    return redirect(url_for('approval.history'))


@approval_bp.route('/export-excel')
@login_required
@department_required
def export_excel():
    """Export pending review list to Excel with timestamp."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    phong_name = ROLE_TO_PHONG.get(current_user.role, '')
    nam_hoc_filter = request.args.get('nam_hoc', '')

    q = PheDuyet.query.filter_by(phong_duyet=phong_name, ket_qua=KetQuaDuyet.CHO_DUYET.value)
    if nam_hoc_filter:
        from app.models.nomination import DeXuat as _DeXuat
        q = q.join(_DeXuat, PheDuyet.de_xuat_id == _DeXuat.id).filter(_DeXuat.nam_hoc == nam_hoc_filter)
    pending_reviews = q.order_by(PheDuyet.created_at.desc()).all()

    # Build out-of-scope set
    out_of_scope_ct_ids = set()
    if current_user.role in (Role.BAN_QUANLUC, Role.BAN_CANBO):
        for pd in pending_reviews:
            for ct in pd.de_xuat.chi_tiets:
                if not _is_in_dept_scope(current_user.role, ct.doi_tuong):
                    out_of_scope_ct_ids.add(ct.id)

    # Build item results
    all_item_results = {}
    for pd in pending_reviews:
        all_item_results[pd.id] = {kq.chi_tiet_id: kq for kq in pd.chi_tiet_duyet}

    # Get criteria fields for current department
    phong_fields_map = get_phong_fields()
    field_labels = get_field_labels()
    criteria_fields = phong_fields_map.get(current_user.role, [])

    wb = Workbook()
    ws = wb.active
    ws.title = 'Phê duyệt khen thưởng'

    # Timestamp header
    ts = datetime.now().strftime('%d/%m/%Y %H:%M')
    total_cols = 8 + len(criteria_fields)
    last_col = get_column_letter(total_cols)
    ws.merge_cells(f'A1:{last_col}1')
    ws['A1'] = f'DANH SÁCH PHÊ DUYỆT KHEN THƯỞNG - {phong_name}'
    ws['A1'].font = Font(bold=True, size=13)
    ws['A1'].alignment = Alignment(horizontal='center')

    ws.merge_cells(f'A2:{last_col}2')
    ws['A2'] = f'Năm học: {nam_hoc_filter or "Tất cả"} | Xuất lúc: {ts}'
    ws['A2'].alignment = Alignment(horizontal='center')
    ws['A2'].font = Font(italic=True, size=10)

    # Header row
    base_headers = ['STT', 'Đơn vị', 'Họ tên', 'Cấp bậc', 'Chức vụ', 'Đối tượng', 'Danh hiệu', 'Kết quả']
    criteria_headers = [field_labels.get(f, f) for f in criteria_fields]
    headers = base_headers + criteria_headers

    header_fill = PatternFill('solid', fgColor='1B3A6B')
    header_font = Font(bold=True, color='FFFFFF', size=10)
    thin = Side(style='thin', color='CCCCCC')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=4, column=col, value=h)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        cell.border = border

    row_num = 5
    stt = 0
    for pd in pending_reviews:
        results = all_item_results.get(pd.id, {})
        don_vi = pd.de_xuat.don_vi.ten_don_vi
        for ct in pd.de_xuat.chi_tiets:
            if ct.id in out_of_scope_ct_ids:
                continue
            stt += 1
            kq = results.get(ct.id)
            ket_qua_str = ''
            if kq:
                if kq.ket_qua == 'Đồng ý':
                    ket_qua_str = 'Nhất trí'
                elif kq.ket_qua == 'Từ chối':
                    ket_qua_str = f'Không NT: {kq.ly_do or ""}'
                else:
                    ket_qua_str = 'Chờ duyệt'
            ho_ten = ct.quan_nhan.ho_ten if ct.quan_nhan else don_vi
            cap_bac = ct.quan_nhan.cap_bac if ct.quan_nhan else ''
            chuc_vu = ct.quan_nhan.chuc_vu if ct.quan_nhan else ''

            # Build criteria values (individual: getattr; tap_the: tap_the_dict)
            is_tap_the = ct.quan_nhan_id is None
            if is_tap_the:
                tt_dict = ct.tap_the_dict or {}
                criteria_vals = [tt_dict.get(f, '') for f in criteria_fields]
            else:
                criteria_vals = [getattr(ct, f, '') or '' for f in criteria_fields]

            row_data = [stt, don_vi, ho_ten, cap_bac or '', chuc_vu or '',
                        ct.doi_tuong or '', ct.loai_danh_hieu or '', ket_qua_str] + criteria_vals
            for col, val in enumerate(row_data, 1):
                cell = ws.cell(row=row_num, column=col, value=val)
                cell.border = border
                cell.alignment = Alignment(vertical='center', wrap_text=True)
                if ket_qua_str == 'Nhất trí':
                    cell.fill = PatternFill('solid', fgColor='D4EDDA')
                elif ket_qua_str.startswith('Không NT'):
                    cell.fill = PatternFill('solid', fgColor='F8D7DA')
            row_num += 1

    # Column widths
    base_widths = [6, 30, 25, 16, 20, 18, 22, 28]
    criteria_widths = [18] * len(criteria_fields)
    col_widths = base_widths + criteria_widths
    for i, w in enumerate(col_widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w
    ws.row_dimensions[4].height = 28

    # Page setup: A4 landscape, fit to 1 page wide
    ws.page_setup.paperSize = 9
    ws.page_setup.orientation = 'landscape'
    ws.page_setup.fitToPage = True
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0
    ws.sheet_properties.pageSetUpPr.fitToPage = True

    # Sheet protection
    ws.protection.sheet = True
    ws.protection.password = 'hktd@2025'

    output = BytesIO()
    wb.save(output)
    output.seek(0)
    ts_file = datetime.now().strftime('%Y%m%d_%H%M')
    filename = f'phe_duyet_{phong_name.replace(" ", "_")}_{ts_file}.xlsx'
    return send_file(output, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                     as_attachment=True, download_name=filename)

@approval_bp.route('/export-word')
@login_required
@department_required
def export_word():
    from io import BytesIO
    from datetime import date, datetime
    import zipfile
    import os
    import binascii
    import hashlib
    from sqlalchemy.orm import joinedload, selectinload

    phong_name     = ROLE_TO_PHONG.get(current_user.role, '')
    nam_hoc_filter = request.args.get('nam_hoc', '')
    q = (
    PheDuyet.query
    .filter_by(phong_duyet=phong_name)
    .join(DeXuat)  # Bắt buộc join DeXuat để có thể dùng DeXuat trong filter và order_by bên dưới
    .options(
        selectinload(PheDuyet.chi_tiet_duyet),
        joinedload(PheDuyet.de_xuat).options(
            joinedload(DeXuat.don_vi),
            selectinload(DeXuat.chi_tiets).joinedload(DeXuatChiTiet.quan_nhan),
        ),
    )
    )

    # Áp dụng các bộ lọc tương tự như mẫu
    if nam_hoc_filter:
        q = q.filter(DeXuat.nam_hoc == nam_hoc_filter)

    q = q.join(DonVi, DeXuat.don_vi_id == DonVi.id)

    # Sắp xếp và lấy kết quả
    q = q.order_by(
        DonVi.thu_tu.asc(), DeXuat.nam_hoc.desc(), DeXuat.ngay_gui.desc()
    ).all()
    # ── 1. Query với eager loading đầy đủ ────────────────────────────────────
    
    pending_reviews = q

    # ── 2. Out-of-scope (không query thêm) ───────────────────────────────────
    out_of_scope_ct_ids = set()
    if current_user.role in (Role.BAN_QUANLUC, Role.BAN_CANBO):
        for pd in pending_reviews:
            for ct in pd.de_xuat.chi_tiets:
                if not _is_in_dept_scope(current_user.role, ct.doi_tuong):
                    out_of_scope_ct_ids.add(ct.id)

    all_item_results = {
        pd.id: {kq.chi_tiet_id: kq for kq in pd.chi_tiet_duyet}
        for pd in pending_reviews
    }

    phong_fields_map = get_phong_fields()
    field_labels     = get_field_labels()
    criteria_fields  = phong_fields_map.get(current_user.role, [])

    # ── 3. Batch load TieuChi cho tap_the_dict ───────────────────────────────
    all_ma_truong = set()
    for pd in pending_reviews:
        for ct in pd.de_xuat.chi_tiets:
            if ct.quan_nhan_id is None:
                all_ma_truong.update((ct.tap_the_dict or {}).keys())

    tieu_chi_map: dict[str, str] = {}
    if all_ma_truong:
        from app.models.nomination import TieuChi as _TieuChi
        tieu_chi_map = {
            tc.ma_truong: tc.ten
            for tc in _TieuChi.query.filter(
                _TieuChi.ma_truong.in_(list(all_ma_truong))
            ).all()
        }
    
    # ── 4. Phân loại chi tiết theo danh hiệu ─────────────────────────────────
    ds_don_vi_qt  = []
    ds_don_vi_tt  = []
    ds_ca_nhan_td = []
    ds_ca_nhan_tt = []
    seen_ids      = set()
    title_nam_hoc = nam_hoc_filter
    for pd in pending_reviews:
        results = all_item_results.get(pd.id, {})
        don_vi  = pd.de_xuat.don_vi.ten_don_vi if pd.de_xuat and pd.de_xuat.don_vi else ''
        for ct in pd.de_xuat.chi_tiets:
            if ct.id in out_of_scope_ct_ids or ct.id in seen_ids:
                continue
            seen_ids.add(ct.id)

            kq  = results.get(ct.id)
            if kq:
                if kq.ket_qua == KetQuaDuyet.DONG_Y.value:
                    ket_qua_str = 'Nhất trí'
                elif kq.ket_qua == KetQuaDuyet.TU_CHOI.value:
                    ket_qua_str = f'Không NT: {kq.ly_do or ""}'
                else:
                    ket_qua_str = 'Chờ duyệt'
            else:
                ket_qua_str = 'Chờ duyệt'

            item = dict(ct=ct, pd=pd, don_vi=don_vi,
                        kq=kq, ket_qua_str=ket_qua_str)

            dh = (ct.loai_danh_hieu or '').strip()
            if dh == 'Đơn vị quyết thắng':
                ds_don_vi_qt.append(item)
            elif dh == 'Đơn vị tiên tiến':
                ds_don_vi_tt.append(item)
            elif dh == 'Chiến sĩ thi đua':
                ds_ca_nhan_td.append(item)
            elif dh == 'Chiến sĩ tiên tiến':
                ds_ca_nhan_tt.append(item)


    # ═══════════════════════════════════════════════════════════════════════════
    # BUILD DOCX — XML template approach (nhanh hơn python-docx ~700x)
    # ═══════════════════════════════════════════════════════════════════════════
    from app.utils.docx_fast import (
        cm_to_twips, _para, _build_table, _data_row,
        _build_document_xml, build_docx,
    )
    import datetime as _dt

    # ── Columns: cá nhân ────────────────────────────────────────────────────────
    CN_WIDTHS = [
        cm_to_twips(0.9),   # STT
        cm_to_twips(3.5),   # Họ tên
        cm_to_twips(1.8),   # Cấp bậc
        cm_to_twips(2.2),   # Chức vụ
        cm_to_twips(2.5),   # Đơn vị
        cm_to_twips(5.5),   # Tóm tắt
        cm_to_twips(1.0),
        
    ]
    CN_HEADERS = ['STT', 'Họ và tên', 'Cấp bậc', 'Chức vụ', 'Đơn vị', 'Tóm tắt thành tích', 'Kết quả duyệt']

    # ── Columns: đơn vị ─────────────────────────────────────────────────────────
    DV_WIDTHS = [
        cm_to_twips(0.7),   # STT
        cm_to_twips(2.5),   # Tên đơn vị
        cm_to_twips(3.0),   # Đề xuất DV
        cm_to_twips(10.0),  # Ghi chú
    ]
    DV_HEADERS = ['STT', 'Tên đơn vị', 'Đề xuất của đơn vị', 'Ghi chú tiêu chí']

    def _cn_rows(items):
        rows_xml = []
        stt =0
        for i, item in enumerate(items, 1):
            if item['ct'].phong_loai == 'Tuyên huấn':
                continue
            stt += 1
            ct = item['ct']; qn = ct.quan_nhan
            
            row_cells = [
                (str(stt), False, 'center'),
                (qn.ho_ten if qn else '', True, 'justify'),
                (qn.cap_bac if qn and qn.cap_bac else '', False, 'left'),
                (qn.chuc_vu if qn and qn.chuc_vu else '', False, 'left'),
                (item['don_vi'], False, 'justify'),
                (build_tom_tat(item['ct']) or '', False, 'justify'),
                (item['ket_qua_str'], False, 'justify')
            ]
            shade = None if i % 2 == 0 else None
            rows_xml.append(_data_row(row_cells, CN_WIDTHS, size_pt=9, shade=shade))
        return rows_xml
    def build_tom_tat(ct) -> list[str]:
        parts = []
        if ct.muc_do_hoan_thanh:
            parts.append(ct.muc_do_hoan_thanh)
        if ct.diem_tong_ket:
            parts.append(f'Kết quả học tập: {ct.diem_tong_ket}')
        if ct.ket_qua_ren_luyen:
            parts.append(f'Rèn luyện: {ct.ket_qua_ren_luyen}')
        if ct.hinh_thuc_tot_nghiep:
            tn = [f'TN: {ct.hinh_thuc_tot_nghiep}']
            for attr, label in (
                ('diem_tn_ctd',         'CTĐ-CT'),
                ('diem_tn_ct',          'CT'),
                ('diem_tn_ta',          'TA'),
                ('diem_tn_mon4',        'Môn 4'),
                ('diem_tn_chuyennganh', 'Chuyên ngành'),
                ('diem_tn_baove',       'Bảo vệ'),
            ):
                val = getattr(ct, attr, None)
                if val:
                    tn.append(f'{label}: {val}')
            parts.append(', '.join(tn))
        if ct.mo_ta_khoa_hoc and phong_name != PhongDuyet.PHONG_KHOAHOC.value:
            parts.append(f'NCKH: {ct.mo_ta_khoa_hoc}')
        if ct.thanh_tich_ca_nhan_khac:
            parts.append(ct.thanh_tich_ca_nhan_khac)
      
            
        for f in criteria_fields:
            if f in ('nckh_noi_dung',): # Thêm dấu phẩy để tạo một tuple đúng chuẩn nếu chỉ có 1 phần tử
                continue
            val = getattr(ct, f, None) or ''
            if val:
                parts.append(f'{field_labels.get(f, f)}: {val}')
        return '\n'.join(parts)
    def _dv_rows(items):
        rows_xml = []
        stt = 0
        for i, item in enumerate(items, 1):
            ct = item['ct']
            if ct.phong_loai == 'Tuyên huấn':
                continue
            stt += 1
            td = ct.tap_the_dict or {}
            criteria_lines = []
            for k, v in td.items():
                if v and str(v).strip() not in ('', '0', 'None'):
                    label = tieu_chi_map.get(k, k)
                    criteria_lines.append(f'- {label}: {v}')
            if ct.muc_do_hoan_thanh:
                criteria_lines.insert(0, f'- Mức độ HT: {ct.muc_do_hoan_thanh}')
            row_cells = [
                (str(stt), False, 'center'),
                (ct.ten_don_vi_de_xuat or item['don_vi'] or '-', True, 'left'),
                (item['don_vi'] or '-', False, 'center'),
                ('\n'.join(criteria_lines), False, 'justify'),
            ]
            shade = None if i % 2 == 0 else None
            rows_xml.append(_data_row(row_cells, DV_WIDTHS, size_pt=9, shade=shade))
        return rows_xml

    today_str = _dt.date.today().strftime('%d/%m/%Y')
    body = []
    now = _dt.datetime.now()
      # ─────────────────────────────────────────────────────────────────────────
    # [TÍCH HỢP] BẢNG QUỐC HIỆU TIÊU NGỮ & TIÊU ĐỀ CHUẨN ĐÚNG THEO XML MẪU
    # ─────────────────────────────────────────────────────────────────────────
    header_table_xml = f"""<w:tbl xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
    <w:tblPr>
    <w:tblW w:type="dxa" w:w="9070"/>
    <w:jc w:val="left"/>
    <w:tblLayout w:type="fixed"/>
    <w:tblLook w:firstColumn="1" w:firstRow="1" w:lastColumn="0" w:lastRow="0" w:noHBand="0" w:noVBand="1" w:val="04A0"/>
    </w:tblPr>
    <w:tblGrid>
    <w:gridCol w:w="4535"/>
    <w:gridCol w:w="4535"/>
    </w:tblGrid>
    <w:tr>
    <w:tc>
    <w:tcPr>
    <w:tcW w:type="dxa" w:w="4535"/>
    <w:tcBorders>
    <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>
    <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>
    <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>
    <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>
    <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>
    <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>
    </w:tcBorders>
    </w:tcPr>
    <w:p>
    <w:pPr>
    <w:spacing w:before="0" w:after="0" w:line="240" w:lineRule="auto"/>
    <w:jc w:val="center"/>
    </w:pPr>
    <w:r>
    <w:rPr>
    <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>
    <w:b w:val="0"/>
    <w:i w:val="0"/>
    <w:sz w:val="24"/>
    </w:rPr>
    <w:t>TRƯỜNG SĨ QUAN CHÍNH TRỊ</w:t>
    </w:r>
    </w:p>
    <w:p>
    <w:pPr>
    <w:spacing w:before="0" w:after="0" w:line="240" w:lineRule="auto"/>
    <w:jc w:val="center"/>
    </w:pPr>
    <w:r>
    <w:rPr>
    <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>
    <w:b/>
    <w:i w:val="0"/>
    <w:sz w:val="24"/>
    <w:u w:val="single"/>
    </w:rPr>
    <w:t>{phong_name.upper()}</w:t>
    </w:r>
    </w:p>
    <w:p>
    <w:pPr>
    <w:spacing w:before="0" w:after="0" w:line="240" w:lineRule="auto"/>
    </w:pPr>
    <w:r>
    <w:rPr>
    <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>
    <w:b w:val="0"/>
    <w:i w:val="0"/>
    <w:sz w:val="22"/>
    </w:rPr>
    </w:r>
    </w:p>
    </w:tc>
    <w:tc>
    <w:tcPr>
    <w:tcW w:type="dxa" w:w="4535"/>
    <w:tcBorders>
    <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>
    <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>
    <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>
    <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>
    <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>
    <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>
    </w:tcBorders>
    </w:tcPr>
    <w:p>
    <w:pPr>
    <w:spacing w:before="0" w:after="0" w:line="240" w:lineRule="auto"/>
    <w:jc w:val="center"/>
    </w:pPr>
    <w:r>
    <w:rPr>
    <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>
    <w:b/>
    <w:i w:val="0"/>
    <w:sz w:val="21"/>
    </w:rPr>
    <w:t>CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM</w:t>
    </w:r>
    </w:p>
    <w:p>
    <w:pPr>
    <w:spacing w:before="0" w:after="0" w:line="240" w:lineRule="auto"/>
    <w:jc w:val="center"/>
    </w:pPr>
    <w:r>
    <w:rPr>
    <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>
    <w:b/>
    <w:i w:val="0"/>
    <w:sz w:val="21"/>
    <w:u w:val="single"/>
    </w:rPr>
    <w:t>Độc lập - Tự do - Hạnh phúc</w:t>
    </w:r>
    </w:p>
    <w:p>
    <w:pPr>
    <w:spacing w:before="0" w:after="0" w:line="240" w:lineRule="auto"/>
    <w:jc w:val="center"/>
    </w:pPr>
    <w:r>
    <w:rPr>
    <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>
    <w:b w:val="0"/>
    <w:i/>
    <w:sz w:val="22"/>
    </w:rPr>
    <w:t>Hà Nội, ngày {now.day} tháng {now.month} năm {now.year}</w:t>
    </w:r>
    </w:p>
    </w:tc>
    </w:tr>
    </w:tbl>"""

    title_xml = f"""<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
    <w:pPr>
    <w:jc w:val="center"/>
    <w:spacing w:before="180" w:after="40"/>
    </w:pPr>
    <w:r>
    <w:rPr>
    <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>
    <w:b/>
    <w:i w:val="0"/>
    <w:sz w:val="26"/>
    </w:rPr>
    <w:t>DANH SÁCH ĐỀ NGHỊ KHEN THƯỞNG NĂM HỌC {title_nam_hoc.upper()}</w:t>
    </w:r>
    </w:p>"""

    today_str = _dt.date.today().strftime('%d/%m/%Y')
    body.append(header_table_xml)
    body.append("<w:p/>")
    body.append(title_xml)
    body.append(_para(f'(Xuất lúc {now.strftime("%H:%M")} ngày {today_str})', italic=True, size_pt=10, align='center', space_before=0, space_after=120))

    def _add_section(label, items, is_tap_the=False):
        if not items:
            return
        body.append(_para(label, bold=True, size_pt=11, space_before=120, space_after=40))
        if is_tap_the:
            rows_xml = _dv_rows(items)
            body.append(_build_table(DV_HEADERS, rows_xml, DV_WIDTHS, total_label=f'Tổng cộng: {len(items)} đơn vị', size_pt=9))
        else:
            rows_xml = _cn_rows(items)
            body.append(_build_table(CN_HEADERS, rows_xml, CN_WIDTHS, total_label=f'Tổng cộng: {len(items)} người', size_pt=9))
        body.append(_para('', space_before=60, space_after=0))

    _add_section('I. DANH HIỆU ĐƠN VỊ QUYẾT THẮNG', ds_don_vi_qt, is_tap_the=True)
    _add_section('II. DANH HIỆU ĐƠN VỊ TIÊN TIẾN', ds_don_vi_tt, is_tap_the=True)
    _add_section('III. DANH HIỆU CHIẾN SĨ THI ĐUA', ds_ca_nhan_td)
    _add_section('IV. DANH HIỆU CHIẾN SĨ TIÊN TIẾN', ds_ca_nhan_tt)

    body.append(_para(f'(Xuất lúc {_dt.datetime.now().strftime("%H:%M ngày %d/%m/%Y")})',
                      italic=True, size_pt=9, align='right', space_before=120, space_after=0))
    chu_ky_xml = f"""<w:tbl>
<w:tblPr>
<w:tblW w:type="auto" w:w="0"/>
<w:jc w:val="center"/>
<w:tblLook w:firstColumn="1" w:firstRow="1" w:lastColumn="0" w:lastRow="0" w:noHBand="0" w:noVBand="1" w:val="04A0"/>
</w:tblPr>
<w:tblGrid>
<w:gridCol w:w="4320"/>
<w:gridCol w:w="4320"/>
</w:tblGrid>
<w:tr>
<w:tc>
<w:tcPr>
<w:tcW w:type="dxa" w:w="4320"/>
<w:tcBorders>
<w:top w:val="none"/>
</w:tcBorders>
<w:tcBorders>
<w:left w:val="none"/>
</w:tcBorders>
<w:tcBorders>
<w:bottom w:val="none"/>
</w:tcBorders>
<w:tcBorders>
<w:right w:val="none"/>
</w:tcBorders>
</w:tcPr>
<w:p>
<w:pPr>
<w:jc w:val="center"/>
</w:pPr>
</w:p>
<w:p/>
<w:p/>
<w:p/>
</w:tc>
<w:tc>
<w:tcPr>
<w:tcW w:type="dxa" w:w="4320"/>
<w:tcBorders>
<w:top w:val="none"/>
</w:tcBorders>
<w:tcBorders>
<w:left w:val="none"/>
</w:tcBorders>
<w:tcBorders>
<w:bottom w:val="none"/>
</w:tcBorders>
<w:tcBorders>
<w:right w:val="none"/>
</w:tcBorders>
</w:tcPr>
<w:p>
<w:pPr>
<w:jc w:val="center"/>
</w:pPr>
<w:r>
<w:rPr>
<w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>
<w:b/>
<w:i w:val="0"/>
<w:sz w:val="22"/>
</w:rPr>
<w:t>THỦ TRƯỞNG ĐƠN VỊ</w:t>
</w:r>
</w:p>
<w:p>
<w:pPr>
<w:jc w:val="center"/>
</w:pPr>
<w:r>
<w:rPr>
<w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>
<w:b w:val="0"/>
<w:i/>
<w:sz w:val="20"/>
</w:rPr>
<w:t>(Ký, ghi rõ họ tên)</w:t>
</w:r>
</w:p>
</w:tc>
</w:tr>
</w:tbl>"""
    body.append(chu_ky_xml)
    doc_xml = _build_document_xml(body, margin_left=2016, margin_right=720, margin_top=1440, margin_bottom=1440)
    buf = build_docx(doc_xml)

    final_buf = BytesIO()
    buf.seek(0)
    # Thuật toán hash mật khẩu Office 2010+ (Agile Encryption)
    password = "bth123" # <--- THAY ĐỔI MẬT KHẨU TẠI ĐÂY
    salt = os.urandom(16)
    salt_b64 = binascii.b2a_base64(salt).strip().decode()

    key = hashlib.sha512(salt + password.encode('utf-16le')).digest()
    spin_count = 10000
    for i in range(spin_count):
        iterator = i.to_bytes(4, byteorder='little')
        key = hashlib.sha512(key + iterator).digest()
        
    hash_b64 = binascii.b2a_base64(key).strip().decode()

    # Tạo chuỗi XML bảo vệ theo chuẩn Agile
    protection_tag = (
        f'<w:documentProtection w:edit="readOnly" w:enforcement="1" '
        f'w:algorithmName="SHA-512" w:spinCount="{spin_count}" '
        f'w:hashValue="{hash_b64}" w:saltValue="{salt_b64}"/>'
    ).encode('utf-8')

    with zipfile.ZipFile(buf, 'r') as zin:
        with zipfile.ZipFile(final_buf, 'w') as zout:
            for item in zin.infolist():
                file_content = zin.read(item.filename)
                
                # Tìm file cấu hình settings.xml
                if item.filename == 'word/settings.xml':
                    # Tiêm mã bảo vệ vào trước khi kết thúc thẻ w:settings
                    if b'</w:settings>' in file_content:
                        # Tránh trùng lặp nếu thẻ đã tồn tại
                        if b'w:documentProtection' not in file_content:
                            file_content = file_content.replace(b'</w:settings>', protection_tag + b'</w:settings>')
                
                # Copy toàn bộ nội dung sang file zip mới
                zout.writestr(item, file_content)
    
    final_buf.seek(0)
    # =========================================================================

    fname_parts = ['DanhSachKhenThuong']
    if nam_hoc_filter:
        fname_parts.append(nam_hoc_filter.replace('-', '_'))
    fname_parts.append(now.strftime('%d%m%Y'))
    filename = '_'.join(fname_parts) + '.docx'

    response = send_file(
        final_buf, as_attachment=True, download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response.set_cookie(
        'export_done', '1',
        max_age=600, httponly=False, samesite='Lax', path='/'
    )
    return response

    # ts_file  = datetime.now().strftime('%Y%m%d_%H%M')
    # filename = f'phe_duyet_{phong_name.replace(" ", "_")}_{ts_file}.docx'
    # return send_file(
    #     buf, as_attachment=True, download_name=filename,
    #     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    # )

def set_fixed_table_widths(tbl, widths_cm):
        """Can thiệp sâu vào XML để khóa chết chiều rộng bảng, Word không thể tự đổi"""
        # 1. Ép kiểu bảng thành Fixed Layout (Không tự co giãn)
        tbl.autofit = False
        tblPr = tbl._tbl.tblPr
        tblLayout = tblPr.find(qn('w:tblLayout'))
        if tblLayout is None:
            tblLayout = OxmlElement('w:tblLayout')
            tblPr.append(tblLayout)
        tblLayout.set(qn('w:type'), 'fixed')

        # 2. Xóa lưới cột cũ và xây lại khung lưới mới theo đúng kích thước cm
        tblGrid = tbl._tbl.find(qn('w:tblGrid'))
        if tblGrid is not None:
            tbl._tbl.remove(tblGrid)
        tblGrid = OxmlElement('w:tblGrid')
        tbl._tbl.insert(1, tblGrid)  # Chèn khung lưới vào đúng vị trí chuẩn XML

        for w in widths_cm:
            gridCol = OxmlElement('w:gridCol')
            # Chuyển đổi Cm sang đơn vị Twips của Word (1 twip = 635 EMUs)
            gridCol.set(qn('w:w'), str(int(Cm(w) / 635)))
            tblGrid.append(gridCol)

        # 3. Khóa cứng chiều rộng ở cấp độ từng Ô (Cell)
        for i, w in enumerate(widths_cm):
            twips_val = str(int(Cm(w) / 635))
            for row in tbl.rows:
                tcPr = row.cells[i]._tc.get_or_add_tcPr()
                tcW = tcPr.find(qn('w:tcW'))
                if tcW is None:
                    tcW = OxmlElement('w:tcW')
                    tcPr.append(tcW)
                tcW.set(qn('w:w'), twips_val)
                tcW.set(qn('w:type'), 'dxa')
def add_corner_logo(doc):
    """Thêm logo nhỏ ở góc phải trên cùng của trang (sau header table hiện tại)."""
    import os
    from flask import current_app

    # ★ Ưu tiên logo nhỏ (19 KB) để giảm kích thước file docx
    logo_path = os.path.join(current_app.root_path, 'static', 'img', 'logo-Si-quan.png')
    if not os.path.exists(logo_path):
        logo_path = os.path.join(current_app.root_path, 'static', 'img', 'watermark.png')
        if not os.path.exists(logo_path):
            return
    
    try:
        for section in doc.sections:
            header = section.header
            
            # Thêm paragraph mới vào cuối header (sau table header hiện tại)
            para = header.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Set paragraph spacing để logo sát lề trên
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            
            # Thêm logo nhỏ căn phải (1.5cm)
            run = para.add_run()
            run.add_picture(logo_path, width=Cm(1.5))
            
    except Exception as e:
        print(f"Warning: Could not add corner logo: {e}")


def protect_document_formatting_only(doc, password: str):
    """
    Khóa tài liệu: chỉ đọc nội dung (readOnly).
    Mật khẩu được hash theo chuẩn Office 2010+ (Agile Encryption).
    """
    # 1. Tạo salt ngẫu nhiên (16 bytes)
    salt = os.urandom(16)
    salt_b64 = binascii.b2a_base64(salt).strip().decode()

    # 2. Hash lần đầu: SHA-512(salt + password)
    # Lưu ý: password bắt buộc encode sang chuẩn UTF-16 Little Endian
    key = hashlib.sha512(salt + password.encode('utf-16le')).digest()
    
    # 3. Lặp N vòng để chống brute-force.
    # NOTE: Đây chỉ là khóa định dạng (readOnly), không phải bảo mật nội dung thật sự
    # (password đã hardcode trong source). 10.000 vòng vẫn tuân thủ chuẩn Agile Encryption
    # nhưng nhanh hơn ~10 lần so với 100.000 vòng, giúp xuất Word nhanh hơn đáng kể.
    spin_count = 10000
    for i in range(spin_count):
        iterator = i.to_bytes(4, byteorder='little')
        # SỬA LỖI: Cần cộng iterator ở PHÍA SAU hash của vòng lặp liền trước
        key = hashlib.sha512(key + iterator).digest()
        
    hash_b64 = binascii.b2a_base64(key).strip().decode()

    # 4. Lấy cấu hình settings của docx
    settings = doc.settings.element

    # Xóa thẻ documentProtection cũ nếu có
    for old in settings.findall(qn('w:documentProtection')):
        settings.remove(old)

    # 5. Tạo thẻ <w:documentProtection> theo chuẩn Office đời mới
    doc_prot = OxmlElement('w:documentProtection')
    doc_prot.set(qn('w:edit'),          'readOnly')
    doc_prot.set(qn('w:enforcement'),   '1')
    
    # BỎ CÁC THẺ CŨ (cryptProviderType, v.v.). SỬ DỤNG CHUẨN AGILE MỚI:
    doc_prot.set(qn('w:algorithmName'), 'SHA-512')
    doc_prot.set(qn('w:spinCount'),     str(spin_count))
    doc_prot.set(qn('w:hashValue'),     hash_b64)     # Đã đổi từ w:hash thành w:hashValue
    doc_prot.set(qn('w:saltValue'),     salt_b64)     # Đã đổi từ w:salt thành w:saltValue

    # Chèn vào đầu <w:settings>
    settings.insert(0, doc_prot)
