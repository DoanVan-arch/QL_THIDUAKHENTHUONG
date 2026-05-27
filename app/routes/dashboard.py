from flask import Blueprint, render_template, send_file
from flask_login import login_required, current_user
import io
from app.models.user import Role
from app.models.unit import DonVi
from app.models.personnel import QuanNhan
from app.models.nomination import DeXuat, TrangThaiDeXuat
from app.models.approval import PheDuyet, KetQuaDuyet

dashboard_bp = Blueprint('dashboard', __name__)


@dashboard_bp.route('/')
@login_required
def index():
    if current_user.is_unit_user:
        return unit_dashboard()
    elif current_user.is_department:
        return department_dashboard()
    elif current_user.is_admin:
        return admin_dashboard()
    return render_template('dashboard/unit_dashboard.html')


@dashboard_bp.route('/huong-dan')
@login_required
def user_guide():
    """Hướng dẫn sử dụng hệ thống - accessible by all authenticated users."""
    return render_template('dashboard/user_guide.html')


@dashboard_bp.route('/huong-dan/tai-word')
@login_required
def download_user_guide_word():
    """Tải file Word hướng dẫn sử dụng."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    title = doc.add_heading('HƯỚNG DẪN SỬ DỤNG HỆ THỐNG', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph('Quản lý Thi đua, Khen thưởng – Trường Sĩ quan Chính trị')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ---- TỔNG QUAN ----
    doc.add_heading('I. TỔNG QUAN HỆ THỐNG', level=1)
    doc.add_paragraph(
        'Hệ thống Quản lý Thi đua, Khen thưởng của Trường Sĩ quan Chính trị giúp số hóa toàn bộ '
        'quy trình đề xuất, xét duyệt và quản lý kết quả thi đua khen thưởng.'
    )

    doc.add_heading('Các nhóm người dùng:', level=2)
    roles = [
        ('Đơn vị (Đề xuất)', 'Quản lý quân nhân, tạo và gửi đề xuất khen thưởng, theo dõi kết quả.'),
        ('Cơ quan (Phê duyệt)', '6 cơ quan xét duyệt song song: P.Chính trị, P.Tham mưu, P.Khoa học, '
                                'P.Đào tạo, B.Cán bộ, B.Quân lực.'),
        ('Ban Tuyên huấn (Quản trị)', 'Theo dõi toàn bộ quy trình, phê duyệt cuối, quản lý danh sách '
                                      'khen thưởng, báo cáo thống kê.'),
    ]
    for name, desc in roles:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(name + ': ').bold = True
        p.add_run(desc)

    # ---- QUY TRÌNH ----
    doc.add_heading('II. QUY TRÌNH PHÊ DUYỆT', level=1)
    steps = [
        ('Bước 1 – Đơn vị tạo đề xuất',
         'Đơn vị chọn quân nhân, điền các tiêu chí đánh giá, đính kèm minh chứng và gửi đề xuất khen thưởng.'),
        ('Bước 2 – 6 cơ quan xét duyệt song song',
         'Mỗi cơ quan xét duyệt các tiêu chí thuộc lĩnh vực phụ trách, đưa ra kết quả Đồng ý hoặc Từ chối '
         'cho từng cá nhân.'),
        ('Bước 3 – Ban Tuyên huấn phê duyệt cuối',
         'Khi cả 6 cơ quan đồng ý, Ban Tuyên huấn xem xét và phê duyệt cuối cùng, lưu kết quả vào danh sách '
         'khen thưởng.'),
    ]
    for title_step, desc in steps:
        p = doc.add_paragraph(style='List Number')
        p.add_run(title_step + ': ').bold = True
        p.add_run(desc)

    # ---- ĐƠN VỊ ----
    doc.add_heading('III. HƯỚNG DẪN DÀNH CHO ĐƠN VỊ', level=1)

    doc.add_heading('1. Quản lý quân nhân', level=2)
    doc.add_paragraph('Vào menu Quân nhân → Danh sách quân nhân để quản lý toàn bộ nhân sự của đơn vị.')
    p = doc.add_paragraph(); p.add_run('Thêm mới từng quân nhân:').bold = True
    for item in [
        'Bấm nút Thêm quân nhân.',
        'Điền đầy đủ thông tin bắt buộc: Họ tên, Cấp bậc, Chức vụ, Đối tượng, Ngày sinh.',
        'Điền thêm: năm nhập ngũ (MM/YYYY), học vị, trình độ, CCCD/CMND, số điện thoại...',
        'Đánh dấu các cờ: Đảng viên, Đoàn viên, Hội viên phụ nữ, Chỉ huy, Bí thư nếu đúng.',
        'Bấm Lưu.',
    ]:
        doc.add_paragraph(item, style='List Number 2')

    p = doc.add_paragraph(); p.add_run('Nhập hàng loạt từ file Excel mẫu:').bold = True
    for item in [
        'Bấm nút Tải Excel mẫu để tải file mẫu về máy.',
        'Điền thông tin quân nhân vào đúng các cột: ho_ten, ngay_sinh (dd/mm/yyyy), ngay_nhap_ngu (MM/YYYY), doi_tuong (tên chính xác), la_dang_vien / la_doan_vien / la_hoi_vien_phu_nu / la_chi_huy / la_bi_thu (1 = có, 0 = không).',
        'Lưu file, quay lại hệ thống, bấm Import Excel, chọn file, bấm Upload.',
        'Hệ thống kiểm tra lỗi và thông báo kết quả import.',
        'Lưu ý: Không thay đổi tên tiêu đề cột. Số CCCD phải là duy nhất toàn hệ thống.',
    ]:
        doc.add_paragraph(item, style='List Number 2')

    doc.add_heading('2. Chỉnh sửa, Chuyển vùng, Chuyển đơn vị', level=2)
    for item in [
        'Chỉnh sửa: Bấm biểu tượng bút chì ở cột Thao tác → cập nhật thông tin → Lưu.',
        'Chuyển vùng: Tích chọn quân nhân → bấm Chuyển vùng → chọn diện quản lý mới (Quân lực / Cán bộ) → Xác nhận.',
        'Chuyển đơn vị: Tích chọn quân nhân → bấm Chuyển đơn vị → chọn đơn vị mới → Xác nhận. (Chức năng này do quản trị viên thực hiện.)',
    ]:
        doc.add_paragraph(item, style='List Number 2')

    doc.add_heading('3. Xóa quân nhân (Soft-delete)', level=2)
    doc.add_paragraph('Hệ thống áp dụng xóa mềm — quân nhân bị xóa vẫn được lưu trong CSDL, có thể khôi phục.')
    for item in [
        'Xóa từng người: Bấm biểu tượng thùng rác ở cột Thao tác → xác nhận.',
        'Xóa hàng loạt: Tích chọn nhiều quân nhân → bấm Xóa trong thanh thao tác hàng loạt → xác nhận.',
        'Xem danh sách đã xóa: Vào menu Quân nhân → Danh sách đã xóa.',
        'Khôi phục: Chỉ quản trị viên mới có thể khôi phục hoặc xóa vĩnh viễn.',
    ]:
        doc.add_paragraph(item, style='List Number 2')

    doc.add_heading('4. Tạo và gửi đề xuất khen thưởng', level=2)
    for item in [
        'Vào menu Đề xuất khen thưởng → Tạo đề xuất mới.',
        'Chọn Năm học, điền Ghi chú nếu cần, bấm Tạo.',
        'Chọn Danh hiệu đề xuất (Chiến sĩ thi đua, Chiến sĩ tiên tiến, Đơn vị quyết thắng, Đơn vị tiên tiến).',
        'Với danh hiệu cá nhân: chọn Quân nhân từ danh sách, hệ thống tự điền đối tượng.',
        'Với danh hiệu tập thể: nhập Tên tập thể (VD: Đại đội 1, Tiểu đoàn 6...).',
        'Điền tiêu chí đánh giá: Chung (Hoàn thành, Phiếu tín nhiệm, KT Chính trị, KT Điều lệnh, Tin học, ĐHQS, Bắn súng, Thể lực, Xếp loại đảng viên nếu là đảng viên, Đoàn viên, Phụ nữ...), Giảng viên, Học viên, NCKH.',
        'Đính kèm minh chứng (file PDF, ảnh) nếu cần.',
        'Bấm Thêm vào đề xuất. Lặp lại cho các cá nhân tiếp theo.',
        'Khi đủ, bấm Gửi duyệt ở góc trên phải. Đề xuất chuyển sang trạng thái Chờ duyệt.',
        'Lưu ý: Sau khi gửi không thể chỉnh sửa. Dùng Thu hồi nếu chưa có cơ quan nào duyệt.',
    ]:
        doc.add_paragraph(item, style='List Number 2')

    doc.add_heading('5. Thông báo', level=2)
    for item in [
        'Biểu tượng chuông trên thanh điều hướng hiển thị số thông báo chưa đọc.',
        'Bấm chuông để xem danh sách gần nhất; bấm Xem tất cả thông báo để vào trang đầy đủ.',
        'Loại thông báo: Từ chối (kèm lý do), Đồng ý, Phê duyệt cuối.',
        'Bấm vào từng thông báo để xem chi tiết và điều hướng đến đề xuất.',
        'Bấm Đánh dấu đã đọc tất cả để xóa số hiển thị trên chuông.',
    ]:
        doc.add_paragraph(item, style='List Number 2')

    doc.add_heading('6. Theo dõi kết quả đề xuất', level=2)
    for item in [
        'Vào menu Đề xuất khen thưởng → Lịch sử đề xuất.',
        'Xem trạng thái và kết quả duyệt của từng cơ quan cho tất cả đề xuất đã tạo.',
        'Lọc theo năm học hoặc trạng thái để tìm nhanh.',
    ]:
        doc.add_paragraph(item, style='List Number 2')

    doc.add_heading('7. Đánh giá xếp loại hằng năm', level=2)
    doc.add_paragraph('Nhập xếp loại năm học cho toàn đơn vị theo 4 tiêu chí: Đảng viên, Cán bộ, Đoàn viên, Phụ nữ.')
    for item in [
        'Vào menu Quân nhân → Đánh giá xếp loại hằng năm.',
        'Chọn Năm học cần nhập (VD: 2024-2025), bấm Tìm.',
        'Nhập xếp loại cho từng quân nhân trong bảng (4 cột xếp loại).',
        'Sử dụng 4 ô Áp dụng cho toàn đơn vị ở phía trên để điền nhanh hàng loạt, sau đó sửa từng dòng riêng lẻ nếu cần.',
        'Bấm Lưu đánh giá ở cuối trang để lưu kết quả.',
        'Lưu ý: Dữ liệu đánh giá hằng năm độc lập với đề xuất khen thưởng.',
    ]:
        doc.add_paragraph(item, style='List Number 2')

    # ---- CƠ QUAN ----
    doc.add_heading('IV. HƯỚNG DẪN DÀNH CHO CƠ QUAN PHÊ DUYỆT', level=1)

    doc.add_heading('1. Duyệt đề xuất', level=2)
    for item in [
        'Vào menu Chờ duyệt để xem danh sách đề xuất cần duyệt.',
        'Bấm vào một đề xuất để xem chi tiết từng cá nhân và các tiêu chí thuộc lĩnh vực quản lý.',
        'Với mỗi cá nhân, chọn Đồng ý hoặc Từ chối. Nếu từ chối, cần nhập lý do.',
        'Có thể duyệt hàng loạt (tất cả đồng ý) hoặc duyệt từng cá nhân.',
    ]:
        doc.add_paragraph(item, style='List Number 2')

    p = doc.add_paragraph()
    p.add_run('Phạm vi duyệt:\n').bold = True
    p.add_run('• Ban Quân lực duyệt: Công nhân viên, Quân nhân chuyên nghiệp, Công chức quốc phòng.\n')
    p.add_run('• Ban Cán bộ duyệt: tất cả đối tượng khác (Sĩ quan, Giảng viên, Học viên...).\n')
    p.add_run('• Các phòng chức năng duyệt theo lĩnh vực chuyên môn.')

    doc.add_heading('2. Thu hồi duyệt và Lịch sử', level=2)
    for item in [
        'Thu hồi: Nếu phát hiện sai sót sau khi duyệt, có thể thu hồi kết quả duyệt (trước khi Ban Tuyên huấn phê duyệt cuối).',
        'Lịch sử duyệt: Vào menu Lịch sử duyệt để xem tất cả đề xuất đã duyệt.',
    ]:
        doc.add_paragraph(item, style='List Number 2')

    # ---- ADMIN ----
    doc.add_heading('V. HƯỚNG DẪN DÀNH CHO BAN TUYÊN HUẤN (QUẢN TRỊ)', level=1)

    doc.add_heading('1. Theo dõi phê duyệt', level=2)
    for item in [
        'Vào menu Theo dõi phê duyệt để xem toàn bộ tiến trình duyệt của tất cả đề xuất.',
        'Bảng hiển thị theo đơn vị, mỗi dòng là một cá nhân, kèm trạng thái duyệt của 6 cơ quan.',
        'Bộ lọc: Lọc theo trạng thái, đơn vị, danh hiệu, diện quản lý, hoặc tìm kiếm theo họ tên.',
        'Chế độ xem: Xem gọn (mặc định) hoặc Xem chi tiết (hiển thị tất cả tiêu chí).',
    ]:
        doc.add_paragraph(item, style='List Number 2')

    doc.add_heading('2. Phê duyệt cuối', level=2)
    for item in [
        'Khi cả 6 cơ quan đều Đồng ý, bấm biểu tượng phê duyệt để phê duyệt cuối cho từng cá nhân.',
        'Hoặc dùng các nút PD cuối ĐV / PD cuối toàn bộ để phê duyệt hàng loạt.',
        'Sau khi phê duyệt cuối, cá nhân sẽ được lưu vào Danh sách khen thưởng.',
    ]:
        doc.add_paragraph(item, style='List Number 2')

    doc.add_heading('3. Danh sách khen thưởng và Xuất Excel', level=2)
    for item in [
        'Vào menu Danh sách khen thưởng để xem tất cả cá nhân đã được phê duyệt cuối.',
        'Lọc theo năm học, đơn vị, danh hiệu hoặc tìm kiếm theo họ tên.',
        'Bấm nút Xuất Excel để tải file với đầy đủ thông tin và tiêu chí đánh giá.',
        'Có thể thu hồi phê duyệt cuối nếu phát hiện sai sót.',
    ]:
        doc.add_paragraph(item, style='List Number 2')

    doc.add_heading('4. Quản lý tài khoản và đơn vị', level=2)
    for item in [
        'Quản lý tài khoản: Tạo, vô hiệu hóa, đặt lại mật khẩu cho các tài khoản đơn vị và cơ quan.',
        'Quản lý đơn vị: Thêm, sửa, kích hoạt/ngừng hoạt động các đơn vị trong trường.',
        'Báo cáo thống kê: Xem tổng quan số liệu thi đua khen thưởng theo đơn vị, danh hiệu, cơ quan duyệt.',
    ]:
        doc.add_paragraph(item, style='List Number 2')

    # ---- CHÚ GIẢI ----
    doc.add_heading('VI. CHÚ GIẢI TRẠNG THÁI ĐỀ XUẤT', level=1)
    statuses = [
        ('Nháp', 'Đề xuất chưa gửi, đang soạn.'),
        ('Chờ duyệt', 'Đã gửi, chờ các cơ quan xét duyệt.'),
        ('Đang duyệt', 'Một số cơ quan đã duyệt, còn lại đang xem xét.'),
        ('Đã duyệt', 'Cả 6 cơ quan đã đồng ý.'),
        ('Phê duyệt cuối', 'Ban Tuyên huấn đã phê duyệt, lưu khen thưởng.'),
        ('Từ chối', 'Đề xuất bị từ chối.'),
    ]
    for status, desc in statuses:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(status + ': ').bold = True
        p.add_run(desc)

    # ---- LIÊN HỆ ----
    doc.add_heading('VII. LIÊN HỆ HỖ TRỢ', level=1)
    doc.add_paragraph('Nếu gặp sự cố hoặc cần hỗ trợ kỹ thuật, vui lòng liên hệ:')
    doc.add_paragraph('Ban Tuyên huấn – Trường Sĩ quan Chính trị', style='List Bullet')
    doc.add_paragraph('Bộ phận quản trị hệ thống', style='List Bullet')

    # Save to buffer
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    return send_file(
        buf,
        as_attachment=True,
        download_name='Huong_dan_su_dung_He_thong_TDKT.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


def unit_dashboard():
    don_vi = current_user.don_vi
    if not don_vi:
        return render_template('dashboard/unit_dashboard.html', don_vi=None,
                               total_personnel=0, total_nominations=0,
                               pending_nominations=0, approved_nominations=0)

    total_personnel = QuanNhan.query.filter_by(don_vi_id=don_vi.id, is_active=True).count()
    total_nominations = DeXuat.query.filter_by(don_vi_id=don_vi.id).count()
    pending_nominations = DeXuat.query.filter_by(
        don_vi_id=don_vi.id, trang_thai=TrangThaiDeXuat.CHO_DUYET.value
    ).count()
    approved_nominations = DeXuat.query.filter_by(
        don_vi_id=don_vi.id, trang_thai=TrangThaiDeXuat.PHE_DUYET_CUOI.value
    ).count()

    recent_nominations = DeXuat.query.filter_by(don_vi_id=don_vi.id)\
        .order_by(DeXuat.ngay_tao.desc()).limit(5).all()

    return render_template('dashboard/unit_dashboard.html',
                           don_vi=don_vi,
                           total_personnel=total_personnel,
                           total_nominations=total_nominations,
                           pending_nominations=pending_nominations,
                           approved_nominations=approved_nominations,
                           recent_nominations=recent_nominations)


def department_dashboard():
    role_phong_map = {
        Role.PHONG_KHOAHOC: 'Phòng Khoa học',
        Role.PHONG_DAOTAO: 'Phòng Đào tạo',
        Role.THU_TRUONG_PHONG_TMHC: 'Thủ trưởng Phòng TM-HC',
        Role.BAN_CANBO: 'Ban Cán bộ',
        Role.BAN_TOCHUC: 'Ban Tổ chức',
        Role.BAN_TUYENHUAN: 'Ban Tuyên huấn',
        Role.BAN_CTCQ: 'Ban Công tác quần chúng',
        Role.BAN_CNTT: 'Ban Công nghệ thông tin',
        Role.BAN_TAC_HUAN: 'Ban Tác huấn',
        Role.BAN_KHAOTHI: 'Ban Khảo thí',
        Role.BAN_BAOVE_ANNINH: 'Ban Bảo vệ an ninh',
        Role.UY_BAN_KIEMTRA: 'Ủy ban Kiểm tra',
        Role.BAN_QUANLUC: 'Ban Quân lực',
    }
    phong_name = role_phong_map.get(current_user.role, '')

    pending_count = PheDuyet.query.filter_by(
        phong_duyet=phong_name,
        ket_qua=KetQuaDuyet.CHO_DUYET.value
    ).count()

    approved_count = PheDuyet.query.filter_by(
        phong_duyet=phong_name,
        ket_qua=KetQuaDuyet.DONG_Y.value
    ).count()

    rejected_count = PheDuyet.query.filter_by(
        phong_duyet=phong_name,
        ket_qua=KetQuaDuyet.TU_CHOI.value
    ).count()

    recent_pending = PheDuyet.query.filter_by(
        phong_duyet=phong_name,
        ket_qua=KetQuaDuyet.CHO_DUYET.value
    ).order_by(PheDuyet.created_at.desc()).limit(10).all()

    return render_template('dashboard/department_dashboard.html',
                           phong_name=phong_name,
                           pending_count=pending_count,
                           approved_count=approved_count,
                           rejected_count=rejected_count,
                           recent_pending=recent_pending)


def admin_dashboard():
    total_units = DonVi.query.filter_by(is_active=True).count()
    total_personnel = QuanNhan.query.filter_by(is_active=True).count()
    total_nominations = DeXuat.query.count()

    awaiting_final = DeXuat.query.filter_by(
        trang_thai=TrangThaiDeXuat.HOI_DONG.value
    ).count()

    final_approved = DeXuat.query.filter_by(
        trang_thai=TrangThaiDeXuat.PHE_DUYET_CUOI.value
    ).count()

    recent_nominations = DeXuat.query.order_by(DeXuat.ngay_tao.desc()).limit(10).all()

    return render_template('dashboard/admin_dashboard.html',
                           total_units=total_units,
                           total_personnel=total_personnel,
                           total_nominations=total_nominations,
                           awaiting_final=awaiting_final,
                           final_approved=final_approved,
                           recent_nominations=recent_nominations)
