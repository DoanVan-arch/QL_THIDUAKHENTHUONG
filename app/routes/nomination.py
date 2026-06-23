from pydoc import doc

from flask import Blueprint, render_template, redirect, url_for, flash, request, jsonify
from flask_login import login_required, current_user
from app.extensions import db
from app.models.user import User
from app.models.personnel import QuanNhan, DoiTuong
from app.models.nomination import DeXuat, DeXuatChiTiet, MinhChung, LoaiDanhHieu, TrangThaiDeXuat, DanhHieu, TieuChi
from app.models.evaluation import NhomTieuChi, DanhGiaHangNam
from app.models.evaluation import DiemQuyDinhDanhHieu
from app.models.approval import PheDuyet, PhongDuyet, KetQuaDuyet, KetQuaDuyetChiTiet
from app.models.reward import KhenThuong
from app.models.notification import ThongBao
from app.models.catalog import DoiTuongOption
from app.utils.decorators import unit_user_required
from app.utils.file_upload import save_upload
from app.utils.activity_logger import log_action
from datetime import datetime
from sqlalchemy.exc import ProgrammingError, OperationalError
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
from flask import send_file
from datetime import date

import hashlib, binascii, os
# The six reviewing departments (excluding admin)
DEPT_NAMES = [
    'Phòng Khoa học', 'Phòng Đào tạo',
    'Thủ trưởng Phòng TM-HC',
    'Ban Cán bộ', 'Ban Tổ chức', 'Ban Tuyên huấn', 'Ban Công tác quần chúng',
    'Ban Công nghệ thông tin', 'Ban Tác huấn', 'Ban Khảo thí', 'Ban Bảo vệ an ninh',
    'Ủy ban Kiểm tra', 'Ban Quân lực', 'Phòng Hậu cần - Kỹ thuật', 'Ban Sau đại học',
]

nomination_bp = Blueprint('nomination', __name__)


def _get_nam_hoc_options():
    """Return a list of academic year strings centred around the current school year.

    School year N–(N+1) starts in September of year N.
    In April 2026 the current school year is 2025-2026 (started Sep 2025).
    We return: (current-1), current, (current+1), (current+2).
    """
    now = datetime.now()
    year = now.year
    # If before August, the school year that started last September is "current"
    start = (year - 1) if now.month < 8 else year
    return [f'{start + i}-{start + i + 1}' for i in range(-1, 3)]


def _get_evidence_required_fields():
    rows = TieuChi.query.filter_by(is_active=True, co_minh_chung=True).order_by(TieuChi.thu_tu, TieuChi.ten).all()
    return [{'ma_truong': tc.ma_truong, 'ten': tc.ten, 'nhom': tc.nhom} for tc in rows]


def _get_doi_tuong_list():
    db_values = [x.ten for x in DoiTuongOption.query.filter_by(is_active=True).order_by(DoiTuongOption.thu_tu, DoiTuongOption.ten).all()]
    return db_values if db_values else [e.value for e in DoiTuong]


def _evidence_input_names_for_field(field_key):
    mapping = {
        'ket_qua_doan_the': ['minh_chung_doan_the', 'minh_chung_ket_qua_doan_the'],
        'thanh_tich_ca_nhan_khac': ['minh_chung_thanh_tich_khac', 'minh_chung_thanh_tich_ca_nhan_khac'],
        'nckh_noi_dung': ['nckh_minh_chung'],
    }
    return mapping.get(field_key, [f'minh_chung_{field_key}'])


def _get_tieu_chi_tap_the_by_danh_hieu(danh_hieu_db):
    """Return a dict: {ten_danh_hieu: {nhom: [tc_dicts]}} for tap_the danh hieus.

    Field names stored in DanhHieu.tieu_chi that have no corresponding active
    TieuChi record are included as fallback entries (using the field name as the
    label) so the nomination form can still render inputs for them.
    """
    result = {}
    for dh in danh_hieu_db:
        if (dh.pham_vi or 'Cá nhân') != 'Đơn vị' or not dh.tieu_chi:
            continue
        tcs = TieuChi.query.filter(
            TieuChi.ma_truong.in_(dh.tieu_chi),
            TieuChi.is_active == True
        ).order_by(TieuChi.thu_tu, TieuChi.ten).all()

        found_ma = {tc.ma_truong for tc in tcs}
        by_nhom = {}
        for tc in tcs:
            nhom = (tc.nhom or '').strip() or 'khac'
            by_nhom.setdefault(nhom, []).append({
                'ma_truong': tc.ma_truong,
                'ten': tc.ten,
                'loai_input': tc.loai_input or 'textbox',
                'gia_tri_chon': tc.gia_tri_chon or [],
                'huong_dan': tc.huong_dan or '',
            })

        # Orphan fields: in dh.tieu_chi but no active TieuChi record found.
        # Add them as plain text inputs so the nomination form can still capture
        # their values (label = field name until admin adds a proper TieuChi entry).
        for ma in dh.tieu_chi:
            if ma not in found_ma:
                by_nhom.setdefault('khac', []).append({
                    'ma_truong': ma,
                    'ten': ma,
                    'loai_input': 'textbox',
                    'gia_tri_chon': [],
                    'huong_dan': '',
                })

        if by_nhom:
            result[dh.ten_danh_hieu] = by_nhom
    return result


@nomination_bp.route('/history')
@login_required
@unit_user_required
def nomination_history():
    """History of submitted nominations for the unit."""
    if not current_user.don_vi:
        flash('Tài khoản chưa được gán đơn vị.', 'warning')
        return redirect(url_for('dashboard.index'))

    page = request.args.get('page', 1, type=int)
    status_filter = request.args.get('status', '')
    nam_hoc_filter = request.args.get('nam_hoc', '')

    query = DeXuat.query.filter(
        DeXuat.don_vi_id == current_user.don_vi_id,
        DeXuat.trang_thai != TrangThaiDeXuat.NHAP.value,
    )

    if status_filter:
        query = query.filter(DeXuat.trang_thai == status_filter)
    if nam_hoc_filter:
        query = query.filter(DeXuat.nam_hoc == nam_hoc_filter)

    nominations = query.order_by(DeXuat.ngay_gui.desc()).paginate(
        page=page, per_page=10, error_out=False
    )

    # Get distinct nam_hoc values for filter
    nam_hoc_list = [row[0] for row in db.session.query(DeXuat.nam_hoc).filter(
        DeXuat.don_vi_id == current_user.don_vi_id,
        DeXuat.trang_thai != TrangThaiDeXuat.NHAP.value,
    ).distinct().order_by(DeXuat.nam_hoc.desc()).all()]

    # Count notifications (will be set up with ThongBao model)
    unread_count = 0
    try:
        from app.models.notification import ThongBao
        unread_count = ThongBao.query.filter_by(
            user_id=current_user.id, da_doc=False
        ).count()
    except Exception:
        pass

    return render_template('nomination/history.html',
                           nominations=nominations,
                           status_filter=status_filter,
                           nam_hoc_filter=nam_hoc_filter,
                           nam_hoc_list=nam_hoc_list,
                           unread_count=unread_count)


@nomination_bp.route('/')
@login_required
@unit_user_required
def list_nominations():
    if not current_user.don_vi:
        flash('Tài khoản chưa được gán đơn vị.', 'warning')
        return redirect(url_for('dashboard.index'))

    page = request.args.get('page', 1, type=int)
    nominations = DeXuat.query.filter_by(don_vi_id=current_user.don_vi_id)\
        .order_by(DeXuat.ngay_tao.desc())\
        .paginate(page=page, per_page=10, error_out=False)

    return render_template('nomination/list.html', nominations=nominations)


@nomination_bp.route('/<int:id>/delete', methods=['POST'])
@login_required
@unit_user_required
def delete_nomination(id):
    de_xuat = DeXuat.query.get_or_404(id)

    if de_xuat.don_vi_id != current_user.don_vi_id:
        flash('Không có quyền thao tác.', 'danger')
        return redirect(url_for('nomination.list_nominations'))

    if de_xuat.trang_thai not in (TrangThaiDeXuat.NHAP.value, TrangThaiDeXuat.TU_CHOI.value):
        flash('Chỉ được xóa đề xuất ở trạng thái Nháp hoặc Từ chối.', 'warning')
        return redirect(url_for('nomination.list_nominations'))

    # Xóa các bản ghi liên quan trước
    chi_tiet_ids = [ct.id for ct in de_xuat.chi_tiets]
    if chi_tiet_ids:
        # Xóa ThongBao liên quan đến chi_tiet
        ThongBao.query.filter(ThongBao.chi_tiet_id.in_(chi_tiet_ids)).delete(synchronize_session=False)
        
        # Xóa YeuCauChinhSua liên quan đến chi_tiet
        from app.models.edit_request import YeuCauChinhSua
        YeuCauChinhSua.query.filter(YeuCauChinhSua.chi_tiet_id.in_(chi_tiet_ids)).delete(synchronize_session=False)
    
    # Xóa ThongBao liên quan trực tiếp đến de_xuat
    ThongBao.query.filter_by(de_xuat_id=de_xuat.id).delete(synchronize_session=False)

    db.session.delete(de_xuat)
    db.session.commit()
    flash('Đã xóa đề xuất.', 'success')
    return redirect(url_for('nomination.list_nominations'))


@nomination_bp.route('/create', methods=['GET', 'POST'])
@login_required
@unit_user_required
def create_nomination():
    if not current_user.don_vi:
        flash('Tài khoản chưa được gán đơn vị.', 'warning')
        return redirect(url_for('dashboard.index'))

    if request.method == 'POST':
        nam_hoc = request.form.get('nam_hoc', '').strip()
        if not nam_hoc:
            flash('Năm học không được để trống.', 'danger')
            return redirect(url_for('nomination.create_nomination'))

        de_xuat = DeXuat(
            don_vi_id=current_user.don_vi_id,
            nam_hoc=nam_hoc,
            nguoi_tao_id=current_user.id,
            trang_thai=TrangThaiDeXuat.NHAP.value,
            ghi_chu=request.form.get('ghi_chu', '').strip() or None,
        )
        db.session.add(de_xuat)
        db.session.commit()
        flash('Đã tạo đề xuất mới. Vui lòng thêm cá nhân vào đề xuất.', 'success')
        return redirect(url_for('nomination.edit_nomination', id=de_xuat.id))

    return render_template('nomination/create.html', nam_hoc_options=_get_nam_hoc_options())


@nomination_bp.route('/<int:id>')
@login_required
def detail_nomination(id):
    de_xuat = DeXuat.query.get_or_404(id)

    # Unit users can only see their own
    if current_user.is_unit_user and de_xuat.don_vi_id != current_user.don_vi_id:
        flash('Không có quyền truy cập.', 'danger')
        return redirect(url_for('dashboard.index'))

    # Build per-individual department approval results for tracking
    # Structure: {ct.id: {dept_name: ket_qua_string}}
    dept_lookup = {}
    for pd in de_xuat.phe_duyets:
        if pd.phong_duyet in DEPT_NAMES:
            dept_lookup[pd.phong_duyet] = {
                'phe_duyet': pd,
                'items': {kq.chi_tiet_id: kq for kq in pd.chi_tiet_duyet},
            }

    chi_tiets_data = []
    # Get all chi_tiet_ids that already have KhenThuong records
    approved_ct_ids = set(
        row[0] for row in db.session.query(KhenThuong.chi_tiet_id).filter(
            KhenThuong.de_xuat_id == de_xuat.id
        ).all()
    )

    for ct in de_xuat.chi_tiets:
        ct_dept_results = {}
        for dept_name in DEPT_NAMES:
            dept_data = dept_lookup.get(dept_name)
            if dept_data:
                kq = dept_data['items'].get(ct.id)
                ct_dept_results[dept_name] = kq.ket_qua if kq else None
            else:
                ct_dept_results[dept_name] = None

        # Parse tap_the_data for collective nominations
        tap_the_parsed = {}
        if ct.tap_the_data:
            try:
                tap_the_parsed = _json.loads(ct.tap_the_data)
            except Exception:
                tap_the_parsed = {}

        chi_tiets_data.append({
            'ct': ct,
            'dept_results': ct_dept_results,
            'is_rewarded': ct.id in approved_ct_ids,
            'is_removed': ct.bi_loai,
            'ly_do_loai': ct.ly_do_loai,
            'phong_loai': ct.phong_loai,
            'tap_the_data': tap_the_parsed,
        })

    return render_template('nomination/detail.html',
                           de_xuat=de_xuat,
                           chi_tiets_data=chi_tiets_data,
                           dept_names=DEPT_NAMES,
                           tieu_chi_map={tc.ma_truong: tc for tc in TieuChi.query.filter_by(is_active=True).all()})


@nomination_bp.route('/<int:id>/edit', methods=['GET', 'POST'])
@login_required
@unit_user_required
def edit_nomination(id):
    de_xuat = DeXuat.query.get_or_404(id)
    if de_xuat.don_vi_id != current_user.don_vi_id:
        flash('Không có quyền truy cập.', 'danger')
        return redirect(url_for('nomination.list_nominations'))

    if not de_xuat.is_editable:
        flash('Đề xuất này không thể sửa (đã gửi duyệt).', 'warning')
        return redirect(url_for('nomination.detail_nomination', id=id))

    if request.method == 'POST':
        de_xuat.nam_hoc = request.form.get('nam_hoc', de_xuat.nam_hoc).strip()
        de_xuat.ghi_chu = request.form.get('ghi_chu', '').strip() or None
        db.session.commit()
        flash('Đã cập nhật thông tin đề xuất.', 'success')
        return redirect(url_for('nomination.edit_nomination', id=id))

    personnel = QuanNhan.query.filter_by(
        don_vi_id=current_user.don_vi_id, is_active=True
    ).order_by(QuanNhan.ho_ten).all()

    # Find personnel already nominated in this nam_hoc (in submitted nominations or current draft)
    # 1. Already in the current draft
    already_in_current = set(
        ct.quan_nhan_id for ct in de_xuat.chi_tiets if ct.quan_nhan_id
    )
    # 2. Already in other submitted nominations for same nam_hoc
    already_in_other_q = db.session.query(DeXuatChiTiet.quan_nhan_id).join(
        DeXuat, DeXuatChiTiet.de_xuat_id == DeXuat.id
    ).filter(
        DeXuatChiTiet.nam_hoc == de_xuat.nam_hoc,
        DeXuatChiTiet.quan_nhan_id.isnot(None),
        DeXuat.id != de_xuat.id,
        DeXuat.trang_thai != TrangThaiDeXuat.NHAP.value,
    ).all()
    already_in_other = set(row[0] for row in already_in_other_q)

    already_nominated_ids = already_in_current | already_in_other

    danh_hieu_db = DanhHieu.query.filter_by(is_active=True).order_by(DanhHieu.thu_tu).all()
    danh_hieu_list = [dh.ten_danh_hieu for dh in danh_hieu_db]
    # Build a mapping of danh_hieu -> criteria fields for JS
    danh_hieu_tieu_chi = {dh.ten_danh_hieu: dh.tieu_chi for dh in danh_hieu_db}
    # Build a mapping of danh_hieu -> pham_vi ('Cá nhân' or 'Đơn vị') for JS
    danh_hieu_pham_vi = {dh.ten_danh_hieu: (dh.pham_vi or 'Cá nhân') for dh in danh_hieu_db}
    doi_tuong_list = _get_doi_tuong_list()
   # muc_do_list = [e.value for e in MucDoHoanThanh]

    # Build tooltips from TieuChi DB: {ma_truong: huong_dan}
    tieu_chi_db = TieuChi.query.filter_by(is_active=True).order_by(TieuChi.thu_tu, TieuChi.ten).all()
    tieu_chi_tooltips = {tc.ma_truong: tc.huong_dan for tc in tieu_chi_db if tc.huong_dan}
    # Full list for dynamic rendering in template (fields not already hardcoded in HTML)
    tieu_chi_list = [{'ma_truong': tc.ma_truong, 'ten': tc.ten, 'nhom': tc.nhom,
                      'loai_input': tc.loai_input or 'textbox',
                      'gia_tri_chon': tc.gia_tri_chon or []} for tc in tieu_chi_db]

    # Map for template rendering: field → {loai_input, gia_tri_chon}
    # Used to render hardcoded fields as combobox when configured so in admin
    tieu_chi_input_map = {
        tc.ma_truong: {
            'loai_input': tc.loai_input or 'textbox',
            'gia_tri_chon': tc.gia_tri_chon or [],
        }
        for tc in tieu_chi_db
    }
    # Inject hardcoded fallback combobox choices for fields not yet in TieuChi DB
    if 'xep_loai_dang_vien' not in tieu_chi_input_map or not tieu_chi_input_map['xep_loai_dang_vien'].get('gia_tri_chon'):
        tieu_chi_input_map['xep_loai_dang_vien'] = {
            'loai_input': 'combobox',
            'gia_tri_chon': DanhGiaHangNam.XEP_LOAI_DANG_VIEN_CHOICES,
        }
   
    if 'xep_loai_tong_ket' not in tieu_chi_input_map or not tieu_chi_input_map['xep_loai_tong_ket'].get('gia_tri_chon'):
        tieu_chi_input_map['xep_loai_tong_ket'] = {
            'loai_input': 'combobox',
            'gia_tri_chon': ['Xuất sắc', 'Giỏi', 'Khá', 'Trung bình khá', 'Trung bình', 'Yếu'],
        }

    diem_field_labels = {
        'diem_kiem_tra_tin_hoc': 'Điểm kỹ năng số',
        'diem_kiem_tra_dieu_lenh': 'Điểm điều lệnh',
        'diem_dia_ly_quan_su': 'Điểm địa hình quân sự',
        'diem_ban_sung': 'Điểm bắn súng',
        'diem_the_luc': 'Điểm thể lực',
        'diem_kiem_tra_chinh_tri': 'Điểm chính trị',
        'diem_tong_ket': 'Điểm tổng kết',
        'diem_nckh': 'Điểm NCKH',
    }

    # Map each score field (and its rating partner) to its nhom group code.
    # This lets fieldVisibility() check NhomTieuChi.doi_tuong_ap_dung for score fields.
    diem_nhom_map = {
        'diem_kiem_tra_tin_hoc':   'chung',
        'kiem_tra_tin_hoc':        'chung',
        'diem_kiem_tra_dieu_lenh': 'chung',
        'kiem_tra_dieu_lenh':      'chung',
        'diem_dia_ly_quan_su':     'chung',
        'dia_ly_quan_su':          'chung',
        'diem_ban_sung':           'chung',
        'ban_sung':                'chung',
        'diem_the_luc':            'chung',
        'the_luc':                 'chung',
        'diem_kiem_tra_chinh_tri': 'chung',
        'kiem_tra_chinh_tri':      'chung',
        'diem_tong_ket':           'hoc_vien',
        'xep_loai_tong_ket':       'hoc_vien',
        'diem_nckh':               'nckh',
    }

    score_rules = {}
    rows = DiemQuyDinhDanhHieu.query.filter_by(is_active=True).all()
    for r in rows:
        score_rules.setdefault(r.loai_danh_hieu, []).append({
            'tieu_chi_field': r.tieu_chi_field,
            'min_diem': r.min_diem,
        })

    # Build criteria/group metadata to keep unit UI aligned with admin group configuration
    nhom_meta = {}
    try:
        nhom_rows = NhomTieuChi.query.filter_by(is_active=True).order_by(NhomTieuChi.thu_tu, NhomTieuChi.ten_nhom).all()
        nhom_meta = {
            row.ma_nhom: {
                'ten_nhom': row.ten_nhom,
                'doi_tuong_ap_dung': row.doi_tuong_ap_dung or [],
            }
            for row in nhom_rows
        }
    except (ProgrammingError, OperationalError):
        db.session.rollback()

    if not nhom_meta:
        nhom_meta = {
            key: {'ten_nhom': label, 'doi_tuong_ap_dung': []}
            for key, label in TieuChi.NHOM_CHOICES.items()
        }

    criteria_meta = {}
    for tc in tieu_chi_db:
        nhom_info = nhom_meta.get(tc.nhom, {'ten_nhom': tc.nhom, 'doi_tuong_ap_dung': []})
        criteria_meta[tc.ma_truong] = {
            'nhom': tc.nhom,
            'nhom_ten': nhom_info.get('ten_nhom') or tc.nhom,
            'doi_tuong_ap_dung': nhom_info.get('doi_tuong_ap_dung') or [],
        }

    return render_template('nomination/edit.html',
                           de_xuat=de_xuat,
                           personnel=personnel,
                           already_nominated_ids=already_nominated_ids,
                           danh_hieu_list=danh_hieu_list,
                           danh_hieu_tieu_chi=danh_hieu_tieu_chi,
                           danh_hieu_pham_vi=danh_hieu_pham_vi,
                           doi_tuong_list=doi_tuong_list,
                       #    muc_do_list=muc_do_list,
                           tieu_chi_tooltips=tieu_chi_tooltips,
                           tieu_chi_list=tieu_chi_list,
                           tieu_chi_input_map=tieu_chi_input_map,
                           criteria_meta=criteria_meta,
                           nhom_meta=nhom_meta,
                           evidence_fields=_get_evidence_required_fields(),
                           score_rules=score_rules,
                           diem_field_labels=diem_field_labels,
                           diem_nhom_map=diem_nhom_map,
                           tieu_chi_tap_the=_get_tieu_chi_tap_the_by_danh_hieu(danh_hieu_db),
                           nam_hoc_options=_get_nam_hoc_options())


@nomination_bp.route('/<int:id>/add-item', methods=['POST'])
@login_required
@unit_user_required
def add_nomination_item(id):
    de_xuat = DeXuat.query.get_or_404(id)
    if de_xuat.don_vi_id != current_user.don_vi_id or not de_xuat.is_editable:
        flash('Không có quyền thao tác.', 'danger')
        return redirect(url_for('nomination.list_nominations'))

    quan_nhan_id = request.form.get('quan_nhan_id', type=int)
    loai_danh_hieu = request.form.get('loai_danh_hieu', '').strip()

    if not loai_danh_hieu:
        flash('Vui lòng chọn danh hiệu đề xuất.', 'danger')
        return redirect(url_for('nomination.edit_nomination', id=id))

    # Determine pham_vi for the selected danh_hieu
    dh_obj = DanhHieu.query.filter_by(ten_danh_hieu=loai_danh_hieu, is_active=True).first()
    is_tap_the = dh_obj and (dh_obj.pham_vi or 'Cá nhân') == 'Đơn vị'

    if is_tap_the:
        ten_don_vi_de_xuat = request.form.get('ten_don_vi_de_xuat', '').strip()
        if not ten_don_vi_de_xuat:
            flash('Vui lòng nhập tên đơn vị đề xuất.', 'danger')
            return redirect(url_for('nomination.edit_nomination', id=id))
    else:
        if not quan_nhan_id:
            flash('Vui lòng chọn quân nhân đề xuất.', 'danger')
            return redirect(url_for('nomination.edit_nomination', id=id))

    # --- Duplicate prevention ---
    if quan_nhan_id:
        # Check if this person is already in the current nomination
        existing_in_current = DeXuatChiTiet.query.filter_by(
            de_xuat_id=de_xuat.id, quan_nhan_id=quan_nhan_id
        ).first()
        if existing_in_current:
            qn_name = QuanNhan.query.get(quan_nhan_id)
            name = qn_name.ho_ten if qn_name else 'Cá nhân'
            flash(f'{name} đã có trong đề xuất này. Không thể thêm trùng.', 'danger')
            return redirect(url_for('nomination.edit_nomination', id=id))

        # Check if this person + nam_hoc already exists in another submitted nomination
        existing_other = db.session.query(DeXuatChiTiet).join(
            DeXuat, DeXuatChiTiet.de_xuat_id == DeXuat.id
        ).filter(
            DeXuatChiTiet.quan_nhan_id == quan_nhan_id,
            DeXuatChiTiet.nam_hoc == de_xuat.nam_hoc,
            DeXuat.id != de_xuat.id,
            DeXuat.trang_thai != TrangThaiDeXuat.NHAP.value,
        ).first()
        if existing_other:
            qn_name = QuanNhan.query.get(quan_nhan_id)
            name = qn_name.ho_ten if qn_name else 'Cá nhân'
            other_unit = existing_other.de_xuat.don_vi.ten_don_vi
            flash(
                f'{name} đã được đề xuất trong năm học {de_xuat.nam_hoc} '
                f'(bởi {other_unit}). Không thể đề xuất trùng.',
                'danger'
            )
            return redirect(url_for('nomination.edit_nomination', id=id))

    # Get personnel info and doi_tuong from form
    qn = QuanNhan.query.get(quan_nhan_id) if quan_nhan_id else None
    doi_tuong = request.form.get('doi_tuong', '').strip() or (qn.doi_tuong if qn else None)

    # Validate score/rating paired fields: must fill both sides together
    pair_rules = [
      
        ('diem_kiem_tra_dieu_lenh', 'kiem_tra_dieu_lenh', 'điều lệnh'),
        ('diem_dia_ly_quan_su', 'dia_ly_quan_su', 'địa hình quân sự'),
        ('diem_ban_sung', 'ban_sung', 'bắn súng'),
       
        ('diem_kiem_tra_chinh_tri', 'kiem_tra_chinh_tri', 'chính trị'),
    ]
    for diem_field, xeploai_field, label in pair_rules:
        diem_val = request.form.get(diem_field, '').strip()
        xeploai_val = request.form.get(xeploai_field, '').strip()
        if (diem_val and not xeploai_val) or (xeploai_val and not diem_val):
            flash(f'Tiêu chí {label}: phải nhập đầy đủ cả Điểm và Xếp loại.', 'danger')
            return redirect(url_for('nomination.edit_nomination', id=id))

    # Validate score threshold rules per award type
    rules = DiemQuyDinhDanhHieu.query.filter_by(loai_danh_hieu=loai_danh_hieu, is_active=True).all()
    below_threshold_fields = []
    for rule in rules:
        field_name = rule.tieu_chi_field
        raw_val = request.form.get(field_name, '').strip()
        if not raw_val:
            continue
        try:
            val = float(raw_val.replace(',', '.'))
            min_val = float(str(rule.min_diem).replace(',', '.'))
        except ValueError:
            continue

        if val < min_val:
            below_threshold_fields.append(f'{field_name} (< {rule.min_diem})')

    combined_score_reason = request.form.get('ly_do_chua_dat_diem', '').strip()
    if below_threshold_fields and not combined_score_reason:
        flash(
            'Có tiêu chí điểm chưa đạt mức quy định. Vui lòng nhập một lý do chung.',
            'danger'
        )
        return redirect(url_for('nomination.edit_nomination', id=id))

    # Evidence upload is optional: do not force validation here.

    # Collect tap_the criteria values if this is a collective nomination
    tap_the_data_dict = None
    if is_tap_the and dh_obj and dh_obj.tieu_chi:
        from app.models.nomination import TieuChi as _TC
        _tc_list = _TC.query.filter(_TC.ma_truong.in_(dh_obj.tieu_chi), _TC.is_active == True).all()
        tap_the_data_dict = {}
        for tc in _tc_list:
            val = request.form.get(tc.ma_truong, '').strip()
            if val:
                tap_the_data_dict[tc.ma_truong] = val

    import json as _json
    chi_tiet = DeXuatChiTiet(
        de_xuat_id=de_xuat.id,
        quan_nhan_id=quan_nhan_id,
        loai_danh_hieu=loai_danh_hieu,
        doi_tuong=doi_tuong,
        nam_hoc=de_xuat.nam_hoc,
        ten_don_vi_de_xuat=request.form.get('ten_don_vi_de_xuat', '').strip() or None if is_tap_the else None,
        tap_the_data=_json.dumps(tap_the_data_dict, ensure_ascii=False) if tap_the_data_dict else None,
        muc_do_hoan_thanh=request.form.get('muc_do_hoan_thanh', '').strip() or None,
        kiem_tra_tin_hoc=request.form.get('kiem_tra_tin_hoc', '').strip() or None,
        diem_kiem_tra_tin_hoc=request.form.get('diem_kiem_tra_tin_hoc', '').strip() or None,
        kiem_tra_dieu_lenh=request.form.get('kiem_tra_dieu_lenh', '').strip() or None,
        diem_kiem_tra_dieu_lenh=request.form.get('diem_kiem_tra_dieu_lenh', '').strip() or None,
        dia_ly_quan_su=request.form.get('dia_ly_quan_su', '').strip() or None,
        diem_dia_ly_quan_su=request.form.get('diem_dia_ly_quan_su', '').strip() or None,
        ban_sung=request.form.get('ban_sung', '').strip() or None,
        diem_ban_sung=request.form.get('diem_ban_sung', '').strip() or None,
        the_luc=request.form.get('the_luc', '').strip() or None,
        diem_the_luc=request.form.get('diem_the_luc', '').strip() or None,
        kiem_tra_chinh_tri=request.form.get('kiem_tra_chinh_tri', '').strip() or None,
        diem_kiem_tra_chinh_tri=request.form.get('diem_kiem_tra_chinh_tri', '').strip() or None,
        phieu_tin_nhiem=request.form.get('phieu_tin_nhiem', '').strip() or None,
        xep_loai_dang_vien=request.form.get('xep_loai_dang_vien', '').strip() or None,
        ket_qua_doan_the=request.form.get('ket_qua_doan_the', '').strip() or None,
        xep_loai_doan_vien=request.form.get('xep_loai_doan_vien', '').strip() or None,
        hinh_thuc_khen_thuong_qc=request.form.get('hinh_thuc_khen_thuong_qc', '').strip() or None,
        ket_qua_phu_nu=request.form.get('ket_qua_phu_nu', '').strip() or None,
        hinh_thuc_khen_thuong_pn=request.form.get('hinh_thuc_khen_thuong_pn', '').strip() or None,
        chu_tri_don_vi_danh_hieu=request.form.get('chu_tri_don_vi_danh_hieu', '').strip() or None,
        # Lecturer fields
        danh_hieu_gv_gioi=request.form.get('danh_hieu_gv_gioi', '').strip() or None,
        tien_do_pgs=request.form.get('tien_do_pgs', '').strip() or None,
        dinh_muc_giang_day=request.form.get('dinh_muc_giang_day', '').strip() or None,
        thoi_gian_lao_dong_kh=request.form.get('thoi_gian_lao_dong_kh', '').strip() or None,
        ket_qua_kiem_tra_giang=request.form.get('ket_qua_kiem_tra_giang', '').strip() or None,
        # Student fields
        danh_hieu_hv_gioi=request.form.get('danh_hieu_hv_gioi', '').strip() or None,
        diem_tong_ket=request.form.get('diem_tong_ket', '').strip() or None,
        ket_qua_thuc_hanh=request.form.get('ket_qua_thuc_hanh', '').strip() or None,
        ket_qua_ren_luyen=request.form.get('ket_qua_ren_luyen', '').strip() or None,
        # Graduation exam fields
        hinh_thuc_tot_nghiep=request.form.get('hinh_thuc_tot_nghiep', '').strip() or None,
        diem_tn_ctd=request.form.get('diem_tn_ctd', '').strip() or None,
        diem_tn_ct=request.form.get('diem_tn_ct', '').strip() or None,
        diem_tn_ta=request.form.get('diem_tn_ta', '').strip() or None,
        diem_tn_mon4=request.form.get('diem_tn_mon4', '').strip() or None,
        diem_tn_chuyennganh=request.form.get('diem_tn_chuyennganh', '').strip() or None,
        diem_tn_baove=request.form.get('diem_tn_baove', '').strip() or None,
        # NCKH
        nckh_noi_dung=(request.form.get('nckh_noi_dung_text', '').strip() or '; '.join([x.strip() for x in request.form.getlist('nckh_noi_dung') if x and x.strip()]) or None),
        diem_nckh=float(request.form.get('diem_nckh')) if request.form.get('diem_nckh', '').strip() else None,
        mo_ta_khoa_hoc=request.form.get('mo_ta_khoa_hoc', '').strip() or None,
        thanh_tich_ca_nhan_khac=request.form.get('thanh_tich_ca_nhan_khac', '').strip() or None,
        ghi_chu=(
            (request.form.get('ghi_chu_item', '').strip() or '') +
            (
                ' | Lý do chưa đạt điểm (' + ', '.join(below_threshold_fields) + '): ' + combined_score_reason
                if below_threshold_fields and combined_score_reason else ''
            )
        ).strip() or None,
    )

    db.session.add(chi_tiet)
    db.session.flush()

    # Handle file uploads for evidence
    evidence_fields = ['minh_chung_chung']
    for field_name in evidence_fields:
        file = request.files.get(field_name)
        if file and file.filename:
            path = save_upload(file, 'evidence')
            if path:
                mc = MinhChung(
                    chi_tiet_id=chi_tiet.id,
                    loai_minh_chung=field_name,
                    duong_dan=path,
                    ten_file_goc=file.filename,
                )
                db.session.add(mc)

    # NCKH evidence supports one or multiple files
    nckh_files = request.files.getlist('nckh_minh_chung')
    for file in nckh_files:
        if file and file.filename:
            path = save_upload(file, 'evidence')
            if path:
                mc = MinhChung(
                    chi_tiet_id=chi_tiet.id,
                    loai_minh_chung='nckh_minh_chung',
                    duong_dan=path,
                    ten_file_goc=file.filename,
                )
                db.session.add(mc)

    # Dynamic evidence files by configured criteria requiring evidence
    for ef in _get_evidence_required_fields():
        field_key = ef['ma_truong']
        # NCKH evidence is handled separately via nckh_minh_chung; skip here to avoid duplication.
        if field_key == 'nckh_noi_dung':
            continue
        files = request.files.getlist(f'minh_chung_{field_key}')
        for file in files:
            if file and file.filename:
                path = save_upload(file, 'evidence')
                if path:
                    mc = MinhChung(
                        chi_tiet_id=chi_tiet.id,
                        loai_minh_chung=f'minh_chung_{field_key}',
                        duong_dan=path,
                        ten_file_goc=file.filename,
                    )
                    db.session.add(mc)

    # Handle evidence files for ket_qua_doan_the
    doan_the_files = request.files.getlist('minh_chung_doan_the')
    for file in doan_the_files:
        if file and file.filename:
            path = save_upload(file, 'evidence')
            if path:
                mc = MinhChung(
                    chi_tiet_id=chi_tiet.id,
                    loai_minh_chung='minh_chung_doan_the',
                    duong_dan=path,
                    ten_file_goc=file.filename,
                )
                db.session.add(mc)

    # Handle evidence files for thanh_tich_ca_nhan_khac
    thanh_tich_files = request.files.getlist('minh_chung_thanh_tich_khac')
    for file in thanh_tich_files:
        if file and file.filename:
            path = save_upload(file, 'evidence')
            if path:
                mc = MinhChung(
                    chi_tiet_id=chi_tiet.id,
                    loai_minh_chung='minh_chung_thanh_tich_khac',
                    duong_dan=path,
                    ten_file_goc=file.filename,
                )
                db.session.add(mc)

    db.session.commit()

    name = qn.ho_ten if qn else 'Đơn vị'
    flash(f'Đã thêm đề xuất cho: {name} - {loai_danh_hieu}', 'success')
    return redirect(url_for('nomination.edit_nomination', id=id))


@nomination_bp.route('/item/<int:id>/delete', methods=['POST'])
@login_required
@unit_user_required
def delete_nomination_item(id):
    chi_tiet = DeXuatChiTiet.query.get_or_404(id)
    de_xuat = chi_tiet.de_xuat

    if de_xuat.don_vi_id != current_user.don_vi_id or not de_xuat.is_editable:
        flash('Không có quyền thao tác.', 'danger')
        return redirect(url_for('nomination.list_nominations'))

    ThongBao.query.filter_by(chi_tiet_id=chi_tiet.id).delete(synchronize_session=False)

    db.session.delete(chi_tiet)
    db.session.commit()
    flash('Đã xóa mục đề xuất.', 'success')
    return redirect(url_for('nomination.edit_nomination', id=de_xuat.id))


@nomination_bp.route('/item/<int:ct_id>/data', methods=['GET'])
@login_required
@unit_user_required
def get_nomination_item_data(ct_id):
    """Return JSON data for a chi_tiet (for edit modal population)."""
    import json as _json
    chi_tiet = DeXuatChiTiet.query.get_or_404(ct_id)
    de_xuat = chi_tiet.de_xuat
    if de_xuat.don_vi_id != current_user.don_vi_id or not de_xuat.is_editable:
        return jsonify({'error': 'Không có quyền'}), 403

    tap_the_data = {}
    if chi_tiet.tap_the_data:
        try:
            tap_the_data = _json.loads(chi_tiet.tap_the_data)
        except Exception:
            tap_the_data = {}

    data = {
        'ct_id': chi_tiet.id,
        'de_xuat_id': de_xuat.id,
        'quan_nhan_id': chi_tiet.quan_nhan_id,
        'ho_ten': (chi_tiet.quan_nhan.ho_ten if chi_tiet.quan_nhan else ''),
        'ten_don_vi_de_xuat': chi_tiet.ten_don_vi_de_xuat or '',
        'loai_danh_hieu': chi_tiet.loai_danh_hieu or '',
        'doi_tuong': chi_tiet.doi_tuong or '',
        'muc_do_hoan_thanh': chi_tiet.muc_do_hoan_thanh or '',
        'kiem_tra_tin_hoc': chi_tiet.kiem_tra_tin_hoc or '',
        'diem_kiem_tra_tin_hoc': chi_tiet.diem_kiem_tra_tin_hoc or '',
        'kiem_tra_dieu_lenh': chi_tiet.kiem_tra_dieu_lenh or '',
        'diem_kiem_tra_dieu_lenh': chi_tiet.diem_kiem_tra_dieu_lenh or '',
        'dia_ly_quan_su': chi_tiet.dia_ly_quan_su or '',
        'diem_dia_ly_quan_su': chi_tiet.diem_dia_ly_quan_su or '',
        'ban_sung': chi_tiet.ban_sung or '',
        'diem_ban_sung': chi_tiet.diem_ban_sung or '',
        'the_luc': chi_tiet.the_luc or '',
        'diem_the_luc': chi_tiet.diem_the_luc or '',
        'kiem_tra_chinh_tri': chi_tiet.kiem_tra_chinh_tri or '',
        'diem_kiem_tra_chinh_tri': chi_tiet.diem_kiem_tra_chinh_tri or '',
        'phieu_tin_nhiem': chi_tiet.phieu_tin_nhiem or '',
        'xep_loai_dang_vien': chi_tiet.xep_loai_dang_vien or '',
        'ket_qua_doan_the': chi_tiet.ket_qua_doan_the or '',
        'xep_loai_doan_vien': chi_tiet.xep_loai_doan_vien or '',
        'hinh_thuc_khen_thuong_qc': chi_tiet.hinh_thuc_khen_thuong_qc or '',
        'ket_qua_phu_nu': chi_tiet.ket_qua_phu_nu or '',
        'hinh_thuc_khen_thuong_pn': chi_tiet.hinh_thuc_khen_thuong_pn or '',
        'chu_tri_don_vi_danh_hieu': chi_tiet.chu_tri_don_vi_danh_hieu or '',
        'danh_hieu_gv_gioi': chi_tiet.danh_hieu_gv_gioi or '',
        'tien_do_pgs': chi_tiet.tien_do_pgs or '',
        'dinh_muc_giang_day': chi_tiet.dinh_muc_giang_day or '',
        'thoi_gian_lao_dong_kh': chi_tiet.thoi_gian_lao_dong_kh or '',
        'ket_qua_kiem_tra_giang': chi_tiet.ket_qua_kiem_tra_giang or '',
        'danh_hieu_hv_gioi': chi_tiet.danh_hieu_hv_gioi or '',
        'diem_tong_ket': chi_tiet.diem_tong_ket or '',
        'ket_qua_thuc_hanh': chi_tiet.ket_qua_thuc_hanh or '',
        'ket_qua_ren_luyen': chi_tiet.ket_qua_ren_luyen or '',
        'hinh_thuc_tot_nghiep': chi_tiet.hinh_thuc_tot_nghiep or '',
        'diem_tn_ctd': chi_tiet.diem_tn_ctd or '',
        'diem_tn_ct': chi_tiet.diem_tn_ct or '',
        'diem_tn_ta': chi_tiet.diem_tn_ta or '',
        'diem_tn_mon4': chi_tiet.diem_tn_mon4 or '',
        'diem_tn_chuyennganh': chi_tiet.diem_tn_chuyennganh or '',
        'diem_tn_baove': chi_tiet.diem_tn_baove or '',
        'nckh_noi_dung': chi_tiet.nckh_noi_dung or '',
        'diem_nckh': str(chi_tiet.diem_nckh) if chi_tiet.diem_nckh is not None else '',
        'mo_ta_khoa_hoc': chi_tiet.mo_ta_khoa_hoc or '',
        'thanh_tich_ca_nhan_khac': chi_tiet.thanh_tich_ca_nhan_khac or '',
        'ghi_chu': chi_tiet.ghi_chu or '',
        'tap_the_data': tap_the_data,
        # Personnel flags (for Bug 3: show/hide sections in edit modal)
        'la_hoi_vien_phu_nu': bool(chi_tiet.quan_nhan.la_hoi_vien_phu_nu) if chi_tiet.quan_nhan else False,
        'la_dang_vien': bool(chi_tiet.quan_nhan.la_dang_vien) if chi_tiet.quan_nhan else False,
        'la_doan_vien': bool(chi_tiet.quan_nhan.la_doan_vien) if chi_tiet.quan_nhan else False,
        'la_chi_huy': bool(chi_tiet.quan_nhan.la_chi_huy) if chi_tiet.quan_nhan else False,
        'la_bi_thu': bool(chi_tiet.quan_nhan.la_bi_thu) if chi_tiet.quan_nhan else False,
    }
    return jsonify(data)


@nomination_bp.route('/item/<int:ct_id>/update', methods=['POST'])
@login_required
@unit_user_required
def update_nomination_item(ct_id):
    """Update an existing chi_tiet (all criteria fields, keep quan_nhan/danh_hieu unchanged)."""
    import json as _json
    chi_tiet = DeXuatChiTiet.query.get_or_404(ct_id)
    de_xuat = chi_tiet.de_xuat

    if de_xuat.don_vi_id != current_user.don_vi_id or not de_xuat.is_editable:
        flash('Không có quyền thao tác.', 'danger')
        return redirect(url_for('nomination.list_nominations'))

    # Allow changing ten_don_vi_de_xuat if tap_the
    dh_obj = DanhHieu.query.filter_by(ten_danh_hieu=chi_tiet.loai_danh_hieu, is_active=True).first()
    is_tap_the = dh_obj and (dh_obj.pham_vi or 'Cá nhân') == 'Đơn vị'
    if is_tap_the:
        new_ten = request.form.get('ten_don_vi_de_xuat', '').strip()
        if new_ten:
            chi_tiet.ten_don_vi_de_xuat = new_ten
        # Update tap_the_data
        if dh_obj and dh_obj.tieu_chi:
            from app.models.nomination import TieuChi as _TC
            _tc_list = _TC.query.filter(_TC.ma_truong.in_(dh_obj.tieu_chi), _TC.is_active == True).all()
            tap_the_data_dict = {}
            for tc in _tc_list:
                # Form field name is 'tap_the_<ma_truong>' in edit modal
                field_name = f'tap_the_{tc.ma_truong}'
                val = request.form.get(field_name, '').strip()
                if val:
                    tap_the_data_dict[tc.ma_truong] = val
            chi_tiet.tap_the_data = _json.dumps(tap_the_data_dict, ensure_ascii=False) if tap_the_data_dict else None

    # Update all standard criteria fields
    simple_fields = [
        'muc_do_hoan_thanh', 'kiem_tra_tin_hoc', 'diem_kiem_tra_tin_hoc',
        'kiem_tra_dieu_lenh', 'diem_kiem_tra_dieu_lenh',
        'dia_ly_quan_su', 'diem_dia_ly_quan_su',
        'ban_sung', 'diem_ban_sung', 'the_luc', 'diem_the_luc',
        'kiem_tra_chinh_tri', 'diem_kiem_tra_chinh_tri',
        'phieu_tin_nhiem', 'xep_loai_dang_vien', 'ket_qua_doan_the',
        'xep_loai_doan_vien', 'hinh_thuc_khen_thuong_qc', 'ket_qua_phu_nu',
        'hinh_thuc_khen_thuong_pn', 'chu_tri_don_vi_danh_hieu',
        'danh_hieu_gv_gioi', 'tien_do_pgs', 'dinh_muc_giang_day',
        'thoi_gian_lao_dong_kh', 'ket_qua_kiem_tra_giang',
        'danh_hieu_hv_gioi', 'diem_tong_ket', 'ket_qua_thuc_hanh', 'ket_qua_ren_luyen',
        'hinh_thuc_tot_nghiep', 'diem_tn_ctd', 'diem_tn_ct', 'diem_tn_ta',
        'diem_tn_mon4', 'diem_tn_chuyennganh', 'diem_tn_baove',
        'thanh_tich_ca_nhan_khac',
    ]
    for field in simple_fields:
        val = request.form.get(field, '').strip()
        setattr(chi_tiet, field, val or None)

    # NCKH
    nckh_text = request.form.get('nckh_noi_dung_text', '').strip()
    if not nckh_text:
        nckh_list = [x.strip() for x in request.form.getlist('nckh_noi_dung') if x and x.strip()]
        nckh_text = '; '.join(nckh_list)
    chi_tiet.nckh_noi_dung = nckh_text or None

    diem_nckh_raw = request.form.get('diem_nckh', '').strip()
    chi_tiet.diem_nckh = float(diem_nckh_raw) if diem_nckh_raw else None
    chi_tiet.mo_ta_khoa_hoc = request.form.get('mo_ta_khoa_hoc', '').strip() or None

    # ghi_chu
    chi_tiet.ghi_chu = request.form.get('ghi_chu_item', '').strip() or None

    db.session.commit()

    name = chi_tiet.quan_nhan.ho_ten if chi_tiet.quan_nhan else (chi_tiet.ten_don_vi_de_xuat or 'Mục')
    flash(f'Đã cập nhật thông tin cho: {name}', 'success')
    return redirect(url_for('nomination.edit_nomination', id=de_xuat.id))


@nomination_bp.route('/<int:id>/submit', methods=['POST'])
@login_required
@unit_user_required
def submit_nomination(id):
    de_xuat = DeXuat.query.get_or_404(id)
    if de_xuat.don_vi_id != current_user.don_vi_id:
        flash('Không có quyền thao tác.', 'danger')
        return redirect(url_for('nomination.list_nominations'))

    if not de_xuat.is_editable:
        flash('Đề xuất này không thể gửi (đã gửi duyệt).', 'warning')
        return redirect(url_for('nomination.detail_nomination', id=id))

    if not de_xuat.chi_tiets:
        flash('Đề xuất phải có ít nhất một cá nhân/đơn vị.', 'danger')
        return redirect(url_for('nomination.edit_nomination', id=id))

    # Validate: commander/secretary must have DON_VI_QUYET_THANG
    # has_unit_award = any(
    #     ct.loai_danh_hieu == 'Đơn vị quyết thắng'
    #     for ct in de_xuat.chi_tiets
    # )
    # for ct in de_xuat.chi_tiets:
    #     if ct.quan_nhan and (ct.quan_nhan.la_chi_huy or ct.quan_nhan.la_bi_thu):
    #         if not has_unit_award:
    #             flash(
    #                 'Có Cấp trưởng/bí thư trong danh sách đề xuất - cần phải có đề xuất "Đơn vị quyết thắng" đi kèm.',
    #                 'danger'
    #             )
    #             return redirect(url_for('nomination.edit_nomination', id=id))

    # Create pending approval records
    ca_nhan_items = [ct for ct in de_xuat.chi_tiets if ct.doi_tuong]  # tập thể có doi_tuong = None
   # has_any_doan_the = any((ct.ket_qua_doan_the or '').strip() for ct in ca_nhan_items)
   # has_any_phu_nu = any((ct.ket_qua_phu_nu or '').strip() for ct in ca_nhan_items)
    # Chỉ auto-approve BAN_CTCQ khi có cá nhân nhưng không ai có kết quả đoàn thể
    # Nếu toàn bộ là tập thể → không auto-approve, để BAN_CTCQ xét bình thường
  #  ctcq_auto = ca_nhan_items and not has_any_doan_the and not has_any_phu_nu

    for phong in [PhongDuyet.PHONG_KHOAHOC, PhongDuyet.PHONG_DAOTAO,
                  PhongDuyet.THU_TRUONG_PHONG_TMHC,
                  PhongDuyet.BAN_CANBO, PhongDuyet.BAN_TOCHUC,
                  PhongDuyet.BAN_TUYENHUAN, PhongDuyet.BAN_CTCQ,
                  PhongDuyet.BAN_BAOVE_ANNINH,
                  PhongDuyet.BAN_CNTT, PhongDuyet.BAN_TAC_HUAN,
                  PhongDuyet.BAN_KHAOTHI, PhongDuyet.UY_BAN_KIEMTRA,
                  PhongDuyet.BAN_QUANLUC,PhongDuyet.PHONG_HAUCANKYTHUAT,PhongDuyet.BAN_SAUDAIHOC]:
        existing = PheDuyet.query.filter_by(de_xuat_id=de_xuat.id, phong_duyet=phong.value).first()
        if not existing:
            initial_ket_qua = KetQuaDuyet.CHO_DUYET.value
            initial_ghi_chu = None
            # if phong == PhongDuyet.BAN_CTCQ and ctcq_auto:
            #     initial_ket_qua = KetQuaDuyet.DONG_Y.value
            #     initial_ghi_chu = 'Tự động duyệt (không có dữ liệu kết quả đoàn thể)'

            pd = PheDuyet(
                de_xuat_id=de_xuat.id,
                phong_duyet=phong.value,
                ket_qua=initial_ket_qua,
                ghi_chu=initial_ghi_chu,
            )
            db.session.add(pd)
            db.session.flush()

            
            for ct in de_xuat.chi_tiets:
                if not ct.ket_qua_doan_the and not ct.ket_qua_doan_the and not ct.ket_qua_phu_nu and not ct.ket_qua_phu_nu:
                    existing = KetQuaDuyetChiTiet.query.filter_by(
                        phe_duyet_id=pd.id,
                        chi_tiet_id=ct.id,
                    ).first()
                    if existing:
                        existing.ket_qua = KetQuaDuyet.DONG_Y.value
                    else:
                        db.session.add(KetQuaDuyetChiTiet(
                            phe_duyet_id=pd.id,
                            chi_tiet_id=ct.id,
                            ket_qua=KetQuaDuyet.DONG_Y.value,
                        ))
                else:
                    existing = KetQuaDuyetChiTiet.query.filter_by(
                        phe_duyet_id=pd.id,
                        chi_tiet_id=ct.id,
                    ).first()
                    if not existing:
                        db.session.add(KetQuaDuyetChiTiet(
                            phe_duyet_id=pd.id,
                            chi_tiet_id=ct.id,
                            ket_qua=KetQuaDuyet.CHO_DUYET.value,
                        ))
                    else:
                        existing.ket_qua = KetQuaDuyet.CHO_DUYET.value
        else:
            # Reset existing approval to pending if re-submitting
            existing.ket_qua = KetQuaDuyet.CHO_DUYET.value
            existing.ly_do = None
            existing.ghi_chu = None
            db.session.flush()
    
    de_xuat.trang_thai = TrangThaiDeXuat.CHO_DUYET.value
    de_xuat.ngay_gui = datetime.utcnow()
    for ct in de_xuat.chi_tiets:
        ct.trang_thai = TrangThaiDeXuat.CHO_DUYET.value
        ct.bi_loai = False  # reset lại nếu có thay đổi sau khi bị loại trước đó
        ct.phong_loai = None
        ct.ly_do_loai = None
        ct.ngay_loai = None
    db.session.flush()

    # Sync per-item trang_thai → DANG_DUYET
    try:
        from app.routes.approval import _recompute_chi_tiet_status
        _recompute_chi_tiet_status(de_xuat)
    except Exception:
        pass

    db.session.commit()
    log_action('submit_nomination', resource_type='de_xuat', resource_id=de_xuat.id,
               detail=f'Năm học {de_xuat.nam_hoc}, {len(de_xuat.chi_tiets)} cá nhân/tập thể')
    db.session.commit()

    # Auto-finalize scope-limited depts (BAN_QUANLUC, BAN_CANBO) immediately
    # so TTPhongTMHC gate works without requiring those depts to visit their page first
    try:
        from app.routes.approval import _auto_finalize_scope_dept
        _auto_finalize_scope_dept(de_xuat.id)
    except Exception:
        pass  # non-critical, will be re-tried when TTPhongTMHC loads their page

    flash('Đã gửi đề xuất lên cấp trên. Chờ các cơ quan phê duyệt.', 'success')
    return redirect(url_for('nomination.detail_nomination', id=id))


@nomination_bp.route('/notifications')
@login_required
@unit_user_required
def notifications():
    """Notification list for unit accounts."""
    from app.models.notification import ThongBao

    page = request.args.get('page', 1, type=int)
    thong_baos = ThongBao.query.filter_by(
        user_id=current_user.id
    ).order_by(ThongBao.created_at.desc()).paginate(
        page=page, per_page=20, error_out=False
    )

    # Map chi_tiet_id -> open edit-request id (CHO_SUA) for this unit, so the
    # notification can show a "Chỉnh sửa ngay" action.
    from app.models.edit_request import YeuCauChinhSua, TrangThaiYeuCauSua
    ct_ids = [tb.chi_tiet_id for tb in thong_baos.items if tb.chi_tiet_id]
    edit_requests = {}
    if ct_ids:
        open_reqs = YeuCauChinhSua.query.filter(
            YeuCauChinhSua.chi_tiet_id.in_(ct_ids),
            YeuCauChinhSua.trang_thai == TrangThaiYeuCauSua.CHO_SUA.value,
        ).order_by(YeuCauChinhSua.created_at.desc()).all()
        for r in open_reqs:
            edit_requests.setdefault(r.chi_tiet_id, r.id)

    return render_template('nomination/notifications.html',
                           thong_baos=thong_baos, edit_requests=edit_requests)


@nomination_bp.route('/notifications/mark-read', methods=['POST'])
@login_required
@unit_user_required
def mark_notifications_read():
    """Mark all notifications as read."""
    from app.models.notification import ThongBao

    ThongBao.query.filter_by(
        user_id=current_user.id, da_doc=False
    ).update({'da_doc': True})
    db.session.commit()
    flash('Đã đánh dấu tất cả thông báo đã đọc.', 'success')
    return redirect(url_for('nomination.notifications'))


@nomination_bp.route('/edit-request/<int:request_id>', methods=['GET', 'POST'])
@login_required
@unit_user_required
def edit_request(request_id):
    """Restricted edit screen for the unit: only the criteria flagged by an approving
    department (via 'Yêu cầu chỉnh sửa') are editable. Everything else stays locked.
    On submit, only the requesting department must re-review."""
    import json as _json
    from app.models.edit_request import YeuCauChinhSua, TrangThaiYeuCauSua

    yc = YeuCauChinhSua.query.get_or_404(request_id)
    chi_tiet = yc.chi_tiet
    de_xuat = yc.de_xuat

    if de_xuat is None or chi_tiet is None or de_xuat.don_vi_id != current_user.don_vi_id:
        flash('Không có quyền thao tác.', 'danger')
        return redirect(url_for('nomination.notifications'))

    if chi_tiet.bi_loai:
        flash('Mục này đã bị loại khỏi đề xuất, không thể chỉnh sửa.', 'warning')
        return redirect(url_for('nomination.notifications'))

    if yc.trang_thai != TrangThaiYeuCauSua.CHO_SUA.value:
        flash('Yêu cầu chỉnh sửa này đã được xử lý.', 'info')
        return redirect(url_for('nomination.notifications'))

    flagged = yc.cac_truong or []
    is_tap_the = chi_tiet.quan_nhan_id is None

    # Field metadata for rendering inputs
    tc_rows = TieuChi.query.filter(
        TieuChi.ma_truong.in_(flagged), TieuChi.is_active == True
    ).all()
    tc_map = {tc.ma_truong: tc for tc in tc_rows}

    if request.method == 'POST':
        tap_the_data = {}
        if chi_tiet.tap_the_data:
            try:
                tap_the_data = _json.loads(chi_tiet.tap_the_data)
            except Exception:
                tap_the_data = {}

        for field in flagged:
            val = request.form.get(field, '').strip()
            if is_tap_the:
                if val:
                    tap_the_data[field] = val
                else:
                    tap_the_data.pop(field, None)
            else:
                setattr(chi_tiet, field, val or None)

        if is_tap_the:
            chi_tiet.tap_the_data = _json.dumps(tap_the_data, ensure_ascii=False) if tap_the_data else None

        yc.trang_thai = TrangThaiYeuCauSua.DA_SUA.value
        yc.ngay_sua = datetime.utcnow()

        # Reset ONLY the requesting department's result for this item to CHO_DUYET,
        # so only that department re-reviews; all other departments keep their result.
        phe_duyet = PheDuyet.query.filter_by(
            de_xuat_id=de_xuat.id, phong_duyet=yc.phong_yeu_cau
        ).first()
        if phe_duyet:
            kq = KetQuaDuyetChiTiet.query.filter_by(
                phe_duyet_id=phe_duyet.id, chi_tiet_id=chi_tiet.id
            ).first()
            if kq:
                kq.ket_qua = KetQuaDuyet.CHO_DUYET.value
                kq.ly_do = None
            if phe_duyet.ket_qua == KetQuaDuyet.DONG_Y.value:
                phe_duyet.ket_qua = KetQuaDuyet.CHO_DUYET.value
                phe_duyet.ngay_duyet = None
            # Keep đề xuất in review state
            if de_xuat.trang_thai not in (TrangThaiDeXuat.DANG_DUYET.value,
                                          TrangThaiDeXuat.HOI_DONG.value):
                de_xuat.trang_thai = TrangThaiDeXuat.DANG_DUYET.value

        # Notify the requesting department's account(s)
        from app.routes.approval import _PHONG_TO_ROLE
        req_role = _PHONG_TO_ROLE.get(yc.phong_yeu_cau)
        if req_role:
            for u in User.query.filter_by(role=req_role).all():
                name = (chi_tiet.quan_nhan.ho_ten if chi_tiet.quan_nhan else
                        (chi_tiet.ten_don_vi_de_xuat or de_xuat.don_vi.ten_don_vi))
                db.session.add(ThongBao(
                    user_id=u.id,
                    de_xuat_id=de_xuat.id,
                    chi_tiet_id=chi_tiet.id,
                    loai='da_sua',
                    tieu_de=f'Đơn vị đã chỉnh sửa: {name}',
                    noi_dung=(f'{de_xuat.don_vi.ten_don_vi} đã chỉnh sửa các tiêu chí theo '
                              f'yêu cầu. Vui lòng duyệt lại.'),
                ))

        db.session.commit()
        flash('Đã chỉnh sửa và gửi lại cho ban yêu cầu.', 'success')
        return redirect(url_for('nomination.notifications'))

    # GET: build current values + render metadata for flagged fields only
    tap_the_data = {}
    if chi_tiet.tap_the_data:
        try:
            tap_the_data = _json.loads(chi_tiet.tap_the_data)
        except Exception:
            tap_the_data = {}

    fields = []
    for field in flagged:
        tc = tc_map.get(field)
        if is_tap_the:
            cur = tap_the_data.get(field, '')
        else:
            cur = getattr(chi_tiet, field, '') or ''
        fields.append({
            'ma_truong': field,
            'ten': tc.ten if tc else field,
            'huong_dan': (tc.huong_dan if tc else '') or '',
            'loai_input': (tc.loai_input if tc else 'textbox') or 'textbox',
            'gia_tri_chon': (tc.gia_tri_chon if tc else []) or [],
            'gia_tri': cur,
        })

    ten_hien_thi = (chi_tiet.quan_nhan.ho_ten if chi_tiet.quan_nhan else
                    (chi_tiet.ten_don_vi_de_xuat or de_xuat.don_vi.ten_don_vi))

    return render_template('nomination/edit_request.html',
                           yc=yc, de_xuat=de_xuat, chi_tiet=chi_tiet,
                           fields=fields, ten_hien_thi=ten_hien_thi)


@nomination_bp.route('/<int:id>/revoke', methods=['POST'])
@login_required
@unit_user_required
def revoke_nomination(id):
    """Revoke (thu hồi) a submitted nomination back to draft status.
    Only allowed if no department has approved/rejected yet (all still CHO_DUYET).
    """
    de_xuat = DeXuat.query.get_or_404(id)
    if de_xuat.don_vi_id != current_user.don_vi_id:
        flash('Không có quyền thao tác.', 'danger')
        return redirect(url_for('nomination.list_nominations'))

    # Only allow revoking submitted nominations (not drafts, not already final-approved)
    if de_xuat.trang_thai not in (
        TrangThaiDeXuat.CHO_DUYET.value,
        TrangThaiDeXuat.DANG_DUYET.value,
        TrangThaiDeXuat.TU_CHOI.value,
    ):
        flash('Không thể thu hồi đề xuất ở trạng thái này.', 'warning')
        return redirect(url_for('nomination.detail_nomination', id=id))

    # Check that no department has finalized their review (all still CHO_DUYET)
    dept_reviews = PheDuyet.query.filter_by(de_xuat_id=de_xuat.id).filter(
        PheDuyet.phong_duyet != 'Tuyên huấn'
    ).all()

    has_any_decided = any(
        pd.ket_qua != KetQuaDuyet.CHO_DUYET.value for pd in dept_reviews
    )
    if has_any_decided:
        flash('Không thể thu hồi - đã có cơ quan hoàn tất duyệt.', 'danger')
        return redirect(url_for('nomination.detail_nomination', id=id))

    # Delete all PheDuyet and their KetQuaDuyetChiTiet records
    for pd in dept_reviews:
        KetQuaDuyetChiTiet.query.filter_by(phe_duyet_id=pd.id).delete()
        db.session.delete(pd)

    # Reset nomination status
    de_xuat.trang_thai = TrangThaiDeXuat.NHAP.value
    de_xuat.ngay_gui = None

    # Reset per-item trang_thai → NHAP
    from app.models.nomination import TrangThaiChiTiet
    for ct in de_xuat.chi_tiets:
        ct.trang_thai = TrangThaiChiTiet.NHAP.value

    db.session.commit()
    flash('Đã thu hồi đề xuất. Đề xuất đã chuyển về trạng thái Nháp.', 'success')
    return redirect(url_for('nomination.detail_nomination', id=id))


# ---------------------------------------------------------------------------
# Export đề xuất ra file Word theo mẫu
# ---------------------------------------------------------------------------
@nomination_bp.route('/<int:id>/export-word')
@login_required
@unit_user_required
def export_nomination_word(id):
    
    # Expire session cache to get fresh data from DB
    db.session.expire_all()

    de_xuat = DeXuat.query.get_or_404(id)
    if de_xuat.don_vi_id != current_user.don_vi_id:
        flash('Không có quyền xuất đề xuất này.', 'danger')
        return redirect(url_for('nomination.list_nominations'))

    don_vi_ten = de_xuat.don_vi.ten_don_vi if de_xuat.don_vi else ''
    nam_hoc = de_xuat.nam_hoc or ''

    # Phân nhóm chi tiết
    ds_quyet_thang = []   # Đơn vị quyết thắng
    ds_tien_tien_dv = []  # Đơn vị tiên tiến
    ds_chien_si_tdcs = [] # Chiến sĩ thi đua cơ sở
    ds_chien_si_tt = []   # Chiến sĩ tiên tiến
    ds_khac = {}          # Danh hiệu khác → list

    for ct in de_xuat.chi_tiets:
        dh = (ct.loai_danh_hieu or '').strip()
        if dh == 'Đơn vị quyết thắng':
            ds_quyet_thang.append(ct)
        elif dh == 'Đơn vị tiên tiến':
            ds_tien_tien_dv.append(ct)
        elif dh == 'Chiến sĩ thi đua':
            ds_chien_si_tdcs.append(ct)
        elif dh == 'Chiến sĩ tiên tiến':
            ds_chien_si_tt.append(ct)
        else:
            ds_khac.setdefault(dh, []).append(ct)

    # ---- Helpers ----
    def set_font(run, bold=False, size=11, italic=False, color=None):
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        run.font.name = 'Times New Roman'
        if color:
            run.font.color.rgb = RGBColor(*color)

    def para_font(para, text, bold=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
        para.alignment = align
        run = para.add_run(text)
        set_font(run, bold=bold, size=size, italic=italic)
        return run

    def set_cell_border(cell, **kwargs):
        """Set borders on a table cell."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ('top', 'left', 'bottom', 'right'):
            val = kwargs.get(edge, 'single')
            sz = kwargs.get(edge + '_sz', 4)
            tag = OxmlElement(f'w:{edge}')
            tag.set(qn('w:val'), val)
            tag.set(qn('w:sz'), str(sz))
            tag.set(qn('w:space'), '0')
            tag.set(qn('w:color'), '000000')
            tcBorders.append(tag)
        tcPr.append(tcBorders)

    def cell_para(cell, text, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = align
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(text)
        set_font(run, bold=bold, size=size, italic=italic)

    def add_cell(cell, text, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.LEFT):
        cell_para(cell, text, bold=bold, size=size, align=align)

    def build_tom_tat(ct):
        """Tóm tắt thành tích từ các trường - trả về list để hiển thị mỗi tiêu chí 1 dòng."""
        parts = []
        qn = ct.quan_nhan
        
        # Mức độ hoàn thành
        if ct.muc_do_hoan_thanh:
            parts.append(ct.muc_do_hoan_thanh)
        if ct.diem_tong_ket:
            parts.append(f'Kết quả học tập: {ct.diem_tong_ket}')
        # Rèn luyện
        if ct.ket_qua_ren_luyen:
            parts.append(f'Rèn luyện: {ct.ket_qua_ren_luyen}')
        
        # Điểm tổng kết
        
        
        # # Xếp loại tổng kết (nếu có)
        # if hasattr(ct, 'xep_loai_tong_ket') and ct.xep_loai_tong_ket:
        #     parts.append(f'Xếp loại TK: {ct.xep_loai_tong_ket}')
        
        # Tốt nghiệp
        if ct.hinh_thuc_tot_nghiep:
            tn_parts = [f'TN: {ct.hinh_thuc_tot_nghiep}']
            if ct.diem_tn_ctd:
                tn_parts.append(f'CTĐ-CT: {ct.diem_tn_ctd}')
            if ct.diem_tn_ct:
                tn_parts.append(f'CT: {ct.diem_tn_ct}')
            if ct.diem_tn_ta:
                tn_parts.append(f'TA: {ct.diem_tn_ta}')
            if ct.diem_tn_mon4:
                tn_parts.append(f'Môn 4: {ct.diem_tn_mon4}')
            if ct.diem_tn_chuyennganh:
                tn_parts.append(f'Chuyên ngành: {ct.diem_tn_chuyennganh}')
            if ct.diem_tn_baove:
                tn_parts.append(f'Bảo vệ: {ct.diem_tn_baove}')
            parts.append(', '.join(tn_parts))
        
        # # Kiểm tra các môn
        # if ct.kiem_tra_tin_hoc and ct.diem_kiem_tra_tin_hoc:
        #     parts.append(f'Tin học: {ct.kiem_tra_tin_hoc} ({ct.diem_kiem_tra_tin_hoc})')
        # elif ct.diem_kiem_tra_tin_hoc:
        #     parts.append(f'Điểm tin học: {ct.diem_kiem_tra_tin_hoc}')
        
        # if ct.kiem_tra_dieu_lenh and ct.diem_kiem_tra_dieu_lenh:
        #     parts.append(f'Điều lệnh: {ct.kiem_tra_dieu_lenh} ({ct.diem_kiem_tra_dieu_lenh})')
        # elif ct.diem_kiem_tra_dieu_lenh:
        #     parts.append(f'Điểm điều lệnh: {ct.diem_kiem_tra_dieu_lenh}')
        
        # if ct.dia_ly_quan_su and ct.diem_dia_ly_quan_su:
        #     parts.append(f'Địa lý QS: {ct.dia_ly_quan_su} ({ct.diem_dia_ly_quan_su})')
        # elif ct.diem_dia_ly_quan_su:
        #     parts.append(f'Điểm địa lý QS: {ct.diem_dia_ly_quan_su}')
        
        # if ct.ban_sung and ct.diem_ban_sung:
        #     parts.append(f'Bắn súng: {ct.ban_sung} ({ct.diem_ban_sung})')
        # elif ct.diem_ban_sung:
        #     parts.append(f'Điểm bắn súng: {ct.diem_ban_sung}')
        
        # if ct.the_luc and ct.diem_the_luc:
        #     parts.append(f'Thể lực: {ct.the_luc} ({ct.diem_the_luc})')
        # elif ct.diem_the_luc:
        #     parts.append(f'Điểm thể lực: {ct.diem_the_luc}')
        
        # if ct.kiem_tra_chinh_tri and ct.diem_kiem_tra_chinh_tri:
        #     parts.append(f'Chính trị: {ct.kiem_tra_chinh_tri} ({ct.diem_kiem_tra_chinh_tri})')
        # elif ct.diem_kiem_tra_chinh_tri:
        #     parts.append(f'Điểm chính trị: {ct.diem_kiem_tra_chinh_tri}')
        
        # # Đảng viên
        # if ct.xep_loai_dang_vien:
        #     parts.append(f'Đảng viên: {ct.xep_loai_dang_vien}')
        # if ct.phieu_tin_nhiem:
        #     parts.append(f'Phiếu tín nhiệm: {ct.phieu_tin_nhiem}')
        
        # # Đoàn viên
        # if ct.xep_loai_doan_vien:
        #     parts.append(f'Đoàn viên: {ct.xep_loai_doan_vien}')
        # if ct.ket_qua_doan_the:
        #     parts.append(f'Đoàn thể: {ct.ket_qua_doan_the}')
        
        # # Phụ nữ
        # if ct.ket_qua_phu_nu:
        #     parts.append(f'Hội Phụ nữ: {ct.ket_qua_phu_nu}')
        
        # # Giảng viên
        # if ct.danh_hieu_gv_gioi:
        #     parts.append(f'GV giỏi: {ct.danh_hieu_gv_gioi}')
        # if ct.tien_do_pgs:
        #     parts.append(f'Tiến độ PGS: {ct.tien_do_pgs}')
        # if ct.dinh_muc_giang_day:
        #     parts.append(f'Định mức giảng dạy: {ct.dinh_muc_giang_day}')
        # if ct.ket_qua_kiem_tra_giang:
        #     parts.append(f'KT giảng: {ct.ket_qua_kiem_tra_giang}')
        
        # # Học viên
        # if ct.danh_hieu_hv_gioi:
        #     parts.append(f'HV giỏi: {ct.danh_hieu_hv_gioi}')
        # if ct.ket_qua_thuc_hanh:
        #     parts.append(f'Thực hành: {ct.ket_qua_thuc_hanh}')
        
        # NCKH
        if ct.mo_ta_khoa_hoc:
            parts.append(f'NCKH: {ct.mo_ta_khoa_hoc}')
        # elif ct.mo_ta_khoa_hoc:
        #     parts.append(f'NCKH: {ct.mo_ta_khoa_hoc}')
        # if ct.diem_nckh:
        #     parts.append(f'Điểm NCKH: {ct.diem_nckh}')
        
        # Khen thưởng
        # if ct.hinh_thuc_khen_thuong_qc:
        #     parts.append(f'Khen thưởng: {ct.hinh_thuc_khen_thuong_qc}')
        # if ct.hinh_thuc_khen_thuong_pn:
        #     parts.append(f'KT Phụ nữ: {ct.hinh_thuc_khen_thuong_pn}')
        
        # # Chủ trì đơn vị
        # if ct.chu_tri_don_vi_danh_hieu:
        #     parts.append(f'Chủ trì: {ct.chu_tri_don_vi_danh_hieu}')
        
        # Thành tích khác
        if ct.thanh_tich_ca_nhan_khac:
            parts.append(ct.thanh_tich_ca_nhan_khac)
        
        return parts  # Return list instead of string

    def get_don_vi_truc_thuoc(ct):
        """Lấy đơn vị trực thuộc (Đại đội / Tiểu đoàn)."""
        if ct.quan_nhan and ct.quan_nhan.don_vi_truc_thuoc:
            return ct.quan_nhan.don_vi_truc_thuoc
        return ''

    def add_personnel_table(doc, chi_tiets, section_label, stt_start=1):
        """Thêm bảng danh sách cá nhân (Chiến sĩ thi đua / tiên tiến)."""
        # Section header
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        para_font(p, section_label, bold=True, size=11)

        if not chi_tiets:
            p2 = doc.add_paragraph()
            para_font(p2, '(Không có)', size=10, italic=True)
            return stt_start

        # Table: STT | Họ tên | Cấp bậc | Chức vụ | Đơn vị | Tóm tắt thành tích | Ghi chú
        tbl = doc.add_table(rows=1, cols=7)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.style = 'Table Grid'
        

        # Tổng độ rộng = 16.0 cm
        widths = [0.8, 3.2, 1.5, 2.0, 2.0, 5.0, 1.5]
        set_fixed_table_widths(tbl, widths)
        for i, w in enumerate(widths):
            for row in tbl.rows:
                row.cells[i].width = Cm(w)

        # Header row
        headers_txt = ['STT', 'Họ và tên', 'Cấp bậc', 'Chức vụ', 'Đơn vị', 'Tóm tắt thành tích', 'Ghi chú']
        hrow = tbl.rows[0]
        for i, h in enumerate(headers_txt):
            add_cell(hrow.cells[i], h, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

        stt = stt_start
        for ct in chi_tiets:
            qn_obj = ct.quan_nhan
            row = tbl.add_row()
            add_cell(row.cells[0], str(stt), align=WD_ALIGN_PARAGRAPH.CENTER)
            add_cell(row.cells[1], qn_obj.ho_ten if qn_obj else '')
            add_cell(row.cells[2], qn_obj.cap_bac if qn_obj else '')
            add_cell(row.cells[3], qn_obj.chuc_vu if qn_obj and qn_obj.chuc_vu else '')
            add_cell(row.cells[4], don_vi_ten if not qn_obj else get_don_vi_truc_thuoc(ct) or '')
            
            # Build tom tat - mỗi tiêu chí 1 dòng
            tom_tat_list = build_tom_tat(ct)
            cell = row.cells[5]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            if tom_tat_list:
                for i, item in enumerate(tom_tat_list):
                    if i > 0:
                        p = cell.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        p.paragraph_format.space_before = Pt(1)
                        p.paragraph_format.space_after = Pt(1)
                    run = p.add_run(f'- {item}')
                    set_font(run, size=10)
            else:
                run = p.add_run('-')
                set_font(run, size=10)
            
            add_cell(row.cells[6], ct.ghi_chu or '')
            stt += 1

        return stt

    def add_unit_table(doc, chi_tiets, section_label):
        """Thêm bảng danh sách tập thể (Đơn vị quyết thắng / Đơn vị tiên tiến)."""
        # Section header
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        para_font(p, section_label, bold=True, size=11)

        if not chi_tiets:
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Cm(1)
            para_font(p2, '(Không có)', size=10, italic=True)
            return

        # Table: STT | Tên đơn vị | Ghi chú (DS tiêu chí)
        tbl = doc.add_table(rows=1, cols=3)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.style = 'Table Grid'
       
       
        # Tổng độ rộng = 16.0 cm
        widths = [0.8, 4.2, 11.0]  # STT | Tên đơn vị | Ghi chú (tiêu chí)
        set_fixed_table_widths(tbl, widths)
        for i, w in enumerate(widths):
            for row in tbl.rows:
                row.cells[i].width = Cm(w)

        # Header row
        headers_txt = ['STT', 'Tên đơn vị', 'Ghi chú']
        hrow = tbl.rows[0]
        for i, h in enumerate(headers_txt):
            add_cell(hrow.cells[i], h, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

        # Data rows
        for idx, ct in enumerate(chi_tiets, 1):
            row = tbl.add_row()
            
            # STT
            add_cell(row.cells[0], str(idx), align=WD_ALIGN_PARAGRAPH.CENTER)
            
            # Tên đơn vị
            ten_dv = ct.ten_don_vi_de_xuat or '-'
            add_cell(row.cells[1], ten_dv)
            
            # Ghi chú - DS tiêu chí (mỗi tiêu chí 1 dòng)
            cell = row.cells[2]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            
            # Build criteria list from tap_the_dict
            criteria_list = []
            td = ct.tap_the_dict or {}
            
            # Get criteria labels from database
            if td:
                # Get TieuChi records to map ma_truong -> ten
                from app.models.nomination import TieuChi as _TieuChi
                ma_truong_list = list(td.keys())
                tieu_chi_map = {}
                if ma_truong_list:
                    tc_rows = _TieuChi.query.filter(_TieuChi.ma_truong.in_(ma_truong_list)).all()
                    tieu_chi_map = {tc.ma_truong: tc.ten for tc in tc_rows}
                
                for key, val in td.items():
                    if val and str(val).strip() and str(val).strip() not in ('0', 'None', ''):
                        # Use friendly name if available, otherwise use ma_truong
                        label = tieu_chi_map.get(key, key)
                        criteria_list.append(f'{label}: {val}')
            
            # Also add other text fields if available
            if ct.muc_do_hoan_thanh:
                criteria_list.insert(0, f'Mức độ hoàn thành: {ct.muc_do_hoan_thanh}')
            if ct.ghi_chu and ct.ghi_chu.strip():
                criteria_list.append(f'Ghi chú: {ct.ghi_chu}')
            
            if criteria_list:
                for i, item in enumerate(criteria_list):
                    if i > 0:
                        p = cell.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        p.paragraph_format.space_before = Pt(1)
                        p.paragraph_format.space_after = Pt(1)
                    run = p.add_run(f'- {item}')
                    set_font(run, size=10)
            else:
                run = p.add_run('-')
                set_font(run, size=10)

    # ---- Build document ----
    doc = Document()

    # Page margins
    for section in doc.sections:
        
        
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3.5)
        section.right_margin = Cm(1.5)
        
        # --- BẬT TÍNH NĂNG: Bỏ qua trang đầu tiên ---
        section.different_first_page_header_footer = True
        
        # --- Đánh số trang căn giữa ở Header (Sẽ tự động chạy từ trang 2) ---
        header = section.header
        
        # Lấy paragraph đầu tiên của header hoặc tạo mới nếu chưa có
        p_header = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Nếu tái sử dụng template, nên xóa nội dung cũ (nếu có) trước khi chèn
        p_header.clear() 
        
        run_header = p_header.add_run()
        add_page_number(run_header)

    # --- Header 2 cột: đơn vị bên trái, quốc hiệu bên phải ---
    tbl_header = doc.add_table(rows=1, cols=2)
    tbl_header.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Remove borders
    for cell in tbl_header.rows[0].cells:
        for edge in ('top','left','bottom','right'):
            tc = cell._tc; tcPr = tc.get_or_add_tcPr()
            b = OxmlElement('w:tcBorders')
            tag = OxmlElement(f'w:{edge}')
            tag.set(qn('w:val'), 'none')
            b.append(tag)
            tcPr.append(b)

    left_cell = tbl_header.rows[0].cells[0]
    right_cell = tbl_header.rows[0].cells[1]
    left_width = Cm(8)
    right_width = Cm(12)
    left_cell.width = left_width
    right_cell.width = right_width

    # Left: TRƯỜNG SĨ QUAN CHÍNH TRỊ / ĐƠN VỊ
    p_l1 = left_cell.paragraphs[0]
    p_l1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_l1.add_run('TRƯỜNG SĨ QUAN CHÍNH TRỊ')
    set_font(r, bold=False, size=13)

    p_l2 = left_cell.add_paragraph()
    p_l2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p_l2.add_run(don_vi_ten.upper())
    set_font(r2, bold=True, size=13)
    r2.underline = True

    # Right: CỘNG HÒA...
    p_r1 = right_cell.paragraphs[0]
    p_r1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_r1 = p_r1.add_run('CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM')
    set_font(r_r1, bold=True, size=13)

    p_r2 = right_cell.add_paragraph()
    p_r2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    r_r2 = p_r2.add_run('Độc lập - Tự do - Hạnh phúc')
    set_font(r_r2, bold=True, size=13)
    r_r2.underline = True
    p_r4 = right_cell.add_paragraph()
    r_r4 = p_r4.add_run('')
    set_font(r_r4, size=11, italic=True)
    p_r3 = right_cell.add_paragraph()
    p_r3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    today = date.today()
   # r_r3 = p_r3.add_run('')
    r_r3 = p_r3.add_run(f'Hà Nội, ngày {today.day} tháng {today.month} năm {today.year}')
    set_font(r_r3, size=11, italic=True)

    # ---------------------------------------------------------
    # THÊM MỚI: Thiết lập dãn dòng Single cho toàn bộ bảng Header
    # ---------------------------------------------------------
    for row in tbl_header.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    doc.add_paragraph()  # spacer

    # --- Tiêu đề ---
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run(f'DANH SÁCH ĐỀ NGHỊ KHEN THƯỞNG NĂM HỌC {nam_hoc}')
    set_font(r_title, bold=True, size=13)

    doc.add_paragraph()  # spacer

    # --- I. Đơn vị quyết thắng ---
    add_unit_table(doc, ds_quyet_thang, 'I. Danh hiệu Đơn vị quyết thắng')

    # --- II. Đơn vị tiên tiến ---
    add_unit_table(doc, ds_tien_tien_dv, 'II. Danh hiệu Đơn vị tiên tiến')

    # --- III. Chiến sĩ thi đua cơ sở ---
    stt = add_personnel_table(doc, ds_chien_si_tdcs, 'III. Danh hiệu Chiến sĩ thi đua cơ sở')

    # --- IV. Chiến sĩ tiên tiến ---
    add_personnel_table(doc, ds_chien_si_tt, 'IV. Danh hiệu Chiến sĩ tiên tiến')

    # --- Các danh hiệu khác ---
    roman = ['V', 'VI', 'VII', 'VIII', 'IX', 'X']
    for idx, (dh_name, chi_tiets) in enumerate(ds_khac.items()):
        label = f'{roman[idx] if idx < len(roman) else str(idx+5)}. {dh_name}'
        add_personnel_table(doc, chi_tiets, label)

    # --- Ngày giờ in ---
    now_str = datetime.now().strftime('%H:%M ngày %d tháng %m năm %Y')
    p_print = doc.add_paragraph()
    p_print.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_print = p_print.add_run(f'(In lúc {now_str})')
    run_print.font.size = Pt(9)
    run_print.font.italic = True
    run_print.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # --- Ký tên ---
    doc.add_paragraph()
    tbl_sign = doc.add_table(rows=1, cols=2)
    tbl_sign.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell in tbl_sign.rows[0].cells:
        for edge in ('top','left','bottom','right'):
            tc = cell._tc; tcPr = tc.get_or_add_tcPr()
            b = OxmlElement('w:tcBorders')
            tag = OxmlElement(f'w:{edge}')
            tag.set(qn('w:val'), 'none')
            b.append(tag)
            tcPr.append(b)

    left_sign = tbl_sign.rows[0].cells[0]
    right_sign = tbl_sign.rows[0].cells[1]

    # Ký xác nhận bên trái
    p_sl = left_sign.paragraphs[0]
    p_sl.alignment = WD_ALIGN_PARAGRAPH.CENTER
   # para_font(p_sl, 'XÁC NHẬN CỦA CẤP TRÊN', bold=True, size=11)
    left_sign.add_paragraph()
    left_sign.add_paragraph()
    left_sign.add_paragraph()

    # Ký đơn vị bên phải
    # Ký đơn vị bên phải
    p_sr = right_sign.paragraphs[0]
    # Sửa ở đây: Truyền thẳng tham số align vào para_font
    para_font(p_sr, 'THỦ TRƯỞNG ĐƠN VỊ', bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    
    p_sr2 = right_sign.add_paragraph()
    # Sửa ở đây: Truyền thẳng tham số align vào para_font
    para_font(p_sr2, '(Ký, ghi rõ họ tên)', size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    

    # --- Khóa tài liệu: chỉ cho sửa định dạng, không cho sửa nội dung ---
    protect_document_formatting_only(doc, 'bth123')
    
    # --- Thêm watermark ---
    # Chọn 1 trong 3 phương án bên dưới (uncomment để sử dụng):
    
    # Phương án 1: Text watermark xéo 45 độ (đơn giản, không che nội dung)
   # add_text_watermark(doc, "TRƯỜNG SĨ QUAN CHÍNH TRỊ")
    
    # Phương án 2: Logo nhỏ ở footer cuối trang + text
    #add_logo_footer(doc)
    
    # Phương án 3: Logo nhỏ ở góc phải header (có thể conflict với header hiện tại)
    add_corner_logo(doc)

    # --- Stream to response ---
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_name = don_vi_ten.replace(' ', '_').replace('/', '-')
    filename = f'DeXuat_KhenThuong_{safe_name}_{nam_hoc}.docx'.replace(' ', '_')

    return send_file(
        buf,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
def set_fixed_table_widths(tbl, widths_cm):
        """Can thiệp sâu vào XML để khóa chết chiều rộng bảng, Word không thể tự đổi"""
        # 1. Ép kiểu bảng thành Fixed Layout (Không tự co giãn)
        tbl.autofit = False
        tblPr = tbl._tbl.tblPr
        tblLayout = tblPr.find(qn('w:tblLayout'))
        if tblLayout is None:
            tblLayout = OxmlElement('w:tblLayout')
            tblPr.append(tblLayout)
        tblLayout.set(qn('w:type'), 'fixed')

        # 2. Xóa lưới cột cũ và xây lại khung lưới mới theo đúng kích thước cm
        tblGrid = tbl._tbl.find(qn('w:tblGrid'))
        if tblGrid is not None:
            tbl._tbl.remove(tblGrid)
        tblGrid = OxmlElement('w:tblGrid')
        tbl._tbl.insert(1, tblGrid)  # Chèn khung lưới vào đúng vị trí chuẩn XML

        for w in widths_cm:
            gridCol = OxmlElement('w:gridCol')
            # Chuyển đổi Cm sang đơn vị Twips của Word (1 twip = 635 EMUs)
            gridCol.set(qn('w:w'), str(int(Cm(w) / 635)))
            tblGrid.append(gridCol)

        # 3. Khóa cứng chiều rộng ở cấp độ từng Ô (Cell)
        for i, w in enumerate(widths_cm):
            twips_val = str(int(Cm(w) / 635))
            for row in tbl.rows:
                tcPr = row.cells[i]._tc.get_or_add_tcPr()
                tcW = tcPr.find(qn('w:tcW'))
                if tcW is None:
                    tcW = OxmlElement('w:tcW')
                    tcPr.append(tcW)
                tcW.set(qn('w:w'), twips_val)
                tcW.set(qn('w:type'), 'dxa')
def add_page_number(run):
        """Thêm mã trường đếm số trang tự động (PAGE) vào run"""
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')

        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)
def add_text_watermark(doc, text="TRƯỜNG SĨ QUAN CHÍNH TRỊ"):
    """Thêm text watermark xéo góc 45 độ, mờ, ở giữa trang."""
    try:
        for section in doc.sections:
            header = section.header
            # Tạo paragraph trong header
            para = header.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Thêm run với text
            run = para.add_run(text)
            run.font.size = Pt(48)
            run.font.name = 'Times New Roman'
            run.font.bold = True
            run.font.color.rgb = RGBColor(192, 192, 192)  # Light gray
            
            # Access paragraph XML to add rotation and positioning
            p_element = para._element
            pPr = p_element.get_or_add_pPr()
            
            # Create text box frame with rotation
            # Note: This is a simplified approach - true watermark needs more complex XML
            # For better watermark effect, we'll use the shape approach below
            
    except Exception as e:
        print(f"Warning: Could not add text watermark: {e}")


def add_logo_footer(doc):
    """Thêm logo nhỏ và số trang vào footer cuối trang."""
    import os
    from flask import current_app
    
    logo_path = os.path.join(current_app.root_path, 'static', 'img', 'watermark.png')
    
    if not os.path.exists(logo_path):
        # Fallback to main logo if watermark doesn't exist
        logo_path = os.path.join(current_app.root_path, 'static', 'img', 'logo-Si-quan.png')
        if not os.path.exists(logo_path):
            return
    
    try:
        for section in doc.sections:
            footer = section.footer
            
            # Thêm số trang ở giữa
            para_page = footer.add_paragraph()
            para_page.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para_page.paragraph_format.space_before = Pt(2)
            para_page.paragraph_format.space_after = Pt(2)
            
            # Add page number field
            run_page = para_page.add_run()
            run_page.font.size = Pt(10)
            run_page.font.name = 'Times New Roman'
            
            # Insert page number field code
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = 'PAGE'
            
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            
            run_page._element.append(fldChar1)
            run_page._element.append(instrText)
            run_page._element.append(fldChar2)
            
            # Thêm logo nhỏ bên dưới số trang
            para = footer.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Thêm logo nhỏ (1.2cm)
            run = para.add_run()
            run.add_picture(logo_path, width=Cm(1.2))
            
            # Thêm text bên dưới logo
            para2 = footer.add_paragraph()
            para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = para2.add_run('Trường Sĩ quan Chính trị')
            run2.font.size = Pt(7)
            run2.font.name = 'Times New Roman'
            run2.font.italic = True
            run2.font.color.rgb = RGBColor(128, 128, 128)
            
    except Exception as e:
        print(f"Warning: Could not add footer logo: {e}")


def add_corner_logo(doc):
    """Thêm logo nhỏ ở góc phải trên cùng của trang (sau header table hiện tại)."""
    import os
    from flask import current_app
    
    logo_path = os.path.join(current_app.root_path, 'static', 'img', 'watermark.png')
    
    if not os.path.exists(logo_path):
        # Fallback to main logo if watermark doesn't exist
        logo_path = os.path.join(current_app.root_path, 'static', 'img', 'logo-Si-quan.png')
        if not os.path.exists(logo_path):
            return
    
    try:
        for section in doc.sections:
            header = section.header
            
            # Thêm paragraph mới vào cuối header (sau table header hiện tại)
            para = header.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Set paragraph spacing để logo sát lề trên
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            
            # Thêm logo nhỏ căn phải (1.5cm)
            run = para.add_run()
            run.add_picture(logo_path, width=Cm(1.5))
            
    except Exception as e:
        print(f"Warning: Could not add corner logo: {e}")


def add_logo_watermark(doc):
    """Thêm watermark logo vào tất cả các trang - DEPRECATED, use specific functions above."""
    import os
    from flask import current_app
    
    # Đường dẫn logo
    logo_path = os.path.join(current_app.root_path, 'static', 'img', 'logo-Si-quan.png')
    
    if not os.path.exists(logo_path):
        return  # Skip if logo not found
    
    try:
        for section in doc.sections:
            # Thêm watermark vào header của section
            header = section.header
            
            # Tạo paragraph cho watermark - đặt ở cuối header
            para = header.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Thêm hình ảnh watermark với opacity thấp
            run = para.add_run()
            picture = run.add_picture(logo_path, width=Cm(6))
            
            # Note: Đơn giản hóa - chỉ thêm logo vào header
            # Không cần chuyển thành anchor phức tạp, vì có thể gây lỗi
            # Logo sẽ xuất hiện ở đầu mỗi trang trong header
            
    except Exception as e:
        # Nếu watermark fails, không làm crash toàn bộ export
        print(f"Warning: Could not add watermark: {e}")
        pass


def protect_document_formatting_only(doc, password: str):
    """
    Khóa tài liệu: chỉ đọc nội dung (readOnly).
    Mật khẩu được hash theo chuẩn Office 2010+ (Agile Encryption).
    """
    # 1. Tạo salt ngẫu nhiên (16 bytes)
    salt = os.urandom(16)
    salt_b64 = binascii.b2a_base64(salt).strip().decode()

    # 2. Hash lần đầu: SHA-512(salt + password)
    # Lưu ý: password bắt buộc encode sang chuẩn UTF-16 Little Endian
    key = hashlib.sha512(salt + password.encode('utf-16le')).digest()
    
    # 3. Lặp 100.000 vòng để chống brute-force
    spin_count = 100000
    for i in range(spin_count):
        iterator = i.to_bytes(4, byteorder='little')
        # SỬA LỖI: Cần cộng iterator ở PHÍA SAU hash của vòng lặp liền trước
        key = hashlib.sha512(key + iterator).digest()
        
    hash_b64 = binascii.b2a_base64(key).strip().decode()

    # 4. Lấy cấu hình settings của docx
    settings = doc.settings.element

    # Xóa thẻ documentProtection cũ nếu có
    for old in settings.findall(qn('w:documentProtection')):
        settings.remove(old)

    # 5. Tạo thẻ <w:documentProtection> theo chuẩn Office đời mới
    doc_prot = OxmlElement('w:documentProtection')
    doc_prot.set(qn('w:edit'),          'readOnly')
    doc_prot.set(qn('w:enforcement'),   '1')
    
    # BỎ CÁC THẺ CŨ (cryptProviderType, v.v.). SỬ DỤNG CHUẨN AGILE MỚI:
    doc_prot.set(qn('w:algorithmName'), 'SHA-512')
    doc_prot.set(qn('w:spinCount'),     str(spin_count))
    doc_prot.set(qn('w:hashValue'),     hash_b64)     # Đã đổi từ w:hash thành w:hashValue
    doc_prot.set(qn('w:saltValue'),     salt_b64)     # Đã đổi từ w:salt thành w:saltValue

    # Chèn vào đầu <w:settings>
    settings.insert(0, doc_prot)
